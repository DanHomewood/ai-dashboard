# ===== SECTION 1: Imports + Page Config =====
from __future__ import annotations

# --- Standard library ---
from pathlib import Path
import os
import re
import time
import uuid
import base64
import platform
from typing import Optional, Dict, Any
from datetime import datetime, timezone
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
import pdfplumber
from pypdf import PdfReader
from rapidfuzz import fuzz
from dateutil import parser as dateparser

# --- Timezone: UK (fallback to UTC if tz data missing) ---
try:
    from zoneinfo import ZoneInfo  # Python 3.9+
    UK_TZ = ZoneInfo("Europe/London")
except Exception:
    UK_TZ = timezone.utc  # consider: pip install tzdata on Windows

# --- Core third-party (lightweight) ---
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go

# --- Expenses master loader (shared) ---
from pathlib import Path
import pandas as pd
import streamlit as st

EXP_MASTER = Path("Expenses/expenses_master.parquet")

@st.cache_data(show_spinner=False)
def load_expense_master_df() -> pd.DataFrame:
    """Return the expenses master if present, else an empty DF with expected columns."""
    if EXP_MASTER.exists():
        try:
            df = pd.read_parquet(EXP_MASTER)
        except Exception:
            df = pd.DataFrame()
    else:
        df = pd.DataFrame()
    # normalise cols
    want = ["Engineer Name","Transaction Date","Amount","Expense Type","Business Purpose",
            "Vendor","City of Purchase","Payment Type","Source File","RowUID"]
    if df.empty:
        return pd.DataFrame(columns=want)

    # safe dtypes
    if "Transaction Date" in df.columns:
        df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
    return df

def _money_to_float(series: pd.Series) -> pd.Series:
    """
    Turn things like '£1,234.56', '(286.84)', '989.93', ' 426.10 ' into floats.
    Ignores stray integers like '2025'.
    """
    s = series.astype(str).str.strip()
    s = s.str.replace(r"\(([^)]+)\)", r"-\1", regex=True)        # (123.45) -> -123.45
    # keep ONE money-looking number per cell
    s = s.str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)", expand=False)
    s = s.str.replace(",", "", regex=False)
    return pd.to_numeric(s, errors="coerce")


# ====== EXPENSES AUTO-INGEST (Step 3A: folders + preview helpers) ======
from pathlib import Path
import pandas as pd
import pdfplumber
from pypdf import PdfReader

EXPENSES_ROOT = Path(__file__).parent / "Expenses"
INBOX_DIR     = EXPENSES_ROOT / "Inbox"
PROCESSED_DIR = EXPENSES_ROOT / "Processed"
REJECTED_DIR  = EXPENSES_ROOT / "Rejected"

def _ensure_expense_dirs():
    INBOX_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSED_DIR.mkdir(parents=True, exist_ok=True)
    REJECTED_DIR.mkdir(parents=True, exist_ok=True)

def list_pdf_inbox() -> list[Path]:
    _ensure_expense_dirs()
    return sorted([p for p in INBOX_DIR.glob("*.pdf")])

def read_pdf_text(path: Path) -> str:
    """
    Try pdfplumber first; if blank, fall back to PyPDF reader.
    This is just for preview in Step 3A.
    """
    # 1) pdfplumber
    try:
        chunks = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    chunks.append(t)
        if chunks:
            return "\n".join(chunks)
    except Exception:
        pass

    # 2) PyPDF
    try:
        reader = PdfReader(str(path))
        chunks = []
        for p in reader.pages:
            t = p.extract_text() or ""
            if t.strip():
                chunks.append(t)
        return "\n".join(chunks)
    except Exception:
        return ""
import re
from dateutil import parser as dateparser


def parse_expense_text(text: str, fallback_name: str = None) -> dict:
    """
    Try to extract Name, Date, Amount from raw expense PDF text.
    Returns a dict: {"Name":.., "Date":.., "Description":.., "Total Value (£)":..}
    """

    out = {"Name": None, "Date": None, "Description": None, "Total Value (£)": None}

    if not text:
        return out

    # --- Name (look for 'Employee Name' or fallback from filename)
    m = re.search(r"Employee Name\s*[:\-]?\s*([A-Za-z'\s]+)", text, flags=re.I)
    if m:
        out["Name"] = m.group(1).strip()
    elif fallback_name:
        out["Name"] = fallback_name

    # --- Date (try dd/mm/yyyy, dd-mm-yyyy, yyyy-mm-dd, or "August 2025")
    m = re.search(r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text)
    if not m:
        m = re.search(r"([A-Za-z]+\s+\d{4})", text)  # e.g. "August 2025"
    if m:
        try:
            out["Date"] = dateparser.parse(m.group(1), dayfirst=True).date()
        except Exception:
            pass

    # --- Amount (look for £ or numbers with decimal)
    m = re.search(r"£\s*([\d,]+(?:\.\d{1,2})?)", text)
    if not m:
        m = re.search(r"(\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)", text)
    if m:
        try:
            out["Total Value (£)"] = float(m.group(1).replace(",", ""))
        except Exception:
            pass

    # --- Description: take first line with "Report Name" or fallback
    m = re.search(r"Report Name\s*[:\-]?\s*(.+)", text, flags=re.I)
    if m:
        out["Description"] = m.group(1).split("\n")[0].strip()
    else:
        out["Description"] = "Expense PDF"

    return out

# NOTE: Heavy/optional deps — import these *inside* the functions that use them:
# import matplotlib.pyplot as plt
# import seaborn as sns
# from mpl_toolkits.mplot3d import Axes3D  # noqa: F401
# import PyPDF2
# import pdfplumber
# import requests
# from st_aggrid import AgGrid, GridOptionsBuilder
# from streamlit_lottie import st_lottie
# from openai import OpenAI, OpenAIError
# from statsmodels.tsa.arima.model import ARIMA
# from langchain_experimental.agents import create_pandas_dataframe_agent
# from langchain.agents.agent_types import AgentType
# from langchain.schema import HumanMessage, SystemMessage
# from langchain_community.chat_models import ChatOpenAI

import base64, os, mimetypes
from pathlib import Path
import base64
# === ORBIT AI: SETUP  =========================================================
# Requires: pip install langchain-openai langchain-experimental openai (if you use the fallback)
# Also requires: an OpenAI key in .streamlit/secrets.toml:
# [openai]
# api_key = "sk-..."

from langchain_openai import ChatOpenAI
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain.agents.agent_types import AgentType

import os, csv, time, datetime
import pandas as pd
import streamlit as st

FILES: Dict[str, str] = {
    "VIP North":  "VIP North Oracle Data.xlsx",
    "VIP South":  "VIP South Oracle Data.xlsx",
    "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
   
}
BUDGET_FILE = "budgets.csv"  # default budget file name
# === Budget loader (robust to filename / column variations) ===
from pathlib import Path
import pandas as pd
import re
import streamlit as st

def _find_first_existing(*names: str) -> str | None:
    for n in names:
        if Path(n).exists():
            return n
    return None

def _normalise_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = (
        df.columns.astype(str)
        .str.replace("\u00A0", " ", regex=False)  # non-breaking spaces
        .str.strip()
    )
    return df


def _find_col(df: pd.DataFrame, *keys: str) -> str | None:
    """Return the first column whose lower name matches any of keys; else fuzzy match."""
    if df.empty:
        return None
    low = {c.lower().strip(): c for c in df.columns}
    for k in keys:
        if k in low:
            return low[k]
    # fuzzy
    for c in df.columns:
        if re.search("|".join(map(re.escape, keys)), c, flags=re.I):
            return c
    return None
TEAM_COL_OVERRIDE = "Stakeholder"   # <-- your budgets.csv uses this

# === Invoices: loader helper (ADD in helpers section) ===
@st.cache_data(show_spinner=False)
def load_invoices(path: str | Path = "Invoices.xlsx") -> pd.DataFrame:
    """
    Robust invoices loader.
    - Reads all readable sheets, concatenates if columns match
    - Normalises headers
    - Coerces common columns: Date, Amount/Total, Supplier/Vendor, Status, Paid
    - Leaves original columns intact as well
    """
    p = Path(path)
    if not p.exists():
        st.warning("⚠️ Invoices.xlsx not found in app folder.")
        return pd.DataFrame()

    try:
        x = pd.ExcelFile(p)
        frames = []
        for s in x.sheet_names:
            try:
                df = x.parse(s)
                if isinstance(df, pd.DataFrame) and not df.empty:
                    frames.append(df)
            except Exception:
                continue
        if not frames:
            st.warning("⚠️ No readable sheets in Invoices.xlsx")
            return pd.DataFrame()

        def _norm(df: pd.DataFrame) -> pd.DataFrame:
            out = df.copy()
            out.columns = (
                out.columns.astype(str)
                .str.replace("\u00A0", " ", regex=False)  # non-breaking spaces
                .str.strip()
            )
            return out

        frames = [_norm(f) for f in frames]
        same_cols = all(set(frames[0].columns) == set(f.columns) for f in frames)
        inv = pd.concat(frames, ignore_index=True) if same_cols else frames[0]
    except Exception as e:
        st.error(f"Failed to read Invoices.xlsx: {e}")
        return pd.DataFrame()
    
    




    # ---- standardise common columns (add canonical columns without losing originals)
    def _first(df: pd.DataFrame, *cands: str) -> str | None:
        low = {c.lower(): c for c in df.columns}
        for c in cands:
            if c.lower() in low:
                return low[c.lower()]
        # fuzzy find
        pat = "|".join(map(re.escape, cands))
        for c in df.columns:
            if re.search(pat, c, flags=re.I):
                return c
        return None

    inv = inv.copy()

    # Date
    c_date = _first(inv, "Date", "Invoice Date", "Raised", "Created", "Posted")
    if c_date:
        inv["Date"] = pd.to_datetime(inv[c_date], errors="coerce")

    # Amount
    c_amt = _first(inv, "Amount", "Total", "Gross", "Net", "Value", "£")
    if c_amt:
        inv["Amount"] = pd.to_numeric(
            inv[c_amt].astype(str)
            .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)   # (123.45) -> -123.45
            .str.replace(r"[£$,]", "", regex=True)
            .str.replace(r"[^0-9.\-]", "", regex=True),
            errors="coerce",
        )

    # Supplier
    c_sup = _first(inv, "Supplier", "Vendor", "Company", "Payee")
    if c_sup:
        inv["Supplier"] = inv[c_sup].astype(str).str.strip()

    # Status / Paid
    c_stat = _first(inv, "Status", "State")
    if c_stat:
        inv["Status"] = inv[c_stat].astype(str).str.strip()

    c_paid = _first(inv, "Paid", "Is Paid", "Cleared")
    if c_paid:
        inv["Paid"] = (
            inv[c_paid]
            .astype(str)
            .str.strip()
            .str.lower()
            .isin(["yes", "y", "true", "1", "paid"])
        )

    # Optional team/area if present
    c_team = _first(inv, "Team", "Area", "Department", "Stakeholder")
    if c_team:
        inv["Team"] = inv[c_team].astype(str).str.strip()

    # Ensure canonical columns exist
    if "Amount" not in inv.columns:
        inv["Amount"] = np.nan
    if "Date" not in inv.columns:
        inv["Date"] = pd.NaT
    if "Supplier" not in inv.columns:
        inv["Supplier"] = None

    return inv
# === Sky Business: loader (sheet 'Main') ===
@st.cache_data(show_spinner=False)
def load_sky_business(path: str | Path = "Sky Business.xlsx") -> pd.DataFrame:
    p = Path(path)
    if not p.exists():
        st.warning("⚠️ Sky Business.xlsx not found in app folder.")
        return pd.DataFrame()

    try:
        df = pd.read_excel(p, sheet_name="Main")
    except Exception:
        df = pd.read_excel(p, sheet_name=0)

    # normalise headers
    df.columns = (
        df.columns.astype(str)
          .str.replace("\u00A0", " ", regex=False)
          .str.strip()
    )

    def pick(*alts):
        for a in alts:
            if a in df.columns:
                return a
        return None

    c_job  = pick("Job Type", "JobType", "Type")
    c_sla  = pick("SLA", "Sla", "S L A")
    c_slot = pick("Preferred time slot", "Prefered time slot", "Preferred Time Slot", "Preferred slot")
    c_date = pick("Date", "Requested Date", "Requested date", "Created Date")

    out = pd.DataFrame()
    if c_job:  out["JobType"]  = df[c_job].astype(str).str.strip()
    if c_sla:  out["SLA"]      = df[c_sla].astype(str).str.strip()
    if c_slot: out["PrefSlot"] = df[c_slot].astype(str).str.strip()

    # date (SBDate) used for filters + monthly grouping
    if c_date:
        out["SBDate"] = pd.to_datetime(df[c_date], errors="coerce")
    else:
        out["SBDate"] = pd.NaT

    for c in ["JobType", "SLA", "PrefSlot"]:
        if c in out.columns:
            out[c] = out[c].replace({"": pd.NA, "nan": pd.NA, "None": pd.NA})

    return out


# —— neat cards & grids for Sky Business ——
from textwrap import shorten

def _stat_card(title: str, value: str, subtitle: str | None = None, color: str = "#0ea5e9"):
    """Small consistent stat card. Uses your card() helper if present, otherwise a minimal fallback."""
    try:
        # If you already have card(title, big, sub, ...), reuse it
        card(title, value, subtitle or "", spark=None, color=color, source="")
    except Exception:
        box = st.container(border=True)
        with box:
            st.markdown(f"<div style='opacity:.75'>{title}</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:26px;font-weight:700;color:{color}'>{value}</div>", unsafe_allow_html=True)
            if subtitle:
                st.caption(subtitle)

def _chunk(items, n):
    for i in range(0, len(items), n):
        yield items[i:i+n]

def render_stat_grid(title: str, counts: dict[str, int], per_row: int = 5, color: str = "#0ea5e9"):
    """Render a responsive grid of even cards from a {label: count} dict."""
    st.markdown(f"#### {title}")
    if not counts:
        st.info("No data.")
        return
    items = [(shorten(str(k), width=28, placeholder='…'), int(v)) for k, v in counts.items()]
    for row in _chunk(items, per_row):
        cols = st.columns(len(row), gap="large")
        for (label, cnt), col in zip(row, cols):
            with col:
                _stat_card(label, f"{cnt:,}", None, color)




# === Sky Business: screen ===
def render_sky_business_screen():
    st.title("🏢 Sky Business")

    sb = load_sky_business()
    if sb.empty or not {"JobType", "SLA", "PrefSlot"}.intersection(sb.columns):
        st.info("No data found on sheet 'Main' (need Job Type, SLA, Preferred time slot).")
        return

    # -------- Filters (Year / Month with "All")
    years = sb["SBDate"].dropna().dt.year.sort_values().unique().tolist() if "SBDate" in sb else []
    col_f1, col_f2, col_f3 = st.columns([1, 1, 2])
    with col_f1:
        year_choice = st.selectbox("Year", ["All"] + [int(y) for y in years], index=0, key="sb_year")
    with col_f2:
        MONTHS = ["All","Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
        month_choice = st.selectbox("Month", MONTHS, index=0, key="sb_month")
    with col_f3:
        show_all_tiles = st.checkbox("Show ALL tiles (disable to show Top 12)", value=False, key="sb_showall")

    # -------- Apply filter to CURRENT view (table & totals)
    cur = sb.copy()
    if year_choice != "All" and "SBDate" in cur:
        cur = cur[cur["SBDate"].dt.year == int(year_choice)]
    if month_choice != "All" and "SBDate" in cur:
        mnum = MONTHS.index(month_choice)  # 1..12
        cur = cur[cur["SBDate"].dt.month == mnum]

    def _fmt_hhmm(td):
    
        if td is None or pd.isna(td):
            return "N/A"
        total_sec = int(td.total_seconds())
        # round to nearest minute so you don’t get :59 noise
        total_min = int(round(total_sec / 60.0))
        hours, minutes = divmod(total_min, 60)
        return f"{hours:02d}:{minutes:02d}"

    # -------- Build base for sparklines (keep monthly history, only year filter)
    spark_base = sb.copy()
    if year_choice != "All" and "SBDate" in spark_base:
        spark_base = spark_base[spark_base["SBDate"].dt.year == int(year_choice)]

    # Cap sparklines at selected month (so deltas work when a single month is chosen)
    end_month = None
    if year_choice != "All" and month_choice != "All" and "SBDate" in spark_base:
        end_month = pd.Timestamp(int(year_choice), MONTHS.index(month_choice), 1)

    # -------- Helper: monthly counts for a column
    def monthly_counts(df: pd.DataFrame, group_col: str) -> pd.DataFrame:
        if df.empty or "SBDate" not in df or group_col not in df:
            return pd.DataFrame(columns=["Month","Group","Count"])
        d = df.dropna(subset=["SBDate", group_col]).copy()
        d["Month"] = d["SBDate"].dt.to_period("M").dt.to_timestamp()
        g = d.groupby(["Month", group_col]).size().reset_index(name="Count")
        g.rename(columns={group_col: "Group"}, inplace=True)
        return g.sort_values(["Group", "Month"])

    # -------- Helper: sparkline values + MoM delta (▲/▼)
    def spark_and_delta(series_df: pd.DataFrame, group_val: str,
                        end: pd.Timestamp | None = None, points: int = 12):
        s = series_df[series_df["Group"] == group_val][["Month","Count"]].sort_values("Month")
        if end is not None:
            s = s[s["Month"] <= pd.Period(end, freq="M").to_timestamp()]
        if s.empty:
            return [0], "—", "—"
        spark_vals = s["Count"].tail(points).tolist()
        if len(s) < 2:
            return spark_vals, "—", "—"
        curr, prev = s["Count"].iloc[-1], s["Count"].iloc[-2]
        if prev == 0:
            return spark_vals, "—", "—"
        pct = (curr - prev) / prev * 100.0
        return spark_vals, f"{pct:+.1f}%", ("▲" if pct >= 0 else "▼")

    # -------- KPI row (unique counts + AM/PM/All Day totals from CURRENT view)
    uniq_jobtypes = cur["JobType"].dropna().nunique() if "JobType" in cur else 0
    uniq_slas     = cur["SLA"].dropna().nunique()     if "SLA" in cur else 0

    am = pm = allday = 0
    if "PrefSlot" in cur:
        s = cur["PrefSlot"].dropna().str.lower().str.strip()
        am     = int(s.str.fullmatch(r"am|am slot|morning").sum())
        pm     = int(s.str.fullmatch(r"pm|pm slot|afternoon").sum())
        allday = int(s.str.contains(r"all\s*day").sum())

    c1, c2, c3 = st.columns(3, gap="large")
    c1.metric("Unique Job Types", f"{uniq_jobtypes:,}")
    c2.metric("Unique SLA Types", f"{uniq_slas:,}")
    c3.metric("Total AM / PM / All Day", f"{am:,} / {pm:,} / {allday:,}")

    st.divider()

    # ==========================
    # Job Types (5 across)
    # ==========================
    st.markdown("### Job Types")
    if "JobType" in cur:
        job_counts = cur["JobType"].dropna().value_counts().sort_values(ascending=False)
        if not show_all_tiles:
            job_counts = job_counts.head(12)
    else:
        job_counts = pd.Series(dtype=int)

    jobs_series = monthly_counts(spark_base, "JobType") if "SBDate" in spark_base else \
                  pd.DataFrame(columns=["Month","Group","Count"])

    PER_ROW = 5
    from textwrap import shorten
    items = list(job_counts.items())
    if not items:
        st.info("No Job Types to display for the selected filters.")
    else:
        for r in range(0, len(items), PER_ROW):
            row = items[r:r+PER_ROW]
            cols = st.columns(len(row), gap="large")
            for c_idx, ((name, cnt), col) in enumerate(zip(row, cols)):
                label = shorten(str(name), width=32, placeholder="…")
                spark, delta_txt, arrow = spark_and_delta(jobs_series, name, end=end_month)
                subtitle = f"{arrow} {delta_txt} vs prev month" if delta_txt != "—" else "No prior month"
                with col:
                    try:
                        card(label, f"{int(cnt):,}", subtitle,
                             spark=spark, color="#0ea5e9", source=f"jobtype_{r}_{c_idx}")
                    except Exception:
                        box = st.container(border=True)
                        with box:
                            st.caption("Month to date"); st.markdown(f"**{label}**")
                            st.markdown(f"<div style='font-size:26px;font-weight:700;'>{int(cnt):,}</div>",
                                        unsafe_allow_html=True)

    st.divider()

    # ==========================
    # SLA Types (5 across)
    # ==========================
    st.markdown("### SLA Types")
    if "SLA" in cur:
        sla_counts = cur["SLA"].dropna().value_counts().sort_values(ascending=False)
        if not show_all_tiles:
            sla_counts = sla_counts.head(12)
    else:
        sla_counts = pd.Series(dtype=int)

    sla_series = monthly_counts(spark_base, "SLA") if "SBDate" in spark_base else \
                 pd.DataFrame(columns=["Month","Group","Count"])

    items = list(sla_counts.items())
    if not items:
        st.info("No SLA Types to display for the selected filters.")
    else:
        for r in range(0, len(items), PER_ROW):
            row = items[r:r+PER_ROW]
            cols = st.columns(len(row), gap="large")
            for c_idx, ((name, cnt), col) in enumerate(zip(row, cols)):
                label = shorten(str(name), width=32, placeholder="…")
                spark, delta_txt, arrow = spark_and_delta(sla_series, name, end=end_month)
                subtitle = f"{arrow} {delta_txt} vs prev month" if delta_txt != "—" else "No prior month"
                with col:
                    try:
                        card(label, f"{int(cnt):,}", subtitle,
                             spark=spark, color="#10b981", source=f"sla_{r}_{c_idx}")
                    except Exception:
                        box = st.container(border=True)
                        with box:
                            st.caption("Month to date"); st.markdown(f"**{label}**")
                            st.markdown(f"<div style='font-size:26px;font-weight:700;'>{int(cnt):,}</div>",
                                        unsafe_allow_html=True)

    st.divider()

    # =========================================
    # Preferred time slot (AM / PM / All Day)
    # =========================================
    st.markdown("### Preferred time slot")

    # Totals from CURRENT view
    def slot_total(df: pd.DataFrame, regex: str, exact: bool = False) -> int:
        if "PrefSlot" not in df:
            return 0
        s = df["PrefSlot"].dropna().str.lower().str.strip()
        return int(s.str.fullmatch(regex).sum() if exact else s.str.contains(regex).sum())

    tot_AM     = slot_total(cur, r"am|am slot|morning", exact=True)
    tot_PM     = slot_total(cur, r"pm|pm slot|afternoon", exact=True)
    tot_AllDay = slot_total(cur, r"all\s*day", exact=False)

    # Canonicalise for spark history
    canon = spark_base.copy()
    if "PrefSlot" in canon.columns:
        canon["PrefSlot"] = (canon["PrefSlot"].fillna("")
                             .str.lower().str.strip()
                             .replace({"am slot":"am","morning":"am","pm slot":"pm","afternoon":"pm"})
                             .replace({r".*\ball\s*day\b.*":"all day"}, regex=True))
    slot_series = monthly_counts(canon, "PrefSlot") if "SBDate" in canon else \
                  pd.DataFrame(columns=["Month","Group","Count"])

    cols = st.columns(3, gap="large")
    for (label, key, total, color, src), col in zip(
        [("AM","am",tot_AM,"#3b82f6","slot_am"),
         ("PM","pm",tot_PM,"#f59e0b","slot_pm"),
         ("All Day","all day",tot_AllDay,"#22c55e","slot_allday")],
        cols
    ):
        spark, delta_txt, arrow = spark_and_delta(slot_series, key, end=end_month)
        subtitle = f"{arrow} {delta_txt} vs prev month" if delta_txt != "—" else "No prior month"
        with col:
            try:
                card(label, f"{int(total):,}", subtitle, spark=spark, color=color, source=src)
            except Exception:
                box = st.container(border=True)
                with box:
                    st.caption("Month to date"); st.markdown(f"**{label}**")
                    st.markdown(f"<div style='font-size:26px;font-weight:700;'>{int(total):,}</div>",
                                unsafe_allow_html=True)

    st.divider()

    # -------- Main table
    st.markdown("#### Main table")
    show_cols = [c for c in ["JobType","SLA","PrefSlot","SBDate"] if c in cur.columns]
    if show_cols:
        try:
            from st_aggrid import AgGrid, GridOptionsBuilder
            gb = GridOptionsBuilder.from_dataframe(cur[show_cols])
            gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=15)
            gb.configure_default_column(editable=False, filter=True, sortable=True)
            AgGrid(cur[show_cols], gridOptions=gb.build(), height=420, fit_columns_on_grid_load=True)
        except Exception:
            st.dataframe(cur[show_cols], use_container_width=True)
    else:
        st.info("No columns to display.")






@st.cache_data(show_spinner=False)
def load_budgets_df() -> pd.DataFrame:
    # 1) find the file
    budget_path = _find_first_existing("budgets.csv", "Budgets.csv", "Budget.csv", "BUDGETS.CSV")
    if not budget_path:
        st.warning("⚠️ budgets.csv file not found in working directory.")
        return pd.DataFrame()

    # 2) read (robust to delimiters)
    try:
        # sep=None lets pandas sniff delimiter (comma vs semicolon)
        bud = pd.read_csv(budget_path, sep=None, engine="python", encoding="utf-8-sig")
    except Exception:
        try:
            bud = pd.read_excel(budget_path)
        except Exception as e:
            st.warning(f"⚠️ Could not read {budget_path}: {e}")
            return pd.DataFrame()

    # 3) normalise header
    bud = _normalise_cols(bud)

    # 4) map team column to 'Team'
    team_col = TEAM_COL_OVERRIDE or _find_col(bud, "team", "area", "team name", "group", "department", "dept", "stakeholder")
    if team_col:
        if team_col != "Team":
            bud.rename(columns={team_col: "Team"}, inplace=True)
    else:
        st.warning(f"⚠️ Could not load {budget_path}: no 'Team' (Stakeholder) column.")
        st.caption("Columns found in budgets file:")
        st.code(list(bud.columns))
        return pd.DataFrame()   

    # 5) normalise Month if present under a different name
    if "Month" not in bud.columns:
        month_col = _find_col(bud, "month", "period", "month name")
        if month_col:
            bud.rename(columns={month_col: "Month"}, inplace=True)

    # tidy Team text
    bud["Team"] = bud["Team"].astype(str).str.strip()
    return bud



# === ORBIT AI: SETUP (clean) =================================================
from langchain_openai import ChatOpenAI
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain.agents.agent_types import AgentType

TEAM_TAGS = {"VIP North", "VIP South", "Tier 2 North", "Tier 2 South"}

def _find_team_col(df: pd.DataFrame) -> str | None:
    if df.empty:
        return None
    for c in df.columns:
        lc = c.lower()
        if lc in ("team", "area", "region") or any(tok in lc for tok in ("team","area","region")):
            return c
    return None

def _scope_df_to_page(df: pd.DataFrame, page_tag: str) -> pd.DataFrame:
    """Return only the rows for an exact team match when page_tag is a team."""
    if df.empty or page_tag not in TEAM_TAGS:
        return df
    team_col = _find_team_col(df)
    if not team_col:
        return df
    # strict match (case-insensitive), avoids VIP picking up both VIP North/South, etc.
    return df[df[team_col].astype(str).str.strip().str.casefold() == page_tag.casefold()]

# --- Currency guardrails ---
CURRENCY_SYMBOL = "£"

def _force_gbp(text: str) -> str:
    """Replace any stray $/USD with GBP symbol. Safe for normal chat output."""
    if not isinstance(text, str):
        return text
    out = re.sub(r'\bUSD\b', 'GBP', text, flags=re.I)
    out = re.sub(r'\bUS\$\b', CURRENCY_SYMBOL, out, flags=re.I)
    out = re.sub(r'(?<![A-Za-z])\$', CURRENCY_SYMBOL, out)  # $5,120 -> £5,120
    return out

# 🔤 Query → DataFrame column/term aliases
# NOTE: keys MUST be lowercase; _alias() lowercases the text before replacing.
ALIASES = {
    # --- Time fields ---
    "activate time": "Activate",
    "activation time": "Activate",
    "activate": "Activate",

    "deactivate time": "Deactivate",
    "deactivation time": "Deactivate",
    "deactivate": "Deactivate",

    "total time": "Total Time",
    "tt": "Total Time",
    "duration": "Total Time",

    "total time (inc travel)": "Total Time (Inc Travel)",
    "total incl travel": "Total Time (Inc Travel)",
    "total including travel": "Total Time (Inc Travel)",
    "tt inc travel": "Total Time (Inc Travel)",

    "total working time": "Total Working Time",
    "avg total working time": "Total Working Time",
    "working time": "Total Working Time",
    "work time": "Total Working Time",
    "twt": "Total Working Time",

    "travel time": "Travel Time",

    # --- Value / Cost fields ---
    "total £": "Total Value",
    "total value": "Total Value",
    "£ value": "Total Value",
    "value £": "Total Value",
    "monetary value": "Total Value",
    "revenue": "Total Value",
    "income": "Total Value",

    "total cost": "Total Cost Inc Travel",
    "total cost inc travel": "Total Cost Inc Travel",
    "cost inc travel": "Total Cost Inc Travel",
    "cost including travel": "Total Cost Inc Travel",
    "spend inc travel": "Total Cost Inc Travel",

    # --- Budgets (exec) ---
    "budget used": "Used",
    "used amount": "Used",
    "spend": "Used",
    "spent": "Used",
    "actuals": "Used",
    "expenses": "Used",
    "expense": "Used",
    "costs": "Used",
    "cost": "Used",

    "budget allocated": "Allocated",
    "allocated budget": "Allocated",
    "allocation": "Allocated",
    "alloc": "Allocated",

    "budget remaining": "Remaining",
    "remaining budget": "Remaining",
    "budget left": "Remaining",
    "left": "Remaining",
    "balance": "Remaining",

    "usage %": "Usage %",
    "usage%": "Usage %",
    "utilisation %": "Usage %",
    "utilization %": "Usage %",
    "use %": "Usage %",

    # --- Overtime ---
    "overtime £": "Overtime",
    "overtime gbp": "Overtime",
    "overtime cost": "Overtime",
    "overtime value": "Overtime",
    "over-time": "Overtime",
    "ot": "Overtime",

    # --- Entities / dimensions ---
    "stakeholder": "Sky Retail Stakeholder",
    "retail stakeholder": "Sky Retail Stakeholder",

    "visit type": "Visit Type",
    "type": "Visit Type",

    "post code": "Postcode",
    "postcode": "Postcode",

    "engineer name": "Engineer",
    "engineer": "Engineer",
    "advisor": "Engineer",
    "agent": "Engineer",

    # --- Team name normalisation (helps the DF agent) ---
    "t2": "Tier 2",
    "tier2": "Tier 2",
    "vip north": "VIP North",
    "vip south": "VIP South",
    "tier 2 north": "Tier 2 North",
    "tier 2 south": "Tier 2 South",
}
# --------- ALIASES (regex-based, safer) ---------
ALIAS_PATTERNS: list[tuple[str, str]] = [
    # Time columns
    (r"\bactivate(?:d| time)?\b", "Activate"),
    (r"\bdeactivate(?:d| time)?\b", "Deactivate"),
    (r"\b(total|overall)\s+time\b", "Total Time"),
    (r"\b(total\s+time\s*(?:inc|including)\s*travel|tt\s*inc\s*travel)\b", "Total Time (Inc Travel)"),
    (r"\b(total|overall)\s+working\s+time\b|\btwt\b", "Total Working Time"),
    (r"\btravel\s*time\b", "Travel Time"),

    # Value / Cost
    (r"\b(total\s*£|£\s*value|total\s*value|monetary\s*value|revenue|income)\b", "Total Value"),
    (r"\b(total\s*cost(?:\s*(inc|including)\s*travel)?|spend\s*inc\s*travel)\b", "Total Cost Inc Travel"),

    # Budgets
    (r"\bbudget\s*used|used\s*amount|spend|spent|actuals|expenses?\b", "Used"),
    (r"\bbudget\s*allocated|allocated\s*budget|allocation|alloc\b", "Allocated"),
    (r"\bbudget\s*remaining|remaining\s*budget|budget\s*left|balance|left\b", "Remaining"),
    (r"\busage\s*%|usage%|utili[sz]ation\s*%|use\s*%\b", "Usage %"),

    # Overtime
    (r"\bovertime\s*(£|gbp|cost|value)?\b|\bo/?t\b", "Overtime"),

    # Dimensions / entities
    (r"\bstakeholder|retail\s*stakeholder\b", "Sky Retail Stakeholder"),
    (r"\bvisit\s*type\b|\btype\b", "Visit Type"),
    (r"\bpost\s*code\b|\bpostcode\b", "Postcode"),
    (r"\bengineer(?:\s*name)?|advisor|agent\b", "Engineer"),

    # Team names (normalise)
    (r"\bt2\b|\btier2\b", "Tier 2"),
    (r"\bvip\s*north\b", "VIP North"),
    (r"\bvip\s*south\b", "VIP South"),
    (r"\btier\s*2\s*north\b", "Tier 2 North"),
    (r"\btier\s*2\s*south\b", "Tier 2 South"),
]

# Activity Status synonyms -> canonical labels (for counting/filters)
STATUS_MAP = {
    r"\b(done|finished|completed|closed)\b": "Completed",
    r"\b(cancel(?:led|ed)?|void(?:ed)?|aborted|no\s*access)\b": "Cancelled",
    r"\b(suspended|on\s*hold)\b": "Suspended",
    r"\bnot\s*(done|completed)\b": "Not Done",
    r"\b(pending|awaiting|open)\b": "Pending",
    r"\b(started|in\s*progress|working)\b": "Started",
}
# --------- PERIOD RESOLVER ---------
MONTH_NAMES = {
    "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
    "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6, "jul": 7, "july": 7,
    "aug": 8, "august": 8, "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12,
}

def _resolve_period(df: pd.DataFrame, q: str) -> tuple[pd.DataFrame, str]:
    """
    Detect common time phrases and filter df:
      - this month / current month / mtd
      - last month / previous month
      - ytd
      - explicit month names like 'July' or 'Jul 2025'
    Falls back to Exec page selection if available, else latest month in data.
    """
    if df.empty or "Date" not in df.columns:
        return df, "(all time)"

    base = df.copy()
    base["__month"] = pd.to_datetime(base["Date"], errors="coerce").dt.to_period("M")
    if base["__month"].isna().all():
        return df, "(all time)"

    ql = (q or "").lower().strip()
    today = datetime.now(UK_TZ)
    this_m = pd.Period(today.date().replace(day=1), "M")
    last_m = (this_m - 1)

    # exec month (if your exec_ctx captured it)
    exec_ctx = st.session_state.get("exec_ctx", {})
    exec_label = exec_ctx.get("month_label")
    exec_period = None
    if exec_label:
        try:
            exec_period = pd.Period(pd.to_datetime(exec_label).date().replace(day=1), "M")
        except Exception:
            exec_period = None

    # Helpers
    def by_period(p: pd.Period, label: str):
        return base[base["__month"] == p].drop(columns="__month"), label
    # in your ALIASES or recognition list
    # phrases that imply money
    MONEY_WORDS = {"£", "gbp", "pound", "pounds", "value", "money", "cost", "spend", "spent", "actuals"}

    # overtime wording variants
    # users often write: "ot", "ot cost", "o/t", "overtime spend"

        # MTD = current month up to today
    if re.search(r"\b(mtd|month\s*to\s*date|this\s*month|current\s*month)\b", ql):
        return by_period(this_m, f"(this month – {this_m.strftime('%b %Y')})")

    if re.search(r"\b(last|previous)\s*month\b", ql):
        return by_period(last_m, f"(last month – {last_m.strftime('%b %Y')})")

    if re.search(r"\bytd|year\s*to\s*date\b", ql):
        mask = pd.to_datetime(base["Date"], errors="coerce").dt.year == today.year
        return base[mask].drop(columns="__month"), f"(YTD {today.year})"

    # Named month (optionally with year)
    m = re.search(r"\b(" + "|".join(MONTH_NAMES.keys()) + r")\b(?:\s+(\d{4}))?", ql)
    if m:
        mm = MONTH_NAMES[m.group(1)]
        yy = int(m.group(2)) if m.group(2) else today.year
        try:
            per = pd.Period(datetime(yy, mm, 1).date().replace(day=1), "M")
            return by_period(per, f"({per.strftime('%b %Y')})")
        except Exception:
            pass

    # Exec-selected month (if any)
    if exec_period is not None:
        return by_period(exec_period, f"({exec_period.strftime('%b %Y')})")

    # Fallback: latest month in data
    latest = base["__month"].dropna().max()
    return by_period(latest, f"(latest – {latest.strftime('%b %Y')})")

def _alias(text: str) -> str:
    """Regex replacements with word boundaries to avoid accidental hits."""
    t = str(text)
    for pat, repl in ALIAS_PATTERNS:
        t = re.sub(pat, repl, t, flags=re.I)
    # Status labels get normalised to their canonical names
    for pat, label in STATUS_MAP.items():
        t = re.sub(pat, label, t, flags=re.I)
    return t

def _alias(text: str) -> str:
    t = text.lower()
    for k, v in ALIASES.items():
        t = t.replace(k, v)
    return t

def _classify_exec_intent(q: str) -> str:
    """
    Return one of: 'chitchat', 'exec_report', 'exec_metric', 'other'.
    - chitchat: hello/thanks/how are you/etc.
    - exec_report: broad overviews/comparisons requiring the full template.
    - exec_metric: specific KPI questions (usage %, variance) -> concise numeric answer, not full template.
    - other: general Qs -> normal DF agent answer.
    """
    t = (q or "").strip().lower()

    # Chit-chat / pleasantries
    chitchat = {"hi", "hello", "hey", "yo", "hiya", "good morning", "good afternoon",
                "good evening", "thanks", "thank you", "cheers", "ok", "okay", "great", "cool"}
    if any(t == w or t.startswith(w + " ") for w in chitchat) or t in chitchat:
        return "chitchat"

    # Keywords that imply KPI/budget analytics
    kpi_kw = [
        "budget", "allocated", "used", "remaining", "usage %", "usage%", "variance",
        "over budget", "under budget", "spend", "actuals", "cost", "expenses",
        "total value", "value per visit", "completion rate", "completed visits",
        "trend", "mom", "m/m", "wow", "w/w", "ytd", "mtd", "by team", "compare",
        "difference", "vip", "tier 2", "percentage", "percent", "%", "vs", "versus"
    ]
    has_kpi = any(kw in t for kw in kpi_kw)

    # Heuristic: detect comparison-style questions
    compare_markers = ["compare", "difference", "vs", "versus"]
    is_compare = any(m in t for m in compare_markers)

    # Broad prompts that usually want an overview
    overview_markers = ["overview", "summary", "executive overview", "exec overview"]

    if has_kpi and (is_compare or any(m in t for m in overview_markers)):
        return "exec_report"

    if has_kpi:
        return "exec_metric"

    return "other"

def _exec_semantic_glossary() -> str:
    """
    Teaches the model the budget & KPI vocabulary we use on the Exec page.
    Keep this short and opinionated.
    """
    return (
        "Vocabulary & Rules (Executive Overview):\n"
        "- 'VIP' means all VIP teams combined; 'Tier 2' means all Tier 2 teams combined.\n"
        "- Budget Allocated: the approved amount for the period.\n"
        "- Budget Used (aka Spend/Actuals/Cost/Expenses): the amount spent so far.\n"
        "- Budget Remaining = Allocated - Used.\n"
        "- Budget Usage % = (Used / Allocated) * 100.\n"
        "- Variance vs Budget = Used - Allocated (positive = over budget, negative = under).\n"
        "- Completion Rate % = Completed Visits / Total Visits * 100.\n"
        "- Value per Visit = Total Value / Total Visits.\n"
        "- % difference between two teams A vs B = (A - B) / B * 100 unless the user specifies another base.\n"
        "- 'This month' defaults to the month selected on the Exec page; if none selected, use the latest month in view.\n"
        "- If the user says 'we/our' on Exec Overview, interpret that as the organisation across VIP + Tier 2 unless they name a team.\n"
    )

def _fewshot_exec_examples() -> str:
    """
    A couple of tight examples that steer answers for budget & comparison questions.
    """
    return (
        "Examples:\n"
        "Q: What's the budget usage % for VIP this month?\n"
        "A: Budget Usage % (VIP) = Used/Allocated*100. Result: 74.2% (Used £X, Allocated £Y).\n"
        "\n"
        "Q: What's the percentage difference in Total Value between VIP and Tier 2 this month?\n"
        "A: (VIP − Tier 2)/Tier 2 * 100. VIP is +12.8% vs Tier 2 (VIP £X vs Tier 2 £Y).\n"
        "\n"
        'Q: Are we over budget?\n'
        "A: Variance = Used − Allocated. We are under budget by £Z (Used £X vs Allocated £Y). Usage: 82.3%.\n"
        "\n"
        "Q: Give me an executive overview for budgets across VIP and Tier 2.\n"
        "A: Bullet the Allocated, Used, Remaining, Usage % for VIP, Tier 2, and Total; then a one-line interpretation.\n"
    )

def _answer_style_exec() -> str:
    return (
        "Answer Style:\n"
        "- Always use this template:\n"
        "  Budget Overview:\n"
        "    - VIP: Allocated, Used, Remaining, Usage %\n"
        "    - Tier 2: Allocated, Used, Remaining, Usage %\n"
        "    - Total: Allocated, Used, Remaining, Usage %\n"
        "  Performance Overview:\n"
        "    - Total Visits\n"
        "    - Completed Visits\n"
        "    - Completion Rate %\n"
        "    - Value per Visit\n"
        "  Takeaway: One decisive line.\n"
        "- Never add extra empty 'Summary:' or unused headers.\n"
        "- Show the formula used where relevant (e.g., Usage % = Used / Allocated * 100).\n"
        "- Be concise and exec-ready.\n"
    )


def _orbit_pick_df() -> pd.DataFrame:
    for name in ("user_df", "df4", "df_all", "combined_oracle_df"):
        if name in st.session_state:
            df = st.session_state[name]
            if isinstance(df, pd.DataFrame) and not df.empty:
                return df
        df = globals().get(name)
        if isinstance(df, pd.DataFrame) and not df.empty:
            return df
    return pd.DataFrame()

def _orbit_clean(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    df = df.copy()
    if "Date" in df.columns:
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    for td in ("Activate","Deactivate","Total Time","Total Working Time","Travel Time","Total Time (Inc Travel)"):
        if td in df.columns:
            df[td] = pd.to_timedelta(df[td].astype(str), errors="coerce")
    for c in df.select_dtypes("object").columns:
        df[c] = df[c].astype(str).str.strip()
    return df

def _persona_prefix(page_tag: str) -> str:
    """Give the LLM context about who is speaking and what 'I' means."""
    user = (st.session_state.get("username") or "").strip()
    role = (st.session_state.get("role") or "").strip()
    who = f"{user}" if user else "a user"
    scope = "the Executive Overview for all VIP + Tier 2 teams" if page_tag == "exec_overview" else f"the {page_tag} team"
    return (
        "You are Orbit, the analysis assistant for Sky VIP.\n"
        f"The current speaker is **{who}** ({role or 'unknown role'}).\n"
        "When they use words like 'I', 'my', or 'our', interpret them from that person's perspective.\n"
        f"Restrict answers to **{scope}** unless they explicitly ask otherwise.\n"
        "Prefer numbers from the page’s data. If uncertain, ask a precise follow-up."
    )

def _ensure_orbit_agent(page_tag: str = "global"):
    base = _orbit_clean(_orbit_pick_df())
    scoped = _scope_df_to_page(base, page_tag)

    st.session_state.setdefault("orbit_agent_map", {})
    st.session_state.setdefault("orbit_df_map", {})
    st.session_state.setdefault("orbit_agent_sig_map", {})

    sig = (scoped.shape, tuple(scoped.columns), page_tag)
    if st.session_state["orbit_agent_sig_map"].get(page_tag) != sig:
        llm_stream = ChatOpenAI(
            api_key=st.secrets["openai"]["api_key"],
            model_name="gpt-4o-mini",
            streaming=True,
        )
        agent = create_pandas_dataframe_agent(
            llm=llm_stream,
            df=scoped,
            verbose=False,
            agent_type=AgentType.OPENAI_FUNCTIONS,
            allow_dangerous_code=True,
        )
        st.session_state["orbit_agent_map"][page_tag] = agent
        st.session_state["orbit_df_map"][page_tag] = scoped
        st.session_state["orbit_agent_sig_map"][page_tag] = sig

    return (
        st.session_state["orbit_agent_map"][page_tag],
        st.session_state["orbit_df_map"][page_tag],
    )

# --- Period parsing helpers ---------------------------------------------------
import calendar
from datetime import datetime as _dt
import pandas as _pd

_MONTHS = {m.lower(): i for i, m in enumerate(calendar.month_name) if m}
_MONTHS.update({m.lower(): i for i, m in enumerate(calendar.month_abbr) if m})

def _parse_month_from_text(text: str, now: _dt | None = None):
    """
    Returns: (year, month, label) or (None, None, 'all')
    Understands: 'june', 'jun 2025', 'this month', 'mtd', 'last month', 'ytd'.
    Default if nothing found = this month.
    """
    t = (text or "").lower().strip()
    now = now or now_uk()

    # relative
    if "ytd" in t:
        return (now.year, None, "YTD")
    if "last month" in t:
        prev = (now.replace(day=1) - _pd.offsets.Day(1)).to_pydatetime()
        return (prev.year, prev.month, prev.strftime("%b %Y"))
    if "this month" in t or "mtd" in t:
        return (now.year, now.month, now.strftime("%b %Y"))

    # explicit month (with optional year)
    for token, mnum in _MONTHS.items():
        if token and token in t:
            # try to find a 4-digit year around it
            import re
            m = re.search(r"(19|20)\d{2}", t)
            yr = int(m.group(0)) if m else now.year
            label = f"{calendar.month_abbr[mnum]} {yr}"
            return (yr, mnum, label)

    # nothing matched -> default this month
    return (now.year, now.month, now.strftime("%b %Y"))

def _money_series(s: _pd.Series) -> _pd.Series:
    """Parse strings like '£1,234.50' or '(£200)' into floats; NaN if not money."""
    return _pd.to_numeric(
        s.astype(str)
         .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)
         .str.replace(r"[£,\s]", "", regex=True),
        errors="coerce"
    )

def _overtime_gbp_for(df: _pd.DataFrame, year: int | None, month: int | None) -> float | None:
    """
    Sum 'Overtime' in **£** for the requested period.
    If only hours are present (no £), returns None so we can fall back to hours.
    """
    if df.empty:
        return 0.0

    x = df.copy()
    # date filter
    if "Date" in x.columns and (year or month):
        x["Date"] = _pd.to_datetime(x["Date"], errors="coerce")
        x = x.dropna(subset=["Date"])
        if month:  # a specific month (yr, m)
            start = _pd.Timestamp(year=year, month=month, day=1)
            end   = (start + _pd.offsets.MonthBegin(1))
            x = x[(x["Date"] >= start) & (x["Date"] < end)]
        else:      # YTD
            x = x[x["Date"].dt.year == (year or now_uk().year)]

    # find the overtime column
    cand = [c for c in x.columns if c.strip().lower() == "overtime"]
    if not cand:
        return 0.0
    col = cand[0]

    raw = x[col].astype(str).str.replace("\u00A0", " ", regex=False).str.strip()

    # Try £ first
    money = _money_series(raw)
    if money.notna().any():
        return float(money.fillna(0).sum())

    # No £ → looks like time (hh:mm etc.) – tell caller to fall back to hours path
    return None

def _overtime_hours_for(df: _pd.DataFrame, year: int | None, month: int | None) -> float:
    """Sum overtime hours for the requested period (if only time is stored)."""
    if df.empty:
        return 0.0
    x = df.copy()
    if "Date" in x.columns and (year or month):
        x["Date"] = _pd.to_datetime(x["Date"], errors="coerce")
        x = x.dropna(subset=["Date"])
        if month:
            start = _pd.Timestamp(year=year, month=month, day=1)
            end   = (start + _pd.offsets.MonthBegin(1))
            x = x[(x["Date"] >= start) & (x["Date"] < end)]
        else:
            x = x[x["Date"].dt.year == (year or now_uk().year)]

    if "Overtime" not in x.columns:
        return 0.0

    td = _pd.to_timedelta(x["Overtime"].astype(str), errors="coerce")
    return float((td.dt.total_seconds() / 3600.0).fillna(0).sum())

def _orbit_ai_answer(q: str, page_tag: str) -> str:
    """
    Page-aware Q&A with:
      - small-talk guard (no data dump on "hello")
      - 'my team' resolution -> current team
      - special-case: Overtime in £ from proper KPI calc
      - exec intent handling (report/metric/chit-chat) remains
    """
    agent, _ = _ensure_orbit_agent(page_tag)
    q_raw   = (q or "").strip()
    q_lower = q_raw.lower()

    # 0) Small-talk / greetings => keep it conversational
    if re.search(r'^(hi|hello|hey|morning|afternoon|evening|thanks|thank you)\\b', q_lower):
        name = (st.session_state.get("username") or "").split(" ")[0] or "there"
        team = page_tag if page_tag in TEAM_TAGS else (st.session_state.get("user_team") or "your area")
        return f"Hi {name}! 👋 You’re on **{team}**. Ask me things like “budget remaining”, “overtime in £”, “completion rate last month”, or “visits by stakeholder”."

    # 1) Normalise aliases & resolve "my team" to the scoped team
    q_alias = _alias(q_raw)
    if page_tag in TEAM_TAGS:
        q_alias = re.sub(r'\\bmy team\\b', page_tag, q_alias, flags=re.I)

    # --- inside _orbit_ai_answer(q, page_tag) ---

    # 1) Ensure we have the agent and the page-scoped dataframe
    agent, scoped_df = _ensure_orbit_agent(page_tag)

    q_raw   = (q or "").strip()
    q_lower = q_raw.lower()

    # 2) Greetings / small talk → keep it conversational (no data dump)
    if re.search(r'^(hi|hello|hey|morning|afternoon|evening|thanks|thank you)\b', q_lower):
        name = (st.session_state.get("username") or "").split(" ")[0] or "there"
        team = page_tag if page_tag in TEAM_TAGS else (st.session_state.get("user_team") or "your area")
        return f"Hi {name}! 👋 You’re on **{team}**. Ask me things like “budget remaining”, “overtime in £ for June”, or “completion rate last month”."

    # 3) Normalise wording (+ resolve “my team” → actual team)
    q_alias = _alias(q_raw)
    if page_tag in TEAM_TAGS:
        q_alias = re.sub(r'\bmy team\b', page_tag, q_alias, flags=re.I)

    # 4) SPECIAL CASE — Overtime £ (with month detection)
    #    Triggers on “overtime” + any money word: £, gbp, pound(s), value, money, cost, spend, spent, actuals
    if ("overtime" in q_lower) and any(w in q_lower for w in ["£","gbp","pound","pounds","value","money","cost","spend","spent","actuals"]):
        # parse requested month (defaults to this month if none found)
        yr, mo, label = _parse_month_from_text(q_lower, now_uk())

        # Sum in £ if 'Overtime' column contains money; else fall back to hours
        gbp = _overtime_gbp_for(scoped_df, yr, mo)
        if gbp is not None:
            return f"**Overtime (cost) — {label} · {page_tag}: £{gbp:,.2f}.**"
        else:
            hrs = _overtime_hours_for(scoped_df, yr, mo)
            # format hours into H:MM
            h = int(hrs)
            m = int(round((hrs - h) * 60))
            return (
                f"**Overtime (hours) — {label} · {page_tag}: {h}h {m:02d}m.**  \n"
                "No £ values were found in the 'Overtime' column for that period."
            )

    # 5) OPTIONAL: also answer pure “overtime in June” (no money word) as hours
    if ("overtime" in q_lower) and any(tok in q_lower for tok in ["january","february","march","april","may","june","july","august","september","october","november","december","last month","this month","mtd","ytd","jun","jul","aug","sep","oct","nov","dec"]):
        yr, mo, label = _parse_month_from_text(q_lower, now_uk())
        hrs = _overtime_hours_for(scoped_df, yr, mo)
        h = int(hrs); m = int(round((hrs - h) * 60))
        return f"**Overtime (hours) — {label} · {page_tag}: {h}h {m:02d}m.**"


    # 3) Exec page intent handling (keeps your concise exec style)
    user_name = (st.session_state.get("username") or "").strip()
    user_role = (st.session_state.get("role") or "").strip()
    persona_hdr = (
        f"You are Orbit, answering {user_name or 'a user'} ({user_role or 'Exec'}). "
        "Be concise and board-ready.\n"
    )

    exec_ctx_text = ""
    if page_tag == "exec_overview":
        ctx = st.session_state.get("exec_ctx", {})
        if ctx:
            from pprint import pformat
            exec_ctx_text = "Executive Overview Context:\n" + pformat(ctx, compact=True) + "\n"

    intent   = _classify_exec_intent(q_alias) if page_tag == "exec_overview" else "other"
    glossary = _exec_semantic_glossary() if page_tag == "exec_overview" else ""
    style_ex = _answer_style_exec()       if page_tag == "exec_overview" else ""
    fewshot  = _fewshot_exec_examples()   if page_tag == "exec_overview" else ""

    extras = ""
    if page_tag == "exec_overview":
        if intent == "exec_report":
            extras = glossary + "\n" + fewshot + "\n" + style_ex
        elif intent == "exec_metric":
            extras = glossary + (
                "\nAnswer Style:\n"
                "- Give just the metric requested with the formula and numbers used.\n"
                "- No headings. One-line takeaway.\n"
            )
        elif intent == "chitchat":
            return "Hi! 👋 What would you like to look at on the Exec page (budget, variance, completion rate, VIP vs Tier 2, etc.)?"
        else:
            extras = glossary

    full_prompt = (persona_hdr + exec_ctx_text + extras + "\nUser question:\n" + q_alias).strip()

    # 4) Try the dataframe agent first
    try:
        res = agent.invoke(full_prompt)
        out = res.get("output", "") if isinstance(res, dict) else str(res or "")
        if out.strip():
            return out
    except Exception:
        pass

    # 5) Fallback to a small LLM call, but keep scope if on a team tab
    try:
        scope_hint = ""
        if page_tag in TEAM_TAGS:
            scope_hint = (
                f"Only use information from the **{page_tag}** team subset of the data. "
                "Do not include other teams unless explicitly asked."
            )
        llm = ChatOpenAI(api_key=st.secrets["openai"]["api_key"], model_name="gpt-4o-mini", temperature=0.1)
        return llm.invoke(scope_hint + "\n\n" + full_prompt).content
    except Exception as e:
        return f"Sorry — I hit an error answering that: {e}"
# --------- KPI CALCS (deterministic) ---------
def _completion_rate(df: pd.DataFrame) -> tuple[int, int, float]:
    if df.empty:
        return 0, 0, 0.0
    total = len(df)
    comp = 0
    if "Activity Status" in df.columns:
        comp = df["Activity Status"].astype(str).str.contains("Completed", case=False, na=False).sum()
    pct = (comp / total * 100.0) if total else 0.0
    return total, comp, pct

def _value_per_visit(df: pd.DataFrame) -> tuple[float, int, float]:
    if df.empty:
        return 0.0, 0, 0.0
    tv = pd.to_numeric(df.get("Total Value"), errors="coerce").fillna(0).sum()
    n  = len(df)
    vpv = (tv / n) if n else 0.0
    return float(tv), int(n), float(vpv)

def _count_by_status(df: pd.DataFrame, canon_status: str) -> int:
    if df.empty or "Activity Status" not in df.columns:
        return 0
    return df["Activity Status"].astype(str).str.contains(canon_status, case=False, na=False).sum()

def _detect_status_from_query(q: str) -> str | None:
    for pat, label in STATUS_MAP.items():
        if re.search(pat, q, flags=re.I):
            return label
    return None




def render_orbit_ai(page_tag: str = "global"):
    """Drop this once at the bottom of any page/tab to add the Orbit chat box."""
    _ensure_orbit_agent(page_tag)  # make sure it's ready

    hist_key = f"orbit_ai_history_{page_tag}"
    st.session_state.setdefault(hist_key, [])
    st.markdown(f"### 🤖 Sky Orbit — Ask anything · {page_tag}")

    if st.button("🧹 Clear Chat", key=f"clear_ai_{page_tag}"):
        st.session_state[hist_key] = []

    # render history
    for msg in st.session_state[hist_key]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("timestamp"):
                st.caption(msg["timestamp"])

    # chat input
    q = st.chat_input("Ask me anything (Oracle, charts, summaries, budgets, emails, etc.)",
                      key=f"chat_input_{page_tag}")
    if not q:
        return

    now = datetime.now(UK_TZ).strftime("%Y-%m-%d %H:%M")

    # --- show user question
    st.session_state[hist_key].append({"role": "user", "content": q, "timestamp": now})
    with st.chat_message("user"):
        st.markdown(f"**Q:** {q}")
        st.caption(now)

    # --- generate AI reply
    with st.chat_message("assistant"):
        reply = _orbit_ai_answer(q, page_tag)
        st.markdown(reply)
        st.caption(now)

    # --- save reply after it's defined
    st.session_state[hist_key].append({"role": "assistant", "content": reply, "timestamp": now})

    # --- log to CSV
    try:
        import csv
        with open("chat_logs.csv", "a", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            if f.tell() == 0:
                w.writerow(["timestamp", "page", "question", "answer"])
            w.writerow([datetime.now(UK_TZ).isoformat(timespec="seconds"), page_tag, q, reply])
    except Exception:
        pass

# === END ORBIT AI: SETUP =====================================================



# --- safety: make sure ops subsection is always present ---
st.session_state.setdefault("op_area_section", "menu")
section = st.session_state.get("op_area_section", "menu")
username = (st.session_state.get("username") or "").strip().lower()
is_ops = (st.session_state.get("role") == "operations_manager") or username.startswith("rob")

def img_to_base64(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# Load once at the top of your script
vip_t2_logo_b64 = img_to_base64("sky_vip_logo.png")

def to_data_uri(path: str) -> str | None:
    """Convert an image file to a data URI for embedding."""
    p = Path(path)
    if not p.exists():
        return None
    mime, _ = mimetypes.guess_type(str(p))
    with open(p, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
    return f"data:{mime or 'image/png'};base64,{b64}"

# Build once, only if not already built
if "logo_data_uri" not in globals():
    logo_data_uri = to_data_uri("sky_vip_logo.png")  # 👈 rename to your real filename
# ---- tiny helper to show the exec logo, centered ----
from pathlib import Path
import streamlit as st



# --- Streamlit page config (keep this near the top) ---
st.set_page_config(
    page_title="Sky VIP Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed",
)
st.markdown("""
<style>
/* compact, centered content */
div.block-container{
  max-width: 1200px;   /* tweak 1000–1280 if you like */
  padding-top: 0.75rem;
  padding-bottom: 1.5rem;
  margin: 0 auto;      /* center */
}
</style>
""", unsafe_allow_html=True)
# --- Global compact layout + utility classes ---
st.markdown("""
<style>
/* Center the page and cap width */
div.block-container{max-width:1100px;margin:0 auto;padding-top:.5rem;padding-bottom:1rem;}
/* Headings + spacing */
h2, h3 { margin: .25rem 0 .5rem 0 !important; }
hr { margin: .6rem 0 !important; opacity:.25; }
/* Tighten tab header spacing */
.css-13ejsyy, .stTabs [data-baseweb="tab-list"] { gap: .25rem !important; }
/* KPI band */
.kpi {border:1px solid rgba(148,163,184,.25); border-radius:12px; padding:10px 14px; background:rgba(2,6,23,.02);}
.kpi h3 {font-size:0.9rem; font-weight:600; margin:0 0 .25rem 0;}
.kpi .val {font-size:1.6rem; font-weight:800; line-height:1; margin:0;}
/* Budget pill */
.budget-pill {background:linear-gradient(135deg,#0ea5e9,#2563eb);color:#fff;padding:10px 16px;border-radius:14px;
  font-weight:700; display:inline-block; min-width:120px; text-align:center; box-shadow:0 8px 22px rgba(37,99,235,.16);}
.small {color:#64748b; font-size:.9rem;}
.rowpad {padding:.25rem 0;}
/* Tighter columns default */
section.main > div {padding-top:.25rem;}
</style>
""", unsafe_allow_html=True)


# --- Global display defaults ---
pd.options.display.float_format = "{:,.2f}".format
np.random.seed(42)

def now_uk() -> datetime:
    """Return tz-aware 'now' in UK timezone."""
    return datetime.now(UK_TZ)

# --- Secrets / env ---
OPENAI_API_KEY: str = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY", ""))
from datetime import datetime, timezone
from datetime import datetime as dt  # alias for convenience

# ---------- Teams webhook helpers (fixed) ----------
from typing import Optional, Dict, Any
import requests
from requests.exceptions import RequestException

def _get_teams_webhook() -> Optional[str]:
    """
    Reads the Teams webhook from Streamlit secrets.
    Supports either:
      TEAMS_WEBHOOK_URL = "https://..."
    or:
      [teams]
      webhook_url = "https://..."
    """
    s = st.secrets

    # flat key
    if "TEAMS_WEBHOOK_URL" in s:
        return str(s["TEAMS_WEBHOOK_URL"]).strip()

    # table key
    if "teams" in s:
        for k in ("webhook_url", "WEBHOOK_URL", "url", "URL"):
            if k in s["teams"]:
                return str(s["teams"][k]).strip()

    return None


def send_to_teams(title: str,
                  text: str,
                  facts: Optional[Dict[str, Any]] = None,
                  button_text: Optional[str] = None,
                  button_url: Optional[str] = None,
                  color: str = "0078D4") -> None:
    """Send Teams Incoming Webhook (raise if fails)."""
    url = _get_teams_webhook()
    if not url:
        st.error("⚠️ No Teams webhook configured in secrets.")
        return

    payload: Dict[str, Any] = {
        "@type": "MessageCard",
        "@context": "https://schema.org/extensions",
        "summary": title,
        "themeColor": color,
        "title": title,
        "text": text,
    }
    if facts:
        payload.setdefault("sections", []).append({
            "facts": [{"name": k, "value": str(v)} for k, v in facts.items()],
            "markdown": True,
        })
    if button_text and button_url:
        payload["potentialAction"] = [{
            "@type": "OpenUri",
            "name": button_text,
            "targets": [{"os": "default", "uri": button_url}],
        }]

    # Let this raise if it fails so you see it
    requests.post(url, json=payload, timeout=10).raise_for_status()

def send_login_card(user_name: str,
                    user_team: Optional[str] = None,
                    tab_url: Optional[str] = None) -> bool:
    """Always send login notification to Teams. Returns True on success."""
    url = _get_teams_webhook()
    if not url:
        st.error("⚠️ No Teams webhook configured.")
        return False

    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())[:8]
    session_id = st.session_state.session_id

    now_utc = datetime.now(timezone.utc)
    # IMPORTANT: use the *correct* var name you defined earlier
    now_local = now_utc.astimezone(UK_TZ)
    uk_str  = now_local.strftime("%A, %d %B %Y at %H:%M %Z")
    utc_str = now_utc.strftime("%Y-%m-%d %H:%M UTC")

    facts = {
        "User": user_name or "—",
        "Team": user_team or "—",
        "Local time": uk_str,
        "UTC": utc_str,
        "Session": session_id,
        "App": "Visit Dashboard",
    }

    payload: Dict[str, Any] = {
        "@type": "MessageCard",
        "@context": "https://schema.org/extensions",
        "summary": f"Login — {user_name or '—'}",
        "themeColor": "0078D4",
        "title": "🔐 New login",
        "text": "A user has signed in.",
        "sections": [{
            "activityTitle": f"**{user_name or '—'}**",
            "activitySubtitle": "Visit Dashboard",
            "facts": [{"name": k, "value": str(v)} for k, v in facts.items()],
            "markdown": True
        }]
    }
    if tab_url:
        payload["potentialAction"] = [{
            "@type": "OpenUri",
            "name": "Open dashboard",
            "targets": [{"os": "default", "uri": tab_url}],
        }]

    import requests
    requests.post(url, json=payload, timeout=10).raise_for_status()
    return True

DEBUG_TEAMS = True  # flip to True temporarily

if DEBUG_TEAMS:
    with st.sidebar:
        st.markdown("**Debug (Teams)**")
        st.write("Webhook configured:", bool(_get_teams_webhook()))
        if st.button("Send Teams test"):
            send_to_teams("✅ Test", "Manual test from sidebar")
            st.success("Sent")


# ===== SECTION 3: Session State Defaults =====

# Choose ONE start screen for the app.
# Options we’ll use later: "login", "instruction_guide", "area_selection", "team_overview", "ai"
START_SCREEN = "instruction_guide"

DEFAULT_STATE = {
    "screen": START_SCREEN,
    "quick_summary": False,
    "user_df": None,
    "user_file_name": None,
    "selected_dataset": None,
    "authenticated": False,
    "ai_chat": [],                 # chat history for AI page
    "login_notified": False,       # Teams login card sent?
    # helpful extras we’ll use soon:
    "username": None,
    "user_team": None,
    "preferred_team_tab": None,    # e.g. "VIP North"
}

# Apply defaults once
for k, v in DEFAULT_STATE.items():
    st.session_state.setdefault(k, v)

# One short, stable session id (useful for logging/Teams, etc.)
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]

# Validate the current screen to avoid typos from previous runs
VALID_SCREENS = {"login", "instruction_guide", "area_selection", "team_overview", "team_graphs","ai","instructions_page", "operational_area", "highlands_islands", "exec_overview", "sky_retail", "suggestions", "engineer_kpi", "team_engineers"}
if st.session_state.screen not in VALID_SCREENS:
    st.session_state.screen = START_SCREEN

# ===== SECTION 4: Small Utilities + Cached Loaders =====
from typing import Optional, Dict, Any, Iterable, Union
from pathlib import Path

# --- Lottie loader (cached, safe) ---
@st.cache_data(show_spinner=False)
def load_lottie_url(url: str) -> Optional[Dict[str, Any]]:
    """Fetch and cache a Lottie JSON by URL with basic error handling."""
    try:
        import requests  # local import keeps optional dep light
        resp = requests.get(url, timeout=8)
        resp.raise_for_status()
        data = resp.json()
        # Quick sanity check: many lotties have 'v' and 'layers' keys
        if isinstance(data, dict) and ("layers" in data or "assets" in data):
            return data
        return None
    except Exception:
        return None

# --- CSS injector (use this once on app start if needed) ---
def inject_css(css: str) -> None:
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)

# Handy default styles (optional)
BASE_CSS = """
/* tighten top padding and hide deploy button if desired */
.block-container { padding-top: 1.2rem; }
"""
st.markdown(
    "<style>.block-container{ padding-top: 2rem !important; }</style>",
    unsafe_allow_html=True,
)

# --- Number/time formatters you'll reuse for KPIs ---
def fmt_int(x: Union[int, float, None]) -> str:
    try:
        return f"{int(x):,}"
    except Exception:
        return "—"

def fmt_money(x: Union[int, float, None], symbol: str = "£") -> str:
    try:
        return f"{symbol}{float(x):,.0f}"
    except Exception:
        return "—"

def fmt_hhmm(minutes: Optional[float]) -> str:
    """Format minutes as HH:MM; returns — for None/NaN."""
    try:
        if minutes is None:
            return "—"
        m = int(round(float(minutes)))
        h, mm = divmod(m, 60)
        return f"{h:02d}:{mm:02d}"
    except Exception:
        return "—"

# --- Cache: Load Excel file and tag with source/team ---
@st.cache_data(show_spinner=False)
def load_oracle_file(
    path: Union[str, Path],
    source_label: str,
    *,
    sheet_name: Union[str, int, None] = None,
    usecols: Optional[Iterable[str]] = None,
    nrows: Optional[int] = None,
    dtype: Optional[Dict[str, str]] = None,
) -> pd.DataFrame:
    """
    Load a single Oracle-like Excel into a DataFrame and add 'Source' and 'Team' cols.
    - sheet_name: pass if not the first sheet
    - usecols/nrows: speed up dev/testing
    - dtype: force specific columns if needed
    """
    path = Path(path)
    if not path.exists():
        st.warning(f"File not found: {path}")
        return pd.DataFrame()

    try:
        df = pd.read_excel(
            path,
            sheet_name=sheet_name,
            usecols=usecols,
            nrows=nrows,
            dtype=dtype,
            engine=None,  # let pandas choose
        )
        # Ensure DataFrame (in case of sheet dict)
        if isinstance(df, dict):
            # take the first sheet by default
            df = next(iter(df.values()))
    except Exception as e:
        st.error(f"Failed to read {path.name}: {e}")
        return pd.DataFrame()

    # Normalize columns a bit (optional, safe no-ops if already present)
    if "Activity Status" in df.columns:
        df["Activity Status"] = df["Activity Status"].astype(str)

    df["Source"] = source_label
    df["Team"] = source_label  # keep your current convention
    return df

# --- Cache: Load logo as data URI (ready for <img src="...">) ---
@st.cache_data(show_spinner=False)
def load_logo_data_uri(logo_path: Union[str, Path] = "sky_vip_logo.png") -> Optional[str]:
    p = Path(logo_path)
    if not p.exists():
        return None
    ext = p.suffix.lower()
    mime = "image/png" if ext == ".png" else ("image/jpeg" if ext in {".jpg", ".jpeg"} else "image/svg+xml" if ext == ".svg" else "application/octet-stream")
    try:
        data = p.read_bytes()
        b64 = base64.b64encode(data).decode("ascii")
        return f"data:{mime};base64,{b64}"
    except Exception:
        return None

logo_data_uri = load_logo_data_uri()
# (later, you can use: st.markdown(f'<img src="{logo_data_uri}" height="36">', unsafe_allow_html=True))

# ===== SECTION 5: Oracle Data Loaders (aligned columns) =====
from pathlib import Path
from typing import Dict, Any
import pandas as pd

# ===== SECTION 5: Oracle Data Loaders (fast + cached) =====
from pathlib import Path
from typing import Dict, Tuple
import pandas as pd
import streamlit as st

# === Budget data (budgets.csv + expenses.csv) ===
@st.cache_data(show_spinner=False)
def load_budget_sources(
    budget_path: str = "budgets.csv",
    expenses_path: str = "expenses.csv",
) -> tuple[pd.DataFrame, pd.DataFrame]:
    def _read_any(path):
        try:
            if path.lower().endswith((".xlsx", ".xls")):
                return pd.read_excel(path)
            return pd.read_csv(path)
        except Exception:
            return pd.DataFrame()

    bud = _read_any(budget_path)
    exp = _read_any(expenses_path)

    for df in (bud, exp):
        if not df.empty:
            df.columns = (
                df.columns.astype(str)
                .str.replace("\u00A0", " ", regex=False)
                .str.strip()
            )
            # normalise a Team column if present under a different name
            low = {c.lower(): c for c in df.columns}
            team_col = next((low[k] for k in ("team", "group", "department") if k in low), None)
            if team_col and team_col != "Team":
                df.rename(columns={team_col: "Team"}, inplace=True)
            elif "Team" not in df.columns:
                df["Team"] = "All"
    return bud, exp


def _first_existing_col(df: pd.DataFrame, *names: str) -> str | None:
    low = {c.lower(): c for c in df.columns}
    for n in names:
        if n.lower() in low:
            return low[n.lower()]
    return None


def compute_team_budget_metrics(team_name: str) -> tuple[float, float, float, float | float("nan")]:
    # put this immediately after the compute_team_budget_metrics(...) function
    compute_team_budget_kpis = compute_team_budget_metrics

    
    # Alias for backwards compatibility with earlier instructions
    

    import numpy as np

    bud, exp = load_budget_sources()

    # --- Budgets (Allocated) ---
    alloc = 0.0
    if not bud.empty:
        col_alloc = _first_existing_col(bud, "Budget Allocated", "Allocated", "Budget", "Total Budget")
        b = bud.copy()
        # filter to team (exact then contains)
        sel = b["Team"].astype(str).str.lower() == str(team_name).lower()
        if not sel.any():
            sel = b["Team"].astype(str).str.lower().str.contains(str(team_name).lower(), na=False)
        b = b[sel] if sel.any() else b
        if col_alloc:
            alloc_series = b[col_alloc]
            # use your existing currency cleaner if available
            try:
                nums = _money_to_float(alloc_series)  # defined earlier in your file
            except Exception:
                nums = pd.to_numeric(
                    alloc_series.astype(str).str.replace(r"[£$,]", "", regex=True).str.replace(r"[^\d\.\-]", "", regex=True),
                    errors="coerce"
                )
            alloc = float(nums.sum(skipna=True)) if nums.notna().any() else 0.0

    # --- Expenses (Used) ---
    used = 0.0
    if not exp.empty:
        col_used = _first_existing_col(exp, "Budget Used", "Used", "Spent", "Spend", "Amount", "Cost", "Expenses", "Actuals")
        e = exp.copy()
        sel = e["Team"].astype(str).str.lower() == str(team_name).lower()
        if not sel.any():
            sel = e["Team"].astype(str).str.lower().str.contains(str(team_name).lower(), na=False)
        e = e[sel] if sel.any() else e
        if col_used:
            try:
                nums = _money_to_float(e[col_used])
            except Exception:
                nums = pd.to_numeric(
                    e[col_used].astype(str).str.replace(r"[£$,]", "", regex=True).str.replace(r"[^\d\.\-]", "", regex=True),
                    errors="coerce"
                )
            used = float(nums.sum(skipna=True)) if nums.notna().any() else 0.0

    remaining = alloc - used
    pct_used = (used / alloc * 100.0) if alloc not in (0, None) else float("nan")
    return float(alloc), float(used), float(remaining), (float(pct_used) if pd.notna(pct_used) else float("nan"))


# Filenames (exact names you uploaded)
FILES: Dict[str, str] = {
    "VIP North":  "VIP North Oracle Data.xlsx",
    "VIP South":  "VIP South Oracle Data.xlsx",
    "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
}
budgets_df = load_budgets_df()


def _normalize_cols(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = (
        df.columns.astype(str)
        .str.replace("\u00A0", " ", regex=False)
        .str.strip()
    )
    return df

def _maybe_newer(a: Path, b: Path) -> bool:
    """Return True if file a is newer than file b, or b doesn't exist."""
    if not b.exists(): 
        return True
    try:
        return a.stat().st_mtime > b.stat().st_mtime
    except Exception:
        return True

@st.cache_data(show_spinner=False)
def _read_fast(path_xlsx: str) -> pd.DataFrame:
    """
    Fast reader:
      - If a sibling .parquet is present and up-to-date -> read_parquet
      - Else read Excel once, write parquet for next runs
    Cache ensures we don't re-read on each rerun.
    """
    x_path = Path(path_xlsx)
    if not x_path.exists():
        # Graceful: return empty DF with a note
        st.warning(f"File not found: {x_path.name}")
        return pd.DataFrame()

    p_path = x_path.with_suffix(".parquet")

    # Prefer parquet if it's up-to-date (or Excel hasn't changed)
    if p_path.exists() and not _maybe_newer(x_path, p_path):
        try:
            return pd.read_parquet(p_path)
        except Exception:
            pass  # fall back to excel

    # Read Excel (slower), then save parquet for next time
    try:
        # dtype_backend="pyarrow" can speed up + lower memory on pandas >= 2.0
        df = pd.read_excel(x_path)  # keep simple/compatible
        df = _normalize_cols(df)
        try:
            df.to_parquet(p_path, index=False)
        except Exception:
            # If parquet write fails (missing engine), no worries—still return df
            pass
        return df
    except Exception as e:
        st.error(f"Couldn't read {x_path.name}: {e}")
        return pd.DataFrame()

@st.cache_data(show_spinner=False)
def load_all_frames() -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame]]:
    """
    Loads and returns:
      - df_all: concatenation of all four teams (with Team column)
      - per_team: dict of team_name -> dataframe
    Everything cached so your UI is snappy after first load.
    """
    per_team: Dict[str, pd.DataFrame] = {}
    frames = []
    for team, fn in FILES.items():
        df = _read_fast(fn)
        if not df.empty:
            df = df.copy()
            df["Team"] = team
            per_team[team] = df
            frames.append(df)

    df_all = pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()
    return df_all, per_team

# Public helpers for the rest of your app
def get_df_all() -> pd.DataFrame:
    df_all, _ = load_all_frames()
    return df_all

def get_team_df(team_name: str) -> pd.DataFrame:
    df_all, per_team = load_all_frames()
    return per_team.get(team_name, pd.DataFrame())
# --- logo helper (put near your other small helpers) ---
def show_logo(filename: str, width_px: int = 1100):
    """Show a local image by filename, at a fixed pixel width."""
    import os, streamlit as st
    path = os.path.join(os.getcwd(), filename)
    if os.path.exists(path):
        st.image(path, use_container_width=False, width=width_px)



# Optional: root data directory via secrets, else current working dir
DATA_DIR = Path(st.secrets.get("DATA_DIR", "."))

# Canonical schema we expect/use across the app
CANONICAL_COLS = [
    "Name",
    "Date",
    "Day",                     # may be missing (e.g., Tier 2 South) — we’ll add empty
    "Month",
    "Activate",
    "Deactivate",
    "Total Working Time",
    "Overtime",                # Tier 2 North has "Overtime " — trimmed to this
    "Activity Status",
    "Visit Type",
    "Visit Number",
    "Visit Notes",
    "Start",
    "End",
    "Travel Time",
    "Total Time",
    "Total Time (Inc Travel)",
    "Total Value",
    "Total Cost Inc Travel",
    "Difference",
    "Quarter",
    "week",
    "Sky Retail Stakeholder",
    # plus we add:
    "Team",
    "Source",
]

def _normalize_cols(df: pd.DataFrame) -> pd.DataFrame:
    # Trim whitespace from all column names
    df.columns = [str(c).strip() for c in df.columns]

    # Fix known variant(s)
    rename_map = {
        "Overtime ": "Overtime",   # trailing space found in Tier 2 North
    }
    df = df.rename(columns=rename_map)

    # Ensure required columns exist; if missing, create empty
    for col in CANONICAL_COLS:
        if col not in df.columns:
            # sensible defaults per type
            if col in {"Total Value", "Total Cost Inc Travel"}:
                df[col] = pd.to_numeric([], errors="coerce")
            elif col in {"Date", "Start", "End"}:
                df[col] = pd.to_datetime([], errors="coerce")
            else:
                df[col] = pd.Series(dtype="object")

    # Reorder for consistency (keeping only canonical; extras are dropped here intentionally)
    df = df[CANONICAL_COLS]
    return df

def _coerce_types(df: pd.DataFrame) -> pd.DataFrame:
    # Dates
    for c in ("Date", "Start", "End"):
        if c in df.columns:
            df[c] = pd.to_datetime(df[c], errors="coerce")

    # Numerics
    for c in ("Total Value", "Total Cost Inc Travel"):
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce")

    # Text fields
    for c in (
        "Name", "Day", "Month", "Activate", "Deactivate", "Total Working Time",
        "Overtime", "Activity Status", "Visit Type", "Visit Number", "Visit Notes",
        "Travel Time", "Total Time", "Total Time (Inc Travel)", "Difference",
        "Quarter", "week", "Sky Retail Stakeholder", "Team", "Source",
    ):
        if c in df.columns:
            df[c] = df[c].astype("string").replace({"<NA>": None})

    return df

@st.cache_data(show_spinner=True)
def load_all_oracle(files_map: Dict[str, str], data_dir: Path) -> Dict[str, Any]:
    """
    Load all team workbooks, normalize to a single schema, and return:
      - combined: pd.DataFrame (all teams)
      - per_team: Dict[str, pd.DataFrame]
      - errors:   List[str] of warnings
    """
    per_team: Dict[str, pd.DataFrame] = {}
    errors: list[str] = []

    for team, rel in files_map.items():
        path = (data_dir / rel).resolve()
        if not path.exists():
            errors.append(f"Missing file for {team}: {path.name}")
            continue

        df = load_oracle_file(path, team)  # from Section 4 (adds Team/Source)
        if df.empty:
            errors.append(f"No rows in {path.name} ({team})")
            continue

        # Normalize schema and coerce types
        df = _normalize_cols(df)
        df = _coerce_types(df)
        per_team[team] = df

    if not per_team:
        return {"combined": pd.DataFrame(columns=CANONICAL_COLS), "per_team": {}, "errors": errors}

    combined = pd.concat(per_team.values(), ignore_index=True)
    combined = _normalize_cols(combined)
    combined = _coerce_types(combined)

    return {"combined": combined, "per_team": per_team, "errors": errors}

# --- Load once and expose to the app ---
load_result = load_all_oracle(FILES, DATA_DIR)
combined_oracle_df: pd.DataFrame = load_result["combined"]
oracle_by_team: Dict[str, pd.DataFrame] = load_result["per_team"]
load_errors: list[str] = load_result["errors"]

if load_errors:
    with st.expander("⚠️ Data load warnings", expanded=False):
        for msg in load_errors:
            st.write("•", msg)

if combined_oracle_df.empty:
    st.error("No data loaded from Oracle files. Please check file paths and names.")
else:
    combined_oracle_df = combined_oracle_df.reset_index(drop=True)



    
# --- GLOBAL USERS DICTIONARY ---
users = {
    "Matt Hughes":     {"password": "mattpw", "team": "VIP Team North"},
    "Andy Holmes":     {"password": "andypw", "team": "Tier 2 North"},
    "Steve Paisley":   {"password": "stevepw", "team": "Tier 2 South"},
    "Branks Krneta":   {"password": "brankspw", "team": "TAP's Team"},
    "Chris Woods":     {"password": "chrispw", "team": "Highlands & Islands"},
    "Rachel Wylie":    {"password": "rachelpw", "team": "Sky Retail"},
    "Darryl Fuller":   {"password": "darrylpw", "team": "Highlands & Islands"},
    "Christian Boyce": {"password": "chrisbpw", "team": "VIP Team South"},
    "Rob Mangar":      {"password": "robpw", "team": "Operational Team"},
    "Dan Homewood":    {"password": "danpw", "team": "Finance Team"},
}



# ===== SECTION 6: Styling + Single Login (with named users & passwords) =====

# ---- LIGHT THEME (matches your old app look, full width, responsive) ----
LIGHT_CSS = """
/* App background + base typography */
html, body, .stApp {
  background: #f5f7fb !important;
  color: #111827 !important;
  font-family: 'Inter','Segoe UI', system-ui, -apple-system, sans-serif;
}

/* Make the main area full width with nice side gutters */
.block-container {
  max-width: 100%;
  padding: 0 24px;
}

/* Hide Streamlit's default top gap */
header[data-testid="stHeader"] { background: transparent; }

/* ---------- Uniform 'menu tiles' (buttons) ---------- */
.menu-row { margin-top: 8px; }
.stButton { margin: 6px 8px !important; } /* consistent gaps */
.stButton > button {
  min-width: 200px;
  height: 72px;
  border-radius: 18px;
  padding: .6rem 1rem;
  white-space: normal;
  line-height: 1.15;
  word-break: keep-all;
  background: linear-gradient(90deg,#0ea5e9,#2563eb) !important;
  color:#fff !important; border:0;
  box-shadow:0 6px 18px rgba(2,32,71,.18);
  font-weight:700;
}
.stButton > button:hover {
  filter:brightness(1.06);
  box-shadow:0 10px 24px rgba(2,32,71,.22);
}

/* Responsive sizing for tiles so zooming doesn't break rows */
@media (max-width: 1280px) {
  .stButton > button { min-width: 170px; height: 64px; }
}
@media (max-width: 992px) {
  .stButton > button { min-width: 150px; height: 60px; }
}
@media (max-width: 640px) {
  .stButton > button { min-width: 100%; height: auto; }
}

/* ---------- Login hero (logo + strapline) ---------- */
.logo { display:flex; justify-content:center; margin:12px 0 4px 0; }
.logo img { max-width: 560px; width: 100%; height:auto; }
.login-header-wrapper { display:flex; justify-content:center; margin:8px 0 12px; }
.login-header { font-size:20px; font-weight:700; color:#111827; }

/* Centered white login card */
.login-card {
  max-width: 860px; margin: 8px auto 0 auto;
  background:#fff; border:1px solid #e6e8eb; border-radius:16px;
  padding:20px 22px; box-shadow:0 18px 40px rgba(2,32,71,.06);
}

/* Inputs: light fields */
div[data-baseweb="select"] > div { background:#f8fafc; border-radius:10px; }
.stTextInput > div > div > input { background:#f8fafc; border-radius:10px; }
.stPassword > div > div > input { background:#f8fafc; border-radius:10px; }

/* Headings */
h1,h2,h3,h4 { color:#0f172a !important; font-weight:800; }

/* Sidebar */
section[data-testid="stSidebar"] { background:#ffffff !important; border-right:1px solid #eef1f5; }

/* ---------- KPI 'pill' cards: WHITE background + DARK text ---------- */
.kpi-pill {
  background:#ffffff !important;
  color:#0f172a !important;
  border-radius:24px;
  padding:18px 20px;
  border:1px solid #e6e8eb;
  box-shadow:0 14px 28px rgba(2,32,71,.10);
}
/* Ensure all text inside is dark */
.kpi-pill h1, .kpi-pill h2, .kpi-pill h3, .kpi-pill h4,
.kpi-pill p, .kpi-pill li, .kpi-pill strong, .kpi-pill span, .kpi-pill em {
  color:#0f172a !important;
}
.kpi-pill strong { font-weight:800; }

/* ---------- Expander styling (light) ---------- */
div[data-testid="stExpander"] {
  background:#ffffff; border:1px solid #e6e8eb; border-radius:12px;
}
div[data-testid="stExpander"] > div[role="button"] {
  background:#ffffff; border-bottom:1px solid #e6e8eb; border-radius:12px 12px 0 0;
}
"""

# Inject the light theme (make sure any earlier dark injection is removed/commented out)
inject_css(LIGHT_CSS)




# ---- Users: EXACT names, passwords, and teams you provided ----
USERS = {
    "Matt Hughes":     {"password": "mattpw",  "team": "VIP Team North"},
    "Andy Holmes":     {"password": "andypw",  "team": "Tier 2 North"},
    "Steve Paisley":   {"password": "stevepw", "team": "Tier 2 South"},
    "Branks Krneta":   {"password": "brankspw","team": "TAP's Team"},
    "Chris Woods":     {"password": "chrispw", "team": "Highlands & Islands"},
    "Rachel Wylie":    {"password": "rachelpw","team": "Sky Retail"},
    "Darryl Fuller":   {"password": "darrylfpw","team": "Highlands & Islands"},  # keep as given
    "Christian Boyce": {"password": "chrisbpw","team": "VIP Team South"},
    "Rob Mangar":      {"password": "robpw",   "team": "Operational Team"},
    "Mark Wilson":     {"password": "markpw",  "team": "Operational Leadership"},
    "Dylan Cleverly":  {"password": "dylancpw","team": "Sky Business"},
    "Dan Homewood":    {"password": "danpw",   "team": "Finance Team"},
}
NAME_LIST = list(USERS.keys())

def _normalize_team(label: str | None) -> str | None:
    if not label: return None
    label = label.strip()
    mapping = {
        "VIP Team North": "VIP North",
        "VIP Team South": "VIP South",
        "Tier 2 North": "Tier 2 North",
        "Tier 2 South": "Tier 2 South",
        "Operational Team": "Operational Team",
        "Highlands & Islands": "Highlands & Islands",
        "Sky Retail": "Sky Retail",
        "TAP's Team": "TAP's Team",
        "Finance Team": "Finance Team",
        "Sky Business": "Sky Business"
    }
    return mapping.get(label, label)

def _render_login_hero() -> None:
    """Top logo + animated strapline + Lottie art."""
    # Prefer data-URI logo (from Section 4)
    logo_src = logo_data_uri if "logo_data_uri" in globals() else None

    if logo_src:
        st.markdown(
            f'<div class="logo"><img src="{logo_src}" alt="Sky VIP"></div>',
            unsafe_allow_html=True,
        )

    st.markdown(
        '<div class="login-header-wrapper"><div class="login-header">'
        '🔐 Login to Visit Insights Dashboard'
        '</div></div>',
        unsafe_allow_html=True,
    )

    # Lottie animation
    lottie = load_lottie_url("https://assets2.lottiefiles.com/packages/lf20_jcikwtux.json")
    if lottie:
        try:
            from streamlit_lottie import st_lottie
            st_lottie(lottie, height=260, key="login_anim")
        except Exception:
            pass


def _show_welcome_modal(name: str, team: str | None) -> None:
    try:
        from streamlit_modal import Modal
    except Exception:
        return
    import random
    quotes = [
        "Success is the sum of small efforts, repeated day in and day out.",
        "Your hard work is making a difference!",
        "Teamwork divides the task and multiplies the success.",
        "Keep up the amazing work!",
    ]
    modal = Modal(key="welcome_modal", title=f"👋 Welcome, {name}!", padding=28)
    if modal.is_open():
        st.write(f"Welcome to the dashboard, **{name}**!")
        if team:
            st.write(f"You're viewing **{team}**. {random.choice(quotes)}")
        if st.button("Let's go!"):
            modal.close()

def login_screen() -> None:
    """Single login UI. Name select outside the form so password appears immediately."""
    _render_login_hero()
    st.markdown("### 🔐 Sign in")
    st.markdown("<p style='text-align:center;'>Please choose your name to continue.</p>", unsafe_allow_html=True)

    # Name is outside the form so changing it reruns the script
    selected_name = st.selectbox("Choose Your Name", ["-- Select --"] + NAME_LIST, index=0, key="login_name")

    if selected_name and selected_name != "-- Select --":
        hint_team = USERS[selected_name]["team"]
        st.caption(f"Team: {hint_team}")

    # The ONLY form (password + submit)
    with st.form("login_form", clear_on_submit=False):
        pwd = st.text_input(
            "Enter your password",
            type="password",
            key="user_pw",
            disabled=(selected_name in (None, "", "-- Select --"))
        )
        submitted = st.form_submit_button("Sign in")

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
    b1, b2 = st.columns([1, 1], gap="large")

    with b1:
        if st.button("📘 Instructions", key="login_instructions", use_container_width=True):
            st.session_state["screen"] = "instructions_page"
            st.rerun()

    with b2:
        if st.button("💡 Suggestion Box (no login)", key="login_suggestions", use_container_width=True):
            st.session_state["screen"] = "suggestions"
            st.rerun()


    if submitted:
        # Validate selection & password
        if not selected_name or selected_name == "-- Select --":
            st.error("Please select your name.")
            return
        expected = USERS[selected_name]["password"]
        if not pwd or pwd.strip() != expected:
            st.error("Incorrect password for this user.")
            return

        # Success: remember identity + team (raw + normalized)
        st.session_state.authenticated = True
        st.session_state.username = selected_name
        raw_team = USERS[selected_name]["team"]              # e.g. "VIP Team North"
        st.session_state.team = raw_team                     # original label
        st.session_state.user_team = _normalize_team(raw_team)  # e.g. "VIP North"

        # Preselect Team Overview tab if it's one of the 4 core teams
        core_tabs = {"VIP North", "VIP South", "Tier 2 North", "Tier 2 South"}
        st.session_state.preferred_team_tab = (
            st.session_state.user_team if st.session_state.user_team in core_tabs else None
        )

        # Send Teams notification on EVERY successful login (before rerun)
        ok = send_login_card(
            user_name=st.session_state.username,
            user_team=st.session_state.user_team,
            tab_url=None,
        )
        if ok:
            st.toast("Teams login notification sent ✅", icon="✅")
        # Under (or next to) your Sign in button in login_screen()
        c_s1, c_s2, c_s3 = st.columns([1,1,1])
        with c_s2:
            if st.button("💡 Suggestion Box (no login)", key="goto_suggestions_public", use_container_width=True):
                st.session_state.screen = "suggestions"
                st.rerun()

        # Route to home and rerun once
        u = (st.session_state.get("username") or "").strip().lower()
        team_norm = (st.session_state.get("user_team") or "").strip().lower()

        # 1) Rob -> Ops Area
        if u.startswith("rob"):
            st.session_state.role = "operations_manager"
            st.session_state.screen = "operational_area"
        
            import pandas as pd

        # --- inside login_screen(), after you have validated the password ---
        u = (st.session_state.get("username") or "").strip().lower()
        team_norm = (st.session_state.get("user_team") or "").strip().lower()

        # Rob -> Ops menu
        if u.startswith("rob"):
            st.session_state.role = "operations_manager"
            st.session_state.screen = "operational_area"
            st.session_state.op_area_section = "menu"        # start on the ops menu

        elif u == "dylan cleverly":
            st.session_state.role = "operations_manager"
            st.session_state.screen = "operational_area"
            st.session_state.op_area_section = "menu"        # start on the ops menu


        elif u == "mark wilson":
            st.session_state.role = "executive"
            st.session_state.screen = "operational_area"
            st.session_state.op_area_section = "exec_overview"

        # Highlands & Islands
        elif team_norm.startswith("highlands") or u in ("chris woods", "darryl fuller"):
            st.session_state.role = "islands_manager"
            st.session_state.screen = "highlands_islands"

        # ✅ Rachel -> Sky Retail
        elif u == "rachel wylie":
            st.session_state.role = "retail_lead"
            st.session_state.screen = "sky_retail"
        



        # Everyone else -> normal flow
        else:
            st.session_state.role = "manager"
            st.session_state.screen = "instruction_guide"

        st.rerun()

    



# --- Router guard for login (allow public pages without auth) ---
PUBLIC_SCREENS = {"login", "instructions_page", "suggestions"}  # ← added "suggestions"
if not st.session_state.get("authenticated") and st.session_state.get("screen") not in PUBLIC_SCREENS:
    st.session_state["screen"] = "login"



if st.session_state.screen == "login":
    login_screen()
    st.stop()






# Show welcome modal once on the instruction page
if st.session_state.get("authenticated") and st.session_state.screen == "instruction_guide":
    if st.session_state.get("welcome_pending"):
        _show_welcome_modal(st.session_state.username, st.session_state.get("user_team"))
        st.session_state["welcome_pending"] = False



# ===== SECTION 7: Home – big buttons (Instructions / Team Overview / Engineer) =====
if st.session_state.screen == "instruction_guide":
    t1, t2, t3 = st.columns([1, 2, 1])
    with t2:
        # Show VIP/Tier 2 logo above title (only on this page)
        if logo_data_uri:
            st.markdown(
                f"<div style='text-align:center;margin-top:6px;'>"
                f"<img src='{logo_data_uri}' alt='VIP/Tier2 Logo' "
                f"style='max-width:950px;width:100%;height:auto;margin-bottom:10px;'>"
                f"</div>",
                unsafe_allow_html=True,
            )
        st.markdown(
        """
        <div style="text-align:center; margin-top: 6px; white-space: nowrap;">
          <h1 style="display:inline-block; margin:0; font-size:2.2em;">
            Welcome to the Sky Orbit Dashboard 
          </h1>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Big button styles
    st.markdown("""
    <style>
      #home-choices .stButton>button{
        width:100%; height:230px; border-radius:24px; font-size:2rem; font-weight:800; color:#fff;
        box-shadow:0 10px 28px rgba(0,0,0,.25); transition:transform .2s, box-shadow .2s, filter .2s;
      }
      #home-choices .stButton>button:hover{
        transform:translateY(-6px); box-shadow:0 16px 34px rgba(0,0,0,.35); filter:brightness(1.05);
      }
      #home-choices .graphs>button{
        background:linear-gradient(135deg,#00c6ff,#0072ff) !important;
      }
      #home-choices .team>button{
        background:linear-gradient(135deg,#ff7e5f,#feb47b) !important;
      }
      #home-choices .engineer>button{
        background:linear-gradient(135deg,#43cea2,#185a9d) !important;
      }
    </style>
    """, unsafe_allow_html=True)

    # who is logged in?
    raw_name = (
        st.session_state.get("user_full_name")
        or st.session_state.get("user_name")
        or st.session_state.get("selected_user")
        or st.session_state.get("user")
        or ""
    )
    name_norm = str(raw_name).strip().casefold()
    role = (st.session_state.get("role") or "").lower()

    st.markdown("<div id='home-choices'>", unsafe_allow_html=True)

    # Show two buttons for everyone, add third if manager
    if role == "manager":
        col1, col2, col3 = st.columns(3, gap="large")
    else:
        col1, col2 = st.columns(2, gap="large")
        col3 = None

    # LEFT → Team Overview — Graphs
    with col1:
        st.markdown("<div class='graphs'>", unsafe_allow_html=True)
        if st.button("📊 Team Overview — Graphs", key="btn_team_graphs", use_container_width=True):
            st.session_state.screen = "team_graphs"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # MIDDLE → Team Overview
    with col2:
        st.markdown("<div class='team'>", unsafe_allow_html=True)
        if st.button("👥 Team Overview", key="go_team", use_container_width=True):
            st.session_state.preferred_team_tab = st.session_state.get("user_team")
            st.session_state.screen = "team_overview"
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # RIGHT (managers only) → Engineer
    if role == "manager" and col3 is not None:
        with col3:
            st.markdown("<div class='engineer'>", unsafe_allow_html=True)
            if st.button("🧑‍🔧 Engineer", key="btn_team_engineers", use_container_width=True):
                st.session_state.screen = "team_engineers"
                st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)




if st.session_state.screen == "instruction_guide":
  

    # --- Back to Login button (top-left or bottom-left, your choice)
    if st.button("🔑 Back to Login"):
        st.session_state.authenticated = False
        st.session_state.username = None
        st.session_state.team = None
        st.session_state.screen = "login"
        st.rerun()




# ===== SECTION 7a: Instructions page =====
if st.session_state.screen == "instructions_page":

    # --- Logo (same as home) ---
    # If you have logo_data_uri, use it; otherwise point to a file like 'vip_tier2_logo.png'
    if "logo_data_uri" in globals() and logo_data_uri:
        st.markdown(
            f"""
            <div style="text-align:center;">
              <img src="{logo_data_uri}" alt="VIP/Tier2 Logo"
                   style="max-width:1150px;height:auto;margin-bottom:10px;">
            </div>
            """,
            unsafe_allow_html=True,
        )

    # --- One-line heading (no wrapping) ---
    st.markdown(
        """
        <div style="text-align:center; margin-top: 6px; white-space: nowrap;">
          <h1 style="display:inline-block; margin:0; font-size:2.2em;">
            Welcome to the Sky Orbit Dashboard 
          </h1>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Your existing instructions content follows...
    st.subheader("📘 Instructions Guide")

    # --- your card grid content exactly as you had it ---
    sections = [
        {"title":"What is Sky Orbit?","content":"Sky Orbit is an interactive dashboard built especially for our teams that brings together Oracle visit data from multiple teams (VIP North, VIP South, Tier 2 North, Tier 2 South). It helps you explore, analyse, and visualize this data easily with features like secure login, AI-powered chat, detailed KPIs, trends, and forecasting."},
        {"title":"Main Menu","content":"From here, pick the area you want to explore — like Operational Area, Dashboards, AI Chat (Sky Orbit AI), Suggestions, Forecasts, Sky Retail, Sky Business, and more. Just click the button to jump in."},
        {"title":"Upload Your Data (Optional)","content":"You can upload your own Excel or CSV file with visit data to explore dynamically. If you don’t upload anything, the app uses default Oracle data combined from the four main teams."},
        {"title":"Sky Retail Area","content":"View detailed KPIs and trends filtered by stakeholders such as Currys, Sky Retail, and EE. See totals for visits and value, average times, monthly trends, and day-of-week breakdowns. Explore charts and tables by stakeholder, team, or engineer."},
        {"title":"Sky Business Area","content":"This filters the data to “Sky Business” visit types, showing KPIs for total visits, values, and completion rates. You get activity breakdown charts, monthly trends, sunburst visuals, forecasts by team, and detailed tables with heatmaps."},
        {"title":"VIP - SB Standby Section","content":"Specialized KPIs and charts for “VIP - SB Standby” visits. Summaries of completed visits, total values, average start/end times, activity status distributions, and monthly counts. Forecasts are included here too."},
        {"title":"SLA Dashboard","content":"Track tickets against SLA buckets like 2h, 4h, 5 day, and 8h targets. View KPIs for total tickets, SLA met/missed counts, and percentages. Visualize ticket volumes by SLA and monthly trends, plus get forecasts for upcoming months."},
        {"title":"Suggestion Box","content":"Submit or delete suggestions through an Excel-based interface. It uses OneDrive-safe temp files with unique keys to avoid data conflicts."},
        {"title":"AI Chat Assistant (“Sky Orbit”)","content":"Ask natural language questions about the visit data. The AI also generates charts when relevant. All chats are logged with timestamps and password-protected access."},
    ]

    if "open_card" not in st.session_state:
        st.session_state.open_card = None

    st.markdown("""
    <style>
      .card-button{background:linear-gradient(135deg,#0099ff 30%,#004488 100%); color:white; border:none; border-radius:20px;
        width:100%; height:110px; font-size:1.05em; font-weight:bold; box-shadow:0 4px 16px rgba(0,0,0,.07);
        margin-bottom:12px; transition:box-shadow .18s, background .22s; cursor:pointer;}
      .card-button:hover{box-shadow:0 8px 30px rgba(0,153,255,.18); background:linear-gradient(135deg,#00c6ff 10%,#004488 90%);}
      .card-content{background:#151d2c; color:#fff; border-radius:18px; padding:16px 20px 10px 20px; margin-top:-8px; min-height:85px; font-size:.99em; box-shadow:0 4px 14px rgba(0,153,255,.07); animation:fadeIn .38s;}
      @keyframes fadeIn{from{opacity:0}to{opacity:1}}
    </style>
    """, unsafe_allow_html=True)

    for row in range(3):
        cols = st.columns(3)
        for col in range(3):
            idx = row*3 + col
            if idx < len(sections):
                with cols[col]:
                    if st.button(sections[idx]["title"], key=f"cardbtn_{idx}", help="Click to expand"):
                        st.session_state.open_card = None if st.session_state.open_card == idx else idx
                    st.markdown(f"""
                        <script>
                          const b=document.querySelector('button[data-testid="baseButton-cardbtn_{idx}"]'); 
                          if(b) b.className += " card-button";
                        </script>
                    """, unsafe_allow_html=True)
                    if st.session_state.open_card == idx:
                        st.markdown(f'<div class="card-content">{sections[idx]["content"]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-header">Important Notes</div>', unsafe_allow_html=True)
    st.markdown("""
    <ul class="important-notes">
      <li>All data and views refresh on app start and can be filtered dynamically.</li>
      <li>Forecasts use simple linear trends based on recent months.</li>
      <li>Some KPIs depend on available columns—missing data will show warnings.</li>
      <li>Uploading your own file replaces the default Oracle data for analysis.</li>
      <li>Time-related metrics require valid time columns.</li>
    </ul>
    """, unsafe_allow_html=True)

    # --- Back button to return home ---
    st.markdown("<br>", unsafe_allow_html=True)  # little spacing

    # inside the Instructions page footer
    if st.button("⬅️ Back to Login", key="back_from_instructions"):
        st.session_state["screen"] = "login"
        st.rerun()

# ✅ Optimized Dataset Loading Block

import pandas as pd
import streamlit as st
import re

# --- Oracle Files ---
oracle_files = {
    "VIP North": "VIP North Oracle Data.xlsx",
    "VIP South": "VIP South Oracle Data.xlsx",
    "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
   
}

@st.cache_data(show_spinner=False)
def load_oracle_files(oracle_files):
    dataframes = []
    for team_name, file in oracle_files.items():
        try:
            df = pd.read_excel(file)
            df["Team"] = team_name
            df.columns = df.columns.str.strip()
            dataframes.append(df)
        except Exception as e:
            st.warning(f"⚠️ Could not load {file}: {e}")
    return dataframes

oracle_dataframes = load_oracle_files(oracle_files)

combined_oracle_df = pd.concat(oracle_dataframes, ignore_index=True) if oracle_dataframes else pd.DataFrame()

# --- Subsets from Oracle ---
vip_north_df = combined_oracle_df[combined_oracle_df["Team"] == "VIP North"]
vip_south_df = combined_oracle_df[combined_oracle_df["Team"] == "VIP South"]
tier2_north_df = combined_oracle_df[combined_oracle_df["Team"] == "Tier 2 North"]
tier2_south_df = combined_oracle_df[combined_oracle_df["Team"] == "Tier 2 South"]

sky_retail_df = combined_oracle_df[
    combined_oracle_df.get("Sky Retail Stakeholder", "").astype(str).str.contains("Sky Retail", case=False, na=False)
] if "Sky Retail Stakeholder" in combined_oracle_df.columns else pd.DataFrame()

highlands_df = combined_oracle_df[
    combined_oracle_df["Visit Type"].astype(str).str.contains("Highlands|Islands", case=False, na=False)
] if "Visit Type" in combined_oracle_df.columns else pd.DataFrame()

# --- Other Datasets ---
@st.cache_data(show_spinner=False)
def load_excel(path):
    return pd.read_excel(path)

try:
    sky_business_df = load_excel("Sky Business.xlsx")
    sky_business_df.columns = sky_business_df.columns.str.strip()
except Exception as e:
    st.warning(f"⚠️ Could not load Sky Business file: {e}")
    sky_business_df = pd.DataFrame()

try:
    call_log_df = load_excel("Call Log Data.xlsx")
    call_log_df.columns = call_log_df.columns.str.strip()
except Exception as e:
    st.warning(f"⚠️ Could not load Call Log Data.xlsx: {e}")
    call_log_df = pd.DataFrame()

# --- Highlands File ---
try:
    highlands_file = pd.ExcelFile("Highlands Islands.xlsx")
    yearly_sheets = [s for s in highlands_file.sheet_names if "Year" in s]
    company_sheets = [s for s in highlands_file.sheet_names if "Company" in s]
    dfs = {sheet: highlands_file.parse(sheet) for sheet in yearly_sheets + company_sheets}
except Exception as e:
    st.error(f"⚠️ Error loading Highlands & Islands data: {e}")
    highlands_file = None
    yearly_sheets = []
    dfs = {}

    # --- Invoices File ---


# --- Combined Dictionary ---
all_tabs = {
    "VIP North": vip_north_df,
    "VIP South": vip_south_df,
    "Tier 2 North": tier2_north_df,
    "Tier 2 South": tier2_south_df,
    "Sky Retail": sky_retail_df,
    "Sky Business": sky_business_df,
    "Call Log": call_log_df,
    "Highlands & Islands": highlands_df,
    
}

# --- Time Formatting Functions ---
def format_time_avg(series):
    clean = series.dropna().astype(str)
    clean = clean[~clean.str.contains("00:00:00|00:00|NaT|nan|None", na=False)]
    times = pd.to_timedelta(clean, errors='coerce').dropna()
    return str((times.mean()).components.hours).zfill(2) + ":" + str((times.mean()).components.minutes).zfill(2) if not times.empty else "00:00"

def format_total_time(series):
    times = pd.to_timedelta(series.dropna().astype(str), errors='coerce')
    total = times.sum()
    return f"{int(total.total_seconds() // 3600):02}:{int((total.total_seconds() % 3600) // 60):02}" if not times.empty else "00:00"
# --- Utility for Highlands weighted averages ---
def weighted_avg(df, value_col, weight_col):
    try:
        valid = df[weight_col] > 0
        return (df.loc[valid, value_col] * df.loc[valid, weight_col]).sum() / df.loc[valid, weight_col].sum()
    except Exception:
        return None

# --- Your other functions above ---

# ==================================
# =========================
# TEAM KPI OVERVIEW (CARDS)
# =========================
# --- imports ---
import uuid
import pandas as pd, numpy as np
from pathlib import Path
import streamlit as st
import plotly.graph_objects as go

# ---------- Keep/assume these exist earlier ----------
# combined_oracle_df  -> your combined Oracle dataframe

# ---------- Utilities ----------
DAY_ORDER = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]

def _to_td(series):
    s = pd.to_datetime(series.astype(str), errors="coerce") if str(series.dtype) == "datetime64[ns]" else pd.to_timedelta(series.astype(str), errors="coerce")
    return s[~s.isna()]

def hhmm(td: pd.Timedelta) -> str:
    secs = int(td.total_seconds()); h = secs // 3600; m = (secs % 3600)//60
    return f"{h:02}:{m:02}"

def avg_hhmm(series) -> str:
    t = _to_td(series)
    return hhmm(t.mean()) if len(t) else "00:00"

def sum_hhmm(series) -> str:
    t = _to_td(series)
    return hhmm(t.sum()) if len(t) else "00:00"

def pct(n, d):
    try: return (n/d*100.0) if d else 0.0
    except: return 0.0

def last_weeks_trend(df, count_col="__count__"):
    if "Date" not in df.columns or df.empty: return 0.0
    x = df.copy()
    x["Date"] = pd.to_datetime(x["Date"], errors="coerce").dropna()
    if x.empty: return 0.0
    x["WeekISO"] = x["Date"].dt.isocalendar().week.astype(int)
    x["Year"]    = x["Date"].dt.year
    w = (x.groupby(["Year","WeekISO"]).size().rename(count_col).reset_index())
    if len(w) < 8: return 0.0
    recent4 = w.sort_values(["Year","WeekISO"]).tail(4)[count_col].sum()
    prior4  = w.sort_values(["Year","WeekISO"]).tail(8).head(4)[count_col].sum()
    return pct(recent4 - prior4, prior4 if prior4 else recent4)

def get_team_df(team_name: str) -> pd.DataFrame:
    mapping = {
        "VIP North": "VIP North Oracle Data.xlsx",
        "VIP South": "VIP South Oracle Data.xlsx",
        "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
        "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
    }
    if 'Team' in combined_oracle_df.columns:
        df_team = combined_oracle_df[combined_oracle_df['Team'] == team_name].copy()
        df_team.columns = df_team.columns.str.replace("\u00A0", " ", regex=False).str.strip()
        if 'Overtime' in df_team.columns and not df_team['Overtime'].astype(str).str.strip().eq("").all():
            return df_team
        path = mapping.get(team_name)
        if path and Path(path).exists():
            df_file = pd.read_excel(path)
            df_file.columns = df_file.columns.str.replace("\u00A0", " ", regex=False).str.strip()
            df_file['Team'] = team_name
            return df_file
        return df_team
    path = mapping.get(team_name)
    if path and Path(path).exists():
        df = pd.read_excel(path)
        df.columns = df.columns.str.replace("\u00A0", " ", regex=False).str.strip()
        df['Team'] = team_name
        return df
    return pd.DataFrame()

THRESHOLD_TWT = pd.to_timedelta("10:25:00")

def _period_key(dt_series: pd.Series, period: str):
    if period == "Weekly":
        return dt_series.dt.to_period("W").apply(lambda p: p.start_time.date())
    if period == "Monthly":
        return dt_series.dt.to_period("M").apply(lambda p: p.start_time.date())
    return dt_series.dt.date

def _series_by_period(df: pd.DataFrame, period="Daily", reducer="count", col=None):
    if "Date" not in df.columns or df.empty: return [0]
    x = df.copy()
    x["Date"] = pd.to_datetime(x["Date"], errors="coerce")
    x = x.dropna(subset=["Date"])
    if x.empty: return [0]
    key = _period_key(x["Date"], period)
    if reducer == "count":
        s = x.groupby(key).size()
    elif reducer == "nunique" and col:
        s = x.groupby(key)[col].nunique()
    elif reducer == "sum_time" and col:
        td = pd.to_timedelta(x[col].astype(str), errors="coerce")
        s = pd.Series(td.values, index=key).groupby(level=0).sum().dt.total_seconds() / 3600.0
    elif reducer == "mean_time" and col:
        td = pd.to_timedelta(x[col].astype(str), errors="coerce")
        s = pd.Series(td.values, index=key).groupby(level=0).mean().dt.total_seconds() / 60.0
    elif reducer == "sum_num" and col:
        s = pd.to_numeric(x[col], errors="coerce").groupby(key).sum()
    else:
        s = x.groupby(key).size()
    return s.tail(30).to_list()

def _status_mask(df: pd.DataFrame, keyword: str) -> pd.Series:
    if "Activity Status" not in df.columns or df.empty:
        return pd.Series([False]*len(df))
    return df["Activity Status"].astype(str).str.contains(keyword, case=False, na=False)

def _money_to_float(s):
    if s is None: return pd.Series(dtype="float64")
    if not isinstance(s, pd.Series): s = pd.Series(s)
    if pd.api.types.is_numeric_dtype(s): return pd.to_numeric(s, errors="coerce")
    s = s.astype("string").str.strip()
    s = s.str.replace("\u2212", "-", regex=False).str.replace(r"\(([^)]+)\)", r"-\1", regex=True)
    s = s.str.replace(r"[£$,]", "", regex=True).str.replace(r"[^\d\.\-]", "", regex=True)
    return pd.to_numeric(s, errors="coerce")

def overtime_from_df(df, period="Daily"):
    title = "Overtime total (£)"
    ot_col = None
    for c in df.columns:
        if "overtime" in c.lower():
            ot_col = c; break
    if ot_col is None: return title, "—", None, [0]
    money = _money_to_float(df[ot_col])
    if money.isna().all(): return title, "—", ot_col, [0]
    ts_df = df.copy(); ts_df["__ot"] = money
    ts = _series_by_period(ts_df[ts_df["__ot"].notna()], period=period, reducer="sum_num", col="__ot")
    total = float(money.sum(skipna=True)); formatted = f"£{total:,.0f}"
    return title, formatted, ot_col, ts

def _value_per_visit_series(df: pd.DataFrame, period="Daily"):
    if df.empty or "Date" not in df.columns or "Total Value" not in df.columns: return [0]
    x = df.copy(); x["Date"] = pd.to_datetime(x["Date"], errors="coerce").dropna()
    if x.empty: return [0]
    key = _period_key(x["Date"], period)
    val = pd.to_numeric(x["Total Value"], errors="coerce")
    sum_val = pd.Series(val.values, index=key).groupby(level=0).sum()
    cnt = x.groupby(key).size()
    avg = (sum_val / cnt).replace([np.inf, -np.inf], np.nan).fillna(0)
    return avg.tail(30).to_list()

def safe_numeric(df, cols):
    """Convert given columns to numeric safely (coerce errors)."""
    for col in cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce')
    return df

# ---------- Card CSS ----------
# ---------- Card CSS (WHITE) ----------
st.markdown("""
<style>
/* Base white KPI card */
.kpi-card {
  background:#ffffff !important;
  color:#0f172a !important;
  border-radius:16px;
  padding:16px 18px;
  border:1px solid #e6e8eb;
  box-shadow:0 14px 28px rgba(2,32,71,.10);
}

/* Typography inside the card */
.kpi-top   { font-size:.85rem; color:#475569 !important; margin-bottom:8px; }
.kpi-title { font-size:1.25rem; font-weight:800; color:#0f172a !important; margin-bottom:8px; }
.kpi-value { font-size:2.0rem; font-weight:800; color:#0f172a !important; line-height:1.1; }
.kpi-sub   { font-size:.95rem; color:#334155 !important; margin-top:6px; }

/* Footer row + deltas */
.kpi-footer { display:flex; justify-content:space-between; margin-top:10px; font-size:.9rem; color:#334155 !important; }
.kpi-up   { color:#16a34a !important; font-weight:700; }
.kpi-down { color:#dc2626 !important; font-weight:700; }

/* Progress bar */
.kpi-bar { height:8px; background:#eef2f7; border-radius:999px; overflow:hidden; margin-top:10px; }
.kpi-bar > div { height:100%; background:linear-gradient(90deg,#0ea5e9,#2563eb); }
</style>
""", unsafe_allow_html=True)
st.markdown("""
<style>
/* Dark KPI card + force white text inside */
.kpi-card {
  background:#0f172a !important;      /* dark box */
  color:#f8fafc !important;           /* default text color */
  border-radius:16px;
  padding:16px 18px;
  border:1px solid rgba(255,255,255,.06);
  box-shadow:0 14px 28px rgba(2,32,71,.18);
}
/* Make EVERYTHING inside the card white (wins against other styles) */
.kpi-card *, 
.kpi-card h1, .kpi-card h2, .kpi-card h3, .kpi-card h4,
.kpi-card p, .kpi-card li, .kpi-card strong, .kpi-card span, .kpi-card em {
  color:#f8fafc !important;
}
.kpi-card a { color:#93c5fd !important; }          /* links readable */
.kpi-card .muted { color:#cbd5e1 !important; }     /* optional muted */
</style>
""", unsafe_allow_html=True)




# ---------- Tiny sparkline ----------
def _sparkline(y_vals, color="#22c55e"):
    if not y_vals:
        y_vals = [0]
    fig = go.Figure(go.Scatter(y=y_vals, mode="lines", line=dict(width=2), hoverinfo="skip"))
    fig.update_layout(margin=dict(l=0, r=0, t=0, b=0),
                      xaxis=dict(visible=False),
                      yaxis=dict(visible=False),
                      height=60)
    if color:
        fig.data[0].line.color = color
    return fig



# ---------- Card ----------
def card(title, value, sub="", trend=None, bar_pct=None, spark=None, color=None, source=""):
    st.markdown('<div class="kpi-card">', unsafe_allow_html=True)
    st.markdown('<div class="kpi-top">Month to date</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="kpi-title">{title}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="kpi-value">{value}</div>', unsafe_allow_html=True)
    if sub: st.markdown(f'<div class="kpi-sub">{sub}</div>', unsafe_allow_html=True)
    if bar_pct is not None:
        w = max(0, min(int(bar_pct), 100))
        st.markdown(f'<div class="kpi-bar"><div style="width:{w}%"></div></div>', unsafe_allow_html=True)
    if spark is not None:
        def _slug(s: str) -> str: return "".join(ch if ch.isalnum() else "_" for ch in (s or "")).strip("_")
        st.plotly_chart(_sparkline(spark, color=color), use_container_width=True,
                        config={"displayModeBar": False},
                        key=f"spark_{_slug(source)}_{_slug(title)}_{uuid.uuid4().hex[:8]}")
    left = ""
    if trend is not None:
        cls = "kpi-up" if trend >= 0 else "kpi-down"; arrow = "▲" if trend >= 0 else "▼"
        left = f'<span class="{cls}">{arrow} {trend:+.1f}%</span>'
    st.markdown(f'<div class="kpi-footer">{left}<span>{source}</span></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# ---------- MoM helpers ----------
def _pct_change(curr, prev):
    if prev is None or prev == 0 or pd.isna(prev): return 0.0
    return (curr - prev) / prev * 100.0

def _pp_change(curr_rate, prev_rate):
    if pd.isna(prev_rate): return 0.0
    return (curr_rate - prev_rate)

def compute_mom_trends(df_full, df_curr, sel_period):
    out = dict(value_pct=0.0, visits_pct=0.0, active_pct=0.0,
               completed_rate_pp=0.0, cancel_rate_pp=0.0, notdone_rate_pp=0.0,
               vpv_pct=0.0, over_1025_pct=0.0, under_1025_pct=0.0)
    if "_dt" not in df_full.columns: return out
    m_idx = pd.PeriodIndex(df_full["_dt"].dropna().dt.to_period("M")).sort_values().unique()
    if len(m_idx) == 0: return out
    if sel_period is None:
        cur_p = m_idx[-1]; prev_p = m_idx[-2] if len(m_idx) >= 2 else None
    else:
        cur_p = sel_period; prev_candidates = m_idx[m_idx < cur_p]; prev_p = prev_candidates[-1] if len(prev_candidates) else None
    if prev_p is None: return out
    prev_df = df_full[df_full["_dt"].dt.to_period("M") == prev_p].copy()

    cur_visits, prev_visits = len(df_curr), len(prev_df)
    cur_value = pd.to_numeric(df_curr.get("Total Value", pd.Series(dtype=float)), errors="coerce").sum()
    prev_value= pd.to_numeric(prev_df.get("Total Value",  pd.Series(dtype=float)), errors="coerce").sum()
    cur_active= df_curr["Name"].nunique() if "Name" in df_curr.columns else np.nan
    prev_active= prev_df["Name"].nunique() if "Name" in prev_df.columns else np.nan

    def _rate(dfi, kw):
        if "Activity Status" not in dfi.columns: return np.nan
        tot = len(dfi);  s = dfi["Activity Status"].astype(str)
        if tot == 0: return np.nan
        return s.str.contains(kw, case=False, na=False).sum() / tot * 100.0

    cur_completed = _rate(df_curr, "Completed");  prev_completed = _rate(prev_df, "Completed")
    cur_cancel    = _rate(df_curr, "Cancelled");  prev_cancel    = _rate(prev_df, "Cancelled")
    cur_notdone   = _rate(df_curr, "Not Done");   prev_notdone   = _rate(prev_df, "Not Done")

    cur_vpv  = (cur_value / cur_visits) if cur_visits else np.nan
    prev_vpv = (prev_value / prev_visits) if prev_visits else np.nan

    def _over_under(dfi):
        tw_col = next((c for c in ["Total Working Time","Total working time"] if c in dfi.columns), None)
        if tw_col is None or dfi.empty: return 0,0
        tw = pd.to_timedelta(dfi[tw_col].astype(str), errors="coerce")
        valid = tw.notna() & (tw > pd.Timedelta(0))
        over = int((valid & (tw > THRESHOLD_TWT)).sum())
        under= int((valid & (tw <= THRESHOLD_TWT)).sum())
        return over, under

    cur_over, cur_under = _over_under(df_curr)
    prev_over, prev_under = _over_under(prev_df)

    out["value_pct"]         = _pct_change(cur_value,  prev_value)
    out["visits_pct"]        = _pct_change(cur_visits, prev_visits)
    out["active_pct"]        = _pct_change(cur_active, prev_active)
    out["completed_rate_pp"] = _pp_change(cur_completed, prev_completed)
    out["cancel_rate_pp"]    = _pp_change(cur_cancel,    prev_cancel)
    out["notdone_rate_pp"]   = _pp_change(cur_notdone,   prev_notdone)
    out["vpv_pct"]           = _pct_change(cur_vpv,      prev_vpv)
    out["over_1025_pct"]     = _pct_change(cur_over,     prev_over)
    out["under_1025_pct"]    = _pct_change(cur_under,    prev_under)
    return out

# ---------- KPIs builder (unchanged) ----------
def team_kpis(df: pd.DataFrame) -> dict:
    cols = df.columns.str.strip()
    df = df.copy(); df.columns = cols
    total_visits = len(df)
    status = df["Activity Status"].astype(str) if "Activity Status" in cols else pd.Series([], dtype=str)
    completed = status.str.contains("Completed", case=False, na=False).sum() if not status.empty else 0
    cancelled = status.str.contains("Cancelled", case=False, na=False).sum() if not status.empty else 0
    not_done  = status.str.contains("Not Done",  case=False, na=False).sum() if not status.empty else 0
    pending   = status.str.contains("Pending",   case=False, na=False).sum() if not status.empty else 0
    started   = status.str.contains("Started",   case=False, na=False).sum() if not status.empty else 0
    unique_engs = df["Name"].nunique() if "Name" in cols else 0
    tw_col = "Total Working Time" if "Total Working Time" in cols else ("Total working time" if "Total working time" in cols else None)
    ot_col = "Overtime" if "Overtime" in cols else None
    total_value = pd.to_numeric(df.get("Total Value", pd.Series(dtype=float)), errors="coerce").sum() if "Total Value" in cols else np.nan
    if tw_col:
        tw = pd.to_timedelta(df[tw_col].astype(str), errors="coerce")
        valid = tw.notna() & (tw > pd.Timedelta(0))
        avg_work  = hhmm(tw[valid].mean()) if valid.any() else "00:00"
        total_time = hhmm(tw[valid].sum()) if valid.any() else "00:00"
        over_mask  = valid & (tw > THRESHOLD_TWT); under_mask = valid & (tw <= THRESHOLD_TWT)
        over_1025_count  = int(over_mask.sum()); under_1025_count = int(under_mask.sum())
        over_1025_rate   = pct(over_1025_count, valid.sum()); under_1025_rate = pct(under_1025_count, valid.sum())
    else:
        avg_work = total_time = "00:00"; over_1025_count = under_1025_count = 0; over_1025_rate = under_1025_rate = 0.0
    lunch_col = "Total Time" if "Total Time" in cols else ("Total Working Time" if "Total Working Time" in cols else ("Total working time" if "Total working time" in cols else None))
    if "Visit Type" in cols and lunch_col:
        lunch_df = df[df["Visit Type"].astype(str).str.contains("Lunch", case=False, na=False)]
        td = pd.to_timedelta(lunch_df[lunch_col].astype(str), errors="coerce"); td = td[td.notna() & (td > pd.Timedelta(0))]
        avg_lunch = hhmm(td.mean()) if len(td) else "00:00"
    else:
        avg_lunch = "00:00"
    overtime_total = np.nan
    if ot_col:
        raw = df[ot_col].astype(str).str.replace("\u00A0", " ", regex=False).str.strip()
        td  = pd.to_timedelta(raw, errors="coerce"); hrs = td.dt.total_seconds() / 3600.0
        if hrs.notna().any(): num = pd.to_numeric(hrs, errors="coerce")
        else:
            cleaned = raw.str.replace(r"\(([^)]+)\)", r"-\1", regex=True).str.replace(r"[£,\s]", "", regex=True)
            money = pd.to_numeric(cleaned, errors="coerce")
            num = money if money.notna().any() else pd.to_numeric(raw, errors="coerce")
        valid = pd.to_numeric(num, errors="coerce").notna()
        overtime_total = float(pd.to_numeric(num, errors="coerce")[valid].sum()) if valid.any() else 0.0
    comp_rate    = pct(completed,  total_visits)
    cancel_rate  = pct(cancelled,  total_visits)
    notdone_rate = pct(not_done,   total_visits)
    trend_visits = last_weeks_trend(df)
    return dict(total_visits=total_visits, completed=completed, cancelled=cancelled, not_done=not_done,
                pending=pending, started=started, unique_engs=unique_engs, total_value=total_value,
                avg_work=avg_work, total_time=total_time, avg_lunch=avg_lunch, comp_rate=comp_rate,
                cancel_rate=cancel_rate, notdone_rate=notdone_rate, over_1025_count=over_1025_count,
                under_1025_count=under_1025_count, over_1025_rate=over_1025_rate, under_1025_rate=under_1025_rate,
                overtime_total=overtime_total, trend_visits=trend_visits)

# === Invoices: loader tuned to your Excel ===
@st.cache_data(show_spinner=False)
def load_invoices(path: str | Path = "Invoices.xlsx") -> pd.DataFrame:
    p = Path(path)
    if not p.exists():
        st.warning("⚠️ Invoices.xlsx not found in app folder.")
        return pd.DataFrame()

    try:
        df = pd.read_excel(p)  # Sheet1 in your file
    except Exception as e:
        st.error(f"Failed to read Invoices.xlsx: {e}")
        return pd.DataFrame()

    # Normalise headers
    df = df.copy()
    df.columns = (
        df.columns.astype(str)
        .str.replace("\u00A0", " ", regex=False)  # non-breaking space
        .str.strip()
    )

    # Helper: currency/number coercion
    def to_num(s: pd.Series) -> pd.Series:
        return pd.to_numeric(
            s.astype(str)
             .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)  # (123) -> -123
             .str.replace(r"[£$,]", "", regex=True)
             .str.replace(r"[^0-9.\-]", "", regex=True),
            errors="coerce",
        )

    # Map your actual columns to canonical working columns
    col = lambda name: name if name in df.columns else None

    # Rechargeable
    if col("Rechargeable"):
        df["RechargeableRaw"] = df["Rechargeable"].astype(str).str.strip()
        yes_vals = {"yes", "y", "true", "1", "chargeable", "rechargeable", "rc"}
        df["RechargeableFlag"] = df["RechargeableRaw"].str.lower().isin(yes_vals)

    # Dates
    if col("Date of visit"):
        df["VisitDate"] = pd.to_datetime(df["Date of visit"], errors="coerce")

    # People / groupers
    if col("Engineers Name"):
        df["Engineer"] = df["Engineers Name"].astype(str).str.strip()
    if col("Client Name"):
        df["Client"] = df["Client Name"].astype(str).str.strip()
    if col("Team"):
        df["Team"] = df["Team"].astype(str).str.strip()
    if col("Department"):
        df["Department_"] = df["Department"].astype(str).str.strip()
    if col("Invoice Type"):
        df["InvoiceType"] = df["Invoice Type"].astype(str).str.strip()
    if col("Visit Type"):
        df["VisitType"] = df["Visit Type"].astype(str).str.strip()
    if col("Internal Or External"):
        df["InternalExternal"] = df["Internal Or External"].astype(str).str.strip()

    # Values
    if col("Total Value"):
        df["TotalValue"] = to_num(df["Total Value"])
    if col("Equipment Value"):
        df["EquipmentValue"] = to_num(df["Equipment Value"])
    if col("Labour Value"):
        df["LabourValue"] = to_num(df["Labour Value"])
    if col("Additional Costs"):
        df["AdditionalCosts"] = to_num(df["Additional Costs"])
    if col("Hotel/ Food Value"):
        df["HotelFoodValue"] = to_num(df["Hotel/ Food Value"])

    # Time on site (Excel time-of-day used as duration)
    if col("Time On-Site"):
        # Keep original; also provide minutes for optional metrics
        try:
            t = pd.to_datetime(df["Time On-Site"], errors="coerce")
            df["TimeOnSite_minutes"] = (t.dt.hour * 60 + t.dt.minute).astype("Int64")
        except Exception:
            df["TimeOnSite_minutes"] = pd.NA

    return df
# === Invoices: screen renderer (shows your real columns) ===
def render_invoices_screen():
    st.title("📄 Invoices")

    inv = load_invoices()
    if inv.empty:
        st.info("No invoices to show yet.")
        return

    # -------- Filters (match your file)
    f1, f2, f3, f4 = st.columns([1.3, 1, 1, 1])
    with f1:
        min_d = pd.to_datetime(inv.get("VisitDate")).min()
        max_d = pd.to_datetime(inv.get("VisitDate")).max()
        start_default = (min_d.date() if pd.notna(min_d) else pd.Timestamp.today().date())
        end_default = (max_d.date() if pd.notna(max_d) else pd.Timestamp.today().date())
        start, end = st.date_input("Date of visit", value=(start_default, end_default))

    with f2:
        team_vals = sorted(inv.get("Team", pd.Series(dtype=str)).dropna().unique().tolist())
        team = st.multiselect("Team", team_vals)

    with f3:
        invtype_vals = sorted(inv.get("InvoiceType", pd.Series(dtype=str)).dropna().unique().tolist())
        invtypes = st.multiselect("Invoice Type", invtype_vals)

    with f4:
        visittype_vals = sorted(inv.get("VisitType", pd.Series(dtype=str)).dropna().unique().tolist())
        visittype = st.multiselect("Visit Type", visittype_vals)

    g1, g2, g3 = st.columns([1, 1, 1])
    with g1:
        engineer_vals = sorted(inv.get("Engineer", pd.Series(dtype=str)).dropna().unique().tolist())
        engineers = st.multiselect("Engineers Name", engineer_vals)
    with g2:
        client_vals = sorted(inv.get("Client", pd.Series(dtype=str)).dropna().unique().tolist())
        clients = st.multiselect("Client Name", client_vals)
    with g3:
        q = st.text_input("Search (VR / Ticket / ASA / Venue / Notes)")

    # -------- Apply filters
    cur = inv.copy()
    if "VisitDate" in cur.columns:
        cur = cur[
            (cur["VisitDate"] >= pd.to_datetime(start)) &
            (cur["VisitDate"] <= pd.to_datetime(end) + pd.Timedelta(days=1))
        ]
    if team:
        cur = cur[cur["Team"].isin(team)]
    if invtypes:
        cur = cur[cur["InvoiceType"].isin(invtypes)]
    if visittype:
        cur = cur[cur["VisitType"].isin(visittype)]
    if engineers:
        cur = cur[cur["Engineer"].isin(engineers)]
    if clients:
        cur = cur[cur["Client"].isin(clients)]
    if q:
        # free text across common ID/text fields
        cols = [c for c in [
            "VR Number", "Ticket Number", "ASA Number", "Venue Name",
            "Activity", "Additional Info", "PostCode", "Cost Code"
        ] if c in cur.columns]
        if cols:
            mask = pd.Series(False, index=cur.index)
            for c in cols:
                mask = mask | cur[c].astype(str).str.contains(q, case=False, na=False)
            cur = cur[mask]

        # -------- KPIs
        total_val  = float(pd.to_numeric(cur.get("TotalValue"), errors="coerce").fillna(0).sum())
        count      = int(len(cur))
        avg_val    = (total_val / count) if count else 0.0
        equip_tot  = float(pd.to_numeric(cur.get("EquipmentValue"), errors="coerce").fillna(0).sum())
        labour_tot = float(pd.to_numeric(cur.get("LabourValue"), errors="coerce").fillna(0).sum())

        # Rechargeable breakdown
        recharge_val = float(pd.to_numeric(cur.loc[cur.get("RechargeableFlag") == True, "TotalValue"], errors="coerce").fillna(0).sum())
        noncharge_val = float(pd.to_numeric(cur.loc[cur.get("RechargeableFlag") == False, "TotalValue"], errors="coerce").fillna(0).sum())

    # -------- KPI ROW 1 with sparklines + MoM arrows (5 across)
    # Helper: build a monthly series from cur
    def _monthly(cur_df: pd.DataFrame) -> pd.DataFrame:
        d = cur_df.dropna(subset=["VisitDate"]).copy()
        if d.empty:
            return pd.DataFrame(columns=["Month", "TotalValue", "Invoices", "Equip", "Labour", "AvgValue"])
        d["Month"] = d["VisitDate"].dt.to_period("M").dt.to_timestamp()
        g = d.groupby("Month")
        s_val    = pd.to_numeric(d["TotalValue"], errors="coerce").fillna(0)
        s_equip  = pd.to_numeric(d.get("EquipmentValue", 0), errors="coerce").fillna(0)
        s_labour = pd.to_numeric(d.get("LabourValue", 0), errors="coerce").fillna(0)
        dfm = pd.DataFrame({
            "TotalValue": g["TotalValue"].sum(min_count=1),
            "Invoices":   g.size(),
            "Equip":      g["EquipmentValue"].sum(min_count=1) if "EquipmentValue" in d.columns else 0,
            "Labour":     g["LabourValue"].sum(min_count=1) if "LabourValue" in d.columns else 0,
        }).reset_index()
        # Avg invoice value per month
        dfm["AvgValue"] = dfm["TotalValue"] / dfm["Invoices"].replace(0, pd.NA)
        dfm = dfm.fillna(0)
        return dfm

    def _format_money(x: float) -> str:
        return f"£{x:,.2f}"

    def _format_int(x: float) -> str:
        return f"{int(x):,}"

    def _pct_delta(curr: float, prev: float) -> tuple[str, str]:
        # returns (text, arrow) where arrow is ▲/▼/—
        if prev == 0:
            return ("—", "—")
        pct = (curr - prev) / prev * 100
        arrow = "▲" if pct >= 0 else "▼"
        return (f"{pct:+.1f}%", arrow)

    def _sparkline(col, x, y, title, big_text, sublabel, is_money=False):
        import plotly.express as px
        import pandas as pd

        # Minimal sparkline
        fig = px.line(pd.DataFrame({"x": x, "y": y}), x="x", y="y")
        fig.update_traces(mode="lines+markers")
        fig.update_layout(
            height=90, margin=dict(l=0, r=0, t=0, b=0),
            showlegend=False
        )
        fig.update_xaxes(visible=False)
        fig.update_yaxes(visible=False)

        with col:
            st.markdown(f"**{title}**")
            st.markdown(f"<div style='font-size:28px; font-weight:700;'>{big_text}</div>", unsafe_allow_html=True)
            st.caption(sublabel)
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

    # Build monthly series for the sparkline + MoM
    m = _monthly(cur)

    # Current values (already computed earlier)
    total_val  = float(pd.to_numeric(cur.get("TotalValue"), errors="coerce").fillna(0).sum())
    count      = int(len(cur))
    avg_val    = (total_val / count) if count else 0.0
    equip_tot  = float(pd.to_numeric(cur.get("EquipmentValue"), errors="coerce").fillna(0).sum())
    labour_tot = float(pd.to_numeric(cur.get("LabourValue"), errors="coerce").fillna(0).sum())

    # Month-over-month deltas (vs previous month in filtered data)
    if not m.empty:
        m = m.sort_values("Month")
        curr_m = m.iloc[-1]
        prev_m = m.iloc[-2] if len(m) > 1 else None

        tv_delta_txt, tv_arrow   = _pct_delta(curr_m["TotalValue"], prev_m["TotalValue"] if prev_m is not None else 0)
        inv_delta_txt, inv_arrow = _pct_delta(curr_m["Invoices"],   prev_m["Invoices"]   if prev_m is not None else 0)
        av_delta_txt, av_arrow   = _pct_delta(curr_m["AvgValue"],   prev_m["AvgValue"]   if prev_m is not None else 0)
        eq_delta_txt, eq_arrow   = _pct_delta(curr_m["Equip"],      prev_m["Equip"]      if prev_m is not None else 0)
        lb_delta_txt, lb_arrow   = _pct_delta(curr_m["Labour"],     prev_m["Labour"]     if prev_m is not None else 0)
    else:
        tv_delta_txt = inv_delta_txt = av_delta_txt = eq_delta_txt = lb_delta_txt = "—"
        tv_arrow = inv_arrow = av_arrow = eq_arrow = lb_arrow = "—"

    c1, c2, c3, c4, c5 = st.columns(5)

    # 1) Total Value
    _sparkline(
        c1,
        x=m["Month"] if not m.empty else [],
        y=m["TotalValue"] if not m.empty else [],
        title="Total Value",
        big_text=_format_money(total_val),
        sublabel=f"{tv_arrow} {tv_delta_txt} vs prev month"
    )

    # 2) Invoices
    _sparkline(
        c2,
        x=m["Month"] if not m.empty else [],
        y=m["Invoices"] if not m.empty else [],
        title="Invoices",
        big_text=_format_int(count),
        sublabel=f"{inv_arrow} {inv_delta_txt} vs prev month"
    )

    # 3) Avg Value
    _sparkline(
        c3,
        x=m["Month"] if not m.empty else [],
        y=m["AvgValue"] if not m.empty else [],
        title="Avg Value",
        big_text=_format_money(avg_val),
        sublabel=f"{av_arrow} {av_delta_txt} vs prev month"
    )

    # 4) Equipment Total
    _sparkline(
        c4,
        x=m["Month"] if not m.empty else [],
        y=m["Equip"] if not m.empty else [],
        title="Equipment Total",
        big_text=_format_money(equip_tot),
        sublabel=f"{eq_arrow} {eq_delta_txt} vs prev month"
    )

    # 5) Labour Total
    _sparkline(
        c5,
        x=m["Month"] if not m.empty else [],
        y=m["Labour"] if not m.empty else [],
        title="Labour Total",
        big_text=_format_money(labour_tot),
        sublabel=f"{lb_arrow} {lb_delta_txt} vs prev month"
    )


    # -------- KPI ROW 2 with sparklines + MoM arrows (5 across)

    # Monthly Rechargeable split
    def _monthly_recharge(cur_df: pd.DataFrame) -> pd.DataFrame:
        d = cur_df.dropna(subset=["VisitDate"]).copy()
        if d.empty or "RechargeableFlag" not in d.columns:
            return pd.DataFrame(columns=["Month", "Rechargeable", "NonRechargeable", "RecCount", "NonCount"])
        d["Month"] = d["VisitDate"].dt.to_period("M").dt.to_timestamp()
        g = d.groupby(["Month", "RechargeableFlag"])
        out = g.agg(
            TotalValue=("TotalValue", "sum"),
            Count=("TotalValue", "size")
        ).reset_index()
        # Pivot Rechargeable vs Non
        rec = out[out["RechargeableFlag"] == True].set_index("Month")
        non = out[out["RechargeableFlag"] == False].set_index("Month")
        dfm = pd.DataFrame({
            "Rechargeable": rec["TotalValue"],
            "NonRechargeable": non["TotalValue"],
            "RecCount": rec["Count"],
            "NonCount": non["Count"]
        }).fillna(0).reset_index()
        dfm["RecAvg"] = dfm["Rechargeable"] / dfm["RecCount"].replace(0, pd.NA)
        dfm["NonAvg"] = dfm["NonRechargeable"] / dfm["NonCount"].replace(0, pd.NA)
        dfm["PctRec"] = dfm["Rechargeable"] / (dfm["Rechargeable"] + dfm["NonRechargeable"]).replace(0, pd.NA) * 100
        return dfm.fillna(0)

    mr = _monthly_recharge(cur)

    # Current filtered snapshot
    rec_df = cur[cur.get("RechargeableFlag") == True]
    non_df = cur[cur.get("RechargeableFlag") == False]

    recharge_val = float(pd.to_numeric(rec_df["TotalValue"], errors="coerce").fillna(0).sum())
    noncharge_val = float(pd.to_numeric(non_df["TotalValue"], errors="coerce").fillna(0).sum())

    rec_count = len(rec_df)
    non_count = len(non_df)

    rec_avg = (recharge_val / rec_count) if rec_count else 0.0
    non_avg = (noncharge_val / non_count) if non_count else 0.0

    pct_rec = (recharge_val / total_val * 100) if total_val else 0.0


    # MoM deltas
    if not mr.empty:
        mr = mr.sort_values("Month")
        curr_m = mr.iloc[-1]
        prev_m = mr.iloc[-2] if len(mr) > 1 else None

        rec_delta_txt, rec_arrow   = _pct_delta(curr_m["Rechargeable"],   prev_m["Rechargeable"]   if prev_m is not None else 0)
        non_delta_txt, non_arrow   = _pct_delta(curr_m["NonRechargeable"],prev_m["NonRechargeable"]if prev_m is not None else 0)
        pct_delta_txt, pct_arrow   = _pct_delta(curr_m["PctRec"],         prev_m["PctRec"]         if prev_m is not None else 0)
        ravg_delta_txt, ravg_arrow = _pct_delta(curr_m["RecAvg"],         prev_m["RecAvg"]         if prev_m is not None else 0)
        navg_delta_txt, navg_arrow = _pct_delta(curr_m["NonAvg"],         prev_m["NonAvg"]         if prev_m is not None else 0)
    else:
        rec_delta_txt = non_delta_txt = pct_delta_txt = ravg_delta_txt = navg_delta_txt = "—"
        rec_arrow = non_arrow = pct_arrow = ravg_arrow = navg_arrow = "—"

    c6, c7, c8, c9, c10 = st.columns(5)

    # 1) Rechargeable Total
    _sparkline(
        c6,
        x=mr["Month"] if not mr.empty else [],
        y=mr["Rechargeable"] if not mr.empty else [],
        title="Rechargeable Total",
        big_text=_format_money(recharge_val),
        sublabel=f"{rec_arrow} {rec_delta_txt} vs prev month"
    )

    # 2) Non-Rechargeable Total
    _sparkline(
        c7,
        x=mr["Month"] if not mr.empty else [],
        y=mr["NonRechargeable"] if not mr.empty else [],
        title="Non-Rechargeable Total",
        big_text=_format_money(noncharge_val),
        sublabel=f"{non_arrow} {non_delta_txt} vs prev month"
    )

    # 3) Rechargeable %
    _sparkline(
        c8,
        x=mr["Month"] if not mr.empty else [],
        y=mr["PctRec"] if not mr.empty else [],
        title="Rechargeable %",
        big_text=f"{pct_rec:.1f}%",
        sublabel=f"{pct_arrow} {pct_delta_txt} vs prev month"
    )

    # 4) Avg Rechargeable
    _sparkline(
        c9,
        x=mr["Month"] if not mr.empty else [],
        y=mr["RecAvg"] if not mr.empty else [],
        title="Avg Rechargeable",
        big_text=_format_money(rec_avg),
        sublabel=f"{ravg_arrow} {ravg_delta_txt} vs prev month"
    )

    # 5) Avg Non-Rechargeable
    _sparkline(
        c10,
        x=mr["Month"] if not mr.empty else [],
        y=mr["NonAvg"] if not mr.empty else [],
        title="Avg Non-Rechargeable",
        big_text=_format_money(non_avg),
        sublabel=f"{navg_arrow} {navg_delta_txt} vs prev month"
    )






    # -------- Tabs
    # New tab order adds 3: Total Over Time, Sky Business, Sky Business Cover

    tabs = st.tabs([
        "📈 Trend",                # tabs[0]
        "📊 Total Over Time",      # tabs[1]
        "🏢 Sky Business",         # tabs[2]
        "🧑‍💼 Sky Business Cover", # tabs[3]
        "👥 By team",              # tabs[4]
        "🏷️ By client",           # tabs[5]
        "👷 Engineer",             # tabs[6]
        "💸 Rechargeable",         # tabs[7]  <-- NEW
        "🧾 Non Chargeable",       # tabs[8]  <-- NEW
        "📋 Table",                # tabs[9]  (moved)
        "⬇️ Export",              # tabs[10] (moved)
    ])


    # ---------- Trend by month (bar)
    with tabs[0]:
        if "VisitDate" in cur.columns and "TotalValue" in cur.columns:
            ts = (
                cur.dropna(subset=["VisitDate"])
                .assign(Month=lambda d: d["VisitDate"].dt.to_period("M").dt.to_timestamp())
                .groupby("Month")["TotalValue"].sum().reset_index()
            )
            fig = px.bar(ts, x="Month", y="TotalValue", title="Monthly invoice value")
            fig.update_layout(margin=dict(l=0, r=0, b=0, t=40))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Need Date of visit and Total Value for the trend chart.")

    # ---------- 📊 Total Over Time — line by Invoice Type (daily)
    with tabs[1]:
        if "VisitDate" in cur.columns and "TotalValue" in cur.columns:
            df = cur.dropna(subset=["VisitDate", "TotalValue"]).copy()
            if "InvoiceType" not in df.columns:
                st.info("Need 'Invoice Type' and 'Total Value' to build this view.")
            else:
                df["Day"] = df["VisitDate"].dt.floor("D")
                series = (
                    df.groupby(["Day", "InvoiceType"])["TotalValue"]
                    .sum().reset_index().sort_values("Day")
                )
                fig = px.line(
                    series, x="Day", y="TotalValue",
                    color="InvoiceType", markers=True,
                    title="Total value over time by Invoice Type"
                )
                fig.update_layout(margin=dict(l=0, r=0, b=0, t=40), xaxis_title="", yaxis_title="Total Value")
                st.plotly_chart(fig, use_container_width=True)

                st.caption("Table – daily totals by invoice type")
                st.dataframe(series, use_container_width=True)
        else:
            st.info("Need Date of visit and Total Value for this chart.")

    # ---------- 🏢 Sky Business — BAU/NON SLA WORK & SLA Call Out (robust match)
    with tabs[2]:
        need = {"VisitDate", "TotalValue", "InvoiceType"}
        if need.issubset(cur.columns):
            df = cur.dropna(subset=["VisitDate", "TotalValue"]).copy()
            # normalise invoice type for robust matching
            it = df["InvoiceType"].astype(str).str.lower()

            # very tolerant patterns:
            # - BAU/NON SLA WORK (handles spaces, slashes, hyphens, case)
            mask_bau = (
                it.str.contains(r"\bbau\b", regex=True)
                & it.str.contains(r"\bnon\s*[-/]?\s*sla\b", regex=True)
            )
            # - SLA Call Out (handles "callout", "call-out", spacing/case)
            mask_sla = it.str.contains(r"\bsla\b", regex=True) & it.str.contains(r"\bcall\s*-?\s*out\b|\bcallout\b", regex=True)

            sb = df[mask_bau | mask_sla].copy()

            if sb.empty:
                st.info("No rows matching BAU/NON SLA WORK or SLA Call Out in the selected filters/date range.")
            else:
                # bucket into two categories for a clean legend (always 2 lines)
                sb["_SB_Category"] = np.where(mask_bau.loc[sb.index], "BAU/NON SLA WORK", "SLA Call Out")

                sb["Day"] = sb["VisitDate"].dt.floor("D")
                sb_series = (
                    sb.groupby(["Day", "_SB_Category"])["TotalValue"]
                    .sum()
                    .reset_index()
                    .sort_values(["Day", "_SB_Category"])
                )

                fig = px.line(
                    sb_series,
                    x="Day",
                    y="TotalValue",
                    color="_SB_Category",
                    markers=True,
                    title="Sky Business – BAU/NON SLA WORK & SLA Call Out (Total value over time)"
                )
                fig.update_layout(margin=dict(l=0, r=0, b=0, t=48), xaxis_title="", yaxis_title="Total Value")
                st.plotly_chart(fig, use_container_width=True)

                st.caption("Table – filtered rows (BAU/NON SLA WORK & SLA Call Out)")
                preferred_cols = [c for c in [
                    "Date of visit", "Team", "Engineers Name", "Invoice Type", "Visit Type",
                    "Client Name", "VR Number", "Ticket Number", "ASA Number", "Venue Name",
                    "PostCode", "Time On-Site", "Equipment Used", "Additional Info",
                    "Equipment Value", "Labour Value", "Additional Costs", "Hotel/ Food Value",
                    "Total Value"
                ] if c in sb.columns]
                st.dataframe(sb[preferred_cols] if preferred_cols else sb, use_container_width=True)
        else:
            st.info("Need Date of visit, Total Value and Invoice Type for this view.")

    # ---------- 🧑‍💼 Sky Business Cover — only Business Cover
    with tabs[3]:
        if {"VisitDate", "TotalValue", "InvoiceType"}.issubset(cur.columns):
            cover = cur[cur["InvoiceType"].astype(str).str.contains(r"Business\s*Cover", case=False, na=False)].copy()

            if cover.empty:
                st.info("No rows matching Business Cover in the selected filters/date range.")
            else:
                cover["Day"] = cover["VisitDate"].dt.floor("D")
                cover_series = (
                    cover.groupby(["Day"])["TotalValue"]
                        .sum().reset_index().sort_values("Day")
                )
                fig = px.line(
                    cover_series, x="Day", y="TotalValue",
                    markers=True, title="Sky Business Cover – Total value over time"
                )
                fig.update_layout(margin=dict(l=0, r=0, b=0, t=48), xaxis_title="", yaxis_title="Total Value")
                st.plotly_chart(fig, use_container_width=True)

                st.caption("Table – filtered rows (Business Cover)")
                preferred_cols = [c for c in [
                    "Date of visit", "Team", "Engineers Name", "Invoice Type", "Visit Type",
                    "Client Name", "VR Number", "Ticket Number", "ASA Number", "Venue Name",
                    "PostCode", "Time On-Site", "Equipment Used", "Additional Info",
                    "Equipment Value", "Labour Value", "Additional Costs", "Hotel/ Food Value",
                    "Total Value"
                ] if c in cover.columns]
                st.dataframe(cover[preferred_cols] if preferred_cols else cover, use_container_width=True)
        else:
            st.info("Need Date of visit, Total Value and Invoice Type for this view.")

    # ---------- 👥 By team (bar)
    with tabs[4]:
        if "Team" in cur.columns and "TotalValue" in cur.columns:
            byteam = (
                cur.groupby("Team")["TotalValue"]
                .sum().sort_values(ascending=False).reset_index()
            )
            fig = px.bar(byteam, x="Team", y="TotalValue", title="By team (total value)")
            fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Need Team and Total Value for this chart.")

    # ---------- 🏷️ By client (bar)
    with tabs[5]:
        if "Client" in cur.columns and "TotalValue" in cur.columns:
            byclient = (
                cur.groupby("Client")["TotalValue"]
                .sum().sort_values(ascending=False).reset_index().head(25)
            )
            fig = px.bar(byclient, x="Client", y="TotalValue", title="Top clients by value")
            fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Need Client and Total Value for this chart.")

        # ---------- 👷 Engineer (table + line chart)
    with tabs[6]:
        need = {"Engineer", "TotalValue"}
        if need.issubset(cur.columns):
            # summary table per engineer
            eng_summary = (
                cur.groupby("Engineer")
                .agg(Invoices=("Engineer", "size"), TotalValue=("TotalValue", "sum"))
                .reset_index()
                .sort_values(["TotalValue", "Invoices"], ascending=False)
            )

            # format numbers
            eng_summary["TotalValue"] = eng_summary["TotalValue"].apply(lambda x: f"£{x:,.2f}")
            eng_summary["Invoices"]   = eng_summary["Invoices"].apply(lambda x: f"{x:,}")

            st.caption("Table — invoices per engineer")
            try:
                from st_aggrid import AgGrid, GridOptionsBuilder
                gb = GridOptionsBuilder.from_dataframe(eng_summary)
                gb.configure_default_column(editable=False, filter=True, sortable=True)
                # center align all cols
                gb.configure_column("Engineer", cellStyle={"textAlign": "center"})
                gb.configure_column("Invoices", cellStyle={"textAlign": "center"})
                gb.configure_column("TotalValue", cellStyle={"textAlign": "center"})
                AgGrid(eng_summary, gridOptions=gb.build(), height=420, fit_columns_on_grid_load=True)
            except Exception:
                st.dataframe(eng_summary.style.set_properties(**{"text-align": "center"}), use_container_width=True)

            # line chart: total value over time by engineer (top 10 by total value)
            if "VisitDate" in cur.columns:
                top_engs = eng_summary.head(10)["Engineer"].tolist()
                df = cur[cur["Engineer"].isin(top_engs)].dropna(subset=["VisitDate", "TotalValue"]).copy()
                if not df.empty:
                    df["Day"] = df["VisitDate"].dt.floor("D")
                    series = (
                        df.groupby(["Day", "Engineer"])["TotalValue"]
                        .sum().reset_index().sort_values(["Day", "Engineer"])
                    )
                    st.caption("Line — total value over time (top 10 engineers by total value)")
                    fig = px.line(
                        series, x="Day", y="TotalValue",
                        color="Engineer", markers=True,
                        title="Total value over time by Engineer"
                    )
                    fig.update_layout(margin=dict(l=0, r=0, b=0, t=48),
                                    xaxis_title="", yaxis_title="Total Value (£)")
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No rows available to plot after filtering.")
            else:
                st.info("Need Date of visit for the engineer trend chart.")
        else:
            st.info("Need Engineer and Total Value columns for this tab.")
    # ---------- 💸 Rechargeable (filters to RechargeableFlag == True)
    with tabs[7]:
        if "RechargeableFlag" in cur.columns and "TotalValue" in cur.columns:
            r = cur[cur["RechargeableFlag"] == True].copy()  # noqa: E712
            if r.empty:
                st.info("No Rechargeable rows in the current filters/date range.")
            else:
                # KPIs
                total_val = float(pd.to_numeric(r["TotalValue"], errors="coerce").fillna(0).sum())
                n = int(len(r))
                avg_val = (total_val / n) if n else 0.0
                k1, k2, k3 = st.columns(3)
                k1.metric("Rechargeable · Total", f"£{total_val:,.2f}")
                k2.metric("Invoices", f"{n:,}")
                k3.metric("Avg / invoice", f"£{avg_val:,.2f}")

                # Line: total value over time (daily)
                if "VisitDate" in r.columns:
                    rs = (
                        r.dropna(subset=["VisitDate", "TotalValue"])
                        .assign(Day=lambda d: d["VisitDate"].dt.floor("D"))
                        .groupby("Day")["TotalValue"].sum().reset_index()
                    )
                    fig = px.line(rs, x="Day", y="TotalValue", markers=True,
                                title="Rechargeable — Total value over time")
                    fig.update_layout(margin=dict(l=0, r=0, b=0, t=48), xaxis_title="", yaxis_title="Total Value")
                    st.plotly_chart(fig, use_container_width=True)

                # By team
                if "Team" in r.columns:
                    byteam = (
                        r.groupby("Team")["TotalValue"].sum()
                        .sort_values(ascending=False).reset_index()
                    )
                    fig = px.bar(byteam, x="Team", y="TotalValue", title="Rechargeable — by team")
                    fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
                    st.plotly_chart(fig, use_container_width=True)

                # By client (top 25)
                if "Client" in r.columns:
                    byclient = (
                        r.groupby("Client")["TotalValue"].sum()
                        .sort_values(ascending=False).reset_index().head(25)
                    )
                    fig = px.bar(byclient, x="Client", y="TotalValue", title="Rechargeable — top clients")
                    fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
                    st.plotly_chart(fig, use_container_width=True)

                # Table
                st.caption("Rechargeable — table")
                preferred_cols = [c for c in [
                    "Date of visit", "Team", "Engineers Name", "Invoice Type", "Visit Type",
                    "Client Name", "VR Number", "Ticket Number", "ASA Number", "Venue Name",
                    "PostCode", "Time On-Site", "Equipment Used", "Additional Info",
                    "Equipment Value", "Labour Value", "Additional Costs", "Hotel/ Food Value",
                    "Total Value", "Rechargeable"
                ] if c in r.columns]
                table_df = r[preferred_cols].copy() if preferred_cols else r
                try:
                    from st_aggrid import AgGrid, GridOptionsBuilder
                    gb = GridOptionsBuilder.from_dataframe(table_df)
                    gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=12)
                    gb.configure_default_column(editable=False, filter=True, sortable=True)
                    AgGrid(table_df, gridOptions=gb.build(), height=420, fit_columns_on_grid_load=True)
                except Exception:
                    st.dataframe(table_df, use_container_width=True)
        else:
            st.info("Need Rechargeable and Total Value columns for this tab.")


    # ---------- 🧾 Non Chargeable (filters to RechargeableFlag == False)
    with tabs[8]:
        if "RechargeableFlag" in cur.columns and "TotalValue" in cur.columns:
            nc = cur[cur["RechargeableFlag"] == False].copy()  # noqa: E712
            if nc.empty:
                st.info("No Non Chargeable rows in the current filters/date range.")
            else:
                # KPIs
                total_val = float(pd.to_numeric(nc["TotalValue"], errors="coerce").fillna(0).sum())
                n = int(len(nc))
                avg_val = (total_val / n) if n else 0.0
                k1, k2, k3 = st.columns(3)
                k1.metric("Non Chargeable · Total", f"£{total_val:,.2f}")
                k2.metric("Invoices", f"{n:,}")
                k3.metric("Avg / invoice", f"£{avg_val:,.2f}")

                # Line: total value over time (daily)
                if "VisitDate" in nc.columns:
                    ns = (
                        nc.dropna(subset=["VisitDate", "TotalValue"])
                        .assign(Day=lambda d: d["VisitDate"].dt.floor("D"))
                        .groupby("Day")["TotalValue"].sum().reset_index()
                    )
                    fig = px.line(ns, x="Day", y="TotalValue", markers=True,
                                title="Non Chargeable — Total value over time")
                    fig.update_layout(margin=dict(l=0, r=0, b=0, t=48), xaxis_title="", yaxis_title="Total Value")
                    st.plotly_chart(fig, use_container_width=True)

                # By team
                if "Team" in nc.columns:
                    byteam = (
                        nc.groupby("Team")["TotalValue"].sum()
                        .sort_values(ascending=False).reset_index()
                    )
                    fig = px.bar(byteam, x="Team", y="TotalValue", title="Non Chargeable — by team")
                    fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
                    st.plotly_chart(fig, use_container_width=True)

                # By client (top 25)
                if "Client" in nc.columns:
                    byclient = (
                        nc.groupby("Client")["TotalValue"].sum()
                        .sort_values(ascending=False).reset_index().head(25)
                    )
                    fig = px.bar(byclient, x="Client", y="TotalValue", title="Non Chargeable — top clients")
                    fig.update_layout(xaxis={'categoryorder': 'total descending'}, margin=dict(l=0, r=0, b=0, t=40))
                    st.plotly_chart(fig, use_container_width=True)

                # Table
                st.caption("Non Chargeable — table")
                preferred_cols = [c for c in [
                    "Date of visit", "Team", "Engineers Name", "Invoice Type", "Visit Type",
                    "Client Name", "VR Number", "Ticket Number", "ASA Number", "Venue Name",
                    "PostCode", "Time On-Site", "Equipment Used", "Additional Info",
                    "Equipment Value", "Labour Value", "Additional Costs", "Hotel/ Food Value",
                    "Total Value", "Rechargeable"
                ] if c in nc.columns]
                table_df = nc[preferred_cols].copy() if preferred_cols else nc
                try:
                    from st_aggrid import AgGrid, GridOptionsBuilder
                    gb = GridOptionsBuilder.from_dataframe(table_df)
                    gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=12)
                    gb.configure_default_column(editable=False, filter=True, sortable=True)
                    AgGrid(table_df, gridOptions=gb.build(), height=420, fit_columns_on_grid_load=True)
                except Exception:
                    st.dataframe(table_df, use_container_width=True)
        else:
            st.info("Need Rechargeable and Total Value columns for this tab.")

    # ---------- 📋 Table
    with tabs[9]:
        preferred_cols = [c for c in [
            "Date of visit", "Team", "Engineers Name", "Invoice Type", "Visit Type",
            "Client Name", "VR Number", "Ticket Number", "ASA Number", "Venue Name",
            "PostCode", "Time On-Site", "Equipment Used", "Additional Info",
            "Equipment Value", "Labour Value", "Additional Costs", "Hotel/ Food Value",
            "Total Value"
        ] if c in cur.columns]

        table_df = cur[preferred_cols].copy() if preferred_cols else cur.copy()
        try:
            from st_aggrid import AgGrid, GridOptionsBuilder
            gb = GridOptionsBuilder.from_dataframe(table_df)
            gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=12)
            gb.configure_default_column(editable=False, filter=True, sortable=True)
            AgGrid(table_df, gridOptions=gb.build(), height=420, fit_columns_on_grid_load=True)
        except Exception:
            st.dataframe(table_df, use_container_width=True)

    # ---------- ⬇️ Export
    with tabs[10]:
        csv = cur.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            "Download filtered CSV",
            data=csv,
            file_name="Invoices_filtered.csv",
            mime="text/csv",
        )


def render_team_engineers_page():
    import pandas as pd, numpy as np
    from datetime import datetime
    st.markdown("""
    <style>
    .engineer-card{
    background:#10161d; border:1px solid #263445; border-radius:18px;
    padding:16px 18px; box-shadow:0 8px 18px rgba(0,0,0,.25);
    min-height: 210px;
    color:#f8fafc !important;                    /* <-- make default text white */
    }
    /* Force all descendants to white so bullets/headings also show */
    .engineer-card *, .engineer-card h1, .engineer-card h2, .engineer-card h3, .engineer-card h4,
    .engineer-card p, .engineer-card li, .engineer-card strong, .engineer-card span, .engineer-card em {
    color:#f8fafc !important;
    }

    .engineer-card h4{
    margin:0 0 6px 0; font-weight:700; color:#e6f0ff !important; /* heading a bit brighter */
    }
    .engineer-card .eyebrow{
    color:#9fb3c8 !important; font-size:.8rem; margin-bottom:6px;
    }
    .engineer-card ul{ list-style:none; padding-left:0; margin:0; }
    .engineer-card li{ margin: 3px 0; }
    .engineer-kpi b{ font-weight:700; }

    </style>
    """, unsafe_allow_html=True)
    import base64


    # --- Sky Engineers logo at top of Engineer page ---
    logo_path = "Sky Engineers.png"  # ensure this is in the same folder as operations_dashboard.py
    with open(logo_path, "rb") as f:
        logo_bytes = f.read()
    logo_b64 = base64.b64encode(logo_bytes).decode()

    st.markdown(
        f"""
        <div style="text-align:center; margin-bottom:15px;">
            <img src="data:image/png;base64,{logo_b64}" 
                alt="Sky Engineers Logo" 
                style="max-width:750px; height:auto;">
        </div>
        """,
        unsafe_allow_html=True
    )

    # ... your existing Year/Month selectors etc. below ...


    user_team = (st.session_state.get("user_team") or "VIP North").strip()
    # --- Back button (top of Engineer page) ---
    st.markdown("<div style='margin-bottom:10px;'>", unsafe_allow_html=True)
    if st.button("⬅️ Back to Menu", key="btn_back_home", use_container_width=False):
        st.session_state.screen = "instruction_guide"
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)

    # --- Load source data
    inv = load_invoices()                 # you already have this loader
    vis = get_team_df(user_team).copy()   # per-team Oracle data

    # Normalise dates
    for df in (inv, vis):
        if df is not None and not df.empty and "Date" in df.columns:
            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")

    # --- Year/Month filter (YTD default)
    y = datetime.now(UK_TZ).year
    colY, colM = st.columns(2)
    with colY:
        year_choice = st.selectbox("Year", [y], index=0)
    with colM:
        month_choice = st.selectbox("Month", ["YTD"], index=0)

    # Build YTD window
    start_dt = pd.Timestamp(year_choice, 1, 1)
    end_dt   = pd.Timestamp(year_choice, 12, 31)

    def in_window(df):
        if df is None or df.empty or "Date" not in df.columns: return df
        return df[(df["Date"] >= start_dt) & (df["Date"] <= end_dt)].copy()

    inv = in_window(inv)
    vis = in_window(vis)

    def _to_minutes(x):
        """Return minutes from HH:MM[:SS], datetime-like, or int-like. NaN if invalid/zero."""
        import numpy as np, pandas as pd
        from datetime import datetime, time
        if x is None: return np.nan
        # datetime/time
        if isinstance(x, (pd.Timestamp, datetime)): return x.hour*60 + x.minute
        if isinstance(x, time): return x.hour*60 + x.minute
        s = str(x).strip().lower()
        if s in {"", "-", "—", "na", "n/a", "none", "nat", "0", "00", "00:00", "00:00:00"}:
            return np.nan
        # try datetime parser first (catches '1899-12-30 08:30:00')
        try:
            dt = pd.to_datetime(s, errors="raise")
            return int(dt.hour)*60 + int(dt.minute)
        except Exception:
            pass
        # HH:MM[:SS]
        if ":" in s:
            parts = s.split(":")
            if len(parts) >= 2:
                try:
                    h = int(parts[0]); m = int(parts[1])
                    if h==0 and m==0: return np.nan
                    return h*60 + m
                except Exception:
                    return np.nan
        # raw minutes
        try:
            v = float(s)
            if v == 0: return np.nan
            return int(v)
        except Exception:
            return np.nan

    def _avg_hhmm(series_like):
        """Average of a numeric minutes sequence -> 'HH:MM' or '—'."""
        import pandas as pd
        # NEW: always work with a Series
        s = series_like if isinstance(series_like, pd.Series) else pd.Series(series_like)
        v = pd.to_numeric(s, errors="coerce").dropna()
        if v.empty:
            return "—"
        m = int(v.mean())
        return f"{m//60:02d}:{m%60:02d}"


    def _clean_minutes(series):
        """Map to minutes and drop NaNs/zeros."""
        import pandas as pd
        return pd.Series(series).map(_to_minutes).dropna()

    def _minutes_from_any(series):
        """
        Convert a column of durations (strings like '00:27:33', '0:27', timedeltas,
        Excel-style datetimes, or numeric minutes) into minutes.
        Zeros/invalids become NaN and are excluded.
        """
        import numpy as np, pandas as pd

        s = pd.Series(series)

        # Timedelta -> minutes
        if pd.api.types.is_timedelta64_dtype(s):
            mins = s.dt.total_seconds() / 60.0

        # Datetime/time-of-day -> minutes past midnight
        elif pd.api.types.is_datetime64_any_dtype(s):
            mins = s.dt.hour * 60 + s.dt.minute

        else:
            # Normalise text; remove obvious "zero"/blank tokens
            s2 = (
                s.astype(str)
                .str.strip()
                .replace({
                    "": None, "-": None, "—": None,
                    "na": None, "n/a": None, "None": None, "NaT": None, "nan": None,
                    "0": None, "00": None, "00:00": None, "00:00:00": None
                })
            )

            # Try HH:MM[:SS] via to_timedelta first
            td = pd.to_timedelta(s2, errors="coerce")
            mins = td.dt.total_seconds() / 60.0

            # Where that failed, try as datetime (Excel sometimes stores times as datetimes)
            mask = mins.isna()
            if mask.any():
                dt = pd.to_datetime(s2[mask], errors="coerce")
                mins.loc[mask] = (dt.dt.hour * 60 + dt.dt.minute)

            # Numeric fallback (already in minutes)
            mask = mins.isna()
            if mask.any():
                mins.loc[mask] = pd.to_numeric(s2[mask], errors="coerce")

        mins = pd.to_numeric(mins, errors="coerce")
        # drop zero and NaN
        mins = mins.where(mins != 0).dropna()
        return mins

    def _hhmm_from_minutes(mean_minutes):
        """Format a single numeric minutes value as HH:MM (returns '—' if NaN)."""
        import numpy as np
        if mean_minutes is None or (isinstance(mean_minutes, float) and np.isnan(mean_minutes)):
            return "—"
        m = int(round(float(mean_minutes)))
        return f"{m//60:02d}:{m%60:02d}"
    
    

    # ---------- helpers (from your Team Overview) ----------
    def to_minutes_from_time(x):
        """
        Convert times like '08:30', '08:30:00', Timestamp('1899-12-30 08:30:00'),
        or integer-like minutes to total minutes. Returns NaN if not parseable.
        """
        import numpy as np
        import pandas as pd
        from datetime import datetime, time

        if x is None or (isinstance(x, float) and pd.isna(x)):
            return np.nan

        # pandas / python datetime or time objects
        if isinstance(x, (pd.Timestamp, datetime)):
            return x.hour * 60 + x.minute
        if isinstance(x, time):
            return x.hour * 60 + x.minute

        s = str(x).strip()
        if s == "" or s in {"-", "—", "None", "NaT"}:
            return np.nan

        # Try generic datetime parsing (handles '1899-12-30 08:30:00', '08:30:00', etc.)
        try:
            dt = pd.to_datetime(s, errors="raise")
            return int(dt.hour) * 60 + int(dt.minute)
        except Exception:
            pass

        # Handle HH:MM or HH:MM:SS
        if ":" in s:
            parts = s.split(":")
            if len(parts) >= 2:
                try:
                    h = int(parts[0])
                    m = int(parts[1])
                    return h * 60 + m
                except Exception:
                    return np.nan

        # Plain integer minutes
        try:
            return int(float(s))
        except Exception:
            return np.nan


    def mean_hhmm(series_mins):
        v = pd.to_numeric(series_mins, errors="coerce").dropna()
        if v.empty: return "—"
        m = int(v.mean()); return f"{m//60:02d}:{m%60:02d}"

    # Column detection
    eng_col   = next((c for c in vis.columns if c.lower() in ("engineer","engineers name","engineers","engineer name","name")), None)
    act_col   = next((c for c in vis.columns if "activate"   in c.lower()), None)                # Activate = start
    deact_col = next((c for c in vis.columns if "deactivate" in c.lower()), None)                # Deactivate = finish
    lowd_col  = next((c for c in vis.columns if "total working time" in c.lower()), None)        # LOWD source
    lunch_col = next((c for c in vis.columns if "lunch" in c.lower()), None)                     # lunch mins if present
    ot_col    = next((c for c in vis.columns if "overtime" in c.lower()), None)                  # OT cost column
    date_col  = "Date" if "Date" in vis.columns else next((c for c in vis.columns if c.lower()=="date"), None)
    status_col = next((c for c in vis.columns if c.lower() in ("activity status","status")), None)
    visit_col  = next((c for c in vis.columns if c.lower() in ("visit type","job type","type")), None)
    # Prefer Daniel's column per your rules; fall back to "Total Time" if needed
    total_time_for_lunch_col = None
    for cand in ("Total Time for AI","Total Time","Total time","total time"):
        if cand in vis.columns:
            total_time_for_lunch_col = cand
            break


    def is_holiday_row(r):
        if type_col and isinstance(r.get(type_col,""), str):
            return "holiday" in r[type_col].lower()
        return False

    def is_rr_sb_row(r):
        if type_col and isinstance(r.get(type_col,""), str):
            t = r[type_col].lower()
            return ("rr" in t) or ("sb cover" in t) or ("standby" in t)
        return False
    
    
    # ─────────────────────────────────────────────────────────────
    # RECON A: Visit Number (Oracle) vs VR Number (Invoices)
    # ─────────────────────────────────────────────────────────────
    import pandas as pd, numpy as np, re

    team_name = (st.session_state.get("user_team") or "").strip()

    # --- Column detection (Oracle)
    eng_col   = next((c for c in vis.columns if c.lower() in ("engineer","engineers name","engineers","engineer name","name")), None)
    date_col  = "Date" if "Date" in vis.columns else next((c for c in vis.columns if c.lower()=="date"), None)
    status_col= next((c for c in vis.columns if c.lower() in ("activity status","status")), None)
    visit_col = next((c for c in vis.columns if c.lower() in ("visit type","job type","type")), None)
    visit_no_col = next((c for c in vis.columns if "visit number" in c.lower()), None)

    def _norm_visit(s):
        if pd.isna(s): return ""
        s = str(s).strip().casefold()
        s = re.sub(r"\s+", " ", s)
        s = s.replace("pick-up","pick up").replace("pickup","pick up")
        return s
# Use 1 June of the selected year for KPI windows
    kpi_start = pd.Timestamp(start_dt.year, 6, 1)
    kpi_end   = end_dt  # keep your existing end bound

    v_oracle = vis.copy()
    if date_col and date_col in v_oracle.columns:
        v_oracle[date_col] = pd.to_datetime(v_oracle[date_col], errors="coerce")
        v_oracle = v_oracle[(v_oracle[date_col] >= kpi_start) & (v_oracle[date_col] <= kpi_end)]




    # Completed, not lunch/stock
    mask_req = pd.Series(True, index=v_oracle.index)
    if status_col in v_oracle.columns:
        mask_req &= v_oracle[status_col].astype(str).str.strip().str.casefold().eq("completed")
    if visit_col in v_oracle.columns:
        vt_clean = v_oracle[visit_col].map(_norm_visit)
        mask_req &= ~vt_clean.isin({"lunch (30)","lunch(30)","lunch 30","lunch",
                                    "stock pick up","stock-pick up","stock pick-up"})

    req_visits = v_oracle[mask_req].copy()

    # Oracle visit numbers (only non-empty, non-zero)
    oracle_visit_nums = pd.Series(dtype=str)
    if visit_no_col and visit_no_col in req_visits.columns:
        oracle_visit_nums = (
            req_visits[visit_no_col]
            .astype(str).str.strip()
            .replace({"": np.nan, "0": np.nan, "None": np.nan, "none": np.nan})
            .dropna()
        )

    # Oracle / visits
    visit_col   = visit_col   if 'visit_col'   in locals() and visit_col   else next((c for c in vis.columns if c.lower() in ("visit type","job type","type")), None)
    status_col  = status_col  if 'status_col'  in locals() and status_col  else next((c for c in vis.columns if c.lower() in ("activity status","status")), None)
    date_col    = date_col    if 'date_col'    in locals() and date_col    else ("Date" if "Date" in vis.columns else next((c for c in vis.columns if c.lower()=="date"), None))

    # Invoices
    inv_eng_col  = next((c for c in inv.columns if ("engineer" in c.lower()) or (c.lower()=="name")), None) if (inv is not None and not inv.empty) else None
    inv_date_col = next((c for c in inv.columns if c.lower() in ("date of visit","date","invoice date","created","created date")), None) if (inv is not None and not inv.empty) else None
    inv_type_col = next((c for c in inv.columns if c.lower() in ("invoice type","invoice_type")), None) if (inv is not None and not inv.empty) else None

    # --- Column detection (Invoices)
    v_invoices = inv.copy() if (inv is not None and not inv.empty) else pd.DataFrame()
    team_col   = next((c for c in v_invoices.columns if c.lower()=="team"), None) if not v_invoices.empty else None
    inv_date_col = next((c for c in v_invoices.columns if c.lower() in ("date of visit","date","invoice date","created","created date")), None) if not v_invoices.empty else None
    vr_col = next((c for c in v_invoices.columns if c.lower() in ("vr number","vrnumber","vr no","vr")), None) if not v_invoices.empty else None
    inv_type_col = next((c for c in v_invoices.columns if c.lower()=="invoice type"), None) if not v_invoices.empty else None
    inv_eng_col = next((c for c in v_invoices.columns if "engineer" in c.lower() or c.lower()=="name"), None) if not v_invoices.empty else None

    if not v_invoices.empty:
        if team_col:
            v_invoices = v_invoices[v_invoices[team_col].astype(str).str.strip().str.casefold()==team_name.casefold()]
        if inv_date_col and inv_date_col in v_invoices.columns:
            v_invoices[inv_date_col] = pd.to_datetime(v_invoices[inv_date_col], errors="coerce")
            v_invoices = v_invoices[(v_invoices[inv_date_col] >= kpi_start) & (v_invoices[inv_date_col] <= kpi_end)]


    # Invoice VR numbers (clean)
    invoice_vrs = pd.Series(dtype=str)
    if (vr_col and (vr_col in v_invoices.columns)):
        invoice_vrs = (
            v_invoices[vr_col]
            .astype(str).str.strip()
            .replace({"": np.nan, "0": np.nan, "None": np.nan, "none": np.nan})
            .dropna()
        )

    # Outstanding by number: Oracle Visit Numbers that don't exist as VR Numbers
    outstanding_by_num = pd.Series(dtype=str)
    if not oracle_visit_nums.empty:
        outstanding_by_num = oracle_visit_nums[~oracle_visit_nums.isin(set(invoice_vrs))]

    # Details table for outstanding by number
    out_by_num_df = pd.DataFrame()
    if not outstanding_by_num.empty and visit_no_col and visit_no_col in req_visits.columns:
        out_by_num_df = (
            req_visits[req_visits[visit_no_col].astype(str).str.strip().isin(set(outstanding_by_num))]
            .copy()
        )
        keep_cols = [c for c in (eng_col, date_col, visit_col, visit_no_col, "Visit Notes") if c and c in out_by_num_df.columns]
        if keep_cols:
            out_by_num_df = out_by_num_df[keep_cols].sort_values(by=[eng_col, date_col] if (eng_col and date_col) else keep_cols)


    # KPIs for Recon A
    kpi_required_by_num  = int(oracle_visit_nums.nunique())
    kpi_completed_by_vr  = int(invoice_vrs.nunique())
    kpi_outstanding_by_num = max(0, kpi_required_by_num - kpi_completed_by_vr)

    leftA, rightA = st.columns(2)
    with leftA:
        st.metric("Invoices completed (YTD)", f"{kpi_completed_by_vr:,}")
        if not v_invoices.empty:
            # Show what we matched on (VR list with some useful cols)
            show_cols = [x for x in (vr_col, inv_eng_col, team_col, inv_date_col, inv_type_col) if x and x in v_invoices.columns]
            if show_cols:
                with st.expander("View completed details"):
                    st.dataframe(v_invoices[show_cols].sort_values(by=[team_col, inv_eng_col, inv_date_col] if team_col and inv_eng_col and inv_date_col in v_invoices.columns else show_cols),
                                use_container_width=True, hide_index=True)
                    st.download_button("Download Completed CSV",
                                    v_invoices[show_cols].to_csv(index=False).encode("utf-8"),
                                    "invoices_completed.csv","text/csv")
    with rightA:
        st.metric("Invoices missing (YTD)", f"{kpi_outstanding_by_num:,}")
        if not out_by_num_df.empty:
            with st.expander("View outstanding details"):
                st.dataframe(out_by_num_df, use_container_width=True, hide_index=True)
                st.download_button("Download Outstanding CSV",
                                out_by_num_df.to_csv(index=False).encode("utf-8"),
                                "invoices_outstanding_by_number.csv","text/csv")

    # Spacer
    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────
    # RECON B: VIP – SB Standby (visits) vs Business Cover (invoices)
    # ─────────────────────────────────────────────────────────────
    # Oracle side: count completed VIP – SB Standby visits
    vip_sb_aliases = {
        "vip - sb standby","vip sb standby","vip—sb standby","vip – sb standby","vip- sb standby",
        "vip - sb stby","vip sb stby","vip-sb standby"
    }
    visit_type_clean = v_oracle[visit_col].map(_norm_visit) if (visit_col and visit_col in v_oracle.columns) else pd.Series([""], index=v_oracle.index)
    vip_sb_count = int(v_oracle[(visit_type_clean.isin(vip_sb_aliases)) &
                                (v_oracle[status_col].astype(str).str.strip().str.casefold()=="completed" if status_col in v_oracle.columns else True)
                            ].shape[0])

    vip_sb_details = pd.DataFrame()
    if vip_sb_count and visit_col in v_oracle.columns:
        vip_sb_details = v_oracle[(visit_type_clean.isin(vip_sb_aliases))].copy()
        keep_cols = [c for c in (eng_col, date_col, visit_col, visit_no_col, "Visit Notes") if c and c in vip_sb_details.columns]
        if keep_cols:
            vip_sb_details = vip_sb_details[keep_cols].sort_values(by=[eng_col, date_col] if (eng_col and date_col) else keep_cols)

    # Invoices side: count Business Cover invoices for this team
    biz_cover_count = 0
    biz_cover_details = pd.DataFrame()
    if not v_invoices.empty and inv_type_col and (inv_type_col in v_invoices.columns):
        biz_mask = v_invoices[inv_type_col].astype(str).str.strip().str.casefold() == "business cover"
        biz_cover_details = v_invoices[biz_mask].copy()
        biz_cover_count = int(biz_cover_details.shape[0])

    # Outstanding for this pairing
    outstanding_sb_vs_bc = max(0, vip_sb_count - biz_cover_count)

    leftB, rightB = st.columns(2)
    with leftB:
        st.metric("VIP – SB Standby visits (Completed, YTD)", f"{vip_sb_count:,}")
        if not vip_sb_details.empty:
            with st.expander("View VIP – SB Standby visit details"):
                st.dataframe(vip_sb_details, use_container_width=True, hide_index=True)
                st.download_button("Download VIP–SB Standby CSV",
                                vip_sb_details.to_csv(index=False).encode("utf-8"),
                                "visits_vip_sb_standby.csv","text/csv")
    with rightB:
        st.metric("Business Cover invoices (YTD)", f"{biz_cover_count:,}")
        st.metric("Outstanding (VIP–SB Standby minus Business Cover)", f"{outstanding_sb_vs_bc:,}")
        if not biz_cover_details.empty:
            with st.expander("View Business Cover invoice details"):
                show_cols = [c for c in (vr_col, inv_eng_col, team_col, inv_date_col, inv_type_col) if c and c in biz_cover_details.columns]
                st.dataframe(biz_cover_details[show_cols] if show_cols else biz_cover_details,
                            use_container_width=True, hide_index=True)
                st.download_button("Download Business Cover CSV",
                                (biz_cover_details[show_cols] if show_cols else biz_cover_details).to_csv(index=False).encode("utf-8"),
                                "invoices_business_cover.csv","text/csv")
    # ─────────────────────────────────────────────────────────────


        # ---------- Page header ----------
        st.subheader(f"👷 Engineers — {user_team} (YTD)")
        if vis is None or vis.empty or not eng_col:
            st.info("No visit data with engineer names available.")
            return

        from pathlib import Path
        import pandas as pd
        import re

        EXP_MASTER = Path("Expenses/expenses_master.parquet")

        # ---- helpers (unique, non-duplicated) ------------------------------------
        def _coerce_money(s: pd.Series) -> pd.Series:
            """Robust money parser: handles £, commas, Unicode minus, and (123.45) negatives."""
            return pd.to_numeric(
                s.astype(str)
                .str.replace("’", "'", regex=False)
                .str.replace("−", "-", regex=False)              # Unicode minus → ASCII
                .str.replace(r"\(([^)]+)\)", r"-\1", regex=True) # (123.45) → -123.45
                .str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)", expand=False)
                .str.replace(",", "", regex=False),
                errors="coerce",
            )

        def _canon_eng_name(name: str) -> str:
            t = str(name or "").strip().replace("’", "'").casefold()
            # remove finance boilerplate / junk
            t = re.sub(r"\b(cash|advance|utili[sz]ed|amount|cash\s+advance\s+utili[sz]ed\s+amount)\b", " ", t)
            # keep letters, spaces, apostrophes
            t = re.sub(r"[^a-z' ]+", " ", t)
            t = re.sub(r"\s+", " ", t).strip()
            # normalise O'Leary variants
            t = t.replace("o leary", "o'leary").replace("oleary", "o'leary")

            # explicit mappings
            if "michael webb" in t:  return "Michael Webb"
            if "simon wakelin" in t: return "Simon Wakelin"
            if "david o'leary" in t or "david oleary" in t:
                return "David O'Leary"

            parts = []
            for w in t.split():
                if "'" in w:
                    a, b = w.split("'", 1)
                    parts.append(a.capitalize() + "'" + b.capitalize())
                else:
                    parts.append(w.capitalize())
            return " ".join(parts)

        def _norm_visit(x: str) -> str:
            """Normalise visit type text (spaces, dashes, case)."""
            t = str(x or "").strip().casefold()
            t = re.sub(r"\s+", " ", t)
            t = t.replace("–", "-").replace("—", "-")
            t = t.replace("pick-up", "pick up").replace("pickup", "pick up")
            return t

        def sum_expenses_for_engineer(e_name: str, start_dt=None, end_dt=None) -> float:
            if not EXP_MASTER.exists():
                return 0.0
            df = pd.read_parquet(EXP_MASTER)
            if df.empty or "Engineer Name" not in df.columns:
                return 0.0

            # canonicalise both sides
            canon_e = _canon_eng_name(e_name)
            df["__canon"] = df["Engineer Name"].apply(_canon_eng_name)
            df = df[df["__canon"] == canon_e].copy()
            if df.empty:
                return 0.0

            # date window (coerce just in case)
            if "Transaction Date" in df.columns:
                df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
                if start_dt is not None:
                    df = df[df["Transaction Date"] >= pd.to_datetime(start_dt)]
                if end_dt is not None:
                    df = df[df["Transaction Date"] <= pd.to_datetime(end_dt)]
            if df.empty or "Amount" not in df.columns:
                return 0.0

            return float(_coerce_money(df["Amount"]).fillna(0.0).sum())

        # Show one tile per engineer (same metrics you already compute)
        present = sorted(vis[eng_col].dropna().astype(str).unique().tolist())

        def engineer_metrics(e_name: str):
            import numpy as np
            import pandas as pd

            # Slice by engineer
            v = vis[vis[eng_col].astype(str).str.strip().str.lower() == str(e_name).strip().lower()].copy()
            if v.empty:
                return {
                    "avg_lunch":"—","avg_start":"—","avg_finish":"—","avg_lowd":"—",
                    "days_worked":0,"expenses":"£0","ot_cost":"—",
                    "avg_ot_entry":"—","avg_ot_day":"—",
                    "invoices_completed_march":0,"visits_completed_march":0,
                    "potential_missing_invoices":0,"vip_sb_visits_march":0,
                    "biz_cover_invoices_march":0,"potential_missing_sb_march":0,
                    "since_cutoff": pd.Timestamp(pd.Timestamp.now().year, 6, 1).date().isoformat(),
                    "invoices_completed_since":0,"visits_completed_since":0,
                    "potential_missing_since":0,"vip_sb_visits_since":0,
                    "biz_cover_invoices_since":0,"potential_missing_sb_since":0,
                }

            # ---- YTD window ----
            y = pd.Timestamp.now().year
            ytd_start = pd.Timestamp(y, 1, 1)
            ytd_end   = pd.Timestamp(y, 12, 31)
            if date_col and date_col in v.columns:
                v[date_col] = pd.to_datetime(v[date_col], errors="coerce")
                v = v[(v[date_col] >= ytd_start) & (v[date_col] <= ytd_end)]
            # use these for expense window
            start_dt, end_dt = ytd_start, ytd_end

            # ---- Average start/finish from Activate/Deactivate ----
            avg_start  = _avg_hhmm(_clean_minutes(v[act_col])   if act_col   in v.columns else [])
            avg_finish = _avg_hhmm(_clean_minutes(v[deact_col]) if deact_col in v.columns else [])

            # ---- Average LOWD from Total working time ----
            avg_lowd = _avg_hhmm(_clean_minutes(v[lowd_col]) if lowd_col in v.columns else [])

            # ---- Average Lunch: Completed & Visit Type == 'Lunch (30)' ----
            avg_lunch = "—"
            if all([status_col, visit_col, total_time_for_lunch_col]):
                vc_status = v[status_col].astype(str).str.strip().str.casefold()
                vc_visit  = v[visit_col].astype(str).str.strip().str.casefold().str.replace(r"\s+", " ", regex=True)
                lunch_mask = (vc_status == "completed") & (vc_visit == "lunch (30)")
                if lunch_mask.any():
                    mins = _minutes_from_any(v.loc[lunch_mask, total_time_for_lunch_col])
                    avg_lunch = _hhmm_from_minutes(mins.mean()) if not mins.empty else "—"

            # ---- Days worked ----
            days_worked = int(v[date_col].dt.date.nunique()) if date_col and (date_col in v.columns) else 0

            # ---- Invoice-file expenses (sum of columns if present) ----
            expenses_from_invoices = 0.0
            if inv is not None and not inv.empty:
                # try to pick an engineer column if not predefined
                e_col = None
                for cand in [inv_eng_col] if 'inv_eng_col' in globals() else []:
                    if cand and cand in inv.columns: e_col = cand; break
                if e_col is None:
                    e_col = next((c for c in inv.columns if c.lower() in {"engineer","name"} or "engineer" in c.lower()), None)

                inv_e = inv if e_col is None else inv[inv[e_col].astype(str).str.strip().str.lower() == str(e_name).strip().lower()]
                if not inv_e.empty:
                    for c in ("Additional Costs","Hotel/ Food Value","Hotel/ Food Value ","Hotel/Food Value",
                            "Equipment Value","Expenses","Travel"):
                        if c in inv_e.columns:
                            expenses_from_invoices += pd.to_numeric(inv_e[c], errors="coerce").fillna(0).sum()

            # ---- OT (sum from Feb 15 YTD) ----
            ot_total = None
            ot_count = 0
            avg_ot_per_entry = "—"
            avg_ot_per_day = "—"
            if ot_col and (ot_col in v.columns) and date_col and (date_col in v.columns):
                ot_df = v[pd.to_datetime(v[date_col], errors="coerce") >= pd.Timestamp(y, 2, 15)].copy()
                ot_vals = (ot_df[ot_col].astype(str).str.replace(r"[£,\s]", "", regex=True).replace({"": None}))
                ot_vals = pd.to_numeric(ot_vals, errors="coerce")
                ot_count = int(ot_vals[ot_vals > 0].count())
                total_val = float(ot_vals.fillna(0).sum())
                ot_total = total_val
                avg_ot_per_entry = f"£{(total_val / ot_count):,.2f}" if ot_count > 0 else "—"
            if isinstance(ot_total, (int, float)):
                avg_ot_per_day_val = (ot_total / days_worked) if days_worked else 0.0
                avg_ot_per_day = f"£{avg_ot_per_day_val:,.2f}"

            # ----- Since 1 June (current year) -----
            since_cutoff = pd.Timestamp(y, 6, 1)

            # (A) Visits completed (exclude Lunch & Stock Pick Up)
            visits_comp_since = 0
            if date_col and (date_col in v.columns):
                vm = v.copy()
                vm[date_col] = pd.to_datetime(vm[date_col], errors="coerce")
                mask = vm[date_col] >= since_cutoff
                if status_col and (status_col in vm.columns):
                    mask &= vm[status_col].astype(str).str.strip().str.casefold().eq("completed")
                if visit_col and (visit_col in vm.columns):
                    vt = vm[visit_col].apply(_norm_visit)
                    mask &= ~vt.isin({"lunch (30)", "lunch", "stock pick up"})
                visits_comp_since = int(vm[mask].shape[0])

            # (B) Invoices completed since cutoff (row count)
            inv_count_since = 0
            if (inv is not None) and (not inv.empty):
                # engineer column
                e_col = next((c for c in inv.columns if c.lower() in {"engineer","name"} or "engineer" in c.lower()), None)
                inv_e = inv if e_col is None else inv[inv[e_col].astype(str).str.strip().str.casefold()
                                                    == str(e_name).strip().casefold()].copy()
                if not inv_e.empty:
                    # date column for invoices
                    date_cand = next((c for c in inv_e.columns
                                    if c.lower() in {"date","invoice date","completed date","completion date","job date","raised date"}
                                    or ("date" in c.lower())), None)
                    if date_cand:
                        inv_e[date_cand] = pd.to_datetime(inv_e[date_cand], errors="coerce")
                        inv_e = inv_e[inv_e[date_cand] >= since_cutoff]
                    inv_count_since = int(inv_e.shape[0])

            # (C) Potential missing = visits - invoices (floor at 0)
            potential_missing_since = max(0, int(visits_comp_since) - int(inv_count_since))

            # VIP – SB Standby visits since cutoff (Oracle)
            vip_sb_visits_since = 0
            if date_col and (date_col in v.columns) and visit_col and (visit_col in v.columns):
                vv = v.copy()
                vv[date_col] = pd.to_datetime(vv[date_col], errors="coerce")
                mask = vv[date_col] >= since_cutoff
                if status_col and (status_col in vv.columns):
                    mask &= vv[status_col].astype(str).str.strip().str.casefold().eq("completed")
                vt = vv[visit_col].apply(_norm_visit)
                vip_sb_aliases = {
                    "vip - sb standby","vip sb standby","vip-sb standby","vip – sb standby","vip—sb standby","vip - sb stby","vip sb stby"
                }
                mask &= vt.isin(vip_sb_aliases)
                vip_sb_visits_since = int(vv[mask].shape[0])

            # Business Cover invoices since cutoff
            biz_cover_invoices_since = 0
            if (inv is not None) and (not inv.empty):
                e_col = next((c for c in inv.columns if c.lower() in {"engineer","name"} or "engineer" in c.lower()), None)
                inv_e = inv if e_col is None else inv[inv[e_col].astype(str).str.strip().str.casefold()
                                                    == str(e_name).strip().casefold()].copy()
                if not inv_e.empty:
                    # date filter
                    date_cand = next((c for c in inv_e.columns
                                    if c.lower() in {"date","invoice date","completed date","completion date","job date","raised date"}
                                    or ("date" in c.lower())), None)
                    if date_cand:
                        inv_e[date_cand] = pd.to_datetime(inv_e[date_cand], errors="coerce")
                        inv_e = inv_e[inv_e[date_cand] >= since_cutoff]
                    # type column
                    type_col = next((c for c in inv_e.columns
                                    if c.lower() in {"invoice type","type","invoicetype","job type","visit type","category"}
                                    or ("type" in c.lower())), None)
                    if type_col:
                        t = inv_e[type_col].astype(str).str.strip().str.casefold()
                        biz_cover_invoices_since = int((t.eq("business cover") | t.str.contains("business cover")).sum())

            # Potential missing SB Standby invoices (never below zero)
            potential_missing_sb_since = max(0, vip_sb_visits_since - biz_cover_invoices_since)

            # --- Expenses total for this engineer & KPI date window (from Expenses master) ---
            expenses_total = sum_expenses_for_engineer(e_name, start_dt=start_dt, end_dt=end_dt)


            # ---------- Back-compat aliases (keep old keys used elsewhere) ----------
            visits_comp_march          = visits_comp_since
            inv_count_march            = inv_count_since
            potential_missing_march    = potential_missing_since
            vip_sb_visits_march        = vip_sb_visits_since
            biz_cover_invoices_march   = biz_cover_invoices_since
            potential_missing_sb_march = potential_missing_sb_since

            # ---------- Return ----------
            return {
                "avg_lunch":  avg_lunch,
                "avg_start":  avg_start,
                "avg_finish": avg_finish,
                "avg_lowd":   avg_lowd,
                "days_worked": days_worked,
                "expenses":   f"£{expenses_total:,.0f}",
                "ot_cost":    (f"£{ot_total:,.0f}" if isinstance(ot_total, (int, float)) else "—"),
                "ot_count":   ot_count,
                "avg_ot_entry": avg_ot_per_entry,
                "avg_ot_day":   avg_ot_per_day,

                # --- canonical “since 1 Jun” values ---
                "since_cutoff": since_cutoff.date().isoformat(),
                "visits_completed_since":   int(visits_comp_since),
                "invoices_completed_since": int(inv_count_since),
                "potential_missing_since":  int(potential_missing_since),
                "potential_missing_invoices": int(potential_missing_since),  # <- single source of truth

                "vip_sb_visits_since":      int(vip_sb_visits_since),
                "biz_cover_invoices_since": int(biz_cover_invoices_since),
                "potential_missing_sb_since": int(potential_missing_sb_since),

                # --- back-compat aliases (mapped to the same 'since' numbers) ---
                "visits_completed_march":   int(visits_comp_since),
                "invoices_completed_march": int(inv_count_since),
                "potential_missing_march":  int(potential_missing_since),
                "vip_sb_visits_march":      int(vip_sb_visits_since),
                "biz_cover_invoices_march": int(biz_cover_invoices_since),
                "potential_missing_sb_march": int(potential_missing_sb_since),
            }




    # ===== Engineer tiles as dropdowns (expanders) =====
    # tiles = [(name, engineer_metrics(name)) for name in present]  # you already have this
    tiles = [(name, engineer_metrics(name)) for name in present]
    PER_ROW = 4  # keep the same row width you like

    # (optional) quick control to open all by default
    # expanded_all = st.checkbox("Expand all engineers", value=False)

    for i in range(0, len(tiles), PER_ROW):
        cols = st.columns(min(PER_ROW, len(tiles) - i), gap="large")
        for (name, m), col in zip(tiles[i:i+PER_ROW], cols):
            with col:
                with st.expander(name, expanded=False):  # or expanded=expanded_all
                    st.markdown(
                        f"""
    <div class="engineer-card">
    <div class="eyebrow">Year to date</div>
    <h4>{name}</h4>
    <ul class="engineer-kpi">
        <li>• Avg lunch: <b>{m['avg_lunch']}</b></li>
        <li>• Avg start: <b>{m['avg_start']}</b></li>
        <li>• Avg finish: <b>{m['avg_finish']}</b></li>
        <li>• Avg LOWD: <b>{m['avg_lowd']}</b></li>
        <li>• Days worked: <b>{m['days_worked']}</b></li>
        <li>• Total expenses: <b>{m['expenses']}</b></li>
        <li>• Total OT cost: <b>{m['ot_cost']}</b></li>
        <li>• Total OT entries: <b>{m['ot_count']}</b></li>
        <li>• Avg OT cost (per entry): <b>{m['avg_ot_entry']}</b></li>
        <li>• Avg OT cost (per working day): <b>{m['avg_ot_day']}</b></li>
        <li>• Visits completed since 1st Jun (excl. Lunch/Stock): <b>{m.get('visits_completed_march', 0)}</b></li>
        <li>• Invoices completed since 1st Jun: <b>{m.get('invoices_completed_march', 0)}</b></li>
        <li>• Potential incomplete / missing invoices: <b>{m.get('potential_missing_march', 0)}</b></li>
        <li>• VIP – SB Standby visits since 1st Jun: <b>{m.get('vip_sb_visits_march', 0)}</b></li>
        <li>• Business Cover invoices since 1st Jun: <b>{m.get('biz_cover_invoices_march', 0)}</b></li>
        <li>• Potential missing SB Standby invoices: <b>{m.get('potential_missing_sb_march', 0)}</b></li>
    </ul>
    </div>
    """,
                        unsafe_allow_html=True,
                    )


# ---- Suggestion helpers (PLACE THIS ABOVE render_suggestions_page) ----
from pathlib import Path
import pandas as pd
import streamlit as st

SUG_FILE = Path("data/suggestions.csv")  # adjust if you store suggestions elsewhere

def add_suggestion(row: dict) -> None:
    """Persist a suggestion and notify Teams."""
    # 1) Save/append to CSV (or wherever you store them)
    try:
        df_new = pd.DataFrame([row])
        if SUG_FILE.exists():
            try:
                df_old = pd.read_csv(SUG_FILE)
            except Exception:
                df_old = pd.DataFrame()
            df = pd.concat([df_old, df_new], ignore_index=True)
        else:
            SUG_FILE.parent.mkdir(parents=True, exist_ok=True)
            df = df_new
        df.to_csv(SUG_FILE, index=False)
    except Exception:
        # Don’t crash the UI if disk write fails
        pass

    # 2) Send a Teams card using your existing helper
    try:
        send_to_teams(
            title="💡 New Dashboard Suggestion",
            text=f"**{row.get('name','Anonymous')}** submitted a suggestion ({row.get('tag','—')}).",
            facts={
                "Suggestion": row.get("idea", "—"),
                "When": row.get("timestamp", "—"),
                "ID": row.get("id", "—"),
            },
            button_text="Open Suggestion Register",
            button_url=st.secrets.get("APP_URL"),  # optional
        )
    except Exception:
        # If webhook is missing/unreachable, just skip
        pass

def render_sky_orbit_file_upload():
    import pandas as pd
    import streamlit as st
    import pdfplumber
    import docx
    from langchain_openai import ChatOpenAI
    from langchain.agents.agent_types import AgentType
    from langchain_experimental.agents import create_pandas_dataframe_agent

    st.title("📁 Sky Orbit File Uploader")

    # Back to Ops menu (switch the op_area_section)
    if st.button("⬅️ Back to Ops Menu", use_container_width=True, key="back_sky_orbit_file_upload_ops"):
        st.session_state["op_area_section"] = "engineer"   # or whichever default tab you want
        st.rerun()

    st.markdown("""
        Upload your Excel, CSV, PDF, TXT, or Word (.docx) files below and ask questions about their contents.
        This AI is specialized in understanding your uploaded files only, but you can still chat normally without uploading.
    """)

    uploaded_file = st.file_uploader(
        "Choose a file to upload",
        type=["xlsx", "xls", "csv", "pdf", "txt", "docx"],
        accept_multiple_files=False,
        key="sky_orbit_uploader"
    )

    # Reset uploaded data and chat if new file uploaded
    if uploaded_file is not None:
        if uploaded_file.name != st.session_state.get("last_uploaded_file", None):
            st.session_state.file_ai_chat = []
            if "file_df_agent" in st.session_state:
                del st.session_state["file_df_agent"]
            st.session_state.last_uploaded_file = uploaded_file.name
            st.session_state.uploaded_df = None
            st.session_state.uploaded_text = None

        df = None
        extracted_text = None

        if uploaded_file.type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                  "application/vnd.ms-excel"]:
            try:
                df = pd.read_excel(uploaded_file)
            except Exception as e:
                st.error(f"Failed to read Excel file: {e}")

        elif uploaded_file.type == "text/csv":
            try:
                df = pd.read_csv(uploaded_file)
            except Exception as e:
                st.error(f"Failed to read CSV file: {e}")

        elif uploaded_file.type == "application/pdf":
            try:
                with pdfplumber.open(uploaded_file) as pdf:
                    extracted_text = ""
                    for page in pdf.pages:
                        extracted_text += (page.extract_text() or "") + "\n"
                st.text_area("Preview of extracted PDF text (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to extract PDF text: {e}")

        elif uploaded_file.type == "text/plain":
            try:
                extracted_text = uploaded_file.getvalue().decode("utf-8", errors="ignore")
                st.text_area("Preview of uploaded TXT file (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to read TXT file: {e}")

        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(uploaded_file)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                extracted_text = "\n".join(paragraphs)
                st.text_area("Preview of uploaded DOCX file (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to read DOCX file: {e}")

        st.write(f"**Uploaded file:** {uploaded_file.name} ({round(uploaded_file.size / 1024, 2)} KB)")

        st.session_state.uploaded_df = df
        st.session_state.uploaded_text = extracted_text
    else:
        st.info("Upload a file to begin chatting with your file AI or just ask questions below.")
        st.session_state.uploaded_df = None
        st.session_state.uploaded_text = None

    # Initialize chat history if missing
    if "file_ai_chat" not in st.session_state:
        st.session_state.file_ai_chat = []

    # Build (or reuse) the agent
    def get_agent():
        llm_stream = ChatOpenAI(
            api_key=st.secrets["openai"]["api_key"],
            model_name="gpt-4o-mini",
            streaming=True,
        )
        if st.session_state.get("uploaded_df") is not None:
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=st.session_state.uploaded_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )
        elif st.session_state.get("uploaded_text") is not None:
            dummy_df = pd.DataFrame({"Text": [st.session_state.uploaded_text]})
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=dummy_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )
        else:
            # Fallback df for free chat
            dummy_df = pd.DataFrame({"Info": ["No file uploaded. Ask me anything or upload a file."]})
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=dummy_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )

    if "file_df_agent" not in st.session_state:
        st.session_state.file_df_agent = get_agent()

    agent = st.session_state.file_df_agent

    # Chat input
    question = st.chat_input("Ask your file AI a question:")
    if question:
        with st.spinner("Processing your question..."):
            try:
                response = agent.run(question)
            except Exception as e:
                response = f"⚠️ AI error: {e}"

        st.session_state.file_ai_chat.append({"question": question, "response": response})
        st.rerun()

    # Render chat bubbles
    for chat in st.session_state.file_ai_chat:
        st.chat_message("user").markdown(chat["question"])
        st.chat_message("assistant").markdown(chat["response"])

# ──────────────────────────────────────────────────────────────────────
# 💡 Suggestion Box page renderer
# ──────────────────────────────────────────────────────────────────────
def render_suggestions_page():
    import uuid, datetime
    from pathlib import Path
    import pandas as pd
    from openpyxl import Workbook, load_workbook
    from openpyxl.worksheet.table import Table, TableStyleInfo
    from openpyxl.utils import get_column_letter
    import streamlit as st

    # ---- (keep your Excel helpers exactly as you had them) ----
    COMMENTS_FILE = Path(r"C:\Users\dah47\OneDrive - Sky\Oracle development\suggestion_comments.xlsx")
    SUGG_FILE    = Path(r"C:\Users\dah47\OneDrive - Sky\Oracle development\suggestions log.xlsx")

    SUGG_COLS    = ["id","num_id","timestamp","name","tag","idea","status"]
    COMM_COLS    = ["id","timestamp","comment"]
    NEW_STATUSES = ["Received","Notified"]
    TABLE_NAME   = "Table1" ; SHEET_NAME = "Sheet1" ; TMP_SUFFIX = ".tmp.xlsx"

    def _create_table(ws, df):
        n_rows, n_cols = df.shape
        ref = f"A1:{get_column_letter(n_cols)}{n_rows+1}"
        t = Table(displayName=TABLE_NAME, ref=ref)
        t.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
        ws.add_table(t)

    def _ensure_excel(path: Path, cols):
        if path.exists():
            return
        wb = Workbook(); ws = wb.active; ws.title = SHEET_NAME
        ws.append(cols); _create_table(ws, pd.DataFrame(columns=cols)); wb.save(path)

    def _atomic_overwrite(path: Path, df: pd.DataFrame):
        tmp = path.with_suffix(TMP_SUFFIX)
        with pd.ExcelWriter(tmp, engine="openpyxl") as w:
            df.to_excel(w, sheet_name=SHEET_NAME, index=False)
        wb = load_workbook(tmp); ws = wb[SHEET_NAME]
        if TABLE_NAME in ws.tables: del ws.tables[TABLE_NAME]
        _create_table(ws, df); wb.save(tmp); tmp.replace(path)

    def _load_df(path: Path, cols):
        _ensure_excel(path, cols)
        return pd.read_excel(path, dtype=str).fillna("")

    def _save_df(path: Path, df: pd.DataFrame, cols):
        _ensure_excel(path, cols)
        _atomic_overwrite(path, df)

    load_suggestions = lambda: _load_df(SUGG_FILE, SUGG_COLS)
    save_suggestions = lambda d: _save_df(SUGG_FILE, d, SUGG_COLS)
    load_comments    = lambda: _load_df(COMMENTS_FILE, COMM_COLS)
    save_comments    = lambda d: _save_df(COMMENTS_FILE, d, COMM_COLS)

    # OPTIONAL: Teams webhook post (use the same helper as login)
    try:
        # One nice compact card; tweak facts / text to taste
        send_to_teams(
            title="💡 New Dashboard Suggestion",
            text=f"**{row.get('name','Anonymous')}** submitted a suggestion ({row.get('tag','—')}).",
            facts={
                "Suggestion": row.get("idea", "—"),
                "When": row.get("timestamp", "—"),
                "ID": row.get("id", "—"),
            },
            # Optional: show a button that opens your app (set APP_URL in secrets if you want)
            button_text="Open Suggestion Register",
            button_url=st.secrets.get("APP_URL", None),
        )
    except Exception:
        # Don't crash the UI if the webhook is down
        pass


    def add_comment(sugg_id, text):
        df = load_comments()
        df = pd.concat(
            [
                df,
                pd.DataFrame([{
                    "id": sugg_id,
                    "timestamp": datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
                    "comment": text.strip(),
                }])
            ],
            ignore_index=True,
        )
        save_comments(df)

    # ---- PAGE UI ----
    if st.button("⬅️ Back to Login", use_container_width=True):
        st.session_state.screen = "login"
        st.rerun()

    st.markdown("## 💡 Suggestion Box")
    st.caption("Help improve the dashboard by submitting your ideas.")

    with st.form("suggest_form"):
        name = st.text_input("Your name")
        tag = st.selectbox("What does your suggestion relate to?",
                           ["General","Missing Data","Graphs Required","Sky Stake Holder","Performance"])
        idea = st.text_area("Your suggestion (max 500 chars)", height=150, max_chars=500)
        sent = st.form_submit_button("Submit Suggestion")

    if sent and idea.strip():
        df   = load_suggestions()
        next_id = str(len(df) + 1)
        add_suggestion({
            "id": str(uuid.uuid4())[:8],
            "num_id": next_id,
            "timestamp": datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
            "name": name.strip() or "Anonymous",
            "tag": tag,
            "idea": idea.strip(),
            "status": "Received"
        })
        st.rerun()
    elif sent:
        st.warning("Please write something before submitting.")

    tab_new, tab_prog, tab_done = st.tabs(["📥 New","🚧 In Progress","✅ Completed"])
    df_sugg, df_comm = load_suggestions(), load_comments()

    def render_list(df: pd.DataFrame, allowed_status):
        nonlocal df_sugg, df_comm
        rows = df[df["status"].isin(allowed_status)]
        if rows.empty:
            st.info("No suggestions here yet.")
            return
        for idx, (_, r) in enumerate(rows.iterrows()):
            safe_id = r["id"] or r.get("num_id","") or idx
            with st.expander(f"📝 {r['idea']} ({r['name']})"):
                st.caption(f"📅 {r['timestamp']}  |  🏷️ {r['tag']}")

                # Comments block
                if r["status"] == "In Progress":
                    st.markdown("**Comments**")
                    cts = df_comm[df_comm["id"] == r["id"]]
                    for _, c in cts.iterrows():
                        st.markdown(f"- *{c['timestamp']}*: {c['comment']}")
                    new_c = st.text_area("Add new comment", key=f"c_{safe_id}_{idx}")
                    if st.button("Add", key=f"add_{safe_id}_{idx}") and new_c.strip():
                        add_comment(r["id"], new_c)
                        st.success("Added!")
                        st.rerun()

                # Controls row
                btn_cols = st.columns([1,1,1,5])  # In-Prog / Complete / Delete / Spacer

                # Start
                if r["status"] in ["Received","Notified"] and btn_cols[0].button("Start", key=f"start_{safe_id}_{idx}"):
                    df_sugg.loc[df_sugg["id"] == r["id"], "status"] = "In Progress"
                    save_suggestions(df_sugg)
                    st.rerun()

                # Done
                if r["status"] == "In Progress" and btn_cols[1].checkbox("Done", key=f"done_{safe_id}_{idx}"):
                    df_sugg.loc[df_sugg["id"] == r["id"], "status"] = "Completed"
                    save_suggestions(df_sugg)
                    st.rerun()

                # Delete
                if btn_cols[2].button("Delete", key=f"del_{safe_id}_{idx}"):
                    df_sugg = df_sugg[df_sugg["id"] != r["id"]]
                    save_suggestions(df_sugg)
                    st.success("Deleted.")
                    st.rerun()

    with tab_new:  render_list(df_sugg, ["Received","Notified"])
    with tab_prog: render_list(df_sugg, ["In Progress"])
    with tab_done: render_list(df_sugg, ["Completed"])

# ==============================
# TEAM OVERVIEW — GRAPHS (pro layout + month picker + budgets + tables + scaling)
# ==============================
def render_team_graphs_page():
    import pandas as pd, numpy as np
    import plotly.express as px, plotly.graph_objects as go
    import streamlit as st

    # (recommended globally in your app's main): st.set_page_config(layout="wide")

    # --- team + data ---
    user_team = (st.session_state.get("user_team", "") or "VIP North").strip()
    try:
        df = get_team_df(user_team)
    except Exception:
        df = pd.DataFrame()
    if df.empty:
        st.warning(f"No data found for team: {user_team}")
        return
    df = df.copy(); df.columns = df.columns.str.strip()

    # ---------- header + back ----------
    left, right = st.columns([1, 1])
    with left:
        st.subheader(f"📊 Team Overview — Graphs ({user_team})")
        if st.button("⬅️ Back to Menu"):
            st.session_state.screen = "instruction_guide"
            st.rerun()
    with right:
        # Page scale for screenshots (affects entire page). Works in Chrome/Edge.
        scale = st.select_slider("View scale", options=[60,70,80,90,100,110,120], value=80, key="graphs_scale")
        st.markdown(
            f"""
            <style>
              .block-container {{
                zoom: {scale}%;
                max-width: 1400px; /* keep page from stretching super wide */
              }}
            </style>
            """,
            unsafe_allow_html=True
        )

    # ---------- detect best date column + Month picker with 'All months' ----------
    df_full = df.copy()
    date_candidates = [c for c in df_full.columns if ("date" in c.lower() or "time" in c.lower())]
    best_col, best_non_null = None, -1
    for c in date_candidates:
        dtc = pd.to_datetime(df_full[c], errors="coerce", dayfirst=True)
        if dtc.notna().sum() > best_non_null:
            best_non_null, best_col = dtc.notna().sum(), c

    if best_col is None:
        df_full["_dt"] = pd.NaT
        sel_period, gran = None, "Daily"
        st.warning("Could not detect a date column; showing unfiltered data.")
    else:
        df_full["_dt"] = pd.to_datetime(df_full[best_col], errors="coerce", dayfirst=True)
        month_index = pd.PeriodIndex(df_full["_dt"].dropna().dt.to_period("M")).sort_values().unique()
        if len(month_index) == 0:
            sel_period, gran = None, "Daily"
            st.warning("No valid dates found; showing unfiltered data.")
        else:
            labels = [m.to_timestamp().strftime("%b %Y") for m in month_index]
            labels_all = ["All months"] + labels
            sel_label = st.selectbox(
                "Month",
                labels_all,
                index=len(labels_all) - 1,
                key=f"graphs_month_select_{user_team.replace(' ','_')}"
            )
            if sel_label == "All months":
                sel_period, gran, df = None, "Monthly", df_full.copy()
            else:
                sel_period = month_index[labels.index(sel_label)]
                df = df_full[df_full["_dt"].dt.to_period("M") == sel_period].copy()
                gran = "Daily"

    # ---------- optional hooks ----------
    try:
        mom = compute_mom_trends(df_full, df, sel_period)
    except Exception:
        mom = {}

    # ---------- dark UI CSS ----------
    st.markdown(
        """
        <style>
        .card{
            background:#1e1e1e; border:1px solid rgba(255,255,255,0.08);
            box-shadow:0 4px 16px rgba(0,0,0,0.3);
            border-radius:14px; padding:16px 16px 8px 16px; margin-bottom:14px;
            color:#f3f4f6;
        }
        .card-head{ display:flex; align-items:center; justify-content:space-between; margin-bottom:8px;}
        .card-title{ font-weight:600; font-size:1.05rem; color:#f9fafb;}
        .card-menu{ color:#aaa; font-size:1.2rem; line-height:1; padding:0 6px;}
        .kpi{ padding:14px; text-align:left; color:#f3f4f6;}
        .kpi-label{ color:#9ca3af; font-size:.85rem; }
        .kpi-value{ font-size:1.6rem; font-weight:700; margin-top:2px; color:#ffffff;}
        .kpi-delta{ font-size:.9rem; margin-top:2px; color:#93c5fd;}
        .kpi-help{ font-size:.78rem; color:#9ca3af; margin-top:4px;}
        </style>
        """,
        unsafe_allow_html=True
    )
    def _fmt_hhmm(td):
        """Format a pandas Timedelta as HH:MM, or 'N/A'."""
        if td is None or pd.isna(td):
            return "N/A"
        total_sec = int(td.total_seconds())
        # round to nearest minute so you don’t get :59 noise
        total_min = int(round(total_sec / 60.0))
        hours, minutes = divmod(total_min, 60)
        return f"{hours:02d}:{minutes:02d}"

    # ---------- helpers ----------
    def card_start(title:str, menu:bool=False):
        st.markdown(
            f"""
            <div class="card">
              <div class="card-head">
                <div class="card-title">{title}</div>
                {'<div class="card-menu">•••</div>' if menu else ''}
              </div>
            """,
            unsafe_allow_html=True
        )

    def card_end():
        st.markdown("</div>", unsafe_allow_html=True)

    def fig_style(fig, title=""):
        fig.update_layout(
            title=title, template="plotly_dark",
            height=320,  # consistent compact height for screenshots
            margin=dict(l=20, r=20, t=50, b=20),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        )
        fig.update_xaxes(showgrid=True, gridwidth=1, gridcolor="rgba(255,255,255,0.06)")
        fig.update_yaxes(showgrid=True, gridwidth=1, gridcolor="rgba(255,255,255,0.06)")
        return fig

    # ---------- TOP KPIs ----------
    total_visits = int(len(df))
    vtypes = df["Visit Type"].nunique() if "Visit Type" in df.columns else 0
    statuses = df["Activity Status"].nunique() if "Activity Status" in df.columns else 0
    stakeholders = df["Stakeholder"].nunique() if "Stakeholder" in df.columns else 0
    total_value = df["Total Value"].sum() if "Total Value" in df.columns else np.nan

    if "_dt" in df_full.columns:
        by_month_all = df_full.groupby(pd.Grouper(key="_dt", freq="MS")).size().rename("Visits")
        if len(by_month_all) >= 2:
            last, prev = by_month_all.iloc[-1], by_month_all.iloc[-2]
            delta = last - prev
            pct = (delta / prev * 100) if prev else np.nan
            mom_str = f"{'▲' if delta>=0 else '▼'} {abs(delta):,} ({abs(pct):.1f}%) vs prev. mo."
        else:
            mom_str = "—"
    else:
        mom_str = "—"

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.markdown('<div class="card kpi"><div class="kpi-label">Total Visits</div>'
                    f'<div class="kpi-value">{total_visits:,}</div>'
                    f'<div class="kpi-delta">{mom_str}</div></div>', unsafe_allow_html=True)
    with k2:
        st.markdown('<div class="card kpi"><div class="kpi-label">Unique Visit Types</div>'
                    f'<div class="kpi-value">{vtypes:,}</div></div>', unsafe_allow_html=True)
    with k3:
        st.markdown('<div class="card kpi"><div class="kpi-label">Unique Statuses</div>'
                    f'<div class="kpi-value">{statuses:,}</div></div>', unsafe_allow_html=True)
    with k4:
        st.markdown('<div class="card kpi"><div class="kpi-label">Unique Stakeholders</div>'
                    f'<div class="kpi-value">{stakeholders:,}</div></div>', unsafe_allow_html=True)

    # ---------------- ROW A: Engineer sunburst + Over/Under 10:25 sunburst + Totals ----------------
    a1, a2, a3 = st.columns([1.2, 1.2, 1])
    with a1:
        if "Name" in df.columns:
            try:
                total_engs = combined_oracle_df["Name"].nunique()  # all teams
            except Exception:
                total_engs = df["Name"].nunique()
            team_engs  = df["Name"].nunique()
            other_engs = max(total_engs - team_engs, 0)

            labels = [f"{user_team} ({team_engs})", f"Other Teams ({other_engs})"]
            values = [team_engs, other_engs if other_engs > 0 else 1e-6]  # avoid single-slice edge case
            colors = ["#1f77b4", "#e74c3c"]  # team blue, other red

            card_start("Engineer Coverage", menu=True)
            fig = go.Figure(go.Pie(
                labels=labels,
                values=values,
                hole=0.62,
                marker=dict(colors=colors, line=dict(color="#1e1e1e", width=2)),
                textinfo="label",                # labels on the ring
                textposition="inside",
                insidetextorientation="radial",  # curved along the arc
                sort=False
            ))
            fig.update_layout(
                template="plotly_dark",
                height=320,
                annotations=[dict(
                    text=f"<b>{total_engs}</b>",  # center number only
                    x=0.5, y=0.5, showarrow=False,
                    font=dict(size=22, color="white")
                )]
            )
            st.plotly_chart(fig, use_container_width=True)
            card_end()






    with a2:
        time_col = "Total Working Time" if "Total Working Time" in df.columns else ("Total Time" if "Total Time" in df.columns else None)
        if time_col:
            tw = pd.to_timedelta(df[time_col].astype(str), errors="coerce")
            over = int((tw > pd.Timedelta("10:25:00")).sum())
            under = int((tw <= pd.Timedelta("10:25:00")).sum())
            card_start("Working Time > 10:25 vs ≤ 10:25", menu=True)
            sb2 = go.Figure(go.Sunburst(
                labels=["All Visits", "Over 10:25", "Under 10:25"],
                parents=["", "All Visits", "All Visits"],
                values=[over+under, over, under],
                branchvalues="total",
                hoverinfo="label+value+percent parent"
            ))
            sb2.update_layout(height=320, template="plotly_dark")
            st.plotly_chart(sb2, use_container_width=True)
            card_end()

    with a3:
        card_start("Totals", menu=False)
        tv = f"£{total_value:,.0f}" if total_value == total_value else "N/A"
        st.metric("💷 Total Value", tv)
        st.metric("📈 Total Visits", f"{total_visits:,}")
        # choose a sensible completion-time column (Total Time first, with fallbacks)
        time_col_opts = ["Total Time", "Total Time for AI", "Total Working Time", "Total working time"]
        time_col = next((c for c in time_col_opts if c in df.columns), None)

        if time_col:
            td = pd.to_timedelta(df[time_col].astype(str), errors="coerce")
            # keep only positive, non-null durations (ignore 0, '00:00', NaT etc.)
            valid = td.notna() & (td > pd.Timedelta(0))
            mean_td = td[valid].mean() if valid.any() else pd.NaT
            st.metric("⏱ Avg Completion Time Per Visit", _fmt_hhmm(mean_td))
        else:
            st.metric("⏱ Avg Completion Time", "N/A")

        card_end()

    # ---------------- ROW 1: Monthly trend + Top Visit Types ----------------
    r1c1, r1c2 = st.columns([2, 1])
    if "_dt" in df.columns or "_dt" in df_full.columns:
        if "_dt" not in df.columns and best_col is not None:
            df["_dt"] = pd.to_datetime(df[best_col], errors="coerce", dayfirst=True)
        with r1c1:
            card_start("Visits by Month", menu=True)
            m_df = df.groupby(pd.Grouper(key="_dt", freq="MS")).size().reset_index(name="Visits")
            m_df["Month"] = m_df["_dt"].dt.strftime("%b %Y")
            fig_month = px.bar(m_df, x="Month", y="Visits")
            fig_style(fig_month)
            st.plotly_chart(fig_month, use_container_width=True)
            card_end()

    with r1c2:
        if "Visit Type" in df.columns:
            card_start("Top Visit Types", menu=True)
            vt = df["Visit Type"].value_counts().reset_index()
            vt.columns = ["Visit Type", "Count"]
            fig_vtype = px.bar(vt.head(7), x="Visit Type", y="Count")
            fig_style(fig_vtype)
            st.plotly_chart(fig_vtype, use_container_width=True)
            card_end()

    # ---------------- ROW 2: (NEW) Monthly KPI Table + Activity Status mix ----------------
    r2c1, r2c2 = st.columns([2, 1])
    with r2c1:
        # Fills the “blank” space you circled
        card_start("Monthly KPI Table", menu=True)
        if "_dt" in df.columns:
            def _safe_mean_timedelta(s):
                td = pd.to_timedelta(s.astype(str), errors="coerce")
                td = td[(td.notna()) & (td > pd.Timedelta(0))]
                return (td.mean() if len(td) else pd.NaT)
            tmp = df.copy()
            tmp["Month"] = tmp["_dt"].dt.to_period("M").dt.to_timestamp()
            comp_mask = tmp["Activity Status"].astype(str).str.lower().eq("completed") if "Activity Status" in tmp.columns else False
            tbl = tmp.groupby("Month").agg(
                Visits=("Month","count"),
                Completed=("Activity Status", lambda s: int((s.astype(str).str.lower()=="completed").sum()) if "Activity Status" in tmp.columns else 0),
                TotalValue=("Total Value", "sum") if "Total Value" in tmp.columns else ("Month","count"),
                AvgTotalTime=("Total Time", _safe_mean_timedelta) if "Total Time" in tmp.columns else ("Month","count"),
            ).reset_index()
            if "Total Value" in df.columns:
                tbl["TotalValue"] = tbl["TotalValue"].fillna(0).round(0).astype(int)
            if "AvgTotalTime" in tbl.columns:
                tbl["AvgTotalTime"] = tbl["AvgTotalTime"].astype(str).str.replace("NaT","—")
            tbl["CompletionRate%"] = (tbl["Completed"]/tbl["Visits"]*100).round(1).replace([np.inf,-np.inf], np.nan).fillna(0)
            # show a compact table
            st.dataframe(tbl.rename(columns={
                "Month":"Month",
                "Visits":"Visits",
                "Completed":"Completed",
                "CompletionRate%":"Completion %",
                "TotalValue":"Total £",
                "AvgTotalTime":"Avg Total Time"
            }), use_container_width=True, height=320)
        else:
            st.info("No date column to build the table.")
        card_end()

    with r2c2:
        if "Activity Status" in df.columns:
            card_start("Activity Status mix", menu=True)
            ast = df["Activity Status"].value_counts().reset_index()
            ast.columns = ["Activity Status", "Count"]
            fig_stat = px.pie(ast, names="Activity Status", values="Count", hole=0.6)
            fig_stat.update_traces(textposition="inside", textinfo="percent+label")
            fig_style(fig_stat)
            st.plotly_chart(fig_stat, use_container_width=True)
            card_end()

    # ---------------- ROW 3: Day-of-week + 7-day running ----------------
    r3c1, r3c2 = st.columns([2, 1])
    with r3c1:
        if "_dt" in df.columns:
            card_start("Visits by Day of Week", menu=True)
            day_order = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
            by_day = df.groupby(df["_dt"].dt.day_name()).size().reindex(day_order, fill_value=0).reset_index()
            by_day.columns = ["DayOfWeek","Visits"]
            fig_day = px.bar(by_day, x="DayOfWeek", y="Visits")
            fig_style(fig_day)
            st.plotly_chart(fig_day, use_container_width=True)
            card_end()
    with r3c2:
        if "_dt" in df.columns:
            card_start("7-Day Running Trend", menu=True)
            daily = (df.groupby(df["_dt"].dt.date).size()
                        .rename("Visits").to_frame().reset_index())
            daily["_dt"] = pd.to_datetime(daily["_dt"])
            daily["7d"] = daily["Visits"].rolling(7).mean()
            fig_7 = go.Figure()
            fig_7.add_trace(go.Scatter(x=daily["_dt"], y=daily["Visits"], mode="lines+markers", name="Daily"))
            fig_7.add_trace(go.Scatter(x=daily["_dt"], y=daily["7d"], mode="lines", name="7-day avg"))
            fig_style(fig_7)
            st.plotly_chart(fig_7, use_container_width=True)
            card_end()

    # ---------------- ROW 4: Monthly mix by Visit Type (stacked area) ----------------
    if "_dt" in df.columns and "Visit Type" in df.columns:
        card_start("Monthly Mix by Visit Type", menu=True)
        mix = (
            df.assign(Month=df["_dt"].dt.to_period("M").dt.to_timestamp())
              .groupby(["Month","Visit Type"]).size().reset_index(name="Count")
        )
        top_types = df["Visit Type"].value_counts().nlargest(6).index
        mix["Visit Type"] = np.where(mix["Visit Type"].isin(top_types), mix["Visit Type"], "Other")
        mix = mix.groupby(["Month","Visit Type"])["Count"].sum().reset_index()
        fig_mix = px.area(mix, x="Month", y="Count", color="Visit Type")
        fig_style(fig_mix)
        st.plotly_chart(fig_mix, use_container_width=True)
        card_end()

    # ---------------- ROW 5: Top 5 Completed (excl Lunch) + Visit Ratios ----------------
    r5c1, r5c2 = st.columns([1.4, 1])
    with r5c1:
        if "Activity Status" in df.columns and "Visit Type" in df.columns:
            card_start("Top 5 Completed Visit Types (excl Lunch)", menu=True)
            comp = df[df["Activity Status"].astype(str).str.lower().eq("completed")]
            comp = comp[~comp["Visit Type"].astype(str).str.contains("Lunch", case=False, na=False)]
            top5 = comp["Visit Type"].value_counts().head(5).reset_index()
            top5.columns = ["Visit Type","Count"]
            fig5 = px.bar(top5, x="Visit Type", y="Count")
            fig_style(fig5)
            st.plotly_chart(fig5, use_container_width=True)
            card_end()
    with r5c2:
        if "Visit Type" in df.columns:
            card_start("Visit Ratios (all types)", menu=True)
            vt_counts = df["Visit Type"].value_counts().reset_index()
            vt_counts.columns=["Visit Type","Count"]
            keep = vt_counts["Visit Type"].head(7).tolist()
            vt_counts["Visit Type"] = np.where(vt_counts["Visit Type"].isin(keep), vt_counts["Visit Type"], "Other")
            vt_counts = vt_counts.groupby("Visit Type")["Count"].sum().reset_index()
            fig_ratio = px.pie(vt_counts, names="Visit Type", values="Count", hole=0.45)
            fig_ratio.update_traces(textposition="inside", textinfo="percent+label")
            fig_style(fig_ratio)
            st.plotly_chart(fig_ratio, use_container_width=True)
            card_end()

    # ---------------- 💷 Budget KPIs ----------------
    st.markdown("#### 💷 Budget KPIs")
    try:
        alloc, used, remaining, pct_used = compute_team_budget_metrics(user_team)
    except Exception:
        alloc = used = remaining = pct_used = np.nan

    b1, b2, b3, b4 = st.columns(4)
    with b1:
        try: card("Budget Allocated", fmt_money(alloc), "From budgets.csv", trend=None, spark=None, color="#3b82f6", source="Budget")
        except Exception: st.info("Budget Allocated: helper `card()` not found")
    with b2:
        try: card("Budget Used", fmt_money(used), "From expenses.csv", trend=None, spark=None, color="#ef4444", source="Budget")
        except Exception: st.info("Budget Used: helper `card()` not found")
    with b3:
        try: bar_remaining = max(0, min(int(remaining / alloc * 100), 100)) if (alloc and alloc == alloc) else None
        except Exception: bar_remaining = None
        try: card("Budget Remaining", fmt_money(remaining), "Allocated − Used", trend=None, bar_pct=bar_remaining, spark=None, color="#22c55e", source="Budget")
        except Exception: st.info("Budget Remaining: helper `card()` not found")
    with b4:
        try: bar_used = max(0, min(int(pct_used), 100)) if (pct_used == pct_used) else None
        except Exception: bar_used = None
        pct_txt = f"{pct_used:.1f}%" if pct_used == pct_used else "—"
        try: card("Budget % Used", pct_txt, "", trend=None, bar_pct=bar_used, spark=None, color="#06b6d4", source="Budget")
        except Exception: st.info("Budget % Used: helper `card()` not found")

    # ---------------- 📑 Summary (fixed idxtop -> idxmax) ----------------
    try:
        comp_mask = df["Activity Status"].astype(str).str.lower().eq("completed")
        comp_rate = 100 * comp_mask.mean()
    except Exception:
        comp_rate = np.nan

    top_type = "—"
    if "Visit Type" in df.columns and not df["Visit Type"].empty:
        counts = df["Visit Type"].value_counts()
        if len(counts):
            top_type = counts.idxmax()  # <-- fixed

    busiest_day = "—"
    if "_dt" in df.columns:
        try: busiest_day = df["_dt"].dt.day_name().value_counts().idxmax()
        except Exception: pass

    avg_tt_txt = "N/A"
    if "Total Time" in df.columns:
        tt = pd.to_timedelta(df["Total Time"].astype(str), errors="coerce")
        mean_tt = tt[tt.notna() & (tt > pd.Timedelta(0))].mean()
        if pd.notna(mean_tt): avg_tt_txt = str(mean_tt).split()[0]

    card_start("📑 Summary", menu=False)
    st.markdown(
        f"""
        - **Total Visits:** {total_visits:,}  **Total Value:** {('£%s' % format(total_value, ',.0f')) if total_value==total_value else 'N/A'}  
        - **Completion Rate:** {('%.1f%%' % comp_rate) if comp_rate==comp_rate else '—'}  **Top Visit Type:** {top_type}  
        - **Busiest Day:** {busiest_day}  **Avg Completion Time:** {avg_tt_txt}
        """.strip()
    )
    card_end()

    st.caption("Use the View scale in the header to shrink the whole page for clean snapshots. Charts reflect the selected month range and your signed-in team.")



import pandas as pd
import numpy as np
from datetime import datetime, date

def render_engineer_kpi():
    st.title("🧑‍🔧 Engineer KPI")

    # ---------- Load data ----------
    inv = load_invoices()
    vis = load_visits()

    # Defensive normalisation for dates
    for df, col in ((inv, "Date"), (vis, "Date")):
        if df is not None and not df.empty:
            if col not in df.columns:
                # Best-effort find a date-like column
                for c in df.columns:
                    if "date" in c.lower():
                        df.rename(columns={c: "Date"}, inplace=True)
                        break
            if "Date" in df.columns:
                df["Date"] = pd.to_datetime(df["Date"], errors="coerce")

    # Year/Month filters (YTD by default)
    y = date.today().year
    years = sorted(set([y] +
                       [int(x) for x in pd.Series(pd.to_datetime(inv.get("Date", pd.Series(dtype="datetime64[ns]")), errors="coerce").dt.year.dropna().unique()).tolist()] +
                       [int(x) for x in pd.Series(pd.to_datetime(vis.get("Date", pd.Series(dtype="datetime64[ns]")), errors="coerce").dt.year.dropna().unique()).tolist()]
                      ))
    cY, cM = st.columns([1,1])
    with cY:
        year_choice = st.selectbox("Year", years, index=years.index(y))
    with cM:
        month_choice = st.selectbox("Month", ["YTD","Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"], index=0)

    # Build period window
    if month_choice == "YTD":
        start_dt = pd.Timestamp(year_choice, 1, 1)
        end_dt   = pd.Timestamp(year_choice, 12, 31)
        cap_month = None  # use full YTD
    else:
        mnum = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"].index(month_choice) + 1
        start_dt = pd.Timestamp(year_choice, 1, 1)
        end_dt   = pd.Timestamp(year_choice, mnum, 1) + pd.offsets.MonthEnd(0)
        cap_month = pd.Timestamp(year_choice, mnum, 1).to_period("M").to_timestamp()

    def in_window(df):
        if df is None or df.empty or "Date" not in df.columns:
            return df
        return df[(df["Date"] >= start_dt) & (df["Date"] <= end_dt)].copy()

    inv = in_window(inv)
    vis = in_window(vis)

    # ---------- Helper functions ----------
    def to_minutes_from_time(x):
        # accepts "HH:MM" or a Timestamp
        if pd.isna(x):
            return np.nan
        if isinstance(x, (pd.Timestamp, datetime)):
            return x.hour*60 + x.minute
        s = str(x).strip()
        if ":" in s:
            h, m = s.split(":", 1)
            return int(h)*60 + int(m)
        return np.nan

    def mean_hhmm(series_mins):
        v = pd.to_numeric(series_mins, errors="coerce").dropna()
        if v.empty:
            return "—"
        m = int(v.mean())
        return f"{m//60:02d}:{m%60:02d}"

    # Try to detect relevant columns in visits
    eng_col = next((c for c in vis.columns if c.lower() in ("engineer","engineers name","engineers","engineer name")), None) if vis is not None and not vis.empty else None
    start_col = next((c for c in vis.columns if "start" in c.lower()), None) if vis is not None and not vis.empty else None
    finish_col = next((c for c in vis.columns if ("finish" in c.lower()) or ("end" in c.lower())), None) if vis is not None and not vis.empty else None
    lunch_col = next((c for c in vis.columns if "lunch" in c.lower() and "min" in c.lower()), None) if vis is not None and not vis.empty else None
    workmins_col = next((c for c in vis.columns if ("work" in c.lower() and "min" in c.lower()) or ("duration" in c.lower() and "min" in c.lower())), None) if vis is not None and not vis.empty else None
    type_col = next((c for c in vis.columns if c.lower() in ("visit type","job type","type")), None) if vis is not None and not vis.empty else None

    # Detect “holiday”
    def is_holiday_row(r):
        if type_col and isinstance(r.get(type_col,""), str):
            t = r[type_col].lower()
            return "holiday" in t
        return False

    # Detect RR / SB cover days
    def is_rr_sb_row(r):
        if type_col and isinstance(r.get(type_col,""), str):
            t = r[type_col].lower()
            return ("rr" in t) or ("sb cover" in t) or ("standby" in t)
        return False

    # ---------- Team counters ----------
    # Completed vs outstanding in invoices
    def is_completed_invoice(row):
        for c in ("invoice status","status","completed","iscomplete"):
            if c in inv.columns:
                v = str(row[c]).strip().lower()
                if v in ("completed","complete","done","yes","true","1"):
                    return True
                if v in ("open","outstanding","no","false","0",""):
                    return False
        # Fallback: consider not completed if no Total Value
        tv = None
        for c in ("Total Value","TotalValue","total value","total"):
            if c in inv.columns:
                tv = row[c]
                break
        try:
            return float(tv) > 0
        except Exception:
            return False

    outstanding = 0
    completed   = 0
    inv_team = inv.copy() if inv is not None else pd.DataFrame()
    if not inv_team.empty:
        flags = inv_team.apply(is_completed_invoice, axis=1)
        completed   = int(flags.sum())
        outstanding = int((~flags).sum())

    # Average lunch & LOWD (avg length of working day)
    avg_lunch = "—"
    avg_lowd  = "—"
    if vis is not None and not vis.empty:
        # lunch mins
        if lunch_col:
            avg_lunch = mean_hhmm(pd.to_numeric(vis[lunch_col], errors="coerce"))
        # LOWD: prefer working minutes; else (finish-start-lunch)
        if workmins_col:
            avg_lowd = mean_hhmm(pd.to_numeric(vis[workmins_col], errors="coerce"))
        else:
            mins = []
            for _, r in vis.iterrows():
                s = to_minutes_from_time(r.get(start_col))
                e = to_minutes_from_time(r.get(finish_col))
                l = pd.to_numeric(r.get(lunch_col), errors="coerce")
                if not np.isnan(s) and not np.isnan(e):
                    d = (e - s) - (0 if np.isnan(l) else l)
                    mins.append(max(d, 0))
            avg_lowd = mean_hhmm(pd.Series(mins))

    # Team top counters (5 cards)
    c1,c2,c3,c4,c5 = st.columns(5, gap="large")
    with c1: card("Jobs requiring invoices (YTD)", f"{outstanding:,}", "Current period", color="#3b82f6", source="kpi_jobs_req")
    with c2: card("Invoices completed (YTD)", f"{completed:,}", "Current period", color="#10b981", source="kpi_inv_done")
    with c3:
        # drilldown
        card("Outstanding invoices (YTD)", f"{outstanding:,}", "Click to view list", color="#ef4444", source="kpi_inv_out")
        with st.expander("View outstanding"):
            if outstanding and not inv_team.empty:
                try:
                    from st_aggrid import AgGrid, GridOptionsBuilder
                    inv_team["__completed__"] = inv_team.apply(is_completed_invoice, axis=1)
                    out_df = inv_team[~inv_team["__completed__"]].drop(columns=["__completed__"])
                    gb = GridOptionsBuilder.from_dataframe(out_df)
                    gb.configure_pagination(paginationPageSize=15)
                    gb.configure_default_column(filter=True, sortable=True, resizable=True)
                    AgGrid(out_df, gridOptions=gb.build(), height=360, fit_columns_on_grid_load=True)
                except Exception:
                    st.dataframe(inv_team[~inv_team.apply(is_completed_invoice, axis=1)])
            else:
                st.info("No outstanding invoices in this period.")
    with c4: card("Avg lunch duration", avg_lunch, "YTD", color="#a855f7", source="kpi_avg_lunch")
    with c5: card("Avg LOWD", avg_lowd, "YTD", color="#f59e0b", source="kpi_avg_lowd")

    st.divider()

    # ---------- Per-engineer tiles ----------
    st.markdown("### Engineers (YTD)")
    if vis is None or vis.empty or not eng_col:
        st.info("No visit data with engineer names available.")
        return

    # North→South order (adjust any time)
    engineer_order = [
        "Gordon","Neil","Wayne","Chris Drury","James",
        "Matt H","Phil S","Chris Law","Chris Millward","Dave O",
        "Shaun","New Eng 1","New Eng 2"
    ]
    # Use only engineers present in this window
    present = sorted(vis[eng_col].dropna().astype(str).unique().tolist(), key=lambda n: (engineer_order.index(n) if n in engineer_order else 9999, n))



    # Helper: compute all metrics for one engineer
    def engineer_metrics(e_name: str):
        v = vis[vis[eng_col].astype(str) == e_name].copy()
        if v.empty:
            return {
                "avg_lunch": "—","avg_start":"—","avg_finish":"—","avg_lowd":"—",
                "rr_sb_days": 0,"days_worked":0,"holidays":0,"pct_rr_sb":"—",
                "expenses":"£0","ot_cost":"—"
            }

        # times
        s_mins = v[start_col].map(to_minutes_from_time) if start_col in v.columns else pd.Series(dtype=float)
        e_mins = v[finish_col].map(to_minutes_from_time) if finish_col in v.columns else pd.Series(dtype=float)
        l_mins = pd.to_numeric(v[lunch_col], errors="coerce") if lunch_col in v.columns else pd.Series(dtype=float)
        w_mins = pd.to_numeric(v[workmins_col], errors="coerce") if workmins_col in v.columns else pd.Series(dtype=float)

        avg_start  = mean_hhmm(s_mins)   if not s_mins.empty else "—"
        avg_finish = mean_hhmm(e_mins)   if not e_mins.empty else "—"
        avg_lunchE = mean_hhmm(l_mins)   if not l_mins.empty else "—"

        if not w_mins.empty:
            avg_lowdE = mean_hhmm(w_mins)
        else:
            # compute (finish-start-lunch)
            dur = []
            for _, r in v.iterrows():
                s = to_minutes_from_time(r.get(start_col))
                e = to_minutes_from_time(r.get(finish_col))
                l = pd.to_numeric(r.get(lunch_col), errors="coerce")
                if not np.isnan(s) and not np.isnan(e):
                    dur.append(max((e-s) - (0 if np.isnan(l) else l), 0))
            avg_lowdE = mean_hhmm(pd.Series(dur) if dur else pd.Series(dtype=float))

        # day counters
        days_worked = v["Date"].dt.date.nunique() if "Date" in v.columns else 0
        rr_sb_days  = v[v.apply(is_rr_sb_row, axis=1)]["Date"].dt.date.nunique() if type_col and "Date" in v.columns else 0
        holidays    = v[v.apply(is_holiday_row, axis=1)]["Date"].dt.date.nunique() if type_col and "Date" in v.columns else 0
        pct_rr_sb   = f"{(100*rr_sb_days/days_worked):.1f}%" if days_worked else "—"

        # expenses / OT cost from invoices for this engineer
        expenses = 0.0
        ot_cost  = None
        if inv is not None and not inv.empty:
            # find engineer column in inv
            eng_inv_col = next((c for c in inv.columns if "engineer" in c.lower()), None)
            inv_e = inv if eng_inv_col is None else inv[inv[eng_inv_col].astype(str).str.strip().str.lower() == e_name.lower()]
            if not inv_e.empty:
                # expenses columns (sum what exists)
                for c in ("Additional Costs","Hotel/ Food Value","Hotel/ Food Value ","Hotel/Food Value","Equipment Value","Expenses","Travel"):
                    if c in inv_e.columns:
                        expenses += pd.to_numeric(inv_e[c], errors="coerce").fillna(0).sum()
                # OT cost
                for c in ("OT Cost","OTCost","Overtime Cost"):
                    if c in inv_e.columns:
                        val = pd.to_numeric(inv_e[c], errors="coerce").fillna(0).sum()
                        ot_cost = (ot_cost or 0) + val

        return {
            "avg_lunch": avg_lunchE,
            "avg_start": avg_start,
            "avg_finish": avg_finish,
            "avg_lowd": avg_lowdE,
            "rr_sb_days": int(rr_sb_days),
            "days_worked": int(days_worked),
            "holidays": int(holidays),
            "pct_rr_sb": pct_rr_sb,
            "expenses": f"£{expenses:,.0f}",
            "ot_cost": ("£{0:,.0f}".format(ot_cost) if isinstance(ot_cost,(int,float)) else "—")
        }

    # Layout: tiles 5 across
    PER_ROW = 5
    tiles = []
    for name in present:
        m = engineer_metrics(name)
        tiles.append((name, m))

    # Render grid
    for i in range(0, len(tiles), PER_ROW):
        cols = st.columns(min(PER_ROW, len(tiles)-i), gap="large")
        for (name, m), col in zip(tiles[i:i+PER_ROW], cols):
            with col:
                # Top title + 10-line mini-summary
                st.caption("Month to date")
                st.markdown(f"**{name}**")
                st.markdown(
                    f"""
                    • Avg lunch: **{m['avg_lunch']}**  
                    • Start: **{m['avg_start']}**  |  Finish: **{m['avg_finish']}**  
                    • LOWD: **{m['avg_lowd']}**  
                    • RR/SB days: **{m['rr_sb_days']}** / worked **{m['days_worked']}**  (**{m['pct_rr_sb']}**)  
                    • Holidays: **{m['holidays']}**  
                    • Expenses: **{m['expenses']}**  
                    • OT cost: **{m['ot_cost']}**
                    """.strip()
                )

    st.caption(f"Window: {start_dt:%b %Y} – {end_dt:%b %Y}")


# ---------- The one and only renderer ----------
def render_team_overview(team_name: str, tab_index: int):
    labels_all = [...]  # however you build your month labels

    sel_label = st.selectbox(
        "Month",
        labels_all,
        index=len(labels_all) - 1,
        key=f"month_select_{tab_index}_{team_name.replace(' ', '_')}"
    )

    df = get_team_df(team_name)






# ---------- The one and only renderer ----------
def render_team_overview(team_name: str, tab_index: int):
    # 1) Data
    df = get_team_df(team_name)
    if df.empty:
        st.warning("No data found for this team."); return
    df = df.copy(); df.columns = df.columns.str.strip()

    # 2) Month picker with "All months"
    df_full = df.copy()
    date_candidates = [c for c in df.columns if ("date" in c.lower() or "time" in c.lower())]
    best_col, best_non_null = None, -1
    for c in date_candidates:
        dtc = pd.to_datetime(df[c], errors="coerce", dayfirst=True)
        if dtc.notna().sum() > best_non_null: best_non_null, best_col = dtc.notna().sum(), c

    if best_col is None:
        st.warning("Could not detect a date column; showing unfiltered data.")
        df_full["_dt"] = pd.NaT; sel_period = None; gran = "Daily"
    else:
        df_full["_dt"] = pd.to_datetime(df_full[best_col], errors="coerce", dayfirst=True)
        month_index = pd.PeriodIndex(df_full["_dt"].dropna().dt.to_period("M")).sort_values().unique()
        if len(month_index) == 0:
            st.warning("No valid dates found; showing unfiltered data.")
            sel_period = None; gran = "Daily"
        else:
            labels = [m.to_timestamp().strftime("%b %Y") for m in month_index]
            labels_all = ["All months"] + labels
            sel_label = st.selectbox(
                "Month",
                labels_all,
                index=len(labels_all) - 1,
                key=f"month_select_{tab_index}_{team_name.replace(' ', '_')}"  # ✅ unique key
            )

            if sel_label == "All months":
                sel_period = None; gran = "Monthly"; df = df_full.copy()
            else:
                sel_period = month_index[labels.index(sel_label)]
                df = df_full[df_full["_dt"].dt.to_period("M") == sel_period].copy()
                gran = "Daily"


    # 3) MoM trends (works for both All months and a specific month)
    mom = compute_mom_trends(df_full, df, sel_period)

    # 4) KPIs & series
    cols = set(df.columns)
    k = team_kpis(df)
    comp_m    = _status_mask(df, "Completed")
    cancel_m  = _status_mask(df, "Cancelled")
    notdone_m = _status_mask(df, "Not Done")
    pend_m    = _status_mask(df, "Pending")
    start_m   = _status_mask(df, "Started")

    visits_ts     = _series_by_period(df,           period=gran, reducer="count")
    value_ts      = _series_by_period(df,           period=gran, reducer="sum_num",  col="Total Value") if "Total Value" in cols else [0]
    comp_ts       = _series_by_period(df[comp_m],   period=gran, reducer="count")
    cancel_ts     = _series_by_period(df[cancel_m], period=gran, reducer="count")
    notdone_ts    = _series_by_period(df[notdone_m],period=gran, reducer="count")
    pend_ts       = _series_by_period(df[pend_m],   period=gran, reducer="count")
    start_ts      = _series_by_period(df[start_m],  period=gran, reducer="count")
    active_eng_ts = _series_by_period(df,           period=gran, reducer="nunique", col="Name") if "Name" in cols else [0]
    vpv_ts        = _value_per_visit_series(df, period=gran)

    if "Total Working Time" in cols:
        work_avg_ts = _series_by_period(df, period=gran, reducer="mean_time", col="Total Working Time"); tw_name = "Total Working Time"
    elif "Total working time" in cols:
        work_avg_ts = _series_by_period(df, period=gran, reducer="mean_time", col="Total working time"); tw_name = "Total working time"
    else:
        work_avg_ts = [0]; tw_name = None
    total_time_ts = _series_by_period(df, period=gran, reducer="sum_time", col="Total Time") if "Total Time" in cols else [0]

    if "Total Time" in cols: lunch_col = "Total Time"
    elif "Total Working Time" in cols: lunch_col = "Total Working Time"
    elif "Total working time" in cols: lunch_col = "Total working time"
    else: lunch_col = None
    if "Visit Type" in cols and lunch_col:
        lunch_df = df[df["Visit Type"].astype(str).str.contains("Lunch", case=False, na=False)]
        lunch_avg_ts = _series_by_period(lunch_df, period=gran, reducer="mean_time", col=lunch_col) if not lunch_df.empty else [0]
    else:
        lunch_avg_ts = [0]

    if tw_name:
        tw = pd.to_timedelta(df[tw_name].astype(str), errors="coerce")
        valid = tw.notna() & (tw > pd.Timedelta(0))
        over_df  = df[valid & (tw > THRESHOLD_TWT)]
        under_df = df[valid & (tw <= THRESHOLD_TWT)]
        over_1025_ts  = _series_by_period(over_df,  period=gran, reducer="count")
        under_1025_ts = _series_by_period(under_df, period=gran, reducer="count")
    else:
        over_1025_ts = under_1025_ts = [0]

    ot_title, ot_value, _, overtime_ts = overtime_from_df(df, period=gran)

    # ===== 4 × 4 CARDS =====
    r1c1, r1c2, r1c3, r1c4 = st.columns(4)
    with r1c1:
        card("Total Value (£)", f"£{k['total_value']:,.0f}" if not np.isnan(k['total_value']) else "N/A",
             "Sum of 'Total Value'", trend=mom["value_pct"], spark=value_ts, color="#16a34a", source="Oracle")
    with r1c2:
        card("Completed (%)", f"{k['comp_rate']:.1f}%",
             f"{k['completed']:,} of {k['total_visits']:,}",
             bar_pct=k['comp_rate'], trend=mom["completed_rate_pp"], spark=comp_ts, color="#06b6d4", source="Oracle")
    with r1c3:
        card("Visits", f"{k['total_visits']:,}", "All visits",
             trend=mom["visits_pct"], spark=visits_ts, color="#ef4444", source="Oracle")
    with r1c4:
        card("Active engineers", f"{k['unique_engs']:,}", "Unique per period",
             trend=mom["active_pct"], spark=active_eng_ts, color="#22c55e", source="")

    r2c1, r2c2, r2c3, r2c4 = st.columns(4)
    with r2c1:
        card("Avg Working Time", k['avg_work'], "HH:MM", trend=None, spark=work_avg_ts, color="#22c55e", source="(avg minutes)")
    with r2c2:
        card("Cancelled (%)", f"{k['cancel_rate']:.1f}%", f"{k['cancelled']:,} cases",
             bar_pct=k['cancel_rate'], trend=mom["cancel_rate_pp"], spark=cancel_ts, color="#ef4444", source="")
    with r2c3:
        card("Not Done (%)", f"{k['notdone_rate']:.1f}%", f"{k['not_done']:,} cases",
             bar_pct=k['notdone_rate'], trend=mom["notdone_rate_pp"], spark=notdone_ts, color="#ef4444", source="")
    with r2c4:
        card("Total time", k['total_time'], "HH:MM total", trend=None, spark=total_time_ts, color="#06b6d4", source="(hours per period)")

    r3c1, r3c2, r3c3, r3c4 = st.columns(4)
    with r3c1:
        card("Lunch avg", k['avg_lunch'], "Avg of 'Lunch' visits", trend=None, spark=lunch_avg_ts, color="#22c55e", source="(avg minutes)")
    with r3c2:
        card("Pending (count)", f"{int(np.sum(pend_ts) if len(pend_ts)>0 else 0):,}", "", trend=None, spark=pend_ts, color="#f59e0b", source="")
    with r3c3:
        card("Started (count)", f"{int(np.sum(start_ts) if len(start_ts)>0 else 0):,}", "", trend=None, spark=start_ts, color="#3b82f6", source="")
    with r3c4:
        val_per_visit = (k['total_value'] / k['total_visits']) if (k['total_visits'] and not np.isnan(k['total_value'])) else np.nan
        card("Value per visit", f"£{val_per_visit:,.0f}" if val_per_visit == val_per_visit else "N/A",
             "(avg per period)", trend=mom["vpv_pct"], spark=vpv_ts, color="#8b5cf6", source="")

    r4c1, r4c2, r4c3, r4c4 = st.columns(4)
    with r4c1:
        card("Over 10:25 (count)", f"{k['over_1025_count']:,}", "Total Working Time > 10:25",
             trend=mom["over_1025_pct"], spark=over_1025_ts, color="#f43f5e", source="")
    with r4c2:
        card("Under 10:25 (count)", f"{k['under_1025_count']:,}", "Total Working Time ≤ 10:25",
             trend=mom["under_1025_pct"], spark=under_1025_ts, color="#22c55e", source="")
    with r4c3:
        card(ot_title, ot_value, "", trend=None, spark=overtime_ts, color="#f59e0b", source="sum per period")
    with r4c4:
        card("—", "—", "", trend=None, spark=None, source="")

    # ===== 💷 Budget KPIs (new row of 4) =====
    st.markdown("#### 💷 Budget KPIs")

    alloc, used, remaining, pct_used = compute_team_budget_metrics(team_name)


    b1, b2, b3, b4 = st.columns(4)

    with b1:
        card("Budget Allocated", fmt_money(alloc), "From budgets.csv",
            trend=None, spark=None, color="#3b82f6", source="Budget")

    with b2:
        card("Budget Used", fmt_money(used), "From expenses.csv",
            trend=None, spark=None, color="#ef4444", source="Budget")

    with b3:
        bar_remaining = None
        try:
            bar_remaining = max(0, min(int(remaining / alloc * 100), 100)) if alloc else None
        except Exception:
            bar_remaining = None
        card("Budget Remaining", fmt_money(remaining), "Allocated − Used",
            trend=None, bar_pct=bar_remaining, spark=None, color="#22c55e", source="Budget")

    with b4:
        bar_used = None
        try:
            bar_used = max(0, min(int(pct_used), 100)) if pct_used == pct_used else None
        except Exception:
            bar_used = None
        pct_txt = f"{pct_used:.1f}%" if pct_used == pct_used else "—"
        card("Budget % Used", pct_txt, "",
            trend=None, bar_pct=bar_used, spark=None, color="#06b6d4", source="Budget")

  

# ---------- Screen entry (tabs + back) ----------
if st.session_state.get("screen") == "team_overview":
    # Top row: back buttons
    left, right = st.columns([1, 1])
    with left:
        if st.button("⬅ Back to Instructions", key="team_back_instr", use_container_width=True):
            st.session_state.screen = "instructions_guide"
            st.rerun()
    with right:
        if st.session_state.get("username") in ("Rob Mangar", "Mark Wilson"):
            if st.button("Back to Leadership", key="team_back_exec", use_container_width=True):
                st.session_state.screen = "operational_area"
                st.session_state.op_area_section = "exec_overview"
                st.rerun()

    # --- Actual Team Overview content (this was missing previously)
    user_role = (st.session_state.get("role") or "").lower()
    user_team = st.session_state.get("user_team")

    # Managers see ONLY their own team tab
    if user_role == "manager" and user_team in TEAM_TAGS:
        st.subheader(f"{user_team} — Team Overview")
        render_team_overview(user_team, 0)
        render_orbit_ai(user_team)     # Orbit scoped to their team
        st.stop()

    # Execs / leadership: show the 4 tabs
    team_tabs = st.tabs(["VIP North", "VIP South", "Tier 2 North", "Tier 2 South"])
    with team_tabs[0]:
        render_team_overview("VIP North", 0)
        render_orbit_ai("VIP North")
    with team_tabs[1]:
        render_team_overview("VIP South", 1)
        render_orbit_ai("VIP South")
    with team_tabs[2]:
        render_team_overview("Tier 2 North", 2)
        render_orbit_ai("Tier 2 North")
    with team_tabs[3]:
        render_team_overview("Tier 2 South", 3)
        render_orbit_ai("Tier 2 South")

elif st.session_state.get("screen") == "team_graphs":
    # One-page graphs view for the signed-in team
    render_team_graphs_page()
    st.stop()  # important: do not fall through to any other handlers

elif st.session_state.get("screen") == "team_engineers":
    render_team_engineers_page()
    st.stop()

elif st.session_state.get("screen") == "suggestions":
    render_suggestions_page()


def show_exec_logo(
    img_name: str = "Mark W logo.png",
    ratios=(1, 6, 1),            # <- change these to tweak centering width (e.g., (1,8,1) or (2,5,2))
) -> None:
    """
    Center the 'Mark W logo.png' image using columns.
    Looks in the same folder as this script for a few filename variants.
    Shows a warning if nothing is found.
    """
    folder = Path(__file__).parent
    candidates = [
        img_name,
        img_name.replace("  ", " "),
        img_name.replace(" ", "_"),
        "Mark W logo.jpg", "Mark W logo.jpeg", "Mark W logo.PNG", "Mark W logo.png",
    ]

    for name in candidates:
        p = (folder / name)
        if p.exists():
            left, center, right = st.columns(ratios)
            with center:
                st.image(str(p), use_container_width=True)   # fills the center column nicely
            return

    # If we get here, we didn't find an image
    looked = ", ".join([str(folder / n) for n in candidates])
    st.warning(f"⚠️ Could not find an exec logo. Looked for: {looked}")



# ---------- Screen: Sky Retail (Rachel) ----------
if st.session_state.get("screen") == "sky_retail":
    import re, numpy as np, pandas as pd, plotly.express as px, plotly.graph_objects as go

    # Back row: left = Main Menu (everyone), right = Leadership (Rob/Mark only)
    left, right = st.columns([3, 1])
    with left:
        if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_sky_retail"):
            st.session_state.screen = "instruction_guide"
            st.rerun()
    with right:
        if st.session_state.get("username") in ("Rob Mangar", "Mark Wilson"):
            if st.button("Back to Leadership", key="sky_retail_back_exec", use_container_width=True):
                st.session_state.screen = "operational_area"
                st.session_state.op_area_section = "exec_overview"
                st.rerun()

    st.title("Sky Retail")

    # ===== data guard + source =====
    if "combined_oracle_df" not in globals() or combined_oracle_df is None or combined_oracle_df.empty:
        st.error("Oracle data is not loaded.")
        st.stop()

    data = combined_oracle_df.copy()

    col_name = "Sky Retail Stakeholder"
    if col_name not in data.columns:
        st.error(f"Column '{col_name}' not found in Oracle datasets.")
        st.stop()

    # ===== helper functions (scoped to this screen) =====
    def time_to_minutes(val):
        import datetime
        if pd.isnull(val):
            return 0
        if isinstance(val, (int, float)):
            return val
        if isinstance(val, datetime.timedelta):
            return val.total_seconds() / 60
        if isinstance(val, datetime.time):
            return val.hour * 60 + val.minute
        if isinstance(val, str):
            try:
                t = pd.to_datetime(val).time()
                return t.hour * 60 + t.minute
            except Exception:
                return 0
        return 0

    def minutes_to_hhmm(minutes):
        hours = int(minutes // 60)
        mins = int(minutes % 60)
        return f"{hours}:{mins:02}"

    def clean_stakeholder(val):
        v = str(val).strip().lower()
        if re.search(r"currys?|curry's", v): return "Currys"
        if "ee" in v: return "EE"
        if "sky" in v: return "Sky Retail"
        return "Other"

    def better_forecast(series, months=6):
        series = series.sort_index()
        n = len(series)
        if n == 0:
            return [0] * months
        if n < 3:
            avg = int(series.mean())
            return [avg] * months
        x = np.arange(n) if n < 4 else np.arange(n-4, n)
        y = series.values if n < 4 else series.values[-4:]
        m, b = np.polyfit(x, y, 1)
        fut_x = np.arange(n, n + months)
        y_pred = (m * fut_x + b)
        recent_avg = max(1, int(series.tail(2).mean()))
        y_pred = np.where(y_pred < recent_avg, recent_avg, y_pred)
        return y_pred.round().astype(int)

    # ===== data prep =====
    data["Sky Retail Stakeholder Clean"] = data[col_name].apply(clean_stakeholder)

    stakeholders = ["Currys", "Sky Retail", "EE"]
    tabs = st.tabs(stakeholders)

    # Optional: AgGrid (fallback to st.dataframe if missing)
    try:
        from st_aggrid import AgGrid, GridOptionsBuilder
        def show_aggrid(df, height=300):
            gb = GridOptionsBuilder.from_dataframe(df)
            gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=10)
            gb.configure_default_column(editable=False, filter=True, sortable=True)
            grid_options = gb.build()
            AgGrid(df, gridOptions=grid_options, height=height, fit_columns_on_grid_load=True)
    except Exception:
        def show_aggrid(df, height=300):
            st.dataframe(df, use_container_width=True, height=height)

    for i, stakeholder in enumerate(stakeholders):
        with tabs[i]:
            df = data[data["Sky Retail Stakeholder Clean"] == stakeholder].copy()
            if df.empty:
                st.info(f"No data for {stakeholder}.")
                continue

            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
            df = df[df["Date"].notna()]
            df["Month"] = df["Date"].dt.to_period("M").astype(str)

            if "Total Time" in df.columns:
                df["Total Time (min)"] = df["Total Time"].apply(time_to_minutes)
            else:
                df["Total Time (min)"] = 0

            # ===== KPIs =====
            with st.expander(f"📊 {stakeholder} KPIs", expanded=True):
                count_visits = len(df)
                total_value = df["Total Value"].sum() if "Total Value" in df.columns else 0
                avg_value = df["Total Value"].mean() if "Total Value" in df.columns else 0
                total_time = df["Total Time (min)"].sum()
                avg_time = df["Total Time (min)"].mean()

                col1, col2, col3 = st.columns(3)
                with col1: st.metric(f"Total {stakeholder} Visits", f"{count_visits:,}")
                with col2: st.metric("Total Value (£)", f"£{total_value:,.0f}")
                with col3: st.metric("Avg Value", f"£{avg_value:,.0f}")

                col4, col5, col6 = st.columns(3)
                with col4: st.metric("Total Time", f"{minutes_to_hhmm(total_time)}")
                with col5: st.metric("Average Time", f"{minutes_to_hhmm(avg_time)}")
                with col6: pass

                # Monthly Visits (last 4)
                monthly_visits = (
                    df.groupby(df["Date"].dt.to_period("M").astype(str))
                      .size()
                      .reset_index(name="Visit Count")
                      .rename(columns={"Date": "Month"})
                )
                st.markdown("**Monthly Visits (last 4 months shown)**")
                show_aggrid(monthly_visits.tail(4))

                # Monthly Value (last 4)
                if "Total Value" in df.columns:
                    monthly_value = (
                        df.groupby(df["Date"].dt.to_period("M").astype(str))["Total Value"]
                          .sum()
                          .reset_index(name="Total Value")
                          .rename(columns={"Date": "Month"})
                    )
                    monthly_value["Total Value"] = monthly_value["Total Value"].map("£{0:,.0f}".format)
                    st.markdown("**Monthly Total Value (£) (last 4 months shown)**")
                    show_aggrid(monthly_value.tail(4))

                # Visits by Day of Week
                df["DayOfWeek"] = df["Date"].dt.day_name()
                day_order = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
                visits_by_day = df["DayOfWeek"].value_counts().reindex(day_order, fill_value=0).reset_index()
                visits_by_day.columns = ["DayOfWeek","Visit Count"]
                st.markdown("**Visits by Day of Week**")
                show_aggrid(visits_by_day)

                # Value by Day of Week
                if "Total Value" in df.columns:
                    value_by_day = df.groupby(df["Date"].dt.day_name())["Total Value"].sum().reindex(day_order, fill_value=0).reset_index()
                    value_by_day.columns = ["DayOfWeek","Total Value"]
                    value_by_day["Total Value"] = value_by_day["Total Value"].map("£{0:,.0f}".format)
                    st.markdown("**Total Value (£) by Day of Week**")
                    show_aggrid(value_by_day)

            # ===== Breakdown by original stakeholder label =====
            with st.expander(f"📋 Breakdown by {stakeholder} Stakeholder", expanded=False):
                name_col = "Name" if "Name" in df.columns else df.columns[0]
                by_stake = df.groupby(col_name).agg(
                    Visits=(name_col, 'count'),
                    Value=('Total Value', 'sum') if "Total Value" in df.columns else ('Total Value', 'sum'),
                    Time=('Total Time (min)', 'sum')
                ).sort_values("Visits", ascending=False)
                by_stake["Time (hh:mm)"] = by_stake["Time"].apply(minutes_to_hhmm)
                by_stake["Value (£)"] = by_stake["Value"].apply(lambda x: f"£{x:,.0f}")
                by_stake = by_stake.drop(columns=["Value", "Time"])
                st.dataframe(by_stake, use_container_width=True)

            # ===== Monthly Trend charts =====
            with st.expander(f"📈 Monthly Visit Trends for {stakeholder}", expanded=False):
                by_month = df.groupby(["Month", col_name]).size().reset_index(name="Visits")
                fig_visits = px.bar(by_month, x="Month", y="Visits", color=col_name, barmode="group",
                                   title=f"Visits per Month by Stakeholder for {stakeholder}")
                st.plotly_chart(fig_visits, use_container_width=True)

            with st.expander(f"📈 Monthly Value Trends for {stakeholder}", expanded=False):
                if "Total Value" in df.columns:
                    by_value = df.groupby(["Month", col_name])["Total Value"].sum().reset_index()
                    fig_value = px.line(by_value, x="Month", y="Total Value", color=col_name,
                                       title=f"Value per Month by Stakeholder for {stakeholder}")
                    st.plotly_chart(fig_value, use_container_width=True)

            # ===== Activity Status pie =====
            with st.expander(f"📊 Visit Activity Status Split for {stakeholder}", expanded=False):
                if "Activity Status" in df.columns:
                    status_df = df["Activity Status"].value_counts().reset_index()
                    status_df.columns = ["Activity Status", "Count"]
                    fig_pie = px.pie(status_df, names="Activity Status", values="Count")
                    st.plotly_chart(fig_pie, use_container_width=True)

            # ===== Team breakdown pivot =====
            with st.expander("📋 Team Breakdown by Stakeholder", expanded=False):
                if "Team" in df.columns:
                    team_pivot = pd.pivot_table(
                        df, index="Team", columns="Sky Retail Stakeholder Clean",
                        values="Name" if "Name" in df.columns else df.columns[0],
                        aggfunc="count", fill_value=0
                    )
                    st.dataframe(team_pivot, use_container_width=True)

            # ===== Engineer breakdown (by stakeholder) =====
            eng_pivot = (
                df.groupby(["Sky Retail Stakeholder Clean", "Name" if "Name" in df.columns else df.columns[0]])
                .agg(Visits=("Name" if "Name" in df.columns else df.columns[0], "count"),
                     Value=("Total Value", "sum") if "Total Value" in df.columns else ("Total Value", "sum"),
                     Time=("Total Time (min)", "sum"))
                .reset_index()
            )
            eng_pivot["Time (hh:mm)"] = eng_pivot["Time"].apply(minutes_to_hhmm)
            eng_pivot["Value (£)"] = eng_pivot["Value"].apply(lambda x: f"£{x:,.0f}")

            for eng_stakeholder in eng_pivot["Sky Retail Stakeholder Clean"].unique():
                with st.expander(f"👤 Engineer Breakdown for {eng_stakeholder}", expanded=False):
                    nm = "Name" if "Name" in df.columns else df.columns[0]
                    display = eng_pivot[eng_pivot["Sky Retail Stakeholder Clean"] == eng_stakeholder][
                        [nm, "Visits", "Value (£)", "Time (hh:mm)"]
                    ].sort_values("Visits", ascending=False)
                    st.dataframe(display, use_container_width=True)

            # ===== Overall forecasts =====
            with st.expander(f"🔮 Overall {stakeholder} Forecasts", expanded=False):
                # Visits
                monthly_visits = df.groupby("Month").size().sort_index()
                if len(monthly_visits) >= 2:
                    fc_vals = better_forecast(monthly_visits, months=6)
                    last_month = pd.Period(monthly_visits.index.max(), freq="M")
                    fut_months = [str(last_month + i) for i in range(1, 7)]
                    fc_df = pd.concat([
                        pd.DataFrame({"Month": monthly_visits.index, "Visits": monthly_visits.values, "Type": "Actual"}),
                        pd.DataFrame({"Month": fut_months, "Visits": fc_vals, "Type": "Forecast"})
                    ], ignore_index=True)
                    fig_visits_fc = px.line(fc_df, x="Month", y="Visits", color="Type", markers=True,
                                            title=f"{stakeholder} – Overall Visits Forecast")
                    st.plotly_chart(fig_visits_fc, use_container_width=True)
                else:
                    st.info("Not enough data for an overall visits forecast.")

                # Value
                if "Total Value" in df.columns:
                    monthly_value = df.groupby("Month")["Total Value"].sum().sort_index()
                    if len(monthly_value) >= 2:
                        fc_vals = better_forecast(monthly_value, months=6)
                        last_month = pd.Period(monthly_value.index.max(), freq="M")
                        fut_months = [str(last_month + i) for i in range(1, 7)]
                        fc_df = pd.concat([
                            pd.DataFrame({"Month": monthly_value.index, "Value": monthly_value.values, "Type": "Actual"}),
                            pd.DataFrame({"Month": fut_months, "Value": fc_vals, "Type": "Forecast"})
                        ], ignore_index=True)
                        fig_value_fc = px.line(fc_df, x="Month", y="Value", color="Type", markers=True,
                                               title=f"{stakeholder} – Overall Value Forecast (£)" )
                        st.plotly_chart(fig_value_fc, use_container_width=True)
                    else:
                        st.info("Not enough data for a value forecast.")

            # ===== Month-on-Month change tables =====
            pivot_visits = (
                df.groupby(["Month", "Sky Retail Stakeholder Clean"])
                  .size().unstack(fill_value=0)
                  .reindex(columns=[stakeholder], fill_value=0)
                  .sort_index()
            )
            pivot_value = (
                df.groupby(["Month", "Sky Retail Stakeholder Clean"])["Total Value"]
                  .sum().unstack(fill_value=0)
                  .reindex(columns=[stakeholder], fill_value=0)
                  .sort_index()
            )

            def make_change_table(pivot):
                df_ = pivot.copy()
                for col in df_.columns:
                    df_[f"{col} Δ"] = df_[col].diff().fillna(0).astype(int)
                    max_val = df_[col].max()
                    min_val = df_[col].min()
                    df_[f"{col} ΔMax"] = df_[col] - max_val
                    df_[f"{col} ΔMin"] = df_[col] - min_val
                return df_

            visits_tbl = make_change_table(pivot_visits)
            value_tbl = make_change_table(pivot_value)

            with st.expander("📊 Month-on-Month Change Table (Visits)", expanded=False):
                st.dataframe(
                    visits_tbl.style
                        .format({col: "{:,}" for col in visits_tbl.columns if "Δ" not in col})
                        .format({col: "{:+,}" for col in visits_tbl.columns if "Δ" in col}),
                    use_container_width=True
                )

            with st.expander("📊 Month-on-Month Change Table (Value)", expanded=False):
                value_cols  = [c for c in value_tbl.columns if "Δ" not in c]
                delta_cols  = [c for c in value_tbl.columns if "Δ" in c]
                # format display
                fmt_tbl = value_tbl.copy()
                for col in delta_cols:
                    fmt_tbl[col] = fmt_tbl[col].apply(lambda x: f"£{x:+,}")
                fmt_tbl[value_cols] = fmt_tbl[value_cols].round(0).astype(int)
                for col in value_cols:
                    fmt_tbl[col] = "£" + fmt_tbl[col].map("{:,}".format)
                st.dataframe(fmt_tbl, use_container_width=True)

            # ===== Combined charts (collapsible) =====
            with st.expander("📊 Combined Monthly Trend Charts", expanded=False):
                fig_value_trends = px.line(value_tbl.reset_index(), x="Month", y=[stakeholder],
                                           title=f"Monthly Total Value for {stakeholder} (£)", markers=True)
                st.plotly_chart(fig_value_trends, use_container_width=True)

                fig_visits_trends = px.line(visits_tbl.reset_index(), x="Month", y=[stakeholder],
                                            title=f"Monthly Total Visits for {stakeholder}", markers=True)
                st.plotly_chart(fig_visits_trends, use_container_width=True)

                value_change_cols  = [c for c in value_tbl.columns if 'Δ' in c and 'Max' not in c and 'Min' not in c]
                fig_value_changes  = px.bar(value_tbl.reset_index(), x='Month', y=value_change_cols,
                                            title=f"Month-on-Month Change in Value for {stakeholder}", barmode='group')
                st.plotly_chart(fig_value_changes, use_container_width=True)

                visits_change_cols = [c for c in visits_tbl.columns if 'Δ' in c and 'Max' not in c and 'Min' not in c]
                fig_visits_changes = px.bar(visits_tbl.reset_index(), x='Month', y=visits_change_cols,
                                            title=f"Month-on-Month Change in Visits for {stakeholder}", barmode='group')
                st.plotly_chart(fig_visits_changes, use_container_width=True)

                fig_heatmap_value_max = go.Figure(go.Heatmap(
                    z=value_tbl[[c for c in value_tbl.columns if 'ΔMax' in c]].values.T,
                    x=value_tbl.index,
                    y=[c.replace(' ΔMax', '') for c in value_tbl.columns if 'ΔMax' in c],
                    colorscale='RdBu', colorbar=dict(title='Diff vs Max'), zmid=0))
                fig_heatmap_value_max.update_layout(title=f'Value Difference vs Max for {stakeholder}')
                st.plotly_chart(fig_heatmap_value_max, use_container_width=True)

                fig_heatmap_visits_max = go.Figure(go.Heatmap(
                    z=visits_tbl[[c for c in visits_tbl.columns if 'ΔMax' in c]].values.T,
                    x=visits_tbl.index,
                    y=[c.replace(' ΔMax', '') for c in visits_tbl.columns if 'ΔMax' in c],
                    colorscale='RdBu', colorbar=dict(title='Diff vs Max'), zmid=0))
                fig_heatmap_visits_max.update_layout(title=f'Visits Difference vs Max for {stakeholder}')
                st.plotly_chart(fig_heatmap_visits_max, use_container_width=True)

            # ===== Visits by Day of Week (bar) =====
            with st.expander("📅 Visits by Day of Week", expanded=False):
                day_order = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
                visits_by_day2 = df.groupby(df["Date"].dt.day_name()).size().reindex(day_order, fill_value=0).reset_index()
                visits_by_day2.columns = ["DayOfWeek","Visits"]
                fig_day = px.bar(visits_by_day2, x="DayOfWeek", y="Visits",
                                 title=f"{stakeholder} Visits by Day of Week",
                                 labels={"DayOfWeek":"Day of Week","Visits":"Number of Visits"},
                                 color="Visits", color_continuous_scale="Blues")
                st.plotly_chart(fig_day, use_container_width=True)

            # ===== Raw data =====
            with st.expander(f"🔎 Show Raw Data for {stakeholder}", expanded=False):
                st.dataframe(df.dropna(axis=1, how="all"), use_container_width=True)

# ---- SKY BUSINESS: shared normaliser + patterns + counters ----
import unicodedata
import pandas as pd

def sb__strip_accents(text: str) -> str:
    return unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("utf-8")

def sb__normalise(series: pd.Series) -> pd.Series:
    s = series.astype(str).map(sb__strip_accents).str.lower()
    return (
        s.str.replace(r"[\u2010-\u2015–—]", "-", regex=True)   # unusual dashes → hyphen
         .str.replace(r"[^a-z0-9]+", " ", regex=True)         # drop punctuation, collapse spaces
         .str.replace(r"\s+", " ", regex=True)
         .str.strip()
    )

def sb__all_text(df: pd.DataFrame) -> pd.Series:
    """Concatenate ALL object columns into one searchable string."""
    if df.empty:
        return pd.Series([], dtype=str)
    cols = df.select_dtypes(include="object").columns
    if not len(cols):
        return pd.Series([""] * len(df), index=df.index)
    return df[cols].fillna("").agg(" ".join, axis=1)

# ONE source of truth for the patterns your SLA tiles use
SB_SLA_PATTERNS = {
    "nero_all":  r"\bcaf+\w*\s*nero\b",                    # cafe/caffe/caffè nero
    "nero_2h":   r"\bcaf+\w*\s*nero\b.*\b(2\s*hour|2\s*hr)\b",
    "nero_next": r"\bcaf+\w*\s*nero\b.*\bnext\s*day\b",
    "nero_4h":   r"\bcaf+\w*\s*nero\b.*\b(4\s*hour|4\s*hr)\b",
    "sla_8h":    r"\b8\s*hour\s*sla\b",
}

def sb__build_masks(norm_text: pd.Series) -> dict[str, pd.Series]:
    """Return a dict of boolean masks keyed by pattern name."""
    if norm_text.empty:
        return {k: pd.Series(False, index=norm_text.index) for k in SB_SLA_PATTERNS}
    return {k: norm_text.str.contains(pat, na=False) for k, pat in SB_SLA_PATTERNS.items()}

def sb__count_totals(df: pd.DataFrame) -> dict[str, int]:
    """Counts for current filtered df, using the shared rules."""
    norm = sb__normalise(sb__all_text(df))
    masks = sb__build_masks(norm)
    return {k: int(v.sum()) for k, v in masks.items()}

def sb__series_by_month(df: pd.DataFrame, mask: pd.Series, month_col: str) -> pd.Series:
    """Monthly counts for sparkline on filtered base df."""
    if df.empty or mask.sum() == 0:
        return pd.Series([], dtype=int)
    return (df.loc[mask]
            .groupby(month_col).size()
            .sort_index())

# ==============================
# EXECUTIVE OVERVIEW (Mark Wilson)
# ==============================
def render_exec_overview(embed: bool = False):
    import pandas as pd
    from pathlib import Path

    exec_uri = to_data_uri("SkyExec.png")
    if exec_uri:
        st.markdown(
            f"""
            <div style="text-align:center; margin:10px 0;">
            <img src="{exec_uri}" alt="Sky Executive"
                style="max-width:600px; width:100%; height:auto;">
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.warning("⚠️ Could not find 'SkyExec.png' in the app folder.")




    # Title directly UNDER the logo
    st.markdown(
        "<h1 style='text-align:center; margin-top:0.5rem;'>🧭 Executive Overview</h1>",
        unsafe_allow_html=True
    )
# =========================
    # Invoices — Executive Snapshot
    # =========================
    import pandas as pd
    import numpy as np

    # --- tiny fallbacks if helpers are missing (won't override if you already have them)
    def _safe_pct_delta(curr: float, prev: float):
        if prev in (0, None) or pd.isna(prev):
            return "—", "—"
        pct = (curr - prev) / prev * 100.0
        return (f"{pct:+.1f}%", "▲" if pct >= 0 else "▼")

    def _format_money(x: float) -> str:
        try:
            return f"£{float(x):,.2f}"
        except Exception:
            return "£0.00"

    def _get_month_series(df: pd.DataFrame, col="TotalValue"):
        if df.empty or "VisitDate" not in df.columns:
            return pd.DataFrame(columns=["Month", col])
        d = df.dropna(subset=["VisitDate"]).copy()
        d["Month"] = d["VisitDate"].dt.to_period("M").dt.to_timestamp()
        return d.groupby("Month")[col].sum().reset_index()

    # --- load invoices
    inv_all = load_invoices()
    if "RechargeableFlag" not in inv_all.columns and "Rechargeable" in inv_all.columns:
        # normalise text into boolean
        rf = inv_all["Rechargeable"].astype(str).str.strip().str.lower()
        inv_all["RechargeableFlag"] = rf.isin({"yes","y","true","1","chargeable","rechargeable","rc"})

    # If your Exec page has a selected month/range, you can filter here.
    # As a safe default we build the last 12 months snapshot.
    if not inv_all.empty:
        inv_all = inv_all.copy()
        inv_all["VisitDate"] = pd.to_datetime(inv_all["VisitDate"], errors="coerce")
        inv_all = inv_all.dropna(subset=["VisitDate"])

        # last 12 complete months window
        end_month = pd.Timestamp.today().to_period("M").to_timestamp()
        start_month = (end_month - pd.offsets.MonthBegin(12))
        inv_12 = inv_all[ (inv_all["VisitDate"] >= start_month) & (inv_all["VisitDate"] <= end_month + pd.offsets.MonthEnd(0)) ]

        # split
        rec_df  = inv_12[inv_12.get("RechargeableFlag") == True]
        non_df  = inv_12[inv_12.get("RechargeableFlag") == False]

        # monthly series for sparklines
        m_rec   = _get_month_series(rec_df,  "TotalValue")
        m_non   = _get_month_series(non_df,  "TotalValue")

        # % rechargeable per month
        m_all   = _get_month_series(inv_12,  "TotalValue")
        if not m_all.empty:
            m_pct = m_all.merge(m_rec.rename(columns={"TotalValue":"Rec"}), on="Month", how="left") \
                        .merge(m_non.rename(columns={"TotalValue":"Non"}), on="Month", how="left") \
                        .fillna(0.0)
            m_pct["PctRec"] = (m_pct["Rec"] / (m_pct["Rec"] + m_pct["Non"]).replace(0, np.nan)) * 100.0
            m_pct = m_pct.fillna(0.0)
        else:
            m_pct = pd.DataFrame(columns=["Month","PctRec"])

        # “current” month = last point we have (within the 12m window)
        # if you want this to follow your Exec month picker, replace the next 4 lines with that selection.
        curr_month = m_all["Month"].max() if not m_all.empty else None
        prev_month = m_all["Month"].sort_values().iloc[-2] if len(m_all) > 1 else None

        # current snapshot values
        rec_now  = float(pd.to_numeric(rec_df.loc[rec_df["VisitDate"].dt.to_period("M").dt.to_timestamp() == curr_month, "TotalValue"], errors="coerce").fillna(0).sum()) if curr_month is not None else 0.0
        non_now  = float(pd.to_numeric(non_df.loc[non_df["VisitDate"].dt.to_period("M").dt.to_timestamp() == curr_month, "TotalValue"], errors="coerce").fillna(0).sum()) if curr_month is not None else 0.0
        tot_now  = rec_now + non_now
        pct_now  = (rec_now / tot_now * 100.0) if tot_now else 0.0

                # >>> totals across the selected window (to match the other KPIs)
        rec_total = float(pd.to_numeric(rec_df["TotalValue"], errors="coerce").fillna(0).sum())
        non_total = float(pd.to_numeric(non_df["TotalValue"], errors="coerce").fillna(0).sum())
        tot_total = rec_total + non_total
        pct_total = (rec_total / tot_total * 100.0) if tot_total else 0.0


        rec_prev = float(pd.to_numeric(rec_df.loc[rec_df["VisitDate"].dt.to_period("M").dt.to_timestamp() == prev_month, "TotalValue"], errors="coerce").fillna(0).sum()) if prev_month is not None else 0.0
        non_prev = float(pd.to_numeric(non_df.loc[non_df["VisitDate"].dt.to_period("M").dt.to_timestamp() == prev_month, "TotalValue"], errors="coerce").fillna(0).sum()) if prev_month is not None else 0.0
        tot_prev = rec_prev + non_prev
        pct_prev = (rec_prev / tot_prev * 100.0) if tot_prev else 0.0

        rec_delta_txt,  rec_arrow  = _safe_pct_delta(rec_now, rec_prev)
        non_delta_txt,  non_arrow  = _safe_pct_delta(non_now, non_prev)
        pct_delta_txt,  pct_arrow  = _safe_pct_delta(pct_now, pct_prev)

        # top 3 rechargeable invoice types (across the 12m context)
        top3 = (
            rec_df.groupby("InvoiceType")["TotalValue"]
                .sum().sort_values(ascending=False).head(3).reset_index()
            if not rec_df.empty and "InvoiceType" in rec_df.columns else pd.DataFrame()
        )
        with st.expander("📑 Invoices — Executive Snapshot", expanded=False):
            # ---------- title bar
            st.markdown("### 🧾 Invoices — Executive Snapshot")

            # ---------- 4 cards (use your exec card() helper if present)
            c1, c2, c3, c4 = st.columns(4)

            # helper to push a sparkline into the card
            def _spark_values(df_series, colname):
                if df_series is None or df_series.empty:
                    return None
                return df_series[colname].tolist()

            # 1) Rechargeable
            with c1:
                try:
                    card(
                        "Rechargeable",
                        _format_money(rec_total),
                        f"{rec_arrow} {rec_delta_txt} vs prev month",
                        trend=None,
                        spark=_spark_values(m_rec, "TotalValue"),
                        color="#22c55e",
                        source="last 12 months"
                    )
                except Exception:
                    # fallback (minimal look)
                    import plotly.express as px
                    st.caption("Rechargeable")
                    st.markdown(f"<div style='font-size:28px;font-weight:700'>{_format_money(rec_now)}</div>", unsafe_allow_html=True)
                    if m_rec is not None and not m_rec.empty:
                        fig = px.line(m_rec, x="Month", y="TotalValue")
                        fig.update_layout(height=90, margin=dict(l=0,r=0,t=0,b=0), showlegend=False)
                        fig.update_xaxes(visible=False); fig.update_yaxes(visible=False)
                        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
                    st.caption(f"{rec_arrow} {rec_delta_txt} vs prev month")

            # 2) Non-Chargeable
            with c2:
                try:
                    card(
                        "Non-Chargeable",
                        _format_money(non_total),
                        f"{non_arrow} {non_delta_txt} vs prev month",
                        trend=None,
                        spark=_spark_values(m_non, "TotalValue"),
                        color="#ef4444",
                        source="last 12 months"
                    )
                except Exception:
                    import plotly.express as px
                    st.caption("Non-Chargeable")
                    st.markdown(f"<div style='font-size:28px;font-weight:700'>{_format_money(non_now)}</div>", unsafe_allow_html=True)
                    if m_non is not None and not m_non.empty:
                        fig = px.line(m_non, x="Month", y="TotalValue")
                        fig.update_layout(height=90, margin=dict(l=0,r=0,t=0,b=0), showlegend=False)
                        fig.update_xaxes(visible=False); fig.update_yaxes(visible=False)
                        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
                    st.caption(f"{non_arrow} {non_delta_txt} vs prev month")

            # 3) % Rechargeable
            with c3:
                try:
                    card(
                        "% Rechargeable",
                        f"{pct_total:.1f}%",
                        f"{pct_arrow} {pct_delta_txt} vs prev month",
                        trend=None,
                        spark=_spark_values(m_pct, "PctRec"),
                        color="#06b6d4",
                        source="share of total value"
                    )
                except Exception:
                    import plotly.express as px
                    st.caption("% Rechargeable")
                    st.markdown(f"<div style='font-size:28px;font-weight:700'>{pct_now:.1f}%</div>", unsafe_allow_html=True)
                    if m_pct is not None and not m_pct.empty:
                        fig = px.line(m_pct, x="Month", y="PctRec")
                        fig.update_layout(height=90, margin=dict(l=0,r=0,t=0,b=0), showlegend=False)
                        fig.update_xaxes(visible=False); fig.update_yaxes(visible=False)
                        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
                    st.caption(f"{pct_arrow} {pct_delta_txt} vs prev month")

            # 4) Top 3 Rechargeable invoice types
            from textwrap import shorten

            with c4:
                # build clean, readable lines
                if not top3.empty:
                    def _label(row):
                        name = shorten(str(row["InvoiceType"]), width=32, placeholder="…")
                        return f"{name} – £{row['TotalValue']:,.0f}"

                    lines = [f"{i+1}) {_label(row)}" for i, row in top3.iterrows()]
                    big = lines[0]                          # first line big
                    sub = "<br>".join(lines[1:]) if len(lines) > 1 else ""  # remaining lines below
                else:
                    big, sub = "—", ""

                # try your exec card() first (it will wrap HTML in subtitle)
                try:
                    card(
                        "Top Rechargeable Types",
                        big,
                        sub,                 # contains <br> for line breaks
                        trend=None,
                        spark=None,
                        color="#a855f7",
                        source=f"by total value ({range_label})"
                    )
                except Exception:
                    # fallback rendering if card() isn't available here
                    st.caption("Top Rechargeable Types")
                    st.markdown(f"<div style='font-size:22px;font-weight:700'>{big}</div>", unsafe_allow_html=True)
                    if sub:
                        st.markdown(f"<div style='opacity:.8'>{sub}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown("<div style='opacity:.6'>No data</div>", unsafe_allow_html=True)

    # --- Data guard rails ---
    if "combined_oracle_df" not in globals():
        st.error("combined_oracle_df is not loaded yet.")
        return

    df_all = combined_oracle_df.copy()
    if "Team" not in df_all.columns:
        st.warning("Couldn't find a 'Team' column in the Oracle data.")
        return

    # Combine the 4 core teams (VIP + Tier 2)
    CORE_TEAMS = ["VIP North", "VIP South", "Tier 2 North", "Tier 2 South"]
    df4 = df_all[df_all["Team"].astype(str).isin(CORE_TEAMS)].copy()
    if df4.empty:
        st.info("No Oracle rows found for VIP/Tier 2 yet.")
        return

    # Month selector (now with "All Months")
    df4["_dt"] = pd.to_datetime(df4["Date"], errors="coerce")

    months = (
        df4["_dt"]
        .dropna()
        .dt.to_period("M")
        .sort_values()
        .unique()
        .tolist()
    )

    if not months:
        st.info("No dated rows in the data.")
        return

    # Put "All Months" at the top, default to the most recent month
    options = ["All Months"] + months
    sel = st.selectbox(
        "Month",
        options=options,
        index=len(options) - 1,  # default = latest month
        format_func=lambda x: x if isinstance(x, str) else x.strftime("%b %Y"),
    )

    # Build the working frame for the rest of the page
    if isinstance(sel, str) and sel == "All Months":
        mdf = df4.copy()
    else:
        mdf = df4[df4["_dt"].dt.to_period("M") == sel].copy()

    # Optional: set a sensible sparkline granularity
    if "gran" not in locals():
        gran = "M" if (isinstance(sel, str) and sel == "All Months") else "D"

    # (Optional) A label you can reuse in headings/cards
    sel_label = "All Months" if isinstance(sel, str) else sel.strftime("%b %Y")


    # 👉 Keep all your existing KPI cards / sparkline code here.
    #    (Use `mdf` as the filtered month data.)
    #
    # Example (leave your real card code in place):
    # total_value = pd.to_numeric(mdf.get("Total Value"), errors="coerce").sum()
    # st.metric("Total Value (£)", f"£{total_value:,.0f}")
    #
    # If you had a Highlands & Islands section for Mark, paste it below the KPIs.


    # ---- Helpers ----
    def _mask_kw(kw: str):
        if "Activity Status" not in mdf.columns:
            return mdf.index == -1
        return mdf["Activity Status"].astype(str).str.contains(kw, case=False, na=False)

    visits = len(mdf)
    total_value = pd.to_numeric(mdf.get("Total Value"), errors="coerce").sum()

    completed_cnt = int(_mask_kw("Completed").sum())
    completed_pct = (completed_cnt / visits * 100.0) if visits else 0.0

    active_engs = mdf["Name"].nunique() if "Name" in mdf.columns else 0

    tw_col = (
        "Total Working Time"
        if "Total Working Time" in mdf.columns
        else ("Total working time" if "Total working time" in mdf.columns else None)
    )
    avg_work = avg_hhmm(mdf[tw_col]) if tw_col else "00:00"
    total_time = sum_hhmm(mdf[tw_col]) if tw_col else "00:00"

    cancelled_pct = (int(_mask_kw("Cancelled").sum()) / visits * 100.0) if visits else 0.0
    notdone_pct = (int(_mask_kw("Not Done").sum()) / visits * 100.0) if visits else 0.0
    pending_cnt = int(_mask_kw("Pending").sum())
    started_cnt = int(_mask_kw("Started").sum())

    lunch_col = "Total Time" if "Total Time" in mdf.columns else tw_col
    if "Visit Type" in mdf.columns and lunch_col:
        lunch_df = mdf[mdf["Visit Type"].astype(str).str.contains("Lunch", case=False, na=False)]
        lunch_avg = avg_hhmm(lunch_df[lunch_col]) if not lunch_df.empty else "00:00"
    else:
        lunch_avg = "00:00"

    value_per_visit = (total_value / visits) if visits else 0.0

    # Pass None when "All Months" is chosen so the helper knows not to compare Periods
    sel_period = None if isinstance(sel, str) else sel
    trends = compute_mom_trends(df4, mdf, sel_period)
    # ---- Build Exec context for the AI (so Rob/Mark questions make sense) ----
    def _safe_sum(df, col):
        return float(pd.to_numeric(df.get(col, pd.Series(dtype=float)), errors="coerce").fillna(0).sum()) if col in df.columns else 0.0

    def _avg_td_hhmm(df, col):
        if col not in df.columns: return "00:00"
        td = pd.to_timedelta(df[col].astype(str), errors="coerce").dropna()
        if td.empty: return "00:00"
        avg = td.mean()
        h = int(avg.total_seconds() // 3600)
        m = int((avg.total_seconds() % 3600) // 60)
        return f"{h:02d}:{m:02d}"

    # Core monthly stats
    visits = int(len(mdf))
    completed = int(mdf["Activity Status"].astype(str).str.contains("Completed", case=False, na=False).sum()) if "Activity Status" in mdf.columns else 0
    completion_rate = (completed / visits * 100.0) if visits else 0.0
    total_value = _safe_sum(mdf, "Total Value")
    avg_twt = _avg_td_hhmm(mdf, "Total Working Time")

    # Budgets – pull high-level VIP and Tier 2 figures if available
    vip_alloc, vip_used, vip_rem, vip_pct_used = compute_team_budget_metrics("Sky VIP")
    t2_alloc,  t2_used,  t2_rem,  t2_pct_used  = compute_team_budget_metrics("Tier 2")
    total_alloc = (vip_alloc or 0) + (t2_alloc or 0)
    total_used  = (vip_used  or 0) + (t2_used  or 0)
    total_rem   = max(total_alloc - total_used, 0)

    # Optional: overtime (if you already compute it on the page, reuse the value; else leave None)
    try:
        title, ot_text, _, _ = overtime_from_df(mdf, period="Daily")  # your existing helper on this page
        # parse "£123,456" to number for the AI
        _ot_num = float(str(ot_text).replace("£","").replace(",","")) if isinstance(ot_text, str) else None
    except Exception:
        _ot_num = None

    # Label for the month picker (e.g. "All Months" or "Apr 2025")
    if isinstance(sel, str):
        month_label = sel
    else:
        try:
            month_label = sel.strftime("%b %Y")
        except Exception:
            month_label = "Selected period"

    st.session_state["exec_ctx"] = {
        "month_label": month_label,
        "scope_note": "VIP North, VIP South, Tier 2 North, Tier 2 South combined",
        "kpis": {
            "visits": visits,
            "completed": completed,
            "completion_rate_pct": completion_rate,
            "total_value": total_value,
            "avg_total_working_time_hhmm": avg_twt,
            "overtime_total_value": _ot_num,
            "budget_alloc_vip": vip_alloc,
            "budget_used_vip": vip_used,
            "budget_used_pct_vip": vip_pct_used,
            "budget_alloc_t2": t2_alloc,
            "budget_used_t2": t2_used,
            "budget_used_pct_t2": t2_pct_used,
            "budget_total_alloc": total_alloc,
            "budget_total_used": total_used,
            "budget_total_remaining": total_rem,
        }
    }
    with st.expander("💹 Overtime & Budgets", expanded=False):
        # ========= Overtime + Budgets =========
        
        st.subheader("💹 Overtime & Budgets")

        # Overtime (combined)
        title, ot_text, _, ot_ts = overtime_from_df(mdf, period="Daily")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            card("Overtime total (£)", ot_text, "Oracle", spark=ot_ts)

        # Budgets — uses budgets.csv + expenses.csv
        def _budget_cards(area_label: str, col):
            alloc, used, remaining, pct_used = compute_team_budget_metrics(area_label)
            used_txt = f"{pct_used:.1f}%" if pd.notna(pct_used) else "—"
            with col:
                card(
                    f"{area_label} budget used",
                    used_txt,
                    sub=f"£{used:,.0f} of £{alloc:,.0f} • budgets/expenses",
                    bar_pct=pct_used,
                )

        _budget_cards("Sky VIP", c2)   # VIP used %
        _budget_cards("Tier 2", c3)    # Tier 2 used %

        # Simple combined remaining (VIP+T2)
        vip_alloc, vip_used, _, _ = compute_team_budget_metrics("Sky VIP")
        t2_alloc,  t2_used,  _, _ = compute_team_budget_metrics("Tier 2")
        tot_alloc = (vip_alloc or 0) + (t2_alloc or 0)
        tot_used  = (vip_used  or 0) + (t2_used  or 0)
        tot_rem   = max(tot_alloc - tot_used, 0)
        pct_rem   = (tot_rem / tot_alloc * 100.0) if tot_alloc else 0.0
        with c4:
            card("Total remaining (VIP+T2)", f"£{tot_rem:,.0f}", sub=f"of £{tot_alloc:,.0f}", bar_pct=pct_rem)  

    with st.expander("👤VIP / Tier 2 — Executive Snapshot", expanded=False):
    # ---------- title bar
        st.subheader("👤VIP / Tier 2 — Executive Snapshot")
        # ========= KPI GRID (3 rows x 4 cards) =========
        # Row 1
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            card(
                "Total Value (£)",
                f"£{total_value:,.0f}",
                "Sum of 'Total Value' • Oracle",
                trend=trends["value_pct"],
                spark=_series_by_period(mdf, period="Daily", reducer="sum_num", col="Total Value"),
                color="#16a34a",
            )
        with c2:
            card(
                "Completed (%)",
                f"{completed_pct:.1f}%",
                f"{completed_cnt} of {visits} • Oracle",
                bar_pct=completed_pct,
                trend=trends["completed_rate_pp"],
                spark=_series_by_period(mdf[_mask_kw("Completed")], period="Daily", reducer="count"),
                color="#14b8a6",
            )
        with c3:
            card(
                "Visits",
                f"{visits:,}",
                "All visits • Oracle",
                trend=trends["visits_pct"],
                spark=_series_by_period(mdf, period="Daily", reducer="count"),
                color="#ef4444",
            )
        with c4:
            card(
                "Active engineers",
                f"{active_engs:,}",
                "Unique per period • Oracle",
                trend=trends["active_pct"],
                spark=_series_by_period(mdf, period="Daily", reducer="nunique", col="Name"),
                color="#22c55e",
            )

        # Row 2
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            card(
                "Avg Working Time",
                f"{avg_work}",
                "HH:MM • avg minutes",
                spark=_series_by_period(mdf, period="Daily", reducer="mean_time", col=tw_col or "Total Working Time"),
            )
        with c2:
            card(
                "Cancelled (%)",
                f"{cancelled_pct:.1f}%",
                f"{int(_mask_kw('Cancelled').sum())} cases • Oracle",
                trend=trends["cancel_rate_pp"],
                spark=_series_by_period(mdf[_mask_kw("Cancelled")], period="Daily", reducer="count"),
                color="#fb7185",
            )
        with c3:
            card(
                "Not Done (%)",
                f"{notdone_pct:.1f}%",
                f"{int(_mask_kw('Not Done').sum())} cases • Oracle",
                trend=trends["notdone_rate_pp"],
                spark=_series_by_period(mdf[_mask_kw("Not Done")], period="Daily", reducer="count"),
                color="#ef4444",
            )
        with c4:
            card(
                "Total time",
                f"{total_time}",
                "HH:MM total • hours per period",
                spark=_series_by_period(mdf, period="Daily", reducer="sum_time", col=tw_col or "Total Working Time"),
            )

        # Row 3
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            card(
                "Lunch avg",
                lunch_avg,
                "Avg of 'Lunch' visits",
                spark=_series_by_period(
                    mdf[mdf.get("Visit Type", "").astype(str).str.contains("Lunch", case=False, na=False)],
                    period="Daily",
                    reducer="mean_time",
                    col=lunch_col or "Total Time",
                ),
            )
        with c2:
            card("Pending (count)", f"{pending_cnt:,}", "", spark=_series_by_period(mdf[_mask_kw("Pending")], period="Daily", reducer="count"))
        with c3:
            card("Started (count)", f"{started_cnt:,}", "", spark=_series_by_period(mdf[_mask_kw("Started")], period="Daily", reducer="count"))
        with c4:
            card(
                "Value per visit",
                f"£{value_per_visit:,.0f}",
                "(avg per period)",
                trend=trends["vpv_pct"],
                spark=_value_per_visit_series(mdf),
                color="#8b5cf6",
            )
        # Default granularity for sparklines if not already set elsewhere
        if "gran" not in locals():
            gran = "M"   # "M" = monthly, "W" = weekly, "D" = daily


    # === Exec: Sky Business (Caffe Nero) snapshot ===
    with st.expander("🏢 Sky Business", expanded=False):
        import pandas as pd

        sb = load_sky_business()
        if sb.empty or "SBDate" not in sb.columns or "SLA" not in sb.columns:
            st.info("No Sky Business data (need 'SBDate' and 'SLA').")
        else:
            sb = sb.copy()
            sb["SBDate"] = pd.to_datetime(sb["SBDate"], errors="coerce")
            sb = sb.dropna(subset=["SBDate"])

            # --- Local Year / Month selectors (same behaviour as Sky Business page)
            years = sb["SBDate"].dt.year.sort_values().unique().tolist()
            MONTHS = ["All","Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]

            cY, cM = st.columns([1,1])
            with cY:
                sb_year  = st.selectbox("Year", ["All"] + [int(y) for y in years], index=0, key="exec_sb_year")
            with cM:
                sb_month = st.selectbox("Month", MONTHS, index=0, key="exec_sb_month")

            # --- Build current selection (cur) to match the SB page counts
            cur = sb.copy()
            if sb_year != "All":
                cur = cur[cur["SBDate"].dt.year == int(sb_year)]
            if sb_month != "All":
                mnum = MONTHS.index(sb_month)  # 1..12
                cur  = cur[cur["SBDate"].dt.month == mnum]

            # --- Build sparkline base: year-only filter (keeps history for MoM)
            spark_base = sb.copy()
            if sb_year != "All":
                spark_base = spark_base[spark_base["SBDate"].dt.year == int(sb_year)]
            spark_base["Month"] = spark_base["SBDate"].dt.to_period("M").dt.to_timestamp()

            # --- Cap month for the sparkline (end of year / selected month / latest)
            if sb_year != "All" and sb_month != "All":
                end_month = pd.Timestamp(int(sb_year), MONTHS.index(sb_month), 1).to_period("M").to_timestamp()
            elif sb_year != "All":
                end_month = pd.Timestamp(int(sb_year), 12, 1).to_period("M").to_timestamp()
            else:
                end_month = spark_base["Month"].max() if not spark_base.empty else None

            # --- SLA normaliser: lower, unify dashes, strip punctuation -> words only
            def normalise_sla(series: pd.Series) -> pd.Series:
                return (
                    series.astype(str)
                        .str.lower()
                        .str.replace(r"[\u2010-\u2015–—]", "-", regex=True)   # unicode dashes -> '-'
                        .str.replace(r"[^a-z0-9]+", " ", regex=True)         # remove punctuation & collapse
                        .str.replace(r"\s+", " ", regex=True)
                        .str.strip()
                )

            # Normalised SLA for cur and spark_base independently (so totals match cur)
            cur_sla   = normalise_sla(cur["SLA"])    if not cur.empty        else pd.Series([], dtype=str)
            base_sla  = normalise_sla(spark_base["SLA"]) if not spark_base.empty else pd.Series([], dtype=str)

            # Build masks for cur (totals) and spark_base (series)
            def masks_for(df_norm: pd.Series):
                if df_norm.empty:
                    # empty masks with correct index
                    return (
                        pd.Series(False, index=df_norm.index),
                        pd.Series(False, index=df_norm.index),
                        pd.Series(False, index=df_norm.index),
                        pd.Series(False, index=df_norm.index),
                        pd.Series(False, index=df_norm.index),
                    )
                m_all_nero  = df_norm.str.contains(r"\b(caffe|cafe|caffe)\s*nero\b", na=False)
                m_nero_2h   = df_norm.str.contains(r"\b(caffe|cafe|caffe)\s*nero\b.*\b(2\s*hour|2\s*hr)\b", na=False)
                m_nero_next = df_norm.str.contains(r"\b(caffe|cafe|caffe)\s*nero\b.*\bnext\s*day\b", na=False)
                m_nero_4h   = df_norm.str.contains(r"\b(caffe|cafe|caffe)\s*nero\b.*\b(4\s*hour|4\s*hr)\b", na=False)
                m_8h        = df_norm.str.contains(r"\b8\s*hour\s*sla\b", na=False)
                return m_all_nero, m_nero_2h, m_nero_next, m_nero_4h, m_8h

            cur_masks  = masks_for(cur_sla)
            base_masks = masks_for(base_sla)

            # Helper: spark + MoM for a mask on spark_base
            def spark_and_mom(mask_on_base: pd.Series):
                if spark_base.empty or mask_on_base.sum() == 0:
                    return [0], "—", "—"
                ser = (
                    spark_base.loc[mask_on_base]
                            .groupby("Month")
                            .size()
                            .reset_index(name="Count")
                            .sort_values("Month")
                )
                if end_month is not None:
                    ser = ser[ser["Month"] <= end_month]
                spark_vals = ser["Count"].tail(12).tolist()

                if len(ser) < 2:
                    return spark_vals, "—", "—"
                curr, prev = ser["Count"].iloc[-1], ser["Count"].iloc[-2]
                if prev == 0:
                    return spark_vals, "—", "—"
                pct = (curr - prev) / prev * 100.0
                return spark_vals, f"{pct:+.1f}%", ("▲" if pct >= 0 else "▼")

            # Totals come from cur to match the Sky Business page exactly
            totals = [int(mask.sum()) for mask in cur_masks]

            # Build cards
            tiles = [
                ("All Caffe Nero – Requests", totals[0], spark_and_mom(base_masks[0]), "#0ea5e9", "sb_nero_all"),
                ("Caffe Nero 2 hour",        totals[1], spark_and_mom(base_masks[1]), "#0ea5e9", "sb_nero_2h"),
                ("Caffe Nero (Next Day)",    totals[2], spark_and_mom(base_masks[2]), "#0ea5e9", "sb_nero_next"),
                ("Caffe Nero (4 Hour)",      totals[3], spark_and_mom(base_masks[3]), "#0ea5e9", "sb_nero_4h"),
                ("8-Hour SLA",               totals[4], spark_and_mom(base_masks[4]), "#10b981", "sb_8h"),
            ]

            c1, c2, c3, c4, c5 = st.columns(5, gap="large")
            for (title, total, (spark, mom_txt, mom_arrow), color, src), col in zip(tiles, [c1, c2, c3, c4, c5]):
                subtitle = f"{mom_arrow} {mom_txt} vs prev month" if mom_txt != "—" else "No prior month"
                try:
                    with col:
                        card(title, f"{total:,}", subtitle, spark=spark, color=color, source=src)
                except Exception:
                    with col.container(border=True):
                        st.caption("Month to date")
                        st.markdown(f"**{title}**")
                        st.markdown(f"<div style='font-size:26px;font-weight:700;'>{total:,}</div>", unsafe_allow_html=True)

            # Window label (for clarity)
            if spark_base.empty:
                st.caption("Window: —")
            else:
                start_lbl = spark_base["Month"].min().strftime("%b %Y")
                end_lbl   = (end_month or spark_base["Month"].max()).strftime("%b %Y")
                st.caption(f"Window: {start_lbl} – {end_lbl}")








    with st.expander("🏬 Sky Retail — Executive Snapshot", expanded=False):
        # ===== Sky Retail — Executive Snapshot (FIXED) =====
        st.subheader("🏬 Sky Retail — Executive Snapshot")
        # Pick a source (same one you used for the top KPIs)
        src = df4 if "df4" in locals() else df_all

        _dt = pd.to_datetime(src["Date"], errors="coerce")

        # IMPORTANT: handle "All Months"
        mdf = src.copy() if isinstance(sel, str) else src[_dt.dt.to_period("M") == sel].copy()

        # Build month-filtered dataframe from the full combined source so all columns exist
        

        # Normalise column names: collapse spaces, strip, and unify unicode
        import re, unicodedata
        def _norm(s: str) -> str:
            s = unicodedata.normalize("NFKC", str(s))
            s = re.sub(r"\s+", " ", s).strip()
            return s

        mdf.columns = [ _norm(c) for c in mdf.columns ]

        # Find the stakeholder column robustly (case-insensitive, ignores spacing)
        stake_col = None
        for c in mdf.columns:
            lc = c.lower()
            if ("sky" in lc) and ("retail" in lc) and ("stakeholder" in lc):
                stake_col = c
                break

        if stake_col is None:
            st.info("No 'Sky Retail Stakeholder' column in source data.")
        else:
            # continue with the rest of the Sky Retail snapshot logic (SR3, Completed, etc.)
        



            def _sr_bucket(x: str) -> str:
                v = str(x or "").strip().lower()
                if re.search(r"currys?|curry's", v): return "Currys"
                if "ee" in v: return "EE"
                if "sky" in v: return "Sky Retail"
                return ""  # treat everything else as Other/blank

            # Work on the month-filtered dataframe you already call mdf
            mdf = mdf.copy()
            stake_col = next((c for c in ["Sky Retail Stakeholder","Sky Retail Stakeholder Clean"] if c in mdf.columns), None)
            if stake_col is None:
                st.info("No 'Sky Retail Stakeholder' column in source data.")
            else:
                if stake_col == "Sky Retail Stakeholder":
                    mdf["SR_Bucket"] = mdf[stake_col].apply(_sr_bucket)
                else:
                    # already cleaned
                    mdf["SR_Bucket"] = mdf[stake_col].astype(str)

                # WE ONLY WANT these three categories in Retail counts
                SR3 = mdf[mdf["SR_Bucket"].isin(["Currys","Sky Retail","EE"])].copy()

                # Completed mask (for value and completion rate)
                SR3["__completed__"] = SR3.get("Activity Status","").astype(str).str.contains("completed", case=False, na=False)

                # Total Retail visits (sum of the 3 categories ONLY)
                total_sr_visits = len(SR3)

                # Completed % out of those visits
                completed_cnt = int(SR3["__completed__"].sum())
                completed_pct = (completed_cnt / total_sr_visits * 100.0) if total_sr_visits else 0.0

                # Retail value: ONLY Completed + one of the 3 categories
                retail_value = pd.to_numeric(SR3.loc[SR3["__completed__"], "Total Value"], errors="coerce").sum()

                # Visit splits
                counts = SR3.groupby("SR_Bucket").size()
                currys_v   = int(counts.get("Currys", 0))
                skyret_v   = int(counts.get("Sky Retail", 0))
                ee_v       = int(counts.get("EE", 0))

                # Average working time PER VISIT from Total Time (fallback to Total Working Time)
                if "Total Time" in SR3.columns:
                    tcol = "Total Time"
                elif "Total Working Time" in SR3.columns:
                    tcol = "Total Working Time"
                else:
                    tcol = None

                def _to_minutes(v):
                    try:
                        td = pd.to_timedelta(v)
                        if pd.isna(td): return 0
                        return int(td.total_seconds() // 60)
                    except Exception:
                        try:
                            s = str(v)
                            if ":" in s:
                                h, m = s.split(":")[:2]
                                return int(h) * 60 + int(m)
                        except Exception:
                            pass
                        return 0

                if tcol:
                    SR3["__mins__"] = SR3[tcol].apply(_to_minutes)
                    avg_all = SR3["__mins__"].mean() if len(SR3) else 0
                else:
                    avg_all = 0

                # Time-series (sparklines)
                vis_ts   = _series_by_period(SR3,                         period=gran, reducer="count")
                comp_ts  = _series_by_period(SR3[SR3["__completed__"]],   period=gran, reducer="count")
                value_ts = _series_by_period(SR3[SR3["__completed__"]],   period=gran, reducer="sum_num", col="Total Value")
                time_ts  = _series_by_period(SR3 if tcol else SR3.head(0), period=gran, reducer="mean_time", col=tcol) if tcol else [0]

                # Cancelled% / Not done% within Retail (the 3 buckets only)
                cancel_pct  = (SR3["Activity Status"].astype(str).str.contains("cancel",  case=False, na=False).mean() * 100.0) if len(SR3) else 0.0
                notdone_pct = (SR3["Activity Status"].astype(str).str.contains("not.?done",case=False, na=False).mean() * 100.0) if len(SR3) else 0.0

                # ===== 4 × 4 CARDS (Other removed) =====
                r1c1, r1c2, r1c3, r1c4 = st.columns(4)
                with r1c1:
                    card("Retail visits", f"{total_sr_visits:,}", "All retail visits • Oracle",
                        spark=vis_ts, color="#0ea5e9", source="Oracle")
                with r1c2:
                    card("Retail completed (%)", f"{completed_pct:.1f}%",
                        f"{completed_cnt:,} of {total_sr_visits:,}",
                        bar_pct=completed_pct, spark=comp_ts, color="#10b981", source="Oracle")
                with r1c3:
                    card("Retail value (£)", f"£{retail_value:,.0f}",
                        "Sum of 'Total Value' • Oracle (Completed only)",
                        spark=value_ts, color="#16a34a", source="Oracle")
                with r1c4:
                    # Avg Total Time per visit (HH:MM)
                    hh = int(avg_all // 60); mm = int(avg_all % 60)
                    card("Avg working time", f"{hh:02}:{mm:02}",
                        "HH:MM • avg minutes per visit", spark=time_ts, color="#60a5fa", source="Oracle")

                r2c1, r2c2, r2c3, _ = st.columns(4)
                with r2c1:
                    card("Currys visits", f"{currys_v:,}", "Month to date", spark=_series_by_period(SR3[SR3["SR_Bucket"]=="Currys"], period=gran, reducer="count"), color="#f59e0b", source="Oracle")
                with r2c2:
                    card("Sky Retail visits", f"{skyret_v:,}", "Month to date", spark=_series_by_period(SR3[SR3["SR_Bucket"]=="Sky Retail"], period=gran, reducer="count"), color="#3b82f6", source="Oracle")
                with r2c3:
                    card("EE visits", f"{ee_v:,}", "Month to date", spark=_series_by_period(SR3[SR3["SR_Bucket"]=="EE"], period=gran, reducer="count"), color="#22c55e", source="Oracle")
                # (r2c4 intentionally unused — “Other visits” removed)

                r3c1, r3c2, r3c3, r3c4 = st.columns(4)
                with r3c1:
                    # retail total time spark is still useful if you want it; keep as before if you had it
                    tot_time_ts = _series_by_period(SR3 if tcol else SR3.head(0), period=gran, reducer="sum_time", col="Total Time" if "Total Time" in SR3.columns else tcol) if tcol else [0]
                    # total time shown as HH:MM
                    total_mins = int(SR3["__mins__"].sum()) if tcol else 0
                    th, tm = (total_mins // 60, total_mins % 60)
                    card("Retail total time", f"{th:02}:{tm:02}", "HH:MM total", spark=tot_time_ts, color="#64748b", source="Oracle")
                with r3c2:
                    # Value per visit (Completed only) to match the fixed logic
                    vpv = (retail_value / completed_cnt) if completed_cnt else 0.0
                    card("Retail value/visit", f"£{vpv:,.0f}", "(avg per period)", spark=_series_by_period(SR3[SR3['__completed__']], period=gran, reducer='sum_num', col='Total Value'), color="#8b5cf6", source="Oracle")
                with r3c3:
                    card("Retail cancelled (%)", f"{cancel_pct:.1f}%", "Month to date", spark=_series_by_period(SR3[SR3['Activity Status'].astype(str).str.contains('cancel', case=False, na=False)], period=gran, reducer="count"), color="#ef4444", source="Oracle")
                with r3c4:
                    card("Retail not done (%)", f"{notdone_pct:.1f}%", "Month to date", spark=_series_by_period(SR3[SR3['Activity Status'].astype(str).str.contains('not.?done', case=False, na=False)], period=gran, reducer="count"), color="#ef4444", source="Oracle")
            
            # ----- Individual stakeholder metrics: % Cancelled + Avg Time (6 cards) -----
            def _pct_cancelled_for(sub_df: pd.DataFrame) -> float:
                if sub_df.empty:
                    return 0.0
                status = (sub_df.get("Activity Status","")
                        .astype(str).str.normalize("NFKC").str.strip().str.lower())
                is_completed = status.str.fullmatch(r"completed")
                is_cancelled = status.str.contains(r"\bcancel", regex=True)
                is_notdone   = status.str.contains(r"\bnot\s*done\b", regex=True)
                den = int((is_completed | is_cancelled | is_notdone).sum())
                if den == 0:
                    return 0.0
                return (int(is_cancelled.sum()) / den) * 100.0

            def _avg_time_hhmm_for(sub_df: pd.DataFrame) -> tuple[int,int]:
                if not tcol or sub_df.empty or tcol not in sub_df.columns:
                    return (0, 0)
                mins = sub_df[tcol].apply(_to_minutes)
                if len(mins) == 0 or mins.mean() != mins.mean():  # NaN guard
                    return (0, 0)
                avg_m = float(mins.mean())
                return int(avg_m // 60), int(avg_m % 60)

            def _spark_cancel(sub_df: pd.DataFrame):
                # cancelled count spark
                if sub_df.empty: return [0]
                status = sub_df.get("Activity Status","").astype(str)
                return _series_by_period(
                    sub_df[status.str.contains("cancel", case=False, na=False)],
                    period=gran, reducer="count"
                )

            def _spark_avg_time(sub_df: pd.DataFrame):
                if not tcol or sub_df.empty: return [0]
                return _series_by_period(sub_df, period=gran, reducer="mean_time", col=tcol)

            # Split the three stakeholders
            cur_df = SR3[SR3["SR_Bucket"] == "Currys"]
            sky_df = SR3[SR3["SR_Bucket"] == "Sky Retail"]
            ee_df  = SR3[SR3["SR_Bucket"] == "EE"]

            # Row A: % Cancelled (Currys / Sky Retail / EE)
            cA1, cA2, cA3 = st.columns(3)
            with cA1:
                pct = _pct_cancelled_for(cur_df)
                card("Currys – cancelled (%)", f"{pct:.1f}%", "Month to date",
                    spark=_spark_cancel(cur_df), color="#f59e0b", source="Oracle")
            with cA2:
                pct = _pct_cancelled_for(sky_df)
                card("Sky Retail – cancelled (%)", f"{pct:.1f}%", "Month to date",
                    spark=_spark_cancel(sky_df), color="#3b82f6", source="Oracle")
            with cA3:
                pct = _pct_cancelled_for(ee_df)
                card("EE – cancelled (%)", f"{pct:.1f}%", "Month to date",
                    spark=_spark_cancel(ee_df), color="#22c55e", source="Oracle")

            # Row B: Avg time per visit (Currys / Sky Retail / EE)
            cB1, cB2, cB3 = st.columns(3)
            with cB1:
                hh, mm = _avg_time_hhmm_for(cur_df)
                card("Currys – avg time", f"{hh:02}:{mm:02}", "HH:MM per visit",
                    spark=_spark_avg_time(cur_df), color="#f59e0b", source="Oracle")
            with cB2:
                hh, mm = _avg_time_hhmm_for(sky_df)
                card("Sky Retail – avg time", f"{hh:02}:{mm:02}", "HH:MM per visit",
                    spark=_spark_avg_time(sky_df), color="#3b82f6", source="Oracle")
            with cB3:
                hh, mm = _avg_time_hhmm_for(ee_df)
                card("EE – avg time", f"{hh:02}:{mm:02}", "HH:MM per visit",
                    spark=_spark_avg_time(ee_df), color="#22c55e", source="Oracle")
            # ----- end individual stakeholder metrics -----




        # ---------------------------------------------------------
    # Highlands & Islands — Executive KPIs (simple yearly cards)
    # ---------------------------------------------------------
    with st.expander("🏔️ Highlands & Islands — Executive KPIs", expanded=False):
        
        st.subheader("🏔️Highlands & Islands — Executive KPIs")

        @st.cache_data(show_spinner=False)
        def load_hi_workbook(file_path: str) -> dict[str, pd.DataFrame]:
            from pathlib import Path
            if not Path(file_path).exists():
                return {}
            # Try explicit sheets first; if any missing, fall back to “all sheets”
            wanted = [
                "2022", "2023", "2024", "2025",
                "Company 2022", "Company 2023", "Company 2024", "Company 2025",
            ]
            try:
                dfs = pd.read_excel(file_path, sheet_name=wanted)  # dict of DataFrames
            except Exception:
                xls = pd.ExcelFile(file_path)
                dfs = {s: xls.parse(s) for s in xls.sheet_names}

            # Normalise + numeric coercion
            for k, df in list(dfs.items()):
                df = df.copy()
                df.columns = df.columns.str.strip()
                df = safe_numeric(df, [
                    "Issued Visits", "Completed Visits", "Average Complete Job Time (Min)",
                    "7 Day Revisits", "7 Day Revisits %", "30 Day Revisits", "30 Day Revisits %",
                    "Surveys", "NPS%", "NPS"
                ])
                dfs[k] = df
            return dfs

        dfs = load_hi_workbook("Highlands Islands.xlsx")

        if not dfs:
            st.info("Highlands workbook not loaded. Place 'Highlands Islands.xlsx' next to the app to enable this section.")
        else:
            import re

            # Build the list of YEAR sheets directly from the loaded dict (no globals)
            year_sheets = [s for s in dfs.keys() if re.search(r"\b20\d{2}\b", s) and not s.lower().startswith("company")]
            if not year_sheets:
                st.info("No yearly sheets found in Highlands workbook.")
            else:
                # Extract numeric years and offer a picker
                years = sorted({int(re.search(r"(20\d{2})", s).group(1)) for s in year_sheets})
                sel_year = st.selectbox("H&I year", years, index=len(years)-1, key="exec_hi_year")
                sheet_name = next(s for s in year_sheets if str(sel_year) in s)

                hi = dfs.get(sheet_name, pd.DataFrame())
                if hi.empty:
                    st.info("Selected Highlands sheet is empty.")
                else:
                    # allow for multiple column-name variants
                    def pick(*names):
                        for n in names:
                            if n in hi.columns:
                                return n
                        return None

                    c_issued    = pick("Total Issued Visits", "Issued Visits", "Issued")
                    c_completed = pick("Total Completed Visits", "Completed Visits", "Completed")
                    c_rate      = pick("Completion Rate (%)", "Completion %", "Completed %")
                    c_notdone   = pick("Total Not Done Visits", "Not Done Visits", "Not Done")
                    c_7day      = pick("Total 7 Day Revisits", "7 Day Revisits")
                    c_7day_w    = pick("Weighted 7 Day Revisits %", "7 Day Revisits %")
                    c_nps       = pick("Weighted Average NPS", "NPS")

                    tot_issued    = int(hi[c_issued].sum()) if c_issued else 0
                    tot_completed = int(hi[c_completed].sum()) if c_completed else 0
                    comp_pct      = float(hi[c_rate].mean()) if c_rate else (100 * tot_completed / tot_issued if tot_issued else 0)
                    notdone_cnt   = int(hi[c_notdone].sum()) if c_notdone else 0
                    seven_cnt     = int(hi[c_7day].sum()) if c_7day else 0
                    seven_w       = float(hi[c_7day_w].mean()) if c_7day_w else None
                    nps_val       = float(hi[c_nps].mean()) if c_nps else None

                    a, b, c, d = st.columns(4)
                    with a: card("H&I — Issued", f"{tot_issued:,}", "Year total")
                    with b: card("H&I — Completed", f"{tot_completed:,}", "Year total")
                    with c: card("H&I — Completion (%)", f"{comp_pct:.2f}%", "Year average")
                    with d: card("H&I — Not Done", f"{notdone_cnt:,}", "Year total", color="#ef4444")

                    a, b, c, d = st.columns(4)
                    with a: card("H&I — 7-Day revisits", f"{seven_cnt:,}", "Year total")
                    with b: card("H&I — 7-Day revisits (weighted)", "N/A" if seven_w is None else f"{seven_w:.2f}%", "")
                    with c: card("H&I — Weighted NPS", "N/A" if nps_val is None else f"{nps_val:.1f}", "")
                    # (d left blank intentionally)

if st.session_state.screen == "operational_area":
    import streamlit as st
    import pandas as pd
    import plotly.express as px
    # --- Mark sees ONLY the Executive page, unless he opts to view Ops ---
    u = (st.session_state.get("username") or "").strip().lower()
    view_ops = bool(st.session_state.get("mark_view_ops", False))

    if u == "mark wilson" and not view_ops:
        # small right-aligned button row (doesn't shrink the page)
        _sp, right = st.columns([8, 2])
        with right:
            if st.button("➡ Go to Operations View", key="mark_to_ops", use_container_width=True):
                st.session_state["mark_view_ops"] = True   # allow Ops on next run
                st.session_state.screen = "operational_area"
                st.session_state.op_area_section = "menu"
                st.rerun()

        # render Exec page full width (Mark default)
        render_exec_overview(embed=True)
        st.stop()

    @st.cache_data
    def load_data():
        df_vip_south = pd.read_excel("VIP South Oracle Data.xlsx")
        df_vip_south["Team"] = "VIP South"
        df_vip_north = pd.read_excel("VIP North Oracle Data.xlsx")
        df_vip_north["Team"] = "VIP North"
        df_t2_south = pd.read_excel("Tier 2 South Oracle Data.xlsx")
        df_t2_south["Team"] = "Tier 2 South"
        df_t2_north = pd.read_excel("Tier 2 North Oracle Data.xlsx")
        df_t2_north["Team"] = "Tier 2 North"
        df_all = pd.concat([df_vip_south, df_vip_north, df_t2_south, df_t2_north], ignore_index=True)
        return df_all, df_vip_south, df_vip_north, df_t2_south, df_t2_north

    df_all, df_vip_south, df_vip_north, df_t2_south, df_t2_north = load_data()

    def fmt_td(x):
        if pd.isnull(x): return ""
        if isinstance(x, pd.Timedelta):
            s = str(x)
            if "days" in s:
                s = s.split(" ")[-1]
            return s.split(".")[0]
        return str(x)

    def fix_time_col(series):
        import datetime
        return series.apply(
            lambda x: x.strftime("%H:%M:%S") if isinstance(x, datetime.time) else x
        )

    


   


# —————— SHOW OPERATIONS TEAM LOGO AT TOP ——————
    ops_uri = to_data_uri("SkyOps.png")
    if ops_uri:
        st.markdown(
            f"""
            <div style="text-align:center; margin:10px 0;">
            <img src="{ops_uri}" alt="Sky Operations Team"
                style="max-width:600px; width:100%; height:auto;">
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.warning("⚠️ Could not find SkyOps.png in the app directory.")

    

    st.markdown("### Select an Operations View")

    # ——— First row: 4 buttons centered ———
    sp1, c1, c2, c3, c4, sp2 = st.columns([1, 1, 1, 1, 1, 1])
    with c1:
        if st.button("Engineer", key="op_btn_eng", use_container_width=True):
            st.session_state["op_area_section"] = "engineer"
            st.rerun()
    with c2:
        if st.button("Time analysis", key="op_btn_time", use_container_width=True):
            st.session_state["op_area_section"] = "time"
            st.rerun()
    with c3:
        if st.button("Visits", key="op_btn_visits", use_container_width=True):
            st.session_state["op_area_section"] = "visits"
            st.rerun()
    with c4:
        if st.button("Activity Status", key="op_btn_status", use_container_width=True):
            st.session_state["op_area_section"] = "activity_status"
            st.rerun()

    # ——— Second row: 5 buttons centered ———
    sp3, d1, d2, d3, d4, d5, d6, sp4 = st.columns([1, 1, 1, 1, 1, 1, 1, 1])

    with d1:
        if st.button("MEWP Hires", key="op_btn_mewp", use_container_width=True):
            st.session_state["op_area_section"] = "mewp_hires"
            st.rerun()

    with d2:
        if st.button("💷 Budget", key="op_btn_budget", use_container_width=True):
            st.session_state["op_area_section"] = "budget"
            st.rerun()

    with d3:
        if st.button("📄 Invoices", key="op_btn_invoices", use_container_width=True):
            st.session_state["op_area_section"] = "invoices"
            st.rerun()

    with d4:
        if st.button("🧭 Exec Overview", key="op_btn_exec_overview", use_container_width=True):
            st.session_state["op_area_section"] = "exec_overview"
            st.rerun()

    with d5:
        if st.button("🏢 Sky Business", key="op_btn_sky_business", use_container_width=True):
            st.session_state["op_area_section"] = "sky_business"
            st.rerun()
    with d6:
        if st.button("📁 Sky Orbit Uploader", key="op_btn_sky_orbit", use_container_width=True):
            st.session_state["op_area_section"] = "sky_orbit_upload"
            st.rerun()



    # ——— Back button row (centered) ———
    sp5, back_col, sp6 = st.columns([2, 1, 2])
    with back_col:
        if st.button("⬅ Back to Instructions", key="op_btn_back_instructions", use_container_width=True):
            st.session_state.screen = "instructions_guide"  # exact, valid name
            st.rerun()
    # --- Only for Mark: a way back to Leadership (Exec) ---
    if (st.session_state.get("username") or "").strip().lower() == "mark wilson":
        sp5, back_col, sp6 = st.columns([2, 1, 2])
        with back_col:
            if st.button("⬅ Back to Leadership", key="mark_back_to_exec", use_container_width=True):
                st.session_state["mark_view_ops"] = False   # <-- disable Ops view
                st.session_state.op_area_section = "exec_overview"
                st.rerun()

    # … now dispatch on st.session_state["op_area_section"] below …


    section = st.session_state.get("op_area_section", "engineer")
        # Exec Overview (Mark’s page)
    if section == "exec_overview":
        render_exec_overview(embed=True)
        render_orbit_ai("exec_overview")
        st.stop()
    # Invoices screen
    if section == "invoices":
        render_invoices_screen()
        st.stop()
    elif section == "sky_business":
        render_sky_business_screen()
        st.stop()
    elif section == "sky_orbit_upload":
        render_sky_orbit_file_upload()
        st.stop()

    elif st.session_state.get("screen") == "suggestions":
        render_suggestions_page()

    elif st.session_state.get("screen") == "engineer_kpi":
        render_engineer_kpi()


    # ---- ENGINEER SECTION ----
    from st_aggrid import AgGrid, GridOptionsBuilder

    def show_aggrid(df, height=300):
        gb = GridOptionsBuilder.from_dataframe(df)
        gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=10)
        gb.configure_default_column(editable=False, filter=True, sortable=True)
        grid_options = gb.build()
        AgGrid(df, gridOptions=grid_options, height=height, fit_columns_on_grid_load=True)
        
    if section == "engineer":
        st.title("Engineer Dashboard (All Oracle Data)")
        st.subheader("All Engineers & Visit Counts")
        engineer_counts = df_all["Name"].value_counts().reset_index()
        engineer_counts.columns = ["Engineer", "Visit Count"]
        st.dataframe(engineer_counts, use_container_width=True)

        with st.expander("📊 Bar Chart: Visits per Engineer"):
            fig = px.bar(engineer_counts, x="Engineer", y="Visit Count", title="Visits per Engineer")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🥧 Pie Chart: Visit Share (Top 10 Engineers)"):
            top10 = engineer_counts.head(10)
            fig = px.pie(top10, names="Engineer", values="Visit Count", title="Top 10 Engineers by Visit Share")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🌞 Sunburst: Engineer → Visit Type"):
            if "Visit Type" in df_all.columns and "Name" in df_all.columns:
                sunburst_df = df_all[["Name", "Visit Type"]].dropna()
                fig = px.sunburst(sunburst_df, path=["Name", "Visit Type"], title="Engineer → Visit Type Breakdown")
                st.plotly_chart(fig, use_container_width=True)

        st.subheader("👷‍♂️ Engineer Visit Summary Stats")
        stats = {
            "Average Visits": int(engineer_counts["Visit Count"].mean()),
            "Min Visits": int(engineer_counts["Visit Count"].min()),
            "Max Visits": int(engineer_counts["Visit Count"].max()),
            "Total Engineers": len(engineer_counts)
        }
        st.table(pd.DataFrame(stats, index=["Value"]).T)

        with st.expander("🏢 Visits per Engineer by Team (Stacked Bar)"):
            if "Team" in df_all.columns and "Name" in df_all.columns:
                grouped = df_all.groupby(["Name", "Team"]).size().unstack(fill_value=0)
                fig = px.bar(
                    grouped,
                    barmode="stack",
                    title="Visits per Engineer by Team"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🔎 Top 10 Engineers with Most Visit Type Diversity"):
            if "Name" in df_all.columns and "Visit Type" in df_all.columns:
                n_types = df_all.groupby("Name")["Visit Type"].nunique().sort_values(ascending=False).head(10)
                st.bar_chart(n_types)
                st.dataframe(n_types)

        with st.expander("🍰 Engineer Visit Status Breakdown (Pie)"):
            if "Name" in df_all.columns and "Activity Status" in df_all.columns:
                engineer_status = df_all.groupby("Name")["Activity Status"].value_counts().unstack(fill_value=0)
                top_eng = engineer_status.sum(axis=1).sort_values(ascending=False).head(5).index
                pie_data = df_all[df_all["Name"].isin(top_eng)]
                fig = px.pie(
                    pie_data, names="Activity Status", title="Status Breakdown (Top 5 Engineers)",
                    color="Activity Status"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📈 Engineer Visits Over Time"):
            if "Date" in df_all.columns and "Name" in df_all.columns:
                df_all["Date"] = pd.to_datetime(df_all["Date"], errors="coerce")
                by_month = df_all.groupby([pd.Grouper(key="Date", freq="M"), "Name"]).size().unstack(fill_value=0)
                top_engs = engineer_counts["Engineer"].head(5)
                fig = px.line(
                    by_month[top_engs], 
                    title="Monthly Visit Count - Top 5 Engineers",
                    labels={"value": "Visits", "Date": "Month"}
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🏷️ Visits per Engineer by Visit Type (Stacked Bar)"):
            if "Visit Type" in df_all.columns and "Name" in df_all.columns:
                grouped = df_all.groupby(["Name", "Visit Type"]).size().unstack(fill_value=0)
                top_engs = engineer_counts["Engineer"].head(5)
                fig = px.bar(
                    grouped.loc[top_engs],
                    barmode="stack",
                    title="Visits per Engineer by Visit Type (Top 5 Engineers)"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("⏳ Average & Total Time per Engineer"):
            if "Name" in df_all.columns and "Total Time" in df_all.columns:
                df_time = df_all.copy()
                df_time["Total Time"] = pd.to_timedelta(fix_time_col(df_time["Total Time"]), errors="coerce")
                df_time = df_time[~df_time["Total Time"].isna()]
                time_stats = df_time.groupby("Name")["Total Time"].agg(["count", "mean", "sum"])
                time_stats = time_stats.sort_values("count", ascending=False).head(10)
                time_stats["mean"] = time_stats["mean"].apply(fmt_td)
                time_stats["sum"] = time_stats["sum"].apply(fmt_td)
                st.dataframe(time_stats)
                st.bar_chart(time_stats["count"])

        with st.expander("📋 Engineer vs Visit Type Matrix"):
            if "Name" in df_all.columns and "Visit Type" in df_all.columns:
                matrix = pd.crosstab(df_all["Name"], df_all["Visit Type"])
                st.dataframe(matrix)

        with st.expander("👥 Engineer Visits Per Team (Table)"):
            if "Name" in df_all.columns and "Team" in df_all.columns:
                team_table = pd.crosstab(df_all["Name"], df_all["Team"])
                st.dataframe(team_table)
    import os
    import base64
    import streamlit as st
    import streamlit.components.v1 as components
    import os
    import fitz  # PyMuPDF
    import streamlit as st

    # … inside your operational_area block …

    if section == "mewp_hires":
        st.title("📄 MEWP Hire Receipts")

        folder = os.path.join(os.getcwd(), "Mewp Hires")
        pdf_files = sorted(
            f for f in os.listdir(folder)
            if f.lower().endswith(".pdf")
        )

        if not pdf_files:
            st.info("No MEWP hire receipts found.")
        else:
            for idx, pdf in enumerate(pdf_files):
                pdf_path = os.path.join(folder, pdf)

                # Create a closed expander per file
                with st.expander(pdf, expanded=False):
                    # Open PDF
                    doc = fitz.open(pdf_path)

                    # Render each page as an image
                    for page_number in range(len(doc)):
                        page = doc[page_number]
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img_data = pix.tobytes("png")
                        st.image(
                            img_data,
                            caption=f"Page {page_number+1}",
                            use_container_width=True
                        )

                    # Download button inside the expander
                    with open(pdf_path, "rb") as f:
                        raw = f.read()
                    st.download_button(
                        label="Download receipt",
                        data=raw,
                        file_name=pdf,
                        mime="application/pdf",
                        key=f"dl_mewp_{idx}"
                    )


    # ---- VISITS SECTION ----
    elif section == "visits":
        st.title("Visit Dashboard (All Oracle Data)")

        st.subheader("All Visits (Table View)")
        visit_cols = [c for c in ["Visit Type", "Activity Status", "Date", "Team", "Name"] if c in df_all.columns]
        visit_table = df_all[visit_cols].copy()
        st.dataframe(visit_table.head(200), use_container_width=True)

        with st.expander("📊 Bar Chart: Visit Count by Type"):
            vc_type = df_all["Visit Type"].value_counts().reset_index()
            vc_type.columns = ["Visit Type", "Count"]
            fig = px.bar(vc_type, x="Visit Type", y="Count", title="Visit Count by Type")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🟢 Visit Status Breakdown (Pie Chart)"):
            if "Activity Status" in df_all.columns:
                fig = px.pie(df_all, names="Activity Status", title="Visit Status Breakdown")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("👥 Visits by Team (Bar Chart)"):
            if "Team" in df_all.columns:
                vc_team = df_all["Team"].value_counts().reset_index()
                vc_team.columns = ["Team", "Visit Count"]
                fig = px.bar(vc_team, x="Team", y="Visit Count", title="Visits by Team")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📈 Visits Over Time (Monthly Line Chart)"):
            if "Date" in df_all.columns:
                df_all["Date"] = pd.to_datetime(df_all["Date"], errors="coerce")
                df_all["Month"] = df_all["Date"].dt.to_period("M").dt.to_timestamp()
                by_month = df_all.groupby("Month").size().reset_index(name="Visit Count")
                fig = px.line(by_month, x="Month", y="Visit Count", markers=True, title="Monthly Visit Trend")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🌞 Sunburst: Team → Visit Type"):
            if "Team" in df_all.columns and "Visit Type" in df_all.columns:
                sunburst_df = df_all[["Team", "Visit Type"]].dropna()
                fig = px.sunburst(sunburst_df, path=["Team", "Visit Type"], title="Visits Breakdown: Team → Visit Type")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🔗 Parallel Categories: Team, Status, Visit Type"):
            needed_cols = ["Team", "Activity Status", "Visit Type"]
            if all(c in df_all.columns for c in needed_cols):
                pc_df = df_all[needed_cols].dropna().astype(str)
                fig = px.parallel_categories(pc_df, dimensions=needed_cols, title="Team → Status → Visit Type")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📅 Visits by Day of Week"):
            if "Date" in df_all.columns:
                df_all["Day of Week"] = df_all["Date"].dt.day_name()
                vc_day = df_all["Day of Week"].value_counts().reindex(
                    ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"], fill_value=0
                )
                st.bar_chart(vc_day)

        with st.expander("🏠 Visits by Postcode (Top 10)"):
            if "Postcode" in df_all.columns:
                vc_postcode = df_all["Postcode"].value_counts().head(10)
                st.bar_chart(vc_postcode)

        with st.expander("👷 Visits per Engineer (Top 10)"):
            if "Name" in df_all.columns:
                vc_engineer = df_all["Name"].value_counts().head(10)
                st.bar_chart(vc_engineer)

        with st.expander("🔥 Visits by Month & Team (Heatmap)"):
            if "Date" in df_all.columns and "Team" in df_all.columns:
                df_all["Month"] = df_all["Date"].dt.to_period("M").dt.strftime('%b %Y')
                pivot = pd.pivot_table(df_all, index="Month", columns="Team", values="Visit Type", aggfunc="count", fill_value=0)
                st.dataframe(pivot)

        with st.expander("📋 Visits by Status & Type"):
            if "Activity Status" in df_all.columns and "Visit Type" in df_all.columns:
                pivot = pd.pivot_table(df_all, index="Activity Status", columns="Visit Type", values="Date", aggfunc="count", fill_value=0)
                st.dataframe(pivot)

    # (You can add a 'time analysis' section next, just like above.)
    elif section == "time":
        import datetime

        def fmt_td(x):
            if pd.isnull(x):
                return ""
            total_seconds = int(x.total_seconds())
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            seconds = total_seconds % 60
            return f"{hours:02}:{minutes:02}:{seconds:02}"

        def clean_times(df, time_cols):
            df = df.copy()
            for col in time_cols:
                if col in df.columns:
                    # Convert any datetime.time to string
                    df[col] = df[col].apply(
                        lambda x: x.strftime("%H:%M:%S") if isinstance(x, datetime.time) else x
                    )
                    # Convert to timedelta
                    df[col] = pd.to_timedelta(df[col], errors="coerce")
                    # Remove zeros, blanks, NaT
                    df = df[~df[col].isna()]
                    df = df[df[col] != pd.Timedelta(0)]
            return df

        # All possible time columns
        time_cols = ["Total Time", "Activate", "Deactivate", "Travel Time", "Total Time (Inc Travel)", "Total Working Time"]


        # 1. Summary Table: Mean/Min/Max by Time Column
        st.subheader("⏳ Summary Table: Time Columns Stats")
        summary_rows = []
        for col in time_cols:
            if col in df_all.columns:
                cleaned = clean_times(df_all[[col]], [col])
                if not cleaned.empty:
                    summary_rows.append({
                        "Column": col,
                        "Mean": cleaned[col].mean(),
                        "Min": cleaned[col].min(),
                        "Max": cleaned[col].max(),
                        "Count": cleaned[col].count()
                    })
        if summary_rows:
            stats_df = pd.DataFrame(summary_rows)
            stats_df["Mean"] = stats_df["Mean"].apply(fmt_td)
            stats_df["Min"] = stats_df["Min"].apply(fmt_td)
            stats_df["Max"] = stats_df["Max"].apply(fmt_td)
            st.dataframe(stats_df)
        else:
            st.info("No time columns found.")

        # 2. Average Total Time per Visit Type
        if "Total Time" in df_all.columns and "Visit Type" in df_all.columns:
            with st.expander("⏱️ Average Total Time per Visit Type"):
                cleaned = clean_times(df_all[["Total Time", "Visit Type"]], ["Total Time"])
                avg_time = cleaned.groupby("Visit Type")["Total Time"].mean().sort_values()
                avg_time = avg_time.apply(fmt_td)
                st.bar_chart(avg_time)
                st.dataframe(avg_time)

        # 3. Activate & Deactivate Time Analysis
        for col in ["Activate", "Deactivate"]:
            if col in df_all.columns:
                with st.expander(f"⚡ {col} Time Analysis"):
                    cleaned = clean_times(df_all[[col, "Visit Type"]], [col])
                    if not cleaned.empty:
                        avg = cleaned.groupby("Visit Type")[col].mean().sort_values()
                        avg = avg.apply(fmt_td)
                        st.bar_chart(avg)
                        st.dataframe(avg)
                        st.markdown(f"**Min:** {fmt_td(cleaned[col].min())} &nbsp;&nbsp; **Max:** {fmt_td(cleaned[col].max())}")
                    else:
                        st.info(f"No data for {col} column.")


        # 4. Total Working Time Analysis
        if "Total Working Time" in df_all.columns:
            with st.expander("🛠️ Total Working Time Analysis"):
                cleaned = clean_times(df_all[["Total Working Time", "Visit Type"]], ["Total Working Time"])
                avg = cleaned.groupby("Visit Type")["Total Working Time"].mean().sort_values()
                avg = avg.apply(fmt_td)
                st.bar_chart(avg)
                st.dataframe(avg)
                st.markdown(f"**Min:** {fmt_td(cleaned['Total Working Time'].min())} &nbsp;&nbsp; **Max:** {fmt_td(cleaned['Total Working Time'].max())}")

        # 5. Boxplot of Total Time per Team
        if "Total Time" in df_all.columns and "Team" in df_all.columns:
            with st.expander("📦 Boxplot: Total Time by Team"):
                cleaned = clean_times(df_all[["Total Time", "Team"]], ["Total Time"])
                import plotly.express as px
                fig = px.box(cleaned, x="Team", y="Total Time", title="Total Time Distribution by Team")
                st.plotly_chart(fig, use_container_width=True)

        # 6. Timeline: Mean Total Time Per Month
        if "Total Time" in df_all.columns and "Date" in df_all.columns:
            with st.expander("📈 Avg Total Time Per Month (Line Chart)"):
                cleaned = clean_times(df_all[["Total Time", "Date"]], ["Total Time"])
                cleaned["Month"] = pd.to_datetime(cleaned["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
                monthly = cleaned.groupby("Month")["Total Time"].mean().dropna()
                monthly_fmt = monthly.apply(fmt_td)
                st.line_chart(monthly_fmt)


        # 7. Table: All Time Columns (first 100 rows, cleaned)
        with st.expander("📋 All Time Columns (Sample)"):
            cols_present = [col for col in time_cols if col in df_all.columns]
            sample = clean_times(df_all[cols_present], cols_present)
            # Format columns
            for c in cols_present:
                sample[c] = sample[c].apply(fmt_td)
            st.dataframe(sample.head(100), use_container_width=True)

        # 8. 🔥 Heatmap: Average Total Time by Team & Visit Type
        if all(col in df_all.columns for col in ["Total Time", "Team", "Visit Type"]):
            with st.expander("🔥 Heatmap: Avg Total Time by Team & Visit Type"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Visit Type"]], ["Total Time"])
                pivot = cleaned.pivot_table(index="Team", columns="Visit Type", values="Total Time", aggfunc="mean")
                # Format for display (HH:MM:SS)
                pivot_fmt = pivot.applymap(fmt_td)
                st.dataframe(pivot_fmt)

        # 9. ⏲️ Distribution of Total Time (Histogram)
        if "Total Time" in df_all.columns:
            with st.expander("⏲️ Distribution of Total Time (Histogram)"):
                cleaned = clean_times(df_all[["Total Time"]], ["Total Time"])
                if not cleaned.empty:
                    import matplotlib.pyplot as plt
                    import numpy as np
                    times = cleaned["Total Time"].dt.total_seconds() / 60  # Minutes
                    fig, ax = plt.subplots()
                    ax.hist(times, bins=30, color='skyblue', edgecolor='black')
                    ax.set_title("Distribution of Total Time (Minutes)")
                    ax.set_xlabel("Minutes")
                    ax.set_ylabel("Frequency")
                    st.pyplot(fig)



        # 10. 🕒 Median Time by Visit Type and Team
        if all(col in df_all.columns for col in ["Total Time", "Team", "Visit Type"]):
            with st.expander("🕒 Median Total Time by Visit Type and Team (Table)"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Visit Type"]], ["Total Time"])
                pivot = cleaned.pivot_table(index="Team", columns="Visit Type", values="Total Time", aggfunc="median")
                pivot_fmt = pivot.applymap(fmt_td)
                st.dataframe(pivot_fmt)

        # 11. 📊 Pie Chart: Proportion of Visits with >1 Hour Total Time
        if "Total Time" in df_all.columns:
            with st.expander("📊 Visits > 1 Hour vs <= 1 Hour (Pie Chart)"):
                cleaned = clean_times(df_all[["Total Time"]], ["Total Time"])
                gt1h = (cleaned["Total Time"] > pd.Timedelta(hours=1)).sum()
                le1h = (cleaned["Total Time"] <= pd.Timedelta(hours=1)).sum()
                pie_df = pd.DataFrame({
                    "Category": ["> 1 Hour", "<= 1 Hour"],
                    "Count": [gt1h, le1h]
                })
                fig = px.pie(pie_df, names="Category", values="Count", title="Proportion of Visits > 1 Hour Total Time")
                st.plotly_chart(fig, use_container_width=True)

        # 12. 🏆 Longest & Shortest Total Times (Per Team)
        if "Total Time" in df_all.columns and "Team" in df_all.columns:
            with st.expander("🏆 Longest & Shortest Total Times per Team"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Name", "Visit Type"]], ["Total Time"])
                idxmax = cleaned.groupby("Team")["Total Time"].idxmax()
                idxmin = cleaned.groupby("Team")["Total Time"].idxmin()
                longest = cleaned.loc[idxmax]
                shortest = cleaned.loc[idxmin]
                longest = longest[["Team", "Name", "Visit Type", "Total Time"]]
                shortest = shortest[["Team", "Name", "Visit Type", "Total Time"]]
                longest["Total Time"] = longest["Total Time"].apply(fmt_td)
                shortest["Total Time"] = shortest["Total Time"].apply(fmt_td)
                st.markdown("#### Longest Total Time per Team")
                st.dataframe(longest)
                st.markdown("#### Shortest Total Time per Team")
                st.dataframe(shortest)

        # 13. ⏱️ Total Working Time > 10:25 (Detailed Summary)
    if section == "time":    
        if "Total Working Time" in df_all.columns:
            with st.expander("⏱️ Total Working Time Over 10:25 Summary", expanded=False):

                def convert_mixed_time(val):
                    import datetime
                    if pd.isnull(val): return 0
                    if isinstance(val, datetime.timedelta): return val.total_seconds() / 60
                    if isinstance(val, datetime.time): return val.hour * 60 + val.minute + val.second / 60
                    try: return pd.to_timedelta(val).total_seconds() / 60
                    except:
                        try:
                            if isinstance(val, str) and ":" in val:
                                return pd.to_timedelta("0 days " + val).total_seconds() / 60
                            if isinstance(val, (int, float)):
                                return float(val) * 24 * 60
                        except: return 0
                    return 0

                def mins_to_hhmm(m):
                    h = int(m // 60)
                    mins = int(m % 60)
                    return f"{h}:{mins:02}"

                valid_df = df_all[df_all["Total Working Time"].notna()].copy()
                valid_df["Total Working Time (min)"] = valid_df["Total Working Time"].apply(convert_mixed_time)
                valid_df["Over Minutes"] = valid_df["Total Working Time (min)"] - 625
                valid_df["Over Minutes"] = valid_df["Over Minutes"].apply(lambda x: x if x > 0 else 0)

                total_minutes = valid_df["Total Working Time (min)"].sum()
                total_over_minutes = valid_df["Over Minutes"].sum()
                total_over_cost = (valid_df["Over Minutes"] / 15).apply(lambda x: round(x)).sum() * 5.50

                overall_summary = {
                    "Total Working Time": mins_to_hhmm(total_minutes),
                    "Total Time Over 10:25": mins_to_hhmm(total_over_minutes),
                    "Total Over Time Cost (£)": f"£{total_over_cost:,.2f}"
                }

                st.markdown("#### 🔢 Overall Summary")
                st.dataframe(pd.DataFrame([overall_summary]), use_container_width=True)

                st.markdown("#### 🧑‍🤝‍🧑 Breakdown by Team")
                team_summary = valid_df.groupby("Team").agg(
                    Total_Working_Minutes=("Total Working Time (min)", "sum"),
                    Over_Minutes=("Over Minutes", "sum"),
                    Over_Cost=("Over Minutes", lambda x: (x / 15).round().sum() * 5.50)
                ).reset_index()
                team_summary["Total Working Time"] = team_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
                team_summary["Time Over 10:25"] = team_summary["Over_Minutes"].apply(mins_to_hhmm)
                team_summary["Over Cost (£)"] = team_summary["Over_Cost"].map("£{:,.2f}".format)
                st.dataframe(team_summary[["Team", "Total Working Time", "Time Over 10:25", "Over Cost (£)"]], use_container_width=True)

                st.markdown("#### 👷 Breakdown by Engineer")
                engineer_summary = valid_df.groupby("Name").agg(
                    Total_Working_Minutes=("Total Working Time (min)", "sum"),
                    Over_Minutes=("Over Minutes", "sum"),
                    Over_Cost=("Over Minutes", lambda x: (x / 15).round().sum() * 5.50)
                ).reset_index()
                engineer_summary["Total Working Time"] = engineer_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
                engineer_summary["Time Over 10:25"] = engineer_summary["Over_Minutes"].apply(mins_to_hhmm)
                engineer_summary["Over Cost (£)"] = engineer_summary["Over_Cost"].map("£{:,.2f}".format)
                st.dataframe(engineer_summary[["Name", "Total Working Time", "Time Over 10:25", "Over Cost (£)"]], use_container_width=True)

    

        # 15. 🗓️ Time Over 10:25: Monthly & Quarterly Breakdown
        with st.expander("📅 Time Over 10:25: Monthly & Quarterly Breakdown"):

            df_time = valid_df.copy()
            df_time["Month"] = pd.to_datetime(df_time["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
            df_time["Quarter"] = pd.to_datetime(df_time["Date"], errors="coerce").dt.to_period("Q").astype(str)

            monthly = df_time.groupby("Month").agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            monthly["% Time Over 10:25"] = (monthly["Over_Minutes"] / monthly["Total_Minutes"] * 100).round(2)
            monthly["Month"] = pd.to_datetime(monthly["Month"]).dt.strftime("%B %Y")
            monthly["Total Working Time (hh:mm)"] = monthly["Total_Minutes"].apply(mins_to_hhmm)
            monthly["Time Over 10:25 (hh:mm)"] = monthly["Over_Minutes"].apply(mins_to_hhmm)
            st.markdown("## 🗓️ Monthly Breakdown")
            st.dataframe(monthly[["Month", "Total Working Time (hh:mm)", "Time Over 10:25 (hh:mm)", "% Time Over 10:25"]], use_container_width=True)

            quarterly = df_time.groupby("Quarter").agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            quarterly["% Time Over 10:25"] = (quarterly["Over_Minutes"] / quarterly["Total_Minutes"] * 100).round(2)
            quarterly["Total Working Time (hh:mm)"] = quarterly["Total_Minutes"].apply(mins_to_hhmm)
            quarterly["Time Over 10:25 (hh:mm)"] = quarterly["Over_Minutes"].apply(mins_to_hhmm)
            st.markdown("## 🗓️ Quarterly Breakdown")
            st.dataframe(quarterly[["Quarter", "Total Working Time (hh:mm)", "Time Over 10:25 (hh:mm)", "% Time Over 10:25"]], use_container_width=True)

            st.markdown("### 📈 % Time Over 10:25 by Team")
            team_monthly = df_time.groupby(["Team", "Month"]).agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            team_monthly["% Time Over 10:25"] = (team_monthly["Over_Minutes"] / team_monthly["Total_Minutes"] * 100).round(2)
            team_monthly["Month"] = pd.to_datetime(team_monthly["Month"]).dt.strftime("%b %Y")
            fig_team = px.line(team_monthly, x="Month", y="% Time Over 10:25", color="Team", markers=True)
            fig_team.update_layout(xaxis_tickangle=45)
            st.plotly_chart(fig_team, use_container_width=True)

            st.markdown("### 📊 % Time Over 10:25 by Engineer")
            eng_monthly = df_time.groupby(["Name", "Month"]).agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            eng_monthly["% Time Over 10:25"] = (eng_monthly["Over_Minutes"] / eng_monthly["Total_Minutes"] * 100).round(2)
            eng_monthly["Month"] = pd.to_datetime(eng_monthly["Month"]).dt.strftime("%b %Y")
            fig_eng = px.bar(eng_monthly, x="Month", y="% Time Over 10:25", color="Name", barmode="group")
            fig_eng.update_layout(xaxis_tickangle=45)
            st.plotly_chart(fig_eng, use_container_width=True)


        # 16. 📊 Deep Dive: % Time Over 10:25 by Month and Week
        with st.expander("📊 Time Over 10:25: Monthly & Weekly Insight", expanded=False):

            def mins_to_hhmm(m):
                h = int(m // 60)
                mins = int(m % 60)
                return f"{h}:{mins:02}"

            df_over = df_all[df_all["Total Working Time"].notna()].copy()
            df_over["Total Working Time (min)"] = df_over["Total Working Time"].apply(convert_mixed_time)
            df_over["Over Minutes"] = df_over["Total Working Time (min)"] - 625
            df_over["Over Minutes"] = df_over["Over Minutes"].apply(lambda x: x if x > 0 else 0)

            df_over["Month"] = pd.to_datetime(df_over["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
            df_over["week"] = df_over["week"].astype(str)

            # Monthly Summary
            monthly_summary = df_over.groupby("Month").agg(
                Total_Working_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            monthly_summary["% Time Over 10:25"] = (monthly_summary["Over_Minutes"] / monthly_summary["Total_Working_Minutes"] * 100).round(2)
            monthly_summary["Total Working Time"] = monthly_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
            monthly_summary["Time Over 10:25"] = monthly_summary["Over_Minutes"].apply(mins_to_hhmm)

            st.markdown("### 📅 Monthly Breakdown")
            st.dataframe(monthly_summary[["Month", "Total Working Time", "Time Over 10:25", "% Time Over 10:25"]], use_container_width=True)

            # Max/Min Monthly Delta
            max_row = monthly_summary.loc[monthly_summary["% Time Over 10:25"].idxmax()]
            min_row = monthly_summary.loc[monthly_summary["% Time Over 10:25"].idxmin()]
            delta = round(max_row["% Time Over 10:25"] - min_row["% Time Over 10:25"], 2)
            delta_table = pd.DataFrame([
                {"Metric": "Max Month", "Month": max_row["Month"].strftime("%B %Y"), "% Over": max_row["% Time Over 10:25"]},
                {"Metric": "Min Month", "Month": min_row["Month"].strftime("%B %Y"), "% Over": min_row["% Time Over 10:25"]},
                {"Metric": "Delta", "Month": f"{max_row['Month'].strftime('%B')} vs {min_row['Month'].strftime('%B')}", "% Over": delta}
            ])

            st.markdown("### 🔺 Max vs Min Monthly Comparison")
            st.dataframe(delta_table, use_container_width=True)

            # Weekly Summary
            weekly_summary = df_over.groupby("week").agg(
                Total_Working_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            weekly_summary["% Time Over 10:25"] = (weekly_summary["Over_Minutes"] / weekly_summary["Total_Working_Minutes"] * 100).round(2)
            weekly_summary["Total Working Time"] = weekly_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
            weekly_summary["Time Over 10:25"] = weekly_summary["Over_Minutes"].apply(mins_to_hhmm)

            st.markdown("### 📆 Weekly Breakdown")
            st.dataframe(weekly_summary[["week", "Total Working Time", "Time Over 10:25", "% Time Over 10:25"]], use_container_width=True)

        with st.expander("📊 Combined Charts: Time Over 10:25 (Weekly & Monthly)", expanded=False):
            from plotly.subplots import make_subplots
            import plotly.graph_objects as go

            # Convert hh:mm to minutes for line chart
            def hhmm_to_minutes(hhmm_str):
                if isinstance(hhmm_str, str) and ":" in hhmm_str:
                    h, m = hhmm_str.split(":")
                    return int(h) * 60 + int(m)
                return 0

            # Sort week as int for correct order
            weekly_summary["Week_Num"] = weekly_summary["week"].astype(int)
            weekly_summary = weekly_summary.sort_values("Week_Num")
            weekly_summary["Time Over Minutes"] = weekly_summary["Time Over 10:25"].apply(hhmm_to_minutes)

            fig_week = make_subplots(specs=[[{"secondary_y": True}]])

            fig_week.add_trace(
                go.Bar(x=weekly_summary["Week_Num"], y=weekly_summary["% Time Over 10:25"], name="% Time Over 10:25"),
                secondary_y=False
            )

            fig_week.add_trace(
                go.Scatter(x=weekly_summary["Week_Num"], y=weekly_summary["Time Over Minutes"], name="Time Over 10:25 (min)", mode="lines+markers"),
                secondary_y=True
            )

            fig_week.update_layout(
                title="📊 Weekly Time Over 10:25",
                xaxis_title="Week",
                yaxis_title="% Time Over",
                legend_title="Metrics",
                bargap=0.3,
                xaxis=dict(type="category")
            )

            fig_week.update_yaxes(title_text="% Time Over 10:25", secondary_y=False)
            fig_week.update_yaxes(title_text="Time Over (Minutes)", secondary_y=True)

            st.plotly_chart(fig_week, use_container_width=True)

            # Monthly version
            monthly_summary["Month_Str"] = monthly_summary["Month"].dt.strftime("%b %Y")
            monthly_summary = monthly_summary.sort_values("Month")
            monthly_summary["Time Over Minutes"] = monthly_summary["Time Over 10:25"].apply(hhmm_to_minutes)

            fig_month = make_subplots(specs=[[{"secondary_y": True}]])

            fig_month.add_trace(
                go.Bar(x=monthly_summary["Month_Str"], y=monthly_summary["% Time Over 10:25"], name="% Time Over 10:25"),
                secondary_y=False
            )

            fig_month.add_trace(
                go.Scatter(x=monthly_summary["Month_Str"], y=monthly_summary["Time Over Minutes"], name="Time Over 10:25 (min)", mode="lines+markers"),
                secondary_y=True
            )

            fig_month.update_layout(
                title="📅 Monthly Time Over 10:25",
                xaxis_title="Month",
                yaxis_title="% Time Over",
                legend_title="Metrics",
                bargap=0.3,
                xaxis_tickangle=45
            )

            fig_month.update_yaxes(title_text="% Time Over 10:25", secondary_y=False)
            fig_month.update_yaxes(title_text="Time Over (Minutes)", secondary_y=True)

            st.plotly_chart(fig_month, use_container_width=True)


        # 17. 🧠 Last Visit Type When Total Over 10:25
        with st.expander("🧠 Last Visit Type When Total Working Time is Over 10:25", expanded=False):

                st.markdown("This shows the final visit types on days where total working time exceeds 10:25 (625 mins).")

                def to_minutes(t):
                        try:
                                if isinstance(t, str) and ":" in t:
                                        h, m, s = map(int, t.split(":"))
                                        return h * 60 + m + s / 60
                                elif hasattr(t, 'hour'):
                                        return t.hour * 60 + t.minute + t.second / 60
                                return float(t) * 24 * 60  # Excel float
                        except:
                                return 0

                df_check = df_all.copy()
                df_check["Date"] = pd.to_datetime(df_check["Date"], errors="coerce")
                df_check = df_check[df_check["Activity Status"].str.lower() == "completed"]
                df_check = df_check[df_check["Total Working Time"].notna()]
                df_check = df_check[~df_check["Total Working Time"].isin(["00:00", "0:00", "0", 0])]

                df_check["Total Working Time (min)"] = df_check["Total Working Time"].apply(to_minutes)

                # Step 1: Total time per engineer per day
                total_day = df_check.groupby(["Name", "Date"])["Total Working Time (min)"].sum().reset_index()
                total_day = total_day[total_day["Total Working Time (min)"] > 625]
                total_day = total_day.rename(columns={"Total Working Time (min)": "Total Day Working Time (min)"})

                # Step 2: Merge valid days
                df_valid = pd.merge(df_check, total_day, on=["Name", "Date"], how="inner")

                # Step 3: Create minutes from End time for sorting
                df_valid["End_minutes"] = df_valid["End"].apply(to_minutes)

                # Step 4: Sort by End_minutes to get the last visit
                df_valid.sort_values(by=["Name", "Date", "End_minutes"], inplace=True)
                last_visits = df_valid.groupby(["Name", "Date"]).tail(1)

                display_cols = ["Name", "Date", "Visit Type", "Start", "End", "Total Working Time", "Total Day Working Time (min)"]

                if not last_visits.empty and all(col in last_visits.columns for col in display_cols):
                        st.dataframe(last_visits[display_cols], use_container_width=True)
                else:
                        st.warning("⚠️ Required columns for displaying the table are missing or no data matched the criteria.")
        # 18. 📦 Breakdown of Last Visit Types by Month and Week (Total > 10:25)
        with st.expander("📦 Breakdown of Last Visit Types by Month and Week (Total Working Time > 10:25)", expanded=False):

                st.markdown("This shows a breakdown of the last visit types by **month** and **week number** for days where total working time exceeded 10:25 (625 mins).")

                if not last_visits.empty:

                        last_visits["Date"] = pd.to_datetime(last_visits["Date"], errors="coerce")
                        last_visits["Month"] = last_visits["Date"].dt.strftime("%B")
                        last_visits["Week"] = last_visits["Date"].dt.isocalendar().week

                        # Define proper calendar order
                        month_order = ["January", "February", "March", "April", "May", "June", 
                                    "July", "August", "September", "October", "November", "December"]

                        # Monthly breakdown
                        monthly = last_visits.groupby(["Month", "Visit Type"]).size().reset_index(name="Count")
                        monthly["Month"] = pd.Categorical(monthly["Month"], categories=month_order, ordered=True)
                        monthly = monthly.sort_values(by=["Month", "Count"], ascending=[True, False])
                        months = monthly["Month"].dropna().unique().tolist()

                        st.markdown("### 📅 Monthly Breakdown (Select a Month Tab)")

                        month_tabs = st.tabs(months)
                        for i, month in enumerate(months):
                                with month_tabs[i]:
                                        st.dataframe(monthly[monthly["Month"] == month].reset_index(drop=True), use_container_width=True)

                        # Weekly breakdown
                        st.markdown("### 📆 Weekly Breakdown")
                        weekly = last_visits.groupby(["Week", "Visit Type"]).size().reset_index(name="Count")
                        weekly = weekly.sort_values(by=["Week", "Count"], ascending=[True, False])
                        st.dataframe(weekly, use_container_width=True)

                else:
                        st.warning("No data found for last visits over 10:25. Please check upstream filters.")

    # --- NUMBER 7 ----------------------------------------------------------
    # --- SECTION: ACTIVITY STATUS BREAKDOWN --------------------------------
    from prophet import Prophet
    import plotly.express as px
    import plotly.graph_objects as go

    if section == "activity_status":
        st.markdown("## 📊 Activity Status Breakdown")

        # Prep all 4 datasets as (name, df) pairs
        team_datasets = {
            "VIP North": df_vip_north,
            "VIP South": df_vip_south,
            "Tier 2 North": df_t2_north,
            "Tier 2 South": df_t2_south,
        }
        with st.expander("🌍 All Teams Summary", expanded=True):
            # Combine all team dataframes into one for All Teams analysis
            df = pd.concat([df_vip_north, df_vip_south, df_t2_north, df_t2_south], ignore_index=True)

            # Clean data: remove rows with missing or empty Activity Status or Date
            df = df[df["Activity Status"].notna()]
            df = df[df["Activity Status"].str.strip() != ""]
            df = df[df["Date"].notna()]

            # Normalize Activity Status (lowercase + strip)
            df["Activity Status"] = df["Activity Status"].str.lower().str.strip()

            # Convert Date to datetime and add Month period column for sorting
            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
            df = df[df["Date"].notna()]  # Remove rows with invalid dates if any
            df["Month"] = df["Date"].dt.to_period("M")
            df["Month_str"] = df["Month"].astype(str)

            # 1. Overall counts per Activity Status (for bar chart and KPIs)
            overall_counts = df["Activity Status"].value_counts().sort_index()
            overall_counts_df = overall_counts.reset_index()
            overall_counts_df.columns = ["Activity Status", "Count"]  # Fix column names

            st.markdown("### Overall Visit Counts by Activity Status")
            fig_counts = px.bar(
                overall_counts_df,
                x="Activity Status",
                y="Count",
                title="Overall Visit Counts by Activity Status"
            )
            st.plotly_chart(fig_counts, use_container_width=True)

            # 2. Overall completion rate = Completed / (Completed + Cancelled + Not Done)
            completed = overall_counts.get("completed", 0)
            cancelled = overall_counts.get("cancelled", 0)
            not_done = overall_counts.get("not done", 0)
            denom = completed + cancelled + not_done
            completion_rate = (completed / denom * 100) if denom > 0 else 0
            st.markdown(f"**Overall Completion Rate:** {completion_rate:.2f}% (Completed / (Completed + Cancelled + Not Done))")

            # 3. Monthly counts pivot table (Activity Status x Month)
            monthly_pivot = df.pivot_table(
                index="Month",
                columns="Activity Status",
                values="Date",  # Count of non-null Date as proxy for count
                aggfunc="count",
                fill_value=0
            ).sort_index()

            # 4. Monthly Pending % = (pending visits / total visits per month) * 100
            monthly_totals = monthly_pivot.sum(axis=1)
            monthly_pending = monthly_pivot.get("pending", pd.Series(0, index=monthly_pivot.index))
            monthly_pending_pct = (monthly_pending / monthly_totals * 100).fillna(0)

            # 5. Line chart with all Activity Status monthly counts
            monthly_pivot_str = monthly_pivot.copy()
            monthly_pivot_str.index = monthly_pivot_str.index.astype(str)
            fig_monthly = px.line(
                monthly_pivot_str,
                x=monthly_pivot_str.index,
                y=monthly_pivot_str.columns,
                markers=True,
                title="Monthly Activity Status Counts"
            )
            st.plotly_chart(fig_monthly, use_container_width=True)

            # 6. Monthly Pending % line chart
            fig_pending = px.line(
                x=monthly_pending_pct.index.astype(str),
                y=monthly_pending_pct.values,
                labels={"x": "Month", "y": "Pending %"},
                title="Monthly Pending Percentage"
            )
            st.plotly_chart(fig_pending, use_container_width=True)

            # 7. Display Monthly Pending % table sorted by month ascending
            pending_df = monthly_pending_pct.reset_index()
            pending_df.columns = ["Month", "Pending %"]
            pending_df = pending_df.sort_values("Month")
            st.dataframe(pending_df.style.format({"Pending %": "{:.2f}"}), use_container_width=True)


        with st.expander("🌍 All Teams Summary", expanded=True):  # <-- Collapsible expander added here

                # Combine all 4 datasets for the "All Teams" summary
                df_all_teams = pd.concat([df_vip_north, df_vip_south, df_t2_north, df_t2_south], ignore_index=True)

                # 1. Overall totals by Activity Status (all time)
                overall_counts = df_all_teams["Activity Status"].value_counts().sort_index()
                st.markdown("#### Overall Visit Counts by Activity Status")
                st.bar_chart(overall_counts)

                # 2. Monthly Aggregated Summary
                if "Month" in df_all_teams.columns:
                    monthly_counts = df_all_teams.groupby(["Month", "Activity Status"]).size().unstack(fill_value=0)
                    monthly_counts["Total Visits"] = monthly_counts.sum(axis=1)
                    monthly_pct_change = monthly_counts.pct_change().fillna(0) * 100
                    monthly_pct_change = monthly_pct_change.add_suffix(" % Change")
                    max_visits = monthly_counts.max()
                    max_visits.name = "Max Visits"
                    min_visits = monthly_counts.min()
                    min_visits.name = "Min Visits"

                    monthly_summary = monthly_counts.join(monthly_pct_change).join(max_visits).join(min_visits)
                    monthly_summary = monthly_summary.sort_index()

                    st.markdown("#### Monthly Summary KPIs by Activity Status")
                    st.dataframe(monthly_summary.style.format({
                        col: "{:.2f}" for col in monthly_summary.columns if "% Change" in col
                    }))

                    # Monthly totals bar chart for quick visual
                    st.bar_chart(monthly_counts["Total Visits"])

                else:
                    st.warning("Missing 'Month' column for All Teams summary.")

                # 3. Weekly Aggregated Summary
                if "week" in df_all_teams.columns:
                    weekly_counts = df_all_teams.groupby(["week", "Activity Status"]).size().unstack(fill_value=0)
                    weekly_counts["Total Visits"] = weekly_counts.sum(axis=1)
                    weekly_pct_change = weekly_counts.pct_change().fillna(0) * 100
                    weekly_pct_change = weekly_pct_change.add_suffix(" % Change")
                    max_visits = weekly_counts.max()
                    max_visits.name = "Max Visits"
                    min_visits = weekly_counts.min()
                    min_visits.name = "Min Visits"

                    weekly_summary = weekly_counts.join(weekly_pct_change).join(max_visits).join(min_visits)
                    weekly_summary = weekly_summary.sort_index()

                    st.markdown("#### Weekly Summary KPIs by Activity Status")
                    st.dataframe(weekly_summary.style.format({
                        col: "{:.2f}" for col in weekly_summary.columns if "% Change" in col
                    }))

                    # Weekly totals line chart for quick visual
                    st.line_chart(weekly_counts["Total Visits"])

                else:
                    st.warning("Missing 'week' column for All Teams summary.")

        
    # --- Then continue with your existing individual team breakdowns here ---

        for team_name, df in team_datasets.items():
            # Main expander for each team
            with st.expander(f"🔵 {team_name} — Activity Overview", expanded=False):

                # Check for Activity Status column
                if "Activity Status" not in df.columns:
                    st.warning(f"No 'Activity Status' in {team_name}")
                    continue

                # Use tabs inside the expander instead of nested expanders
                tabs = st.tabs([
                    "📅 Monthly Breakdown",
                    "📆 Weekly Breakdown",
                    "📈 Forecast Next 6 Months",
                    "🌞 Sunburst View",
                    "📅 Gantt Chart",
                    "🔎 Drilldown"
                ])

                # Monthly Breakdown Tab
                with tabs[0]:
                    if "Month" in df.columns:
                        # Count visits by month and activity status
                        monthly_counts = df.groupby(["Month", "Activity Status"]).size().unstack(fill_value=0)

                        st.bar_chart(monthly_counts)

                        # Total visits per month (sum across all statuses)
                        monthly_counts["Total Visits"] = monthly_counts.sum(axis=1)

                        # Month-on-Month % Change per status plus total
                        monthly_pct_change = monthly_counts.pct_change().fillna(0) * 100
                        monthly_pct_change = monthly_pct_change.add_suffix(" % Change")

                        # Max visits by status across all months
                        max_visits = monthly_counts.max()
                        max_visits.name = "Max Visits"

                        # Min visits by status across all months
                        min_visits = monthly_counts.min()
                        min_visits.name = "Min Visits"

                        # Combine all KPI info into one table
                        summary_df = monthly_counts.join(monthly_pct_change)
                        summary_df = summary_df.join(max_visits)
                        summary_df = summary_df.join(min_visits)

                        # Sort by Month if possible (you may want to convert to datetime or categorical first)
                        summary_df = summary_df.sort_index()

                        st.markdown("### Monthly Summary KPIs by Activity Status")
                        st.dataframe(summary_df.style.format({
                            col: "{:.2f}" for col in summary_df.columns if "% Change" in col
                        }))
                    else:
                        st.warning("Missing 'Month' column")

                # Weekly Breakdown Tab
                with tabs[1]:
                    if "week" in df.columns:
                        # Count visits by week and activity status
                        weekly_counts = df.groupby(["week", "Activity Status"]).size().unstack(fill_value=0)

                        st.line_chart(weekly_counts)

                        # Total visits per week (sum across all statuses)
                        weekly_counts["Total Visits"] = weekly_counts.sum(axis=1)

                        # Week-on-Week % Change per status plus total
                        weekly_pct_change = weekly_counts.pct_change().fillna(0) * 100
                        weekly_pct_change = weekly_pct_change.add_suffix(" % Change")

                        # Max visits by status across all weeks
                        max_visits = weekly_counts.max()
                        max_visits.name = "Max Visits"

                        # Min visits by status across all weeks
                        min_visits = weekly_counts.min()
                        min_visits.name = "Min Visits"

                        # Combine all KPI info into one table
                        summary_week_df = weekly_counts.join(weekly_pct_change)
                        summary_week_df = summary_week_df.join(max_visits)
                        summary_week_df = summary_week_df.join(min_visits)

                        summary_week_df = summary_week_df.sort_index()

                        st.markdown("### Weekly Summary KPIs by Activity Status")
                        st.dataframe(summary_week_df.style.format({
                            col: "{:.2f}" for col in summary_week_df.columns if "% Change" in col
                        }))
                    else:
                        st.warning("Missing 'week' column")



                # Forecast Tab
                with tabs[2]:
                    try:
                        forecast_data = df[df["Activity Status"].notna()]
                        forecast_df = forecast_data.groupby("Date").size().reset_index(name="y")
                        forecast_df.columns = ["ds", "y"]

                        m = Prophet()
                        m.fit(forecast_df)
                        future = m.make_future_dataframe(periods=180)
                        forecast = m.predict(future)

                        fig = px.line(forecast, x="ds", y="yhat", title="6-Month Forecast")
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.warning(f"Forecasting failed: {e}")

                # Sunburst View Tab
                with tabs[3]:
                    try:
                        df['Month'] = df['Month'].astype(str)
                        fig = px.sunburst(df, path=["Month", "Activity Status"], title="Status Distribution by Month")
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.warning(f"Sunburst error: {e}")

                # Gantt Chart Tab
                with tabs[4]:
                    try:
                        if "Date" in df.columns and "Start" in df.columns and "End" in df.columns and "Activity Status" in df.columns:
                            df_gantt = df.copy()

                            # Combine Date with Start/End time into full datetime
                            df_gantt["Start"] = pd.to_datetime(df_gantt["Date"].astype(str) + " " + df_gantt["Start"].astype(str), errors="coerce")
                            df_gantt["End"] = pd.to_datetime(df_gantt["Date"].astype(str) + " " + df_gantt["End"].astype(str), errors="coerce")

                            # Remove rows with missing or invalid times or statuses
                            df_gantt.dropna(subset=["Start", "End", "Activity Status"], inplace=True)

                            fig = px.timeline(
                                df_gantt,
                                x_start="Start",
                                x_end="End",
                                y="Activity Status",
                                color="Activity Status",
                                title=f"Activity Gantt Timeline – {team_name}",
                                height=400
                            )
                            fig.update_yaxes(autorange="reversed")
                            st.plotly_chart(fig)
                    except Exception as e:
                        st.error(f"Gantt chart error: {e}")

                # Drilldown Tab
                with tabs[5]:
                    selected_status = st.selectbox(
                        f"Filter {team_name} by Activity Status",
                        options=sorted(df["Activity Status"].dropna().unique()),
                        key=f"drilldown_{team_name.replace(' ', '_')}"
                    )
                    drill = df[df["Activity Status"] == selected_status]
                    st.dataframe(drill, use_container_width=True)



        




    elif section == "budget":
        st.title("💷 Budget (Operations)")
        import os
        import pandas as pd
    # --- BOOTSTRAP: ensure master helpers exist BEFORE any use ---
    from pathlib import Path

    # Single source of truth for master path
    if 'EXP_MASTER' not in globals():
        EXP_MASTER = Path("Expenses/expenses_master.parquet")
        EXP_MASTER.parent.mkdir(parents=True, exist_ok=True)

    # Money parser used by load/save
    if '_sanitize_amount_series' not in globals():
        def _sanitize_amount_series(s: pd.Series) -> pd.Series:
            return pd.to_numeric(
                s.astype(str)
                .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)          # (123.45) -> -123.45
                .str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)")[0]  # one money-looking number
                .str.replace(",", "", regex=False),
                errors="coerce"
            ).fillna(0.0)

    # Safe loader (returns empty schema if file missing)
    if 'load_master' not in globals():
        def load_master() -> pd.DataFrame:
            if EXP_MASTER.exists():
                df = pd.read_parquet(EXP_MASTER)
            else:
                df = pd.DataFrame(columns=[
                    "Engineer Name","Transaction Date","Expense Type","Business Purpose",
                    "Vendor","City of Purchase","Payment Type","Amount","Source File","RowUID"
                ])
            if "Amount" in df.columns:
                df["Amount"] = _sanitize_amount_series(df["Amount"])
            if "Transaction Date" in df.columns:
                df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
            return df

    # Safe saver
    if 'save_master' not in globals():
        def save_master(df: pd.DataFrame) -> None:
            out = df.copy()
            if "Transaction Date" in out.columns:
                out["Transaction Date"] = pd.to_datetime(out["Transaction Date"], errors="coerce")
            if "Amount" in out.columns:
                out["Amount"] = _sanitize_amount_series(out["Amount"])
            out.to_parquet(EXP_MASTER, index=False)
    # --------------------------------------------------------------

        # 0) Initialize override flag
        if "override_alloc" not in st.session_state:
            st.session_state.override_alloc = False

        # —————— ADMIN AUTH ——————
        ADMIN_PW = "Dan"  # change this to something safe!
        if "is_admin" not in st.session_state:
            st.session_state.is_admin = False

        pw = st.text_input("🔒 Admin password to unlock editing:", type="password")
        if pw and pw == ADMIN_PW:
            st.session_state.is_admin = True
            st.success("🔓 Edit mode unlocked!")
        elif pw:
            st.error("❌ Wrong password")
        # ---------- Canonicalise names (single source of truth) ----------
        import re

        def _canon_eng_name(name: str) -> str:
            t = str(name or "").strip().replace("’", "'").casefold()
            t = re.sub(r"\b(cash|advance|utili[sz]ed|amount|cash\s+advance\s+utili[sz]ed\s+amount)\b", " ", t)
            t = re.sub(r"[^a-z' ]+", " ", t)
            t = re.sub(r"\s+", " ", t).strip()
            t = t.replace("o leary", "o'leary").replace("oleary", "o'leary")
            if "michael webb" in t:   return "Michael Webb"
            if "simon wakelin" in t:  return "Simon Wakelin"
            if "david o'leary" in t:  return "David O'Leary"
            parts = []
            for w in t.split():
                if "'" in w:
                    a, b = w.split("'", 1)
                    parts.append(a.capitalize() + "'" + b.capitalize())
                else:
                    parts.append(w.capitalize())
            return " ".join(parts)

        def _canon_set(names: set[str]) -> set[str]:
            return {_canon_eng_name(n) for n in names}

        # ---------- Team rosters (raw names) ----------
        VIP_NORTH = {
            "Christopher Drury","David O'Leary","Neil McPartlin","Christopher Law","Gordon Robertson",
            "Phillip Saunders","Christopher Millward","James Damm","Shaun Lucan","Daniel James",
            "Matthew Hebden","Wayne Hudson",
        }
        VIP_SOUTH = {
            "Dane Thomas","John Ware","Monika Marczynska","Simon Wakelin","Daniel Hakin",
            "Martin Dowdy","Paul Fearne","Garry Stother","Matthew Phillip","Richard Hall",
            "Gregory Sparkes-Smith","Michael Webb","Sean Harris",
        }
        T2_NORTH = {
            "Alan Burns","David Littlewood","Matthew Christie","Richard Wilson","Christopher Evans",
            "Kenneth Thomson","Matthew Foy","David Drew","Leslie Chrisp","Paul Bird","Wilhelm Tiffe",
            "David Forrester","Lorne Lucas","Philip Cliffe",
        }
        T2_SOUTH = {
            "Alex Doman","Mark Fitchett","Samuel Kempson","Steven Edmunds","Andrew Bailey",
            "Oliver Tottman","Sean Voyle","Steven Hutchins","Anna Hesketh","Richard Patterson",
            "Spencer Chalice","Thomas James","David Rowledge","Ryan Upton","Steven Davis","Tony Matharu",
        }

        # ---------- Canonicalised lookups ----------
        TEAM_SETS = {
            "VIP North":    _canon_set(VIP_NORTH),
            "VIP South":    _canon_set(VIP_SOUTH),
            "Tier 2 North": _canon_set(T2_NORTH),
            "Tier 2 South": _canon_set(T2_SOUTH),
            # add Sky Business / Sky Retail when you have rosters
        }

        def engineer_team(name: str) -> str | None:
            cn = _canon_eng_name(name)
            for team, pool in TEAM_SETS.items():
                if cn in pool:
                    return team
            return None
        # === MEWP-aware team inference ===
        def _infer_team_from_row(row) -> str | None:
            """
            1) Try normal engineer→team mapping (VIP/Tier2).
            2) If that fails, detect MEWP rows by filename/purpose and map to Sky Business/Retail.
            """
            # 1) normal mapping
            name = ""
            if isinstance(row, dict):
                name = row.get("Engineer Name", "")
            else:  # pandas Series
                name = row.get("Engineer Name", "")

            t = engineer_team(name)
            if t:
                return t

            # 2) MEWP fallback by file/purpose
            src   = row.get("Source File", "") if hasattr(row, "get") else row.get("Source File", "")
            bp    = row.get("Business Purpose", "") if hasattr(row, "get") else row.get("Business Purpose", "")
            blob  = f"{name} {src} {bp}".lower()

            if "mewp" in blob or "orion access" in blob:
                if "sky retail" in blob:
                    return "Sky Retail"
                if "sky business" in blob:
                    return "Sky Business"

            return None

        # -- Local constants for this screen --
        TOTAL_BUDGET = 280_000
        BUDGET_FILE  = "budgets.csv"
        EXPENSE_FILE = "expenses.csv"

        # -- Header & total budget --
        st.markdown(
            f"<h2 style='text-align:center; color:#fff; font-size:2.5rem;'>"
            f"📅 Total Budget (2025/26): £{TOTAL_BUDGET:,}</h2>",
            unsafe_allow_html=True,
        )
        st.markdown("---")



        # -- Ensure the full expense CSV exists / load it --
        if not os.path.exists(EXPENSE_FILE):
            pd.DataFrame(columns=["Name","Date","Area","Description","Amount"]).to_csv(EXPENSE_FILE, index=False)

        # read + enforce types (prevents silly sums later)
        full_exp = pd.read_csv(
            EXPENSE_FILE,
            parse_dates=["Date"],
            dtype={"Name": "string", "Area": "string", "Description": "string"}
        )
        # coerce Amount once, safely
        if "Amount" in full_exp.columns:
            full_exp["Amount"] = pd.to_numeric(full_exp["Amount"], errors="coerce").fillna(0.0)
        else:
            full_exp["Amount"] = 0.0

        # ensure Area exists (avoid groupby errors)
        if "Area" not in full_exp.columns:
            full_exp["Area"] = pd.Series(dtype="string")

        # -- Initialize the "current" display log if needed --
        if "current_exp" not in st.session_state:
            st.session_state.current_exp = full_exp.copy()

        # also keep a live budgets_df (in case it wasn't set earlier)
        if "budgets_df" not in st.session_state:
            if os.path.exists(BUDGET_FILE):
                _tmp_b = pd.read_csv(BUDGET_FILE)
                if "Stakeholder" not in _tmp_b.columns and "index" in _tmp_b.columns:
                    _tmp_b = _tmp_b.rename(columns={"index": "Stakeholder"})
                elif "Stakeholder" not in _tmp_b.columns:
                    _tmp_b = _tmp_b.rename(columns={_tmp_b.columns[0]: "Stakeholder"})
                _tmp_b["Allocated"] = pd.to_numeric(_tmp_b["Allocated"], errors="coerce").fillna(0)
                st.session_state.budgets_df = _tmp_b.set_index("Stakeholder")
            else:
                st.session_state.budgets_df = pd.DataFrame(columns=["Allocated"])

        budgets_df = st.session_state.budgets_df  # <- convenience alias

        # --- helper to canonicalise whole rosters BEFORE TEAM_SETS is built ---
        def _canon_set(names: set[str]) -> set[str]:
            # Convert every name in a set to the same canonical form your parser uses
            return {_canon_eng_name(n) for n in names}

        

        # 1) Compute & display total remaining budget, honoring override_alloc
        used = 0 if st.session_state.override_alloc else pd.to_numeric(budgets_df.get("Allocated", pd.Series()), errors="coerce").fillna(0).sum()
        exp_sum = pd.to_numeric(st.session_state.current_exp["Amount"], errors="coerce").fillna(0).sum()
        remaining = TOTAL_BUDGET - used - exp_sum

        st.markdown(
            f"<h2 style='text-align:center; color:#62d2a2; font-size:2.5rem;'>"
            f"💷 Budget Remaining: £{remaining:,.0f}</h2>",
            unsafe_allow_html=True,
        )
        st.markdown("---")
        def engineer_team(name: str) -> str | None:
                    cn = _canon_eng_name(name)
                    for team, pool in TEAM_SETS.items():
                        if cn in pool:
                            return team
                    return None
        # 2) Prepare Budget Summary table data (source of truth = Expenses master by team)

        # helpers ----------------------------------------------------
        def _safe_money(s: pd.Series) -> pd.Series:
            return pd.to_numeric(s, errors="coerce").fillna(0.0)

        def _clean_dates(s: pd.Series) -> pd.Series:
            return pd.to_datetime(s, errors="coerce")

        def _canon_set(names: set[str]) -> set[str]:
            # store canonicalised names so engineer_team() never misses
            return { _canon_eng_name(n) for n in names }



        # budgets & current CSV (used for fallback + Adjust tab)
        budgets_df = st.session_state.budgets_df
        cur = st.session_state.current_exp.copy()
        cur["Amount"] = _safe_money(cur.get("Amount", pd.Series(dtype=float)))
        cur["Area"] = cur.get("Area", pd.Series(dtype="string"))

        # try to load the Expenses master and sum by team
        try:
            master_df = load_master()
        except Exception as e:
            st.warning(f"Failed to load expenses master: {e}")
            master_df = pd.DataFrame()

        team_spend = None
        if not master_df.empty:
            m = master_df.copy()
            m["Amount"] = _safe_money(m["Amount"])
            m["Transaction Date"] = _clean_dates(m["Transaction Date"])
            m["Engineer Name"] = m["Engineer Name"].astype(str)

            # existing team mapping
            m["Team"] = m.apply(_infer_team_from_row, axis=1)


            # 👇 NEW: normalise Area and use it when Team is missing (e.g., MEWP rows)
            if "Area" in m.columns:
                m["Area"] = m["Area"].astype(str)
                m.loc[m["Area"].str.contains("sky business", case=False, na=False), "Area"] = "Sky Business"
                m.loc[m["Area"].str.contains("sky retail",   case=False, na=False), "Area"] = "Sky Retail"
                m["Team"] = m["Team"].fillna(m["Area"])

            team_spend = (
                m.dropna(subset=["Team"])
                .groupby("Team", dropna=False)["Amount"].sum()
            )



        # if no master or no teams, fall back to CSV grouped by Area
        if team_spend is None or team_spend.empty:
            team_spend = (
                cur.groupby("Area", dropna=False)["Amount"].sum()
            )

        # align to the stakeholders you actually show in the budget table
        stake_index = [
            "VIP North", "VIP South",
            "Tier 2 North", "Tier 2 South",
            "Sky Business", "Sky Retail",
        ]
        team_spend = team_spend.reindex(stake_index).fillna(0.0)

        # allocated (from file or working copy later)
        allocated_series = budgets_df.get("Allocated", pd.Series(dtype=float))
        allocated_series = pd.to_numeric(allocated_series, errors="coerce").fillna(0.0)
        allocated_series = allocated_series.reindex(stake_index).fillna(0.0)

        summary_df = pd.DataFrame({
            "Total Expense": team_spend,
            "Allocated":     allocated_series,
        })
        summary_df["Remaining"] = summary_df["Allocated"] - summary_df["Total Expense"]
        summary_df.index.name = "Stakeholder"

        pretty = summary_df.applymap(lambda x: f"£{x:,.2f}")


        # 3) Three tabs: Summary / Adjust / Expenses
        tab_sum, tab_adj, tab_exp = st.tabs([
            "🧾 Budget Summary",
            "🔧 Adjust Allocations",
            "💼 Expenses",
        ])


        # ---- TAB 1: Read-only summary ----
        with tab_sum:
            # ---- KPI cards ----
            # totals from the detailed table you show
            alloc_num  = float(pd.to_numeric(summary_df["Allocated"], errors="coerce").fillna(0).sum())
            spent_num  = float(pd.to_numeric(summary_df["Total Expense"], errors="coerce").fillna(0).sum())
            remaining  = max(alloc_num - spent_num, 0.0)
            remaining_pct = (remaining / alloc_num * 100.0) if alloc_num else 0.0

            c1, c2, c3 = st.columns(3)
            c1.metric("💷 Spent so far",       f"£{spent_num:,.0f}")
            c2.metric("📊 Remaining budget",  f"£{remaining:,.0f}")
            c3.metric("📈 % Remaining",       f"{remaining_pct:.1f}%")


        with st.expander("➕ Add Expense (manual)", expanded=False):
            import hashlib
            from datetime import date

            save_target = st.radio(
                "Save to",
                ["Master (recommended — shows in Detailed Table)", "CSV (legacy)"],
                index=0,
                horizontal=True,
            )

            colA, colB = st.columns(2)
            amount = colA.number_input("Amount (£)", min_value=0.0, step=10.0, format="%.2f")
            date_val = colB.date_input("Date", value=pd.Timestamp.today().date())

            if save_target.startswith("Master"):
                # Master path (needs engineer so we can map team)
                eng = st.text_input("Engineer Name (pick exact)", "")
                desc = st.text_input("Business Purpose / Notes", "")
                vendor = st.text_input("Vendor / Merchant (optional)", "")
                city = st.text_input("City (optional)", "")

                # prevent typos: offer a dropdown of known engineers (from rosters)
                known_engs = sorted(set().union(*TEAM_SETS.values()))
                pick = st.selectbox("or pick an engineer from rosters", ["(none)"] + known_engs, index=0)
                if pick != "(none)":
                    eng = pick

                if st.button("Add to Master"):
                    if not eng.strip():
                        st.warning("Engineer is required for Master.")
                    else:
                        # Build one master row
                        eng_canon = _canon_eng_name(eng)
                        # Create a stable uid for this record
                        uid = hashlib.sha1(
                            f"{eng_canon}|{date_val}|{amount:.2f}|{desc}|{vendor}".encode("utf-8")
                        ).hexdigest()

                        new_row = {
                            "Engineer Name": eng_canon,
                            "Transaction Date": pd.to_datetime(date_val),
                            "Expense Type": "",
                            "Business Purpose": desc or "",
                            "Vendor": vendor or "",
                            "City of Purchase": city or "",
                            "Payment Type": "Manual",
                            "Amount": float(amount or 0.0),
                            "Source File": "ManualEntry",
                            "RowUID": uid,
                        }

                        df = load_master()
                        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
                        save_master(df)
                        st.success("Added to Master. It will appear in the Detailed Budget Table.")
                        st.rerun()

            else:
                # Legacy CSV path (stays as-is)
                c1, c2 = st.columns(2)
                name = c1.text_input("Name / Ref (free text)")
                area = c2.selectbox("Stakeholder / Area", [
                    "VIP North","VIP South","Tier 2 North","Tier 2 South","Sky Business","Sky Retail"
                ], index=0)
                desc = st.text_input("Description", "")

                if st.button("Add to CSV"):
                    row = {
                        "Name": name or "",
                        "Date": pd.to_datetime(date_val),
                        "Area": area,
                        "Description": desc or "",
                        "Amount": float(amount or 0.0),
                    }
                    exp_df = pd.read_csv(
                        EXPENSE_FILE,
                        parse_dates=["Date"],
                        dtype={"Name":"string","Area":"string","Description":"string"}
                    )
                    exp_df = pd.concat([exp_df, pd.DataFrame([row])], ignore_index=True)
                    exp_df.to_csv(EXPENSE_FILE, index=False)
                    st.session_state.current_exp = exp_df.copy()
                    st.success("Added to CSV.")
                    st.rerun()



            # ---------- helpers ----------
            def _combine(df, labels, new_label):
                idx = [l for l in labels if l in df.index]
                if not idx:
                    return None
                row = df.loc[idx].sum(numeric_only=True)
                row.name = new_label
                return row

            def _fmt_money(x):
                return f"£{(0 if pd.isna(x) else x):,.2f}"

            def _kpi_block(alloc, used):
                rem = (alloc or 0) - (used or 0)
                pct_used = (used / alloc * 100) if alloc else float("nan")
                pct_rem  = 100 - pct_used if pd.notna(pct_used) else float("nan")
                k1, k2, k3, k4, k5 = st.columns(5)
                k1.metric("Budget Allocated", _fmt_money(alloc))
                k2.metric("Budget Used", _fmt_money(used))
                k3.metric("Budget Remaining", _fmt_money(rem))
                k4.metric("Budget % Used", f"{pct_used:.1f}%" if pd.notna(pct_used) else "—")
                k5.metric("Budget % Remaining", f"{pct_rem:.1f}%" if pd.notna(pct_rem) else "—")

            def _infer_team_from_row(row) -> str | None:
                """
                1) Try engineer → team (VIP/Tier2).
                2) If not found, detect MEWP rows by file/purpose and map to Sky Business/Retail.
                """
                # 1) normal engineer mapping
                t = engineer_team(row.get("Engineer Name", ""))
                if t:
                    return t

                # 2) MEWP detection by filename/purpose blob
                blob = f"{row.get('Engineer Name','')} {row.get('Source File','')} {row.get('Business Purpose','')}"
                low  = blob.lower()
                if "mewp" in low:
                    if "sky business" in low:
                        return "Sky Business"
                    if "sky retail" in low:
                        return "Sky Retail"
                # extra safety: typical vendor appears on these invoices
                if "orion access" in low and "retail" in low:
                    return "Sky Retail"
                if "orion access" in low and "business" in low:
                    return "Sky Business"

                return None

            # Build display copy and append combined totals
            display_df = summary_df.copy()

            vip_combo = _combine(display_df, ["VIP North", "VIP South", "Sky VIP"], "VIP Team")
            if vip_combo is not None:
                display_df.loc["VIP Team"] = vip_combo

            t2_combo = _combine(display_df, ["Tier 2 North", "Tier 2 South", "Tier 2"], "Tier 2 Team")
            if t2_combo is not None:
                display_df.loc["Tier 2 Team"] = t2_combo

            # ---------- Sub-tabs ----------
            tab_overview, tab_vip, tab_t2 = st.tabs(["📊 Overview", "👥 VIP Team", "🛠 Tier 2 Team"])

            # ===== Overview =====
            with tab_overview:
                st.markdown("#### 1️⃣ Budget Usage by Stakeholder")

                desired_order = [
                    "VIP North", "VIP South", "VIP Team",
                    "Tier 2 North", "Tier 2 South", "Tier 2 Team",
                    "Sky Business", "Sky Retail"
                ]
                # keep order where rows exist
                for stakeholder in [x for x in desired_order if x in display_df.index]:
                    row = display_df.loc[stakeholder]
                    used_amt  = float(row.get("Total Expense", 0) or 0)
                    alloc_amt = float(row.get("Allocated", 0) or 0)
                    pct = int(min((used_amt / alloc_amt * 100) if alloc_amt else 0, 100))
                    color = "#62d2a2" if pct < 70 else ("#f0ad4e" if pct < 90 else "#d9534f")

                    st.markdown(f"**{stakeholder}** — £{used_amt:,.2f} / £{alloc_amt:,.2f}")
                    st.markdown(
                        f"""
                        <progress value="{pct}" max="100"
                                style="width:100%; height:1rem; accent-color:{color};">
                        </progress>
                        """,
                        unsafe_allow_html=True,
                    )

                st.markdown("#### 2️⃣ Allocated vs Spent Chart")
                chart_df = display_df.reset_index()[["Stakeholder", "Allocated", "Total Expense"]]
                st.bar_chart(chart_df.set_index("Stakeholder"))

                st.markdown("---")
                pretty_display = display_df.applymap(lambda x: f"£{x:,.2f}")
                st.dataframe(pretty_display, use_container_width=True)

            # ===== VIP Team tab =====
            with tab_vip:
                if "VIP Team" in display_df.index:
                    row = display_df.loc["VIP Team"]
                    _kpi_block(float(row["Allocated"]), float(row["Total Expense"]))
                else:
                    st.info("No VIP Team data to combine (need 'VIP North' + 'VIP South' or 'Sky VIP').")

            # ===== Tier 2 Team tab =====
            with tab_t2:
                if "Tier 2 Team" in display_df.index:
                    row = display_df.loc["Tier 2 Team"]
                    _kpi_block(float(row["Allocated"]), float(row["Total Expense"]))
                else:
                    st.info("No Tier 2 Team data to combine (need 'Tier 2 North' + 'Tier 2 South' or 'Tier 2').")

        def _alloc_series_from_budgets_df(bud: pd.DataFrame, stakeholders: list[str]) -> pd.Series:
            if bud is None or bud.empty:
                return pd.Series([0.0]*len(stakeholders), index=stakeholders, dtype=float)
            team_col  = "Team" if "Team" in bud.columns else ("Stakeholder" if "Stakeholder" in bud.columns else None)
            alloc_col = "Allocated" if "Allocated" in bud.columns else ("QuarterlyBudget" if "QuarterlyBudget" in bud.columns else None)
            if not team_col or not alloc_col:
                return pd.Series([0.0]*len(stakeholders), index=stakeholders, dtype=float)
            return (bud.set_index(team_col)[alloc_col].reindex(stakeholders).fillna(0).astype(float))

        def get_allocations_from_file(stakeholders: list[str]) -> pd.Series:
            bud = st.session_state.get("budgets_df")
            if bud is None or not isinstance(bud, pd.DataFrame) or bud.empty:
                bud = load_budgets_df()              # your function that reads budgets.csv
                if not bud.empty:
                    st.session_state["budgets_df"] = bud
            return _alloc_series_from_budgets_df(bud, stakeholders)


        # === Budget (Operations) – compact, aligned ===
    st.markdown("## 💷 Budget (Operations)")

    # --- KPI bar (one row, two cards) ---
    if False:
        c1, c2 = st.columns([1,1])
        with c1:
            with st.container():  # card
                st.markdown("<div class='kpi'><h3>Total Budget (2025/26)</h3>"
                            f"<div class='val'>£{TOTAL_BUDGET:,.0f}</div></div>", unsafe_allow_html=True)
        with c2:
            # derive current headroom from saved budgets (not the working copy)
            def _alloc_from_budgets(bud):
                col_team = "Team" if "Team" in bud.columns else ("Stakeholder" if "Stakeholder" in bud.columns else None)
                col_alloc = "Allocated" if "Allocated" in bud.columns else ("QuarterlyBudget" if "QuarterlyBudget" in bud.columns else None)
                if not col_team or not col_alloc: return 0.0
                return float(pd.to_numeric(bud[col_alloc], errors="coerce").fillna(0).sum())
            saved_alloc_sum = _alloc_from_budgets(st.session_state.get("budgets_df", budgets_df.reset_index()))
            headroom_now = float(TOTAL_BUDGET - saved_alloc_sum)
            with st.container():
                st.markdown("<div class='kpi'><h3>Budget Remaining</h3>"
                            f"<div class='val'>£{headroom_now:,.0f}</div></div>", unsafe_allow_html=True)


    st.markdown("<hr/>", unsafe_allow_html=True)

    # === Tabs (Summary / Adjust / Expenses) – keep your existing tab names ===
    tab_sum, tab_adj, tab_exp = st.tabs(["Budget Summary", "Adjust Allocations", "Expenses"])

    # ---------------- Adjust Allocations (polished grid) ----------------
# ---------------- Adjust Allocations (polished grid) ----------------
    with tab_adj:
        STAKEHOLDERS = ["VIP North","VIP South","Tier 2 North","Tier 2 South","Sky Business","Sky Retail"]
        STEP = 1_000

        # Seed working copy once (dedupe by team)
        team_col = "Team" if "Team" in budgets_df.columns else (
            "Stakeholder" if "Stakeholder" in budgets_df.columns else budgets_df.columns[0]
        )
        bud = budgets_df.drop_duplicates(subset=[team_col]).set_index(team_col)

        if "alloc_working" not in st.session_state:
            if "Allocated" in bud.columns:
                base = bud["Allocated"].reindex(STAKEHOLDERS).fillna(0).astype(float)
            elif "QuarterlyBudget" in bud.columns:
                base = bud["QuarterlyBudget"].reindex(STAKEHOLDERS).fillna(0).astype(float)
            else:
                base = pd.Series(0.0, index=STAKEHOLDERS)
            st.session_state.alloc_working = base



        work = st.session_state.alloc_working.reindex(STAKEHOLDERS).fillna(0).astype(float)
        total_alloc_now = float(work.sum())
        headroom_live = float(TOTAL_BUDGET - total_alloc_now)

        kc1, kc2 = st.columns(2)
        kc1.markdown(f"**Total Allocated**<br><span class='val'>£{total_alloc_now:,.0f}</span>", unsafe_allow_html=True)
        kc2.markdown(f"**Headroom vs Total**<br><span class='val'>£{headroom_live:,.0f}</span>", unsafe_allow_html=True)

        with st.expander("🔧 Adjust Quarterly Allocations", expanded=True):
            st.caption("Use – / + to tweak by £1,000 (or type a number). Click **Save changes** to write to file.")
            for name in STAKEHOLDERS:
                col1, col2, col3 = st.columns([1.2, 1.2, 1.0])
                with col1:
                    st.markdown(f"<div class='rowpad'><strong>{name}</strong></div>", unsafe_allow_html=True)
                with col2:
                    b1, bmid, b3 = st.columns([0.25, 1.0, 0.25])
                    if b1.button("−", key=f"dec_{name}"):
                        st.session_state.alloc_working[name] = max(0.0, work[name] - STEP)
                        work = st.session_state.alloc_working
                    val = bmid.number_input("", key=f"in_{name}", value=float(work[name]),
                                            step=float(STEP), min_value=0.0, format="%.0f", label_visibility="collapsed")
                    if val != work[name]:
                        st.session_state.alloc_working[name] = float(val); work = st.session_state.alloc_working
                    if b3.button("+", key=f"inc_{name}"):
                        st.session_state.alloc_working[name] = work[name] + STEP; work = st.session_state.alloc_working
                with col3:
                    st.markdown(f"<div class='budget-pill'>£{st.session_state.alloc_working[name]:,.0f}</div>",
                                unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            s_left, s_right = st.columns([1,1])
            with s_left:
                if st.button("💾 Save changes", use_container_width=True):
                    out_df = (
                        st.session_state.alloc_working.rename("Allocated")
                        .reset_index().rename(columns={"index": "Stakeholder"})
                    )
                    out_df.to_csv(BUDGET_FILE, index=False)       # persist to budgets.csv
                    st.session_state.budgets_df = out_df          # in-memory source of truth
                    st.success("Allocations saved to budgets.csv")
                    st.experimental_rerun()




    # ---- TAB 1: Read-only summary ----
    with tab_sum:
        # existing KPIs...
        spent_num = pd.to_numeric(summary_df["Total Expense"], errors="coerce").fillna(0).sum()
        remaining_pct = (remaining / TOTAL_BUDGET * 100) if TOTAL_BUDGET else 0
        c1, c2, c3 = st.columns(3)
        c1.metric("💷 Spent so far", f"£{spent_num:,.0f}")
        c2.metric("📊 Remaining budget", f"£{remaining:,.0f}")
        c3.metric("📈 % Remaining", f"{remaining_pct:.1f}%")
            # 2.3) Detailed budget table (interactive)
        with st.expander("📋 Detailed Budget Table", expanded=True):
            # --- ALWAYS use the Expenses master mapped to teams (not the CSV) ---
            m = load_master()
            if m.empty:
                st.info("No rows in the Expenses master yet.")
            else:
                m = m.copy()
                m["Amount"] = _sanitize_amount_series(m["Amount"]).fillna(0.0)
                m["Engineer Name"] = m["Engineer Name"].astype(str)

                # 1) Existing engineer->team mapping
                m["Team"] = m.apply(_infer_team_from_row, axis=1)


                # 2) 👇 NEW: if Team is missing, use Area (normalised) so MEWP rows land in Sky Business/Retail
                if "Area" in m.columns:
                    m["Area"] = m["Area"].astype(str)
                    # normalise common variants
                    m.loc[m["Area"].str.contains("sky business", case=False, na=False), "Area"] = "Sky Business"
                    m.loc[m["Area"].str.contains("sky retail",   case=False, na=False), "Area"] = "Sky Retail"
                    # fill Team from Area where Team is NaN
                    m["Team"] = m["Team"].fillna(m["Area"])

                # 3) Sum by Team and align to the stakeholders you show
                team_spend_master = (
                    m.dropna(subset=["Team"])
                    .groupby("Team", dropna=False)["Amount"].sum()
                    .reindex(STAKEHOLDERS).fillna(0.0)
                )


                # Use working allocations if present (so the preview matches your Adjust tab)
                work = st.session_state.alloc_working.reindex(STAKEHOLDERS).fillna(0.0).astype(float)

                df = pd.DataFrame({
                    "Stakeholder": STAKEHOLDERS,
                    "Total Expense": team_spend_master.values,
                    "Allocated": work.values,
                })
                df["Remaining"] = df["Allocated"] - df["Total Expense"]
                # % used: cap to [0, 1] for the progress bar
                df["% Used"] = (df["Total Expense"] / df["Allocated"]).replace([pd.NA, pd.NaT, float("inf")], 0.0)
                df["% Used"] = df["% Used"].fillna(0.0).clip(lower=0.0, upper=1.0)

                # quick tools row
                t1, t2, t3 = st.columns([2,2,1])
                q = t1.text_input("Filter stakeholders", "", placeholder="e.g. VIP or Retail")
                sort_by = t2.selectbox("Sort by", ["Stakeholder", "Total Expense", "Allocated", "Remaining", "% Used"], index=0)
                descending = t3.checkbox("Desc", value=False)

                view = df.copy()
                if q:
                    view = view[view["Stakeholder"].str.contains(q, case=False, na=False)]
                view = view.sort_values(sort_by, ascending=not descending, kind="mergesort").reset_index(drop=True)

                # pretty render
                fmt0 = lambda x: f"£{x:,.0f}"
                fmt2 = lambda x: f"£{x:,.2f}"

                render = view.copy()
                render["Total Expense"] = render["Total Expense"].map(fmt2)
                render["Allocated"]     = render["Allocated"].map(fmt0)
                render["Remaining"]     = render["Remaining"].map(fmt0)
                render["% Used"]        = view["% Used"].astype(float)

                st.data_editor(
                    render,
                    hide_index=True,
                    use_container_width=True,
                    disabled=True,
                    column_config={
                        "Stakeholder": st.column_config.TextColumn(),
                        "Total Expense": st.column_config.TextColumn(),
                        "Allocated":     st.column_config.TextColumn(),
                        "Remaining":     st.column_config.TextColumn(),
                        "% Used": st.column_config.ProgressColumn(
                            format="%.0f%%",
                            min_value=0.0,
                            max_value=1.0,
                        ),
                    },
                )

                # totals + export
                tot = {
                    "Stakeholder": "TOTAL",
                    "Total Expense": float(view["Total Expense"].sum()),
                    "Allocated": float(view["Allocated"].sum()),
                    "Remaining": float(view["Remaining"].sum()),
                    "% Used": float((view["Total Expense"].sum() / view["Allocated"].sum()) if view["Allocated"].sum() else 0.0),
                }
                st.markdown(
                    f"**Totals —** "
                    f"Expense: £{tot['Total Expense']:,.2f} &nbsp;|&nbsp; "
                    f"Allocated: £{tot['Allocated']:,.0f} &nbsp;|&nbsp; "
                    f"Remaining: £{tot['Remaining']:,.0f} &nbsp;|&nbsp; "
                    f"% Used: {tot['% Used']*100:.0f}%"
                )

                csv_bytes = view.to_csv(index=False).encode("utf-8")
                st.download_button("⬇️ Export table (CSV)", data=csv_bytes, file_name="budget_detail.csv", mime="text/csv")






            from pathlib import Path
            import pandas as pd
            import re
            import streamlit as st

            EXP_MASTER = Path("Expenses/expenses_master.parquet")

            def _coerce_money(s: pd.Series) -> pd.Series:
                return pd.to_numeric(
                    s.astype(str)
                    .str.replace("’", "'", regex=False)
                    .str.replace("−", "-", regex=False)
                    .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)
                    .str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)", expand=False)
                    .str.replace(",", "", regex=False),
                    errors="coerce"
                )

            def _canon_eng_name(name: str) -> str:
                t = str(name or "").strip().replace("’", "'").casefold()
                t = re.sub(r"\b(cash|advance|utili[sz]ed|amount|cash\s+advance\s+utili[sz]ed\s+amount)\b", " ", t)
                t = re.sub(r"[^a-z' ]+", " ", t)
                t = re.sub(r"\s+", " ", t).strip()
                t = t.replace("o leary", "o'leary").replace("oleary", "o'leary")
                if "michael webb" in t:   return "Michael Webb"
                if "simon wakelin" in t:  return "Simon Wakelin"
                if "david o'leary" in t:  return "David O'Leary"
                parts = []
                for w in t.split():
                    if "'" in w:
                        a, b = w.split("'", 1)
                        parts.append(a.capitalize() + "'" + b.capitalize())
                    else:
                        parts.append(w.capitalize())
                return " ".join(parts)

            @st.cache_data(show_spinner=False)
            def load_expense_master_df() -> pd.DataFrame:
                if not EXP_MASTER.exists():
                    return pd.DataFrame(columns=[
                        "Engineer Name","Transaction Date","Expense Type","Business Purpose",
                        "Vendor","City of Purchase","Payment Type","Amount","Source File","RowUID"
                    ])
                df = pd.read_parquet(EXP_MASTER)
                if "Transaction Date" in df.columns:
                    df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
                return df

            def sum_expenses_for_engineer(e_name: str, start_dt=None, end_dt=None) -> float:
                df = load_expense_master_df()
                if df.empty:
                    return 0.0
                df = df.copy()
                df["__canon"] = df["Engineer Name"].map(_canon_eng_name)
                canon = _canon_eng_name(e_name)
                df = df[df["__canon"] == canon]
                if start_dt is not None:
                    df = df[df["Transaction Date"] >= pd.to_datetime(start_dt)]
                if end_dt is not None:
                    df = df[df["Transaction Date"] <= pd.to_datetime(end_dt)]
                if df.empty or "Amount" not in df.columns:
                    return 0.0
                return float(_coerce_money(df["Amount"]).fillna(0.0).sum())



            # ---- TAB 3: Expenses management ----
            from pathlib import Path
            import re
            import hashlib
            from datetime import date
            import pandas as pd
            import streamlit as st

            # ---------------- Paths ----------------
            EXP_BASE      = Path("Expenses")
            EXP_INBOX     = EXP_BASE / "Inbox"
            EXP_PROCESSED = EXP_BASE / "Processed"
            EXP_MASTER    = EXP_BASE / "expenses_master.parquet"

            for d in (EXP_BASE, EXP_INBOX, EXP_PROCESSED):
                d.mkdir(parents=True, exist_ok=True)

            # ---------------- Helpers (single source of truth) ----------------
            def _sanitize_amount_series(s: pd.Series) -> pd.Series:
                """Extract a money number and convert to float, ignoring years like 2025 or fragments like -17."""
                return pd.to_numeric(
                    s.astype(str)
                    .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)          # (123.45) -> -123.45
                    .str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)")[0]  # one money-looking number
                    .str.replace(",", "", regex=False),
                    errors="coerce"
                )
            def save_master(df: pd.DataFrame) -> None:
                df = df.copy()
                if "Transaction Date" in df.columns:
                    df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
                if "Amount" in df.columns:
                    df["Amount"] = _sanitize_amount_series(df["Amount"]).fillna(0.0)
                df.to_parquet(EXP_MASTER, index=False)

            def _coerce_money(s: pd.Series) -> pd.Series:
                """Turn strings like '£989.93', '(286.84)', '−17.00', '426.10' into floats."""
                return pd.to_numeric(
                    s.astype(str)
                    .str.replace("’", "'", regex=False)
                    .str.replace("−", "-", regex=False)               # unicode minus → ASCII
                    .str.replace(r"\(([^)]+)\)", r"-\1", regex=True)  # (123.45) → -123.45
                    .str.extract(r"(-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?)", expand=False)
                    .str.replace(",", "", regex=False),
                    errors="coerce"
                )



            def import_inbox_pdfs(move_processed: bool = True) -> tuple[pd.DataFrame, pd.DataFrame]:
                master = load_master()
                have_uids = set(master["RowUID"].astype(str)) if not master.empty else set()
                added = []

                inbox_files = [p for p in EXP_INBOX.iterdir() if p.suffix.lower() in {".pdf",".png",".jpg",".jpeg",".tif",".tiff"}]
                inbox_files.sort()

                for p in inbox_files:
                    if "mewp" in p.name.lower():
                        # special MEWP handler
                        total = extract_mewp_total(p)  # new helper you’ll add below
                        if total:
                            team = "Sky Business" if "business" in p.name.lower() else "Sky Retail"
                            uid = _hash_uid("MEWP", team, total, p.name)
                            if uid not in have_uids:
                                added.append({
                                    "Engineer Name": "MEWP Hire",
                                    "Transaction Date": pd.Timestamp.today().date(),
                                    "Expense Type": "MEWP Hire",
                                    "Business Purpose": "MEWP Invoice",
                                    "Vendor": "MEWP",
                                    "City of Purchase": "",
                                    "Payment Type": "",
                                    "Amount": total,
                                    "Source File": p.name,
                                    "RowUID": uid,
                                    "Team": team,  # ← so it maps cleanly
                                })
                    else:
                        rows = extract_rows_from_pdf(p)  # your existing logic
                        new_rows = [r for r in rows if str(r.get("RowUID","")) not in have_uids]
                        if new_rows:
                            added.extend(new_rows)

                    if move_processed:
                        try: p.rename(EXP_PROCESSED / p.name)
                        except Exception: pass

                added_df = pd.DataFrame(added)
                if not added_df.empty:
                    # normalise + merge into master
                    if "Transaction Date" in added_df.columns:
                        added_df["Transaction Date"] = pd.to_datetime(added_df["Transaction Date"], errors="coerce")
                    if "Amount" in added_df.columns:
                        added_df["Amount"] = _sanitize_amount_series(added_df["Amount"]).fillna(0.0)

                    master = pd.concat([master, added_df], ignore_index=True)
                    master = master.drop_duplicates(subset=["RowUID"], keep="first")
                    save_master(master)

                return (added_df, master)


            def extract_mewp_total(pdf_path: Path) -> float | None:
                """Look for 'TOTAL £xxxx.xx' in a MEWP invoice."""
                txt = _read_pdf_text(pdf_path)
                if not txt:
                    return None
                m = re.search(r"TOTAL\s*£?\s*([\d,]+\.\d{2})", txt, flags=re.I)
                if m:
                    try:
                        return float(m.group(1).replace(",", ""))
                    except:
                        return None
                return None



            def _canon_eng_name(name: str) -> str:
                """Canonicalise engineer names and strip junk phrases (cash-advance text, O’Leary variants)."""
                t = str(name or "").strip().replace("’", "'").casefold()
                # remove finance boilerplate / junk
                t = re.sub(r"\b(cash|advance|utili[sz]ed|amount|cash\s+advance\s+utili[sz]ed\s+amount)\b", " ", t)
                t = re.sub(r"[^a-z' ]+", " ", t)    # keep letters, spaces, apostrophes
                t = re.sub(r"\s+", " ", t).strip()
                t = t.replace("o leary", "o'leary").replace("oleary", "o'leary")

                # explicit mappings you had
                if "michael webb" in t:   return "Michael Webb"
                if "simon wakelin" in t:  return "Simon Wakelin"
                if "david o'leary" in t:  return "David O'Leary"

                # fallback: nice title-casing incl. O'X
                parts = []
                for w in t.split():
                    if "'" in w:
                        a, b = w.split("'", 1)
                        parts.append(a.capitalize() + "'" + b.capitalize())
                    else:
                        parts.append(w.capitalize())
                return " ".join(parts)

            def _norm(s: str) -> str:
                return re.sub(r"\s+", " ", str(s or "").strip().lower())

            def _to_amount(s):
                """Parse currency string like '£1,234.56', '(286.84)', '-12.30', '426.10', '989.93 GBP'."""
                try:
                    s = str(s).strip()
                    if not s:
                        return None
                    neg = False
                    if s.startswith("(") and s.endswith(")"):
                        neg = True
                        s = s[1:-1]
                    s = s.replace("£", "").replace(",", "").replace("GBP", "").strip()
                    s = s.replace("−", "-")  # unicode minus
                    val = float(s)
                    return -val if neg else val
                except Exception:
                    return None

            def _to_date(s):
                """Parse dd/mm/yyyy or dd-mm-yy etc."""
                try:
                    s = str(s or "")
                    m = re.search(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})', s)
                    if not m:
                        return None
                    d, mth, y = int(m.group(1)), int(m.group(2)), m.group(3)
                    y = int(y) if len(y) == 4 else 2000 + int(y)  # e.g., 25 -> 2025
                    return date(y, mth, d)
                except Exception:
                    return None

            def _hash_uid(*parts) -> str:
                return hashlib.sha1("|".join([str(p) for p in parts]).encode("utf-8")).hexdigest()

            # --- PDF text / parsing utilities (unchanged logic) ---
            import pdfplumber
            from PyPDF2 import PdfReader

            def _read_pdf_text(p: Path) -> str:
                # 1) pdfplumber text
                try:
                    buf = []
                    with pdfplumber.open(p) as pdf:
                        for page in pdf.pages:
                            t = page.extract_text() or ""
                            if t.strip():
                                buf.append(t)
                    if buf:
                        return "\n".join(buf)
                except Exception:
                    pass
                # 2) PyPDF fallback
                try:
                    reader = PdfReader(str(p))
                    return "\n".join((pg.extract_text() or "") for pg in reader.pages)
                except Exception:
                    return ""

            def _extract_employee_name(full_text: str, fallback: str = None) -> str:
                m = re.search(r"Employee Name\s*[:\-]?\s*([A-Za-z' \-]+)", full_text, flags=re.I)
                if m:
                    return m.group(1).strip()
                return (fallback or "").strip()

            def _extract_rows_from_text(full_text: str, pdf_path: Path, employee: str) -> list[dict]:
                out = []
                if not full_text:
                    return out

                lines = [ln.rstrip() for ln in full_text.splitlines()]
                start_i = 0
                for i, ln in enumerate(lines):
                    low = ln.lower()
                    if "transaction" in low and "amount" in low:
                        start_i = i + 1
                        break

                body = [ln for ln in lines[start_i:] if ln.strip()]
                if not body:
                    return out

                date_re = re.compile(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b')
                currency_re = re.compile(
                    r'(?:£\s*\(?-?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\)?|'
                    r'\(?-?\d{1,3}(?:,\d{3})*(?:\.\d{2})\)?\s*(?:GBP)?)',
                    re.IGNORECASE
                )

                STOP = ("company disbursements", "expenses requiring receipts", "missing receipt declaration")
                PAYMENT_PATTERNS = ["citibank t&e", "citibankt&e", "t&e", "citibank", "sky citibank"]
                CITY_WORDS = {
                    "london","watford","birmingham","brentford","cardiff","caerphilly","hove","yeovil",
                    "leicester","oxford","bedfont","huddersfield","hounslow","leeds","middlesex",
                    "gatwick","basingstoke","canterbury"
                }
                EXPENSE_MAP = {
                    "airfare": ["airfare","flight","flights","airline","klm","ba","british airways","easyjet","ryanair"],
                    "breakfast": ["breakfast","caffe nero","starbucks","costa"],
                    "lunch": ["lunch","pret","subway","kfc","mcdonalds","nando","nando's","tesco"],
                    "dinner": ["dinner","restaurant","wagamama","harvester","pizza","kfc","mcdonalds","nando","nandos"],
                    "hotel": ["hotel","hilton","premier inn","travelodge","marriott","holiday inn"],
                    "parking": ["parking","apcoa","ncp","drop off","drop-off","yourparkingspace","cp plus","groupnexus"],
                    "train": ["train","rail","ticket","underground","tfl","lner","national rail","lul"],
                    "emergency equipment": ["emergency equipment","screwfix","wickes","buyitdirect","amazon","premier farnell","cpc","consumables"],
                    "entertainment - staff": ["entertainment - staff","village hotels","zettle","staff"],
                    "tuition/training reimbursement": ["tuition","training","course","airdat","exam","certificate"],
                    "other": ["install","installation","servicecall","service call","van stock"]
                }

                def detect_expense_type(text_block: str) -> str:
                    t = _norm(text_block)
                    for label in EXPENSE_MAP.keys():
                        if label in t:
                            return label.title()
                    for label, kws in EXPENSE_MAP.items():
                        if any(k in t for k in kws):
                            return label.title()
                    return ""

                def detect_payment(text_block: str) -> str:
                    t = _norm(text_block)
                    for p in PAYMENT_PATTERNS:
                        if p in t:
                            if "citibank" in p and "t&e" in p: return "Citibank T&E"
                            if p == "t&e": return "T&E"
                            if p == "citibank": return "Citibank"
                            return p.title()
                    return ""

                def detect_amount(block_lines: list[str]):
                    best = None
                    best_line_idx = -1
                    for idx, ln in enumerate(block_lines):
                        matches = list(currency_re.finditer(ln))
                        filt = []
                        for m in matches:
                            s = m.group(0)
                            s_low = s.lower()
                            has_pound = "£" in s
                            has_dec   = "." in s
                            has_gbp   = "gbp" in s_low
                            if not has_pound and not has_dec and not has_gbp:
                                continue
                            if re.fullmatch(r'\(?-?(?:19|20)\d{2}\)?', s.strip()):
                                continue
                            filt.append(m)
                        if not filt:
                            continue
                        m = filt[-1]  # rightmost on the line
                        s = m.group(0)
                        if idx >= best_line_idx:
                            best_line_idx = idx
                            best = s
                    return _to_amount(best) if best else None

                starts = [i for i, ln in enumerate(body) if date_re.search(ln)]
                if not starts:
                    return out

                blocks = []
                for j, si in enumerate(starts):
                    ei = starts[j+1] if j+1 < len(starts) else len(body)
                    blocks.append((si, ei))

                def clean(text): return " ".join(text.split())

                for si, ei in blocks:
                    blk_lines = body[si:ei]
                    if any(any(s in ln.lower() for s in STOP) for ln in blk_lines):
                        break

                    blk = " ".join(blk_lines)
                    dt  = _to_date(blk)
                    amt = detect_amount(blk_lines)
                    if not (dt and (amt is not None)):
                        continue

                    pay = detect_payment(blk)
                    typ = detect_expense_type(blk)

                    city = ""
                    for ln in reversed(blk_lines):
                        if re.search(currency_re, ln):
                            tokens = [t.strip(",.") for t in ln.split()]
                            for tkn in reversed(tokens):
                                if tkn.lower() in CITY_WORDS:
                                    city = tkn
                                    break
                            break

                    vendor = ""
                    caps_chunks = re.findall(r'([A-Z][A-Z0-9&\.\'\- ]{2,})', " ".join(blk_lines))
                    ranked = []
                    for c in caps_chunks:
                        t = re.sub(r'\b(19|20)\d{2}\b', '', c).strip()
                        if not t: continue
                        if "T&E" in t or "GBP" in t: continue
                        if not re.search(r'[A-Z]', t): continue
                        ranked.append(t)
                    if ranked:
                        vendor = max(ranked, key=len)

                    purpose = " ".join(ln for ln in blk_lines if not re.search(date_re, ln))
                    for bad in [vendor, city, "Citibank", "T&E"]:
                        if bad:
                            purpose = purpose.replace(bad, " ")
                    purpose = re.sub(r'\b(19|20)\d{2}\b', ' ', purpose)
                    purpose = clean(purpose)

                    uid = _hash_uid(employee, dt, typ, purpose, vendor, amt, pdf_path.name)
                    out.append({
                        "Engineer Name": employee,
                        "Transaction Date": dt,
                        "Expense Type": typ,
                        "Business Purpose": purpose,
                        "Vendor": vendor,
                        "City of Purchase": city,
                        "Payment Type": pay,
                        "Amount": amt,
                        "Source File": pdf_path.name,
                        "RowUID": uid,
                    })

                return out

            def _try_tables_with_settings(page) -> list[list]:
                for settings in [
                    dict(
                        vertical_strategy="lines", horizontal_strategy="lines",
                        snap_tolerance=3, join_tolerance=3,
                        edge_min_length=10, min_words_vertical=1, min_words_horizontal=1
                    ),
                    dict(
                        vertical_strategy="text", horizontal_strategy="text",
                        text_x_tolerance=2, text_y_tolerance=2,
                        intersection_x_tolerance=3, intersection_y_tolerance=3,
                        snap_tolerance=3, join_tolerance=3,
                        keep_blank_chars=False
                    ),
                    None
                ]:
                    try:
                        tbls = page.extract_tables(table_settings=settings) if settings else page.extract_tables()
                        if tbls: return tbls
                    except Exception:
                        pass
                return []

            def extract_rows_from_pdf(pdf_path: Path) -> list[dict]:
                full_text = _read_pdf_text(pdf_path)
                employee = _extract_employee_name(full_text, fallback=pdf_path.stem)

                rows = _extract_rows_from_text(full_text, pdf_path, employee)
                if rows:
                    return rows

                out = []
                try:
                    with pdfplumber.open(pdf_path) as pdf:
                        for page in pdf.pages:
                            tables = _try_tables_with_settings(page)
                            if not tables:
                                continue
                            for tbl in tables:
                                if not tbl or len(tbl) < 2:
                                    continue
                                header = [c.strip().lower() for c in tbl[0]]
                                idx = {
                                    "date":   next((i for i, h in enumerate(header) if "date" in h), None),
                                    "type":   next((i for i, h in enumerate(header) if "type" in h and "payment" not in h), None),
                                    "purpose":next((i for i, h in enumerate(header) if "purpose" in h), None),
                                    "vendor": next((i for i, h in enumerate(header) if "vendor" in h or "merchant" in h), None),
                                    "city":   next((i for i, h in enumerate(header) if "city" in h), None),
                                    "pay":    next((i for i, h in enumerate(header) if "payment" in h), None),
                                    "amount": next((i for i, h in enumerate(header) if "amount" in h or "total" in h), None),
                                }
                                for r in tbl[1:]:
                                    dt   = _to_date(r[idx["date"]])   if idx["date"]   is not None and idx["date"]   < len(r) else None
                                    typ  = (r[idx["type"]]   or "").strip() if idx["type"]   is not None and idx["type"]   < len(r) else ""
                                    bus  = (r[idx["purpose"]]or "").strip() if idx["purpose"]is not None and idx["purpose"]< len(r) else ""
                                    ven  = (r[idx["vendor"]] or "").strip() if idx["vendor"] is not None and idx["vendor"] < len(r) else ""
                                    city = (r[idx["city"]]   or "").strip() if idx["city"]   is not None and idx["city"]   < len(r) else ""
                                    pay  = (r[idx["pay"]]    or "").strip() if idx["pay"]    is not None and idx["pay"]    < len(r) else ""
                                    amt  = _to_amount(r[idx["amount"]])      if idx["amount"] is not None and idx["amount"] < len(r) else None
                                    if not (dt and (amt is not None)):
                                        continue
                                    uid = _hash_uid(employee, dt, typ, bus, ven, amt, pdf_path.name)
                                    out.append({
                                        "Engineer Name": employee,
                                        "Transaction Date": dt,
                                        "Expense Type": typ,
                                        "Business Purpose": bus,
                                        "Vendor": ven,
                                        "City of Purchase": city,
                                        "Payment Type": pay,
                                        "Amount": amt,
                                        "Source File": pdf_path.name,
                                        "RowUID": uid,
                                    })
                except Exception:
                    pass
                return out

            # ---------------- Master I/O ----------------
            def load_master() -> pd.DataFrame:
                if EXP_MASTER.exists():
                    df = pd.read_parquet(EXP_MASTER)
                else:
                    df = pd.DataFrame(columns=[
                        "Engineer Name","Transaction Date","Expense Type","Business Purpose",
                        "Vendor","City of Purchase","Payment Type","Amount","Source File","RowUID"
                    ])
                if "Amount" in df.columns:
                    df["Amount"] = _sanitize_amount_series(df["Amount"])
                if "Transaction Date" in df.columns:
                    df["Transaction Date"] = pd.to_datetime(df["Transaction Date"], errors="coerce")
                return df

            master_df = load_master()
            if not master_df.empty:
                master_df = master_df.copy()
                master_df["Engineer Name"] = master_df["Engineer Name"].astype(str)
                master_df["__canon"] = master_df["Engineer Name"].map(_canon_eng_name)
                master_df["Team"]   = master_df["Engineer Name"].map(engineer_team)


            def import_inbox_pdfs(move_processed: bool = True) -> tuple[pd.DataFrame, pd.DataFrame]:
                master = load_master()
                have_uids = set(master["RowUID"].astype(str)) if not master.empty else set()

                added = []
                pdfs = sorted(EXP_INBOX.glob("*.pdf"))
                for p in pdfs:
                    rows = extract_rows_from_pdf(p)
                    new_rows = [r for r in rows if r["RowUID"] not in have_uids]
                    if new_rows:
                        added.extend(new_rows)
                    if move_processed:
                        try:
                            p.rename(EXP_PROCESSED / p.name)
                        except Exception:
                            pass

                added_df = pd.DataFrame(added)
                if not added_df.empty:
                    added_df["Amount"] = _sanitize_amount_series(added_df["Amount"])
                    master = pd.concat([master, added_df], ignore_index=True)
                    master = master.drop_duplicates(subset=["RowUID"], keep="first")
                    master["Amount"] = _sanitize_amount_series(master["Amount"])
                    save_master(master)

                return (added_df, master)

            # ---------------- UI (tabs; tables behind expanders) ----------------
            tab_import, tab_master, tab_engineer, tab_types = st.tabs(
                ["📥 Import", "📚 Master", "👷 Engineer", "🧩 Type Breakdown"]
            )

            # ========== Import ==========
            with tab_import:
                master_df = load_master()
                with st.expander("📥 Auto-import expenses from PDF (Inbox)", expanded=True):
                    lcol, rcol = st.columns([1, 1], gap="large")

                    with lcol:
                        st.write(f"Inbox: `{EXP_INBOX}`")
                        pdfs = sorted(EXP_INBOX.glob("*.pdf"))
                        st.write(f"Found **{len(pdfs)}** PDF(s).")

                        if st.button("🔎 Scan Inbox (preview)", key="scan_inbox_btn"):
                            parsed = []
                            for p in pdfs:
                                parsed.extend(extract_rows_from_pdf(p))
                            if parsed:
                                prev_df = pd.DataFrame(parsed)
                                if "Amount" in prev_df.columns:
                                    prev_df["Amount"] = _sanitize_amount_series(prev_df["Amount"])
                                try:
                                    prev_df = prev_df.sort_values("Transaction Date", ascending=False)
                                except Exception:
                                    pass
                                with st.expander("Preview parsed rows", expanded=True):
                                    st.dataframe(prev_df, use_container_width=True, hide_index=True)
                            else:
                                st.warning("No rows parsed. Expand **Show raw text** below to inspect the PDF text.")

                        with st.expander("Show raw text (first PDF)"):
                            if pdfs:
                                txt = _read_pdf_text(pdfs[0])
                                st.text("\n".join((txt or "").splitlines()[:200]))
                            else:
                                st.caption("No PDFs in Inbox.")

                        move_after = st.checkbox("Move PDFs to 'Processed' after import", value=True, key="move_after_cb")
                        if st.button("🚀 Import Inbox PDFs", type="primary", key="import_inbox_btn"):
                            added_df, master_df = import_inbox_pdfs(move_processed=move_after)
                            st.success(f"Imported {len(added_df)} new row(s). Master now has {len(master_df)} rows.")
                            if not added_df.empty:
                                with st.expander("Just imported", expanded=False):
                                    st.dataframe(added_df, use_container_width=True, hide_index=True)

                    with rcol:
                        st.markdown("**How it works**")
                        st.caption("• Put your expense PDFs into: `Expenses/Inbox/`")
                        st.caption("• Click **Scan Inbox (preview)** to see parsed rows.")
                        st.caption("• Click **Import Inbox PDFs** to append unique rows to the master table.")
                        st.caption(f"• Master file: `{EXP_MASTER.as_posix()}` (Parquet; needs `pyarrow`).")
                        st.caption("• Duplicate rows are avoided via a stable hash (`RowUID`).")

            # ========== Master ==========
            with tab_master:
                master_df = load_master()
                if master_df.empty:
                    st.info("Master is empty. Go to **Import** to add rows.")
                else:
                    try:
                        master_df = master_df.sort_values("Transaction Date", ascending=False)
                    except Exception:
                        pass

                    with st.expander("📚 Expenses master table", expanded=True):
                        st.dataframe(master_df, use_container_width=True, hide_index=True)
                        st.caption(f"{len(master_df)} row(s) in master — stored at `{EXP_MASTER.as_posix()}`")

                    with st.expander("🔍 Quick filters", expanded=False):
                        c1, c2, c3 = st.columns(3)
                        eng = c1.selectbox(
                            "Engineer", ["(All)"] + sorted(master_df["Engineer Name"].dropna().unique().tolist()),
                            index=0, key="master_eng_sel"
                        )
                        typ = c2.selectbox(
                            "Expense Type", ["(All)"] + sorted(master_df["Expense Type"].dropna().unique().tolist()),
                            index=0, key="master_type_sel"
                        )
                        purpose_q = c3.text_input("Business purpose contains", "", key="master_purpose_q")

                        view = master_df.copy()
                        if eng != "(All)":
                            view = view[view["Engineer Name"] == eng]
                        if typ != "(All)":
                            view = view[view["Expense Type"] == typ]
                        if purpose_q:
                            view = view[view["Business Purpose"].fillna("").str.contains(purpose_q, case=False)]

                        with st.expander("Filtered rows", expanded=True):
                            st.dataframe(view, use_container_width=True, hide_index=True)

            with st.expander("🗑 Delete rows from master", expanded=False):
                mdf = load_master().copy()
                if mdf.empty:
                    st.caption("Master is empty.")
                else:
                    # show a compact view with a selectable checkbox column
                    mdf = mdf.sort_values("Transaction Date", ascending=False)
                    mdf["__delete"] = False
                    edited = st.data_editor(
                        mdf,
                        hide_index=True,
                        use_container_width=True,
                        column_config={
                            "__delete": st.column_config.CheckboxColumn("Delete?", help="Tick to delete this row"),
                        },
                    )
                    to_drop = edited[edited["__delete"] == True]
                    if not to_drop.empty and st.button("⚠️ Permanently delete selected"):
                        keep = edited[edited["__delete"] == False].drop(columns="__delete")
                        save_master(keep)
                        st.success(f"Deleted {len(to_drop)} row(s) from master.")
                        st.rerun()

            # ========== Engineer ==========
            with tab_engineer:
                exp_master_df = load_master().copy()
                exp_master_df["__canon"] = exp_master_df["Engineer Name"].map(_canon_eng_name)

                options = ["(All)"] + sorted(exp_master_df["__canon"].dropna().unique().tolist())
                pick = st.selectbox("Engineer", options, index=0, key="exp_engineer_picker")

                df_view = exp_master_df if pick == "(All)" else exp_master_df[exp_master_df["__canon"] == pick]
                amount_series = _coerce_money(df_view.get("Amount", pd.Series(dtype=object))).fillna(0.0)
                total = float(amount_series.sum())
                st.metric(f"Total Amount for {pick}", f"£{total:,.2f}")

                if "Transaction Date" in df_view.columns:
                    df_view = df_view.sort_values("Transaction Date", ascending=False)

                with st.expander("Rows for selection", expanded=True):
                    st.dataframe(
                        df_view.drop(columns="__canon", errors="ignore"),
                        use_container_width=True,
                        hide_index=True,
                    )

            # ========== Type Breakdown ==========
            with tab_types:
                df = load_master()
                if df.empty:
                    st.info("No data in master yet.")
                else:
                    df = df.copy()
                    df["Amount"] = _sanitize_amount_series(df["Amount"]).fillna(0.0)

                    fc1, fc2 = st.columns(2)
                    start, end = fc1.date_input(
                        "Date range",
                        value=(pd.Timestamp.today().date() - pd.Timedelta(days=90), pd.Timestamp.today().date()),
                        key="types_date_range"
                    )
                    purpose_q = fc2.text_input("Business purpose contains", "", key="types_purpose_q")

                    mask = pd.Series(True, index=df.index)
                    if start:
                        mask &= df["Transaction Date"] >= pd.to_datetime(start)
                    if end:
                        mask &= df["Transaction Date"] <= pd.to_datetime(end)
                    if purpose_q:
                        mask &= df["Business Purpose"].fillna("").str.contains(purpose_q, case=False)

                    view = df.loc[mask].copy()
                    view["Expense Type"] = view["Expense Type"].fillna("Uncategorised")

                    type_sums = (
                        view.groupby("Expense Type", dropna=False)["Amount"].sum()
                            .sort_values(ascending=False)
                    )
                    st.caption("Top expense types")
                    if not type_sums.empty:
                        items = list(type_sums.items())
                        for i in range(0, len(items), 4):
                            cols = st.columns(4)
                            for (t, v), c in zip(items[i:i+4], cols):
                                c.metric(t, f"£{v:,.2f}")
                    else:
                        st.info("No rows for selected filters.")

                    with st.expander("📋 Expense Type totals (table)", expanded=False):
                        tbl = type_sums.reset_index().rename(columns={"Amount": "Total"})
                        tbl["Total"] = tbl["Total"].map(lambda x: f"£{x:,.2f}")
                        st.dataframe(tbl, use_container_width=True, hide_index=True)

                    with st.expander("📊 Expense Type totals (bar chart)", expanded=False):
                        chart_df = type_sums.reset_index().rename(columns={"Amount": "Total"})
                        st.bar_chart(chart_df.set_index("Expense Type"))

                    with st.expander("🔎 Drilldown: Business Purpose within type", expanded=False):
                        t_pick = st.selectbox(
                            "Choose a type to drill into",
                            ["(Select)"] + type_sums.index.tolist(),
                            index=0, key="types_drill_sel"
                        )
                        if t_pick != "(Select)":
                            sub = view[view["Expense Type"] == t_pick]
                            drill = (
                                sub.groupby("Business Purpose", dropna=False)["Amount"].sum()
                                    .sort_values(ascending=False)
                                    .reset_index()
                                    .rename(columns={"Amount": "Total"})
                            )
                            drill["Total"] = drill["Total"].map(lambda x: f"£{x:,.2f}")
                            st.dataframe(drill, use_container_width=True, hide_index=True)




# ============ HIGHLANDS & ISLANDS PAGE (INLINE) ============
if st.session_state.get("screen") == "highlands_islands":

    # 0) Back button (unique key so it never clashes)
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_highlands_page"):
        st.session_state.screen = "instructions_guide"
        st.rerun()

    st.title("🗺️ Highlands & Islands Dashboard")

    import streamlit as st
    import pandas as pd
    import plotly.express as px
    import numpy as np

    # Function to safely convert columns to numeric
    def safe_numeric(df, cols):
        for col in cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col].astype(str).str.replace(',', ''), errors='coerce')
        return df

    @st.cache_data
    def load_data(file_path):
        sheets = [
            "2022", "2023", "2024", "2025",
            "Company 2022", "Company 2023", "Company 2024", "Company 2025"
        ]
        dfs = pd.read_excel(file_path, sheet_name=sheets)
        for sheet, df in dfs.items():
            df.columns = df.columns.str.strip()
            df = safe_numeric(df, ['Issued Visits', 'Completed Visits', 'Average Complete Job Time (Min)',
                                '7 Day Revisits', '7 Day Revisits %', '30 Day Revisits', '30 Day Revisits %',
                                'Surveys', 'NPS%', 'NPS'])
            dfs[sheet] = df
        return dfs

    dfs = load_data("Highlands Islands.xlsx")

    # Separate sheets for tabs
    yearly_sheets = ["2022", "2023", "2024", "2025"]
    company_sheets = ["Company 2022", "Company 2023", "Company 2024", "Company 2025"]

    # Helper functions at top-level (outside tabs)
    def safe_sum(df, col):
        if col in df.columns:
            return pd.to_numeric(df[col], errors='coerce').sum(skipna=True)
        return 0

    def weighted_avg(df, val_col, weight_col):
        if val_col in df.columns and weight_col in df.columns:
            df_clean = df[[val_col, weight_col]].dropna()
            if df_clean.empty:
                return None
            total_weight = df_clean[weight_col].sum()
            if total_weight == 0:
                return None
            return (df_clean[val_col] * df_clean[weight_col]).sum() / total_weight
        return None

    def generate_summary(df, scope="Yearly"):
        total_issued = safe_sum(df, 'Issued Visits')
        total_completed = safe_sum(df, 'Completed Visits')
        completion_rate = (total_completed / total_issued * 100) if total_issued else 0
        total_not_done = safe_sum(df, 'Not Done Vists')
        not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
        total_7day_revisits = safe_sum(df, '7 Day Revisits')
        total_30day_revisits = safe_sum(df, '30 Day Revisits')
        weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
        total_surveys = safe_sum(df, 'Surveys')
        nps_col = 'NPS' if 'NPS' in df.columns else 'NPS%'
        weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col in df.columns else None

        summary = f"In this {scope} dataset, there were a total of {total_issued:,} issued visits, of which {total_completed:,} were completed, resulting in a completion rate of {completion_rate:.2f}%. "
        summary += f"Not done visits accounted for {total_not_done:,} visits ({not_done_pct:.2f}%). "
        summary += f"Revisit rates included {total_7day_revisits:,} 7-day revisits and {total_30day_revisits:,} 30-day revisits. "
        if weighted_job_time:
            summary += f"The average job completion time was approximately {weighted_job_time:.2f} minutes. "
        if total_surveys > 0 and weighted_nps is not None:
            summary += f"Customer satisfaction measured via NPS had an average score of {weighted_nps:.2f} based on {total_surveys:,} surveys. "
        else:
            summary += "NPS data is currently unavailable. "

        summary += f"This data provides valuable insights to guide operational and customer experience improvements."

        return summary

    
    tab1, tab2 = st.tabs(["Yearly Data", "Company Data"])

    with tab1:
        # default to first year so we can display summary immediately
        year_choice = st.selectbox("Select Year", yearly_sheets)
        df = dfs[year_choice]

        # Show advanced summary paragraph before the dropdown
        if not df.empty:
            st.markdown(f"### Advanced Summary for {year_choice}")
            st.info(generate_summary(df, "Yearly"))
        else:
            st.info("No data available for this year.")

        # Yearly KPIs
        with st.expander("Yearly KPIs", expanded=True):
            total_issued = df['Issued Visits'].sum()
            total_completed = df['Completed Visits'].sum()
            completion_rate = (total_completed / total_issued * 100) if total_issued else 0

            total_not_done = df['Not Done Vists'].sum() if 'Not Done Vists' in df.columns else 0
            not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
            total_7day_revisits = df['7 Day Revisits'].sum() if '7 Day Revisits' in df.columns else 0
            weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
            total_surveys = df['Surveys'].sum() if 'Surveys' in df.columns else 0
            weighted_nps = weighted_avg(df, 'NPS', 'Surveys') if 'NPS' in df.columns else None
            weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')

            cols = st.columns(3)

            cols[0].metric("Total Issued Visits", f"{total_issued:,.0f}")
            cols[1].metric("Total Completed Visits", f"{total_completed:,.0f}")
            cols[2].metric("Completion Rate (%)", f"{completion_rate:.2f}%")

            cols[0].metric("Total Not Done Visits", f"{total_not_done:,.0f}")
            cols[1].metric("Not Done Visits %", f"{not_done_pct:.2f}%")
            cols[2].metric("Total 7 Day Revisits", f"{total_7day_revisits:,.0f}")

            cols[0].metric("Weighted 7 Day Revisits %", f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A")
            cols[1].metric("Total Surveys", f"{total_surveys:,.0f}")
            cols[2].metric("Weighted Average NPS", f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A")

        with st.expander("Sunburst Chart - Issued Visits by Field by FRU"):
            if "Field by FRU" in df.columns:
                df_sb = df.copy()
                df_sb['Field by FRU'] = df_sb['Field by FRU'].fillna("Unknown")
                df_sb['Completion %'] = (df_sb['Completed Visits'] / df_sb['Issued Visits']).fillna(0) * 100

                fig = px.sunburst(df_sb, path=['Field by FRU'], values='Issued Visits',
                                color='Completion %', color_continuous_scale='RdYlGn',
                                title="Issued Visits by Field by FRU (colored by Completion %)")
                st.plotly_chart(fig, use_container_width=True, key=f"sunburst_yearly_{year_choice}")
            else:
                st.info("Column 'Field by FRU' not found in this dataset.")

    with tab2:
        comp_choice = st.selectbox("Select Company Year", company_sheets)
        df = dfs[comp_choice]

        if not df.empty:
            st.markdown(f"### Advanced Summary for {comp_choice}")
            st.info(generate_summary(df, "Company"))
        else:
            st.info("No data available for this company dataset.")

        with st.expander("Company KPIs", expanded=True):
            # Group by company and aggregate relevant metrics
            comp_kpis = df.groupby('Company').agg(
                Total_Issued_Visits = ('Issued Visits', 'sum'),
                Total_Completed_Visits = ('Completed Visits', 'sum')
            ).reset_index()

            # Calculate Completion Rate per company
            comp_kpis['Completion Rate (%)'] = (comp_kpis['Total_Completed_Visits'] / comp_kpis['Total_Issued_Visits']) * 100

            # Format numbers for display
            comp_kpis['Total_Issued_Visits'] = comp_kpis['Total_Issued_Visits'].map('{:,.0f}'.format)
            comp_kpis['Total_Completed_Visits'] = comp_kpis['Total_Completed_Visits'].map('{:,.0f}'.format)
            comp_kpis['Completion Rate (%)'] = comp_kpis['Completion Rate (%)'].map('{:.2f}%'.format)

            # Rename columns for nice display
            comp_kpis = comp_kpis.rename(columns={
                'Company': 'Company',
                'Total_Issued_Visits': 'Total Issued Visits',
                'Total_Completed_Visits': 'Total Completed Visits',
                'Completion Rate (%)': 'Completion Rate (%)'
            })

            st.dataframe(comp_kpis)



        with tab1:
            df = dfs[year_choice]

            if not df.empty:
                with st.expander("Section 3: Detailed Metrics & Charts (Yearly)", expanded=False):
                    total_issued = safe_sum(df, 'Issued Visits')
                    total_completed = safe_sum(df, 'Completed Visits')
                    total_not_done = safe_sum(df, 'Not Done Vists')  # spelling as per your header
                    total_7day_revisits = safe_sum(df, '7 Day Revisits')
                    total_30day_revisits = safe_sum(df, '30 Day Revisits')
                    total_surveys = safe_sum(df, 'Surveys')

                    weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
                    weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
                    weighted_30day_revisit_pct = weighted_avg(df, '30 Day Revisits %', 'Completed Visits')

                    nps_col = 'NPS' if 'NPS' in df.columns else None
                    weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col else None

                    # KPI Table
                    kpi_data = {
                        "Metric": [
                            "Total Issued Visits",
                            "Total Completed Visits",
                            "Total Not Done Visits",
                            "Total 7 Day Revisits",
                            "Weighted 7 Day Revisits %",
                            "Total 30 Day Revisits",
                            "Weighted 30 Day Revisits %",
                            "Total Surveys",
                            "Weighted Average NPS",
                            "Weighted Average Job Time (Min)"
                        ],
                        "Value": [
                            f"{total_issued:,.0f}",
                            f"{total_completed:,.0f}",
                            f"{total_not_done:,.0f}",
                            f"{total_7day_revisits:,.0f}",
                            f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A",
                            f"{total_30day_revisits:,.0f}",
                            f"{weighted_30day_revisit_pct:.2f}%" if weighted_30day_revisit_pct is not None else "N/A",
                            f"{total_surveys:,.0f}",
                            f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A",
                            f"{weighted_job_time:.2f}" if weighted_job_time is not None else "N/A",
                        ]
                    }
                    st.table(pd.DataFrame(kpi_data))

                    # Completion vs Not Done Pie Chart
                    pie_df = pd.DataFrame({
                        'Status': ['Completed', 'Not Done'],
                        'Visits': [total_completed, total_not_done]
                    })
                    fig_pie = px.pie(pie_df, names='Status', values='Visits', title="Completion Status Distribution")
                    st.plotly_chart(fig_pie, use_container_width=True, key=f"pie_completion_yearly_{year_choice}")

                    # Revisits Bar Chart (by Field by FRU)
                    revisit_cols = [col for col in ['7 Day Revisits', '30 Day Revisits'] if col in df.columns]
                    if 'Field by FRU' in df.columns and revisit_cols:
                        revisit_df = df.groupby('Field by FRU')[revisit_cols].sum().reset_index()
                        fig_bar = px.bar(revisit_df, x='Field by FRU', y=revisit_cols, barmode='group', title="Revisits by Field by FRU")
                        st.plotly_chart(fig_bar, use_container_width=True, key=f"bar_revisits_yearly_{year_choice}")

        with tab2:
            df = dfs[comp_choice]

            if not df.empty:
                with st.expander("Section 3: Detailed Metrics & Charts (Company)", expanded=False):
                    total_issued = safe_sum(df, 'Issued Visits')
                    total_completed = safe_sum(df, 'Completed Visits')
                    total_not_done = safe_sum(df, 'Not Done Vists')
                    total_7day_revisits = safe_sum(df, '7 Day Revisits')
                    total_30day_revisits = safe_sum(df, '30 Day Revisits')
                    total_surveys = safe_sum(df, 'Surveys')

                    weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
                    weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
                    weighted_30day_revisit_pct = weighted_avg(df, '30 Day Revisits %', 'Completed Visits')

                    nps_col = 'NPS%' if 'NPS%' in df.columns else None
                    weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col else None

                    # KPI Table
                    kpi_data = {
                        "Metric": [
                            "Total Issued Visits",
                            "Total Completed Visits",
                            "Total Not Done Visits",
                            "Total 7 Day Revisits",
                            "Weighted 7 Day Revisits %",
                            "Total 30 Day Revisits",
                            "Weighted 30 Day Revisits %",
                            "Total Surveys",
                            "Weighted Average NPS",
                            "Weighted Average Job Time (Min)"
                        ],
                        "Value": [
                            f"{total_issued:,.0f}",
                            f"{total_completed:,.0f}",
                            f"{total_not_done:,.0f}",
                            f"{total_7day_revisits:,.0f}",
                            f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A",
                            f"{total_30day_revisits:,.0f}",
                            f"{weighted_30day_revisit_pct:.2f}%" if weighted_30day_revisit_pct is not None else "N/A",
                            f"{total_surveys:,.0f}",
                            f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A",
                            f"{weighted_job_time:.2f}" if weighted_job_time is not None else "N/A",
                        ]
                    }
                    st.table(pd.DataFrame(kpi_data))

                    # Completion vs Not Done Pie Chart
                    pie_df = pd.DataFrame({
                        'Status': ['Completed', 'Not Done'],
                        'Visits': [total_completed, total_not_done]
                    })
                    fig_pie = px.pie(pie_df, names='Status', values='Visits', title="Completion Status Distribution")
                    st.plotly_chart(fig_pie, use_container_width=True, key=f"pie_completion_company_{comp_choice}")

                    # Revisits Bar Chart (by Company)
                    revisit_cols = [col for col in ['7 Day Revisits', '30 Day Revisits'] if col in df.columns]
                    if 'Company' in df.columns and revisit_cols:
                        revisit_df = df.groupby('Company')[revisit_cols].sum().reset_index()
                        fig_bar = px.bar(revisit_df, x='Company', y=revisit_cols, barmode='group', title="Revisits by Company")
                        st.plotly_chart(fig_bar, use_container_width=True, key=f"bar_revisits_company_{comp_choice}")


#SECTION 4#

        import matplotlib.pyplot as plt
        import seaborn as sns

        with tab1:
            # Yearly Tab
            with st.expander("Section 4: Advanced Analytics (Yearly)", expanded=False):
                # Completion % Trend over Years
                st.subheader("Completion % Trend Over Years")
                year_data = []
                for year in yearly_sheets:
                    d = dfs[year]
                    total_issued = safe_sum(d, 'Issued Visits')
                    total_completed = safe_sum(d, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    year_data.append({"Year": year, "Completion %": comp_pct})
                trend_df = pd.DataFrame(year_data)
                fig_trend = px.line(trend_df, x='Year', y='Completion %', markers=True)
                st.plotly_chart(fig_trend, use_container_width=True, key=f"comp_trend_yearly")

                # Top 5 Field by FRU by Issued Visits for selected year
                st.subheader(f"Top 5 Field by FRU by Issued Visits ({year_choice})")
                if 'Field by FRU' in df.columns:
                    top5 = df.groupby('Field by FRU')['Issued Visits'].sum().nlargest(5).reset_index()
                    fig_top5 = px.bar(top5, x='Field by FRU', y='Issued Visits', title="Top 5 Field by FRU")
                    st.plotly_chart(fig_top5, use_container_width=True, key=f"top5_frus_{year_choice}")
                else:
                    st.info("No 'Field by FRU' data available")

                # Not Done Visits Pie Chart
                st.subheader("Not Done Visits Proportion")
                total_not_done = safe_sum(df, 'Not Done Vists')
                pie_nd_df = pd.DataFrame({
                    'Status': ['Completed', 'Not Done'],
                    'Visits': [safe_sum(df, 'Completed Visits'), total_not_done]
                })
                fig_nd_pie = px.pie(pie_nd_df, names='Status', values='Visits', title="Not Done Visits Proportion")
                st.plotly_chart(fig_nd_pie, use_container_width=True, key=f"nd_pie_{year_choice}")

                # NPS Distribution (if available)
                if 'NPS' in df.columns:
                    st.subheader("NPS Distribution")
                    fig_nps = px.histogram(df, x='NPS', nbins=20, title="NPS Score Distribution")
                    st.plotly_chart(fig_nps, use_container_width=True, key=f"nps_dist_{year_choice}")
                else:
                    st.info("NPS data not available")

        with tab2:
            # Company Tab
            with st.expander("Section 4: Advanced Analytics (Company)", expanded=False):
                # Completion % Trend over Company Years
                st.subheader("Completion % Trend Over Company Years")
                comp_data = []
                for comp_year in company_sheets:
                    d = dfs[comp_year]
                    total_issued = safe_sum(d, 'Issued Visits')
                    total_completed = safe_sum(d, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    comp_data.append({"Year": comp_year, "Completion %": comp_pct})
                comp_trend_df = pd.DataFrame(comp_data)
                fig_comp_trend = px.line(comp_trend_df, x='Year', y='Completion %', markers=True)
                st.plotly_chart(fig_comp_trend, use_container_width=True, key=f"comp_trend_company")

                # Top 5 Companies by Issued Visits for selected company year
                st.subheader(f"Top 5 Companies by Issued Visits ({comp_choice})")
                if 'Company' in df.columns:
                    top5_comp = df.groupby('Company')['Issued Visits'].sum().nlargest(5).reset_index()
                    fig_top5_comp = px.bar(top5_comp, x='Company', y='Issued Visits', title="Top 5 Companies")
                    st.plotly_chart(fig_top5_comp, use_container_width=True, key=f"top5_company_{comp_choice}")
                else:
                    st.info("No 'Company' data available")

                # Not Done Visits Pie Chart
                st.subheader("Not Done Visits Proportion")
                total_not_done = safe_sum(df, 'Not Done Vists')
                pie_nd_df = pd.DataFrame({
                    'Status': ['Completed', 'Not Done'],
                    'Visits': [safe_sum(df, 'Completed Visits'), total_not_done]
                })
                fig_nd_pie = px.pie(pie_nd_df, names='Status', values='Visits', title="Not Done Visits Proportion")
                st.plotly_chart(fig_nd_pie, use_container_width=True, key=f"nd_pie_company_{comp_choice}")

                # NPS% Distribution (if available)
                if 'NPS%' in df.columns:
                    st.subheader("NPS% Distribution")
                    fig_nps = px.histogram(df, x='NPS%', nbins=20, title="NPS% Score Distribution")
                    st.plotly_chart(fig_nps, use_container_width=True, key=f"nps_dist_company_{comp_choice}")
                else:
                    st.info("NPS% data not available")



        with tab1:
            df = dfs[year_choice]

            if not df.empty:
                with st.expander("Section 5: Interactive Filters & Drilldown (Yearly)", expanded=False):
                    # Filter by Field by FRU
                    fru_options = df['Field by FRU'].dropna().unique()
                    selected_fru = st.multiselect(
                        "Filter by Field by FRU",
                        options=fru_options,
                        default=fru_options,
                        key=f"filter_fru_{year_choice}"
                    )

                    # Filter by Completion Status
                    status_options = ['Completed', 'Not Done']
                    selected_status = st.multiselect(
                        "Filter by Completion Status",
                        options=status_options,
                        default=status_options,
                        key=f"filter_status_{year_choice}"
                    )

                    # Filter dataframe based on selections
                    filtered_df = df[df['Field by FRU'].isin(selected_fru)]

                    # Create completion mask
                    if 'Not Done Vists' in filtered_df.columns and 'Completed Visits' in filtered_df.columns:
                        if set(selected_status) == {'Completed', 'Not Done'}:
                            pass
                        elif selected_status == ['Completed']:
                            filtered_df = filtered_df[filtered_df['Completed Visits'] > 0]
                        elif selected_status == ['Not Done']:
                            filtered_df = filtered_df[filtered_df['Not Done Vists'] > 0]
                        else:
                            filtered_df = filtered_df.iloc[0:0]

                    # Show KPIs for filtered data
                    total_issued = safe_sum(filtered_df, 'Issued Visits')
                    total_completed = safe_sum(filtered_df, 'Completed Visits')
                    total_not_done = safe_sum(filtered_df, 'Not Done Vists')

                    st.write("### Filtered KPIs")
                    st.metric("Total Issued Visits", f"{total_issued:,.0f}")
                    st.metric("Total Completed Visits", f"{total_completed:,.0f}")
                    st.metric("Total Not Done Visits", f"{total_not_done:,.0f}")

                    # Optionally show filtered data table
                    if st.checkbox("Show filtered data table", key=f"show_table_{year_choice}"):
                        st.dataframe(filtered_df)

        with tab2:
            df = dfs[comp_choice]

            if not df.empty:
                with st.expander("Section 5: Interactive Filters & Drilldown (Company)", expanded=False):
                    # Filter by Company
                    comp_options = df['Company'].dropna().unique()
                    selected_comp = st.multiselect(
                        "Filter by Company",
                        options=comp_options,
                        default=comp_options,
                        key=f"filter_company_{comp_choice}"
                    )

                    # Filter by Completion Status
                    status_options = ['Completed', 'Not Done']
                    selected_status = st.multiselect(
                        "Filter by Completion Status",
                        options=status_options,
                        default=status_options,
                        key=f"filter_status_company_{comp_choice}"
                    )

                    filtered_df = df[df['Company'].isin(selected_comp)]

                    if 'Not Done Vists' in filtered_df.columns and 'Completed Visits' in filtered_df.columns:
                        if set(selected_status) == {'Completed', 'Not Done'}:
                            pass
                        elif selected_status == ['Completed']:
                            filtered_df = filtered_df[filtered_df['Completed Visits'] > 0]
                        elif selected_status == ['Not Done']:
                            filtered_df = filtered_df[filtered_df['Not Done Vists'] > 0]
                        else:
                            filtered_df = filtered_df.iloc[0:0]

                    # Calculate KPIs for filtered data
                    total_issued = safe_sum(filtered_df, 'Issued Visits')
                    total_completed = safe_sum(filtered_df, 'Completed Visits')
                    total_not_done = safe_sum(filtered_df, 'Not Done Vists')
                    not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
                    total_7day_revisits = safe_sum(filtered_df, '7 Day Revisits')

                    weighted_7day_revisit_pct = weighted_avg(filtered_df, '7 Day Revisits %', 'Completed Visits')
                    total_surveys = safe_sum(filtered_df, 'Surveys')
                    weighted_nps = weighted_avg(filtered_df, 'NPS%', 'Surveys')

                    completion_rate = (total_completed / total_issued * 100) if total_issued else 0

                    st.write("### Filtered KPIs")
                    cols = st.columns(3)

                    cols[0].metric("Total Issued Visits", f"{total_issued:,.0f}")
                    cols[1].metric("Total Completed Visits", f"{total_completed:,.0f}")
                    cols[2].metric("Completion Rate (%)", f"{completion_rate:.2f}%")

                    cols[0].metric("Total Not Done Visits", f"{total_not_done:,.0f}")
                    cols[1].metric("Not Done Visits %", f"{not_done_pct:.2f}%")
                    cols[2].metric("Total 7 Day Revisits", f"{total_7day_revisits:,.0f}")

                    cols[0].metric("Weighted 7 Day Revisits %", f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A")
                    cols[1].metric("Total Surveys", f"{total_surveys:,.0f}")
                    cols[2].metric("Weighted Average NPS%", f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A")

                    if st.checkbox("Show filtered data table", key=f"show_table_company_{comp_choice}"):
                        st.dataframe(filtered_df)



        from sklearn.linear_model import LinearRegression
        import numpy as np
        import plotly.express as px

        with tab1:
            with st.expander("Section 6: Yearly Trends & Forecasts", expanded=False):
                # Combine all yearly data into one DataFrame
                combined_yearly = []
                for year in yearly_sheets:
                    df_year = dfs[year]
                    total_issued = safe_sum(df_year, 'Issued Visits')
                    total_completed = safe_sum(df_year, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    combined_yearly.append({
                        "Year": int(year),
                        "Issued Visits": total_issued,
                        "Completed Visits": total_completed,
                        "Completion %": comp_pct
                    })
                trend_df = pd.DataFrame(combined_yearly).sort_values("Year")

                # KPIs to plot
                kpis = ["Issued Visits", "Completed Visits", "Completion %"]

                for kpi in kpis:
                    st.subheader(f"{kpi} Trend Over Years")
                    fig = px.line(trend_df, x='Year', y=kpi, markers=True, title=f"{kpi} Trend Over Years")

                    # Forecast next year using simple linear regression
                    X = trend_df['Year'].values.reshape(-1,1)
                    y = trend_df[kpi].values

                    model = LinearRegression()
                    model.fit(X, y)

                    next_year = np.array([[trend_df['Year'].max() + 1]])
                    forecast = model.predict(next_year)[0]

                    # Add forecast point to the chart
                    fig.add_scatter(x=[next_year[0][0]], y=[forecast], mode='markers+text',
                                    marker=dict(color='red', size=10),
                                    text=["Forecast"],
                                    textposition="top center",
                                    name="Forecast")

                    st.plotly_chart(fig, use_container_width=True, key=f"trend_forecast_{kpi.replace(' ', '_')}")

    # 5) Stop here so nothing else renders on this screen
    st.stop()
# ---- tiny helper to show the exec logo, centered ----
from pathlib import Path
import streamlit as st


