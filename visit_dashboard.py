# ===== top of file (replace your lines 1–178 with this) =====
from __future__ import annotations

# --- core + libs you already use elsewhere in the app ---
import os
import re
import time
import uuid
import base64
import platform
import pathlib
from typing import Optional, Dict, Any

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from mpl_toolkits.mplot3d import Axes3D  # noqa: F401 (kept for compatibility)
import plotly.express as px
import plotly.graph_objects as go
import PyPDF2
import pdfplumber

import streamlit as st
import streamlit.components.v1 as components
from st_aggrid import AgGrid, GridOptionsBuilder
from streamlit_lottie import st_lottie

import requests

# dates/times
from datetime import datetime, timezone
from datetime import datetime as dt  # keep alias if you use dt elsewhere

# openai / langchain (kept as in your app)
from openai import OpenAI, OpenAIError
from statsmodels.tsa.arima.model import ARIMA
from langchain_experimental.agents import create_pandas_dataframe_agent
from langchain.agents.agent_types import AgentType
from langchain.schema import HumanMessage, SystemMessage
from collections import defaultdict
from langchain_community.chat_models import ChatOpenAI
if st.sidebar.button("Send Teams test (cloud)"):
    url = (st.secrets.get("TEAMS_WEBHOOK_URL")
           or st.secrets.get("teams", {}).get("webhook_url"))
    import requests
    requests.post(url, json={"text": "✅ Cloud test ping"}, timeout=10).raise_for_status()
    st.sidebar.success("Sent from Cloud")

# ---------- UK timezone setup ----------
try:
    # Python 3.9+ stdlib
    from zoneinfo import ZoneInfo
    _UK_TZ = ZoneInfo("Europe/London")
except Exception:
    # If tz data is missing on Windows, fall back to UTC.
    # (Optional: pip install tzdata)
    _UK_TZ = timezone.utc

# ---------- Teams webhook helpers ----------
def _get_teams_webhook() -> Optional[str]:
    """
    Reads the Teams webhook from Streamlit secrets.
    Supports either:
      TEAMS_WEBHOOK_URL = "https://..."
    or:
      [teams]
      webhook_url = "https://..."
    """
    s = st.secrets
    # flat key
    if "TEAMS_WEBHOOK_URL" in s:
        return str(s["TEAMS_WEBHOOK_URL"]).strip()
    # table key
    if "teams" in s:
        for k in ("webhook_url", "WEBHOOK_URL", "url", "URL"):
            if k in s["teams"]:
                return str(s["teams"][k]).strip()
    return None

def send_to_teams(
    title: str,
    text: str,
    facts: Optional[Dict[str, Any]] = None,
    button_text: Optional[str] = None,
    button_url: Optional[str] = None,
    color: str = "0078D4",
) -> None:
    """
    Generic sender for Teams Incoming Webhook using MessageCard.
    """
    url = _get_teams_webhook()
    if not url:
        return

    payload: Dict[str, Any] = {
        "@type": "MessageCard",
        "@context": "https://schema.org/extensions",
        "summary": title,
        "themeColor": color,
        "title": title,
        "text": text,
    }

    if facts:
        payload.setdefault("sections", []).append({
            "facts": [{"name": k, "value": str(v)} for k, v in facts.items()],
            "markdown": True,
        })

    if button_text and button_url:
        payload["potentialAction"] = [{
            "@type": "OpenUri",
            "name": button_text,
            "targets": [{"os": "default", "uri": button_url}],
        }]

    # Raises if Teams rejects it (use try/except where you call it if desired)
    requests.post(url, json=payload, timeout=10).raise_for_status()

def send_login_card(
    user_name: str,
    user_team: Optional[str] = None,
    tab_url: Optional[str] = None,
) -> None:
    """
    Pretty login notification with UK local time + UTC, session id, and optional button.
    """
    url = _get_teams_webhook()
    if not url:
        return

    # short, stable session id for the browser session
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())[:8]
    session_id = st.session_state.session_id

    now_utc = datetime.now(timezone.utc)
    now_uk  = now_utc.astimezone(_UK_TZ)
    uk_str  = now_uk.strftime("%A, %d %B %Y at %H:%M %Z")   # e.g., Friday, 15 August 2025 at 10:42 BST
    utc_str = now_utc.strftime("%Y-%m-%d %H:%M UTC")

    facts = {
        "User": user_name,
        "Team": user_team or "—",
        "Local time": uk_str,
        "App": "Visit Dashboard",
    }

    payload: Dict[str, Any] = {
        "@type": "MessageCard",
        "@context": "https://schema.org/extensions",
        "summary": f"Login — {user_name}",
        "themeColor": "0078D4",
        "title": "🔐 New login",
        "text": "A user has signed in.",
        "sections": [{
            "activityTitle": f"**{user_name}**",
            "activitySubtitle": "Visit Dashboard",
            "facts": [{"name": k, "value": str(v)} for k, v in facts.items()],
            "markdown": True
        }]
    }

    if tab_url:
        payload["potentialAction"] = [{
            "@type": "OpenUri",
            "name": "Open dashboard",
            "targets": [{"os": "default", "uri": tab_url}],
        }]

    requests.post(url, json=payload, timeout=10).raise_for_status()

# initialize once-per-session flag (prevents duplicate messages on refresh)
if "login_notified" not in st.session_state:
    st.session_state.login_notified = False

# ---------- Optional debug sidebar (toggle on/off) ----------
DEBUG_TEAMS = False  # set True temporarily if you need to test

if DEBUG_TEAMS:
    with st.sidebar:
        st.markdown("**Debug**")
        st.write("CWD:", os.getcwd())
        st.write("Local secrets exists:", os.path.exists(os.path.join(os.getcwd(), ".streamlit", "secrets.toml")))
        st.write("Home secrets exists:", os.path.exists(os.path.join(pathlib.Path.home(), ".streamlit", "secrets.toml")))
        try:
            st.write("Secrets keys:", list(st.secrets.keys()))
        except Exception as e:
            st.write("Secrets error:", str(e))
        st.write("Webhook present:", bool(_get_teams_webhook()))
        if st.button("Send Teams test"):
            send_to_teams("✅ Test", "Manual test from sidebar")
            st.success("Sent")
# ===== end of replacement block =====







import db
db.init_db()





# --- Session State Defaults ---
if "screen" not in st.session_state:
    st.session_state.screen = "instruction_guide"
if "quick_summary" not in st.session_state:
    st.session_state.quick_summary = False
if "screen" not in st.session_state:
    st.session_state.screen = "area_selection"
if "user_df" not in st.session_state:
    st.session_state.user_df = None
if "user_file_name" not in st.session_state:
    st.session_state.user_file_name = None
if "selected_dataset" not in st.session_state:
    st.session_state.selected_dataset = None
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "ai_chat" not in st.session_state:
    st.session_state.ai_chat = []


# --- Load Lottie Animation ---
import streamlit as st
import requests
from streamlit_lottie import st_lottie

def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# Then rest of your code here


# --- Cache: Load Excel file and add source ---
@st.cache_data(show_spinner=False)
def load_oracle_file(path, source_label):
    df = pd.read_excel(path)
    df["Source"] = source_label
    df["Team"] = source_label
    return df

# --- Cache: Load logo image ---
@st.cache_data(show_spinner=False)
def load_logo_base64(logo_path="sky_vip_logo.png"):
    try:
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    except:
        return None

logo_base64 = load_logo_base64()

# --- Files to load ---
files = {
    "VIP North": "VIP North Oracle Data.xlsx",
    "VIP South": "VIP South Oracle Data.xlsx",
    "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
}

# --- Load all data once ---
all_frames = []
for team, path in files.items():
    try:
        df = load_oracle_file(path, team)
        all_frames.append(df)
    except Exception as e:
        st.warning(f"⚠️ Could not load {path}: {e}")

if all_frames:
    data = pd.concat(all_frames, ignore_index=True)
    data.columns = data.columns.str.strip()
else:
    st.error("No data loaded from Oracle files.")
    st.stop()

## --- Full Styling and Login Block with Unified Look ---
import streamlit as st
from datetime import datetime as dt

# --- LIGHT THEME + CUSTOM STYLES ---
st.markdown("""
<style>
/* 🌕 LIGHT MODE THEME */
body, .main, .stApp {
    background-color: #ffffff !important;
    color: #1f1f1f !important;
    font-family: 'Segoe UI', 'Inter', sans-serif;
}

/* 🔵 BLUE BUTTON STYLE */
.stButton > button {
    background: linear-gradient(to right, #007bff, #0051a2) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.6rem 1.2rem !important;
    font-weight: 600 !important;
    margin: 0.4rem !important;
    box-shadow: 0 0 10px rgba(0,123,255,0.4) !important;
    transition: 0.3s ease-in-out !important;
}
.stButton > button:hover {
    background: linear-gradient(to right, #3399ff, #0073e6) !important;
    box-shadow: 0 0 16px rgba(0,123,255,0.6) !important;
}

/* 🖼️ FULL WIDTH LOGO */
.logo {
    width: 100%;
    display: flex;
    justify-content: center;
    margin: 2rem 0 1rem 0;
}
.logo img {
    width: 100%;
    max-width: 600px;
    height: auto;
}

/* 🧭 HEADINGS */
h1, h2, h3, h4 {
    color: #1f1f1f !important;
    font-weight: 700;
    text-align: center;
    margin-bottom: 1rem;
}

/* 📦 MARKDOWN & CONTAINERS */
.stMarkdown, .stText, .stContainer {
    background-color: transparent !important;
    color: #1f1f1f !important;
}

/* ✅ INFO BOX */
.info-box {
    background-color: #f0f4f8;
    color: #1f1f1f;
    border-left: 4px solid #007bff;
    padding: 0.8rem 1.2rem;
    border-radius: 8px;
    margin-top: 2rem;
    font-size: 0.95rem;
    text-align: center;
}

/* 🌀 TYPING ANIMATION */
@keyframes typing {
  0% { width: 0 }
  40%, 60% { width: 100% }
  100% { width: 0 }
}
@keyframes blink {
  50% { border-color: rgba(0,0,0,0.75); }
}
.login-header-wrapper {
    display: flex;
    justify-content: center;
    margin-top: 10px;
    margin-bottom: 10px;
}
.login-header {
    display: inline-block;
    font-size: 20px;
    font-weight: 500;
    border-right: 2px solid rgba(0,0,0,0.75);
    white-space: nowrap;
    overflow: hidden;
    width: 32ch;
    animation: typing 6s steps(32, end) infinite, blink 0.75s step-end infinite;
    color: #1f1f1f;
    text-align: center;
}

/* 🧩 SIDEBAR + FOOTER */
.stSidebar {
    background-color: #f8f9fa !important;
}
footer, #MainMenu, header {
    visibility: hidden;
}
</style>
""", unsafe_allow_html=True)


# --- LOGIN FUNCTION ---
import streamlit as st
from streamlit_lottie import st_lottie
from datetime import datetime as dt
import requests

# Load lottie animation JSON from URL
def load_lottie_url(url: str):
    try:
        response = requests.get(url)
        if response.status_code != 200:
            return None
        return response.json()
    except:
        return None

# --- LOGIN SCREEN FUNCTION ---
import streamlit as st
from streamlit_lottie import st_lottie
import requests
from datetime import datetime as dt

# Your user dictionary
users = {
    "Matt Hughes":     {"password": "mattpw", "team": "VIP Team North"},
    "Andy Holmes":     {"password": "andypw", "team": "Tier 2 North"},
    "Steve Paisley":   {"password": "stevepw", "team": "Tier 2 South"},
    "Branks Krneta":   {"password": "brankspw", "team": "TAP's Team"},
    "Chris Woods":     {"password": "chrispw", "team": "Highlands & Islands"},
    "Rachel Wylie":    {"password": "rachelpw", "team": "Sky Retail"},
    "Darryl Fuller":   {"password": "darrylpw", "team": "Highlands & Islands"},
    "Christian Boyce": {"password": "chrisbpw", "team": "VIP Team South"},
    "Rob Mangar":      {"password": "robpw", "team": "Operational Team"},
    "Dan Homewood":    {"password": "danpw", "team": "Finance Team"},
}
name_list = list(users.keys())

def load_lottie_url(url: str):
    try:
        response = requests.get(url)
        if response.status_code != 200:
            return None
        return response.json()
    except:
        return None

# --- LOGIN SCREEN FUNCTION ---
def login_screen_with_animation(logo_base64):
    lottie = load_lottie_url("https://assets2.lottiefiles.com/packages/lf20_jcikwtux.json")

    # Full-width logo
    st.markdown(f"""
    <div class="logo">
        <img src='data:image/png;base64,{logo_base64}' />
    </div>
    """, unsafe_allow_html=True)

    # Animated login header
    st.markdown("""
    <div class='login-header-wrapper'>
      <div class='login-header'>
        🔐 Login to Visit Insights Dashboard
      </div>
    </div>
    """, unsafe_allow_html=True)

    # Subheader
    st.markdown("<p style='text-align:center;'>Please choose your name to continue.</p>", unsafe_allow_html=True)

    # Name dropdown
    selected_name = st.selectbox("Choose Your Name", ["-- Select --"] + name_list)
    if selected_name and selected_name != "-- Select --":
        password = st.text_input("Enter your password", type="password", key="user_pw")
        if password:
            if password == users[selected_name]["password"]:
                st.session_state.authenticated = True
                st.session_state.username = selected_name
                st.rerun()
            else:
                st.error("Incorrect password for this user.")
    else:
        st.info("Please select your name to proceed.")

    # Optional animation
    if lottie:
        st_lottie(lottie, height=280, key="login_anim")
    else:
        st.info("🌀 Animation failed to load — but login still works.")

    # Footer
    current_time = dt.now().strftime("%A, %d %B %Y | %H:%M:%S")
    st.markdown(f"<div class='footer' style='text-align:center; opacity:0.6;'>{current_time} | v1.0.0</div>", unsafe_allow_html=True)
import streamlit as st
from streamlit_modal import Modal
import random

# Example messages or themes
fun_quotes = [
    "Success is the sum of small efforts, repeated day in and day out.",
    "Your hard work is making a difference!",
    "Teamwork divides the task and multiplies the success.",
    "Keep up the amazing work!"
]

def show_welcome_modal(name, team):
    modal = Modal(key="welcome_modal", title=f"👋 Welcome, {name}!", padding=30)
    if modal.is_open():
        st.write(f"Welcome to the dashboard, {name}!")
        st.write(f"You're leading **{team}**. {random.choice(fun_quotes)}")
        st.button("Let's Go!", on_click=modal.close)

# After successful login:
if st.session_state.get("authenticated") and st.session_state.screen == "instruction_guide":
    user = users[st.session_state.username]
    # Show welcome modal (this pops up on login)
    show_welcome_modal(st.session_state.username, user['team'])

    # Or as a card at the top of the page:
    

# --- AUTH CHECK ---
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if "screen" not in st.session_state:
    st.session_state.screen = "login"

if not st.session_state.authenticated:
    login_screen_with_animation(logo_base64 or "")
    st.stop()

if st.session_state.authenticated and st.session_state.screen == "login":
    st.session_state.screen = "instruction_guide"

if st.session_state.screen == "instruction_guide":
    # Your instruction guide logic here
    
    st.markdown(f"""
    <div class="logo">
        <img src='data:image/png;base64,{logo_base64}' />
    </div>
    """, unsafe_allow_html=True)

    # --- TYPING INTRO TEXT ---
    st.markdown("""
    <style>
    .adv-summary {
        font-size: 20px;
        font-weight: 400;
        border-right: 2px solid rgba(255,255,255,0.75);
        white-space: nowrap;
        overflow: hidden;
        width: fit-content;
        margin: 0 auto 30px;
        animation: typing 6s steps(60, end) infinite, blink 0.75s step-end infinite;
    }
    </style>
    <div class='adv-summary'>
    Welcome to the advanced reporting hub use the options below to explore all areas
    </div>
    """, unsafe_allow_html=True)


    # Show logo centered above title with glow effect
    st.markdown(f"""
    <style>
    .logo {{
        text-align: center;
        margin-bottom: 2rem;
    }}

    .logo img {{
        width: 750x;  /* Adjust width as needed */
        height: auto;
        filter: drop-shadow(0 0 5px rgba(10, 102, 194, 0.7));
    }}
            
    /* Container for header */
    .header-box {{
        text-align: center;
        margin-bottom: 40px;
    }}
    .header-box img {{
        width: 700px;
        filter: drop-shadow(0 0 8px #0a66c2);
        margin-bottom: 20px;
    }}
    .header-box h1 {{
        font-size: 3rem;
        color: #0a66c2;
        font-weight: 900;
        text-shadow: 0 0 12px #2171c7;
        margin-bottom: 0.2rem;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}
    .header-box h3 {{
        font-size: 1.3rem;
        font-style: italic;
        color: #4491ff;
        margin-top: 0;
        font-weight: 600;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}

    /* Section headers */
    .section-header {{
        font-size: 1.8rem;
        font-weight: 700;
        margin-top: 2.5rem;
        margin-bottom: 1rem;
        color: #80c3ff;
        text-shadow: 0 0 6px rgba(0, 115, 230, 0.7);
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}

    /* Expander customization */
    details > summary {{
        cursor: pointer;
        font-weight: 600;
        font-size: 1.25rem;
        padding: 0.7rem 1.2rem;
        border-radius: 12px;
        background: linear-gradient(90deg, #0a66c2, #2171c7);
        color: white;
        box-shadow: 0 5px 10px rgba(10,102,194,0.5);
        margin-bottom: 0.7rem;
        transition: background 0.3s ease;
        user-select: none;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}
    details[open] > summary {{
        background: linear-gradient(90deg, #2171c7, #0a66c2);
        box-shadow: 0 7px 14px rgba(33,113,199,0.7);
    }}
    details > summary:hover {{
        background: #0d4a8f;
    }}
    details > div, details > p {{
        margin-left: 1.2rem;
        margin-right: 1.2rem;
        font-size: 1rem;
        line-height: 1.6;
        color: #dceeff;
        white-space: pre-wrap;
        user-select: text;
        overflow-wrap: break-word;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}

    /* Important Notes list style */
    .important-notes {{
        font-size: 1rem;
        line-height: 1.6;
        margin-top: 1rem;
        margin-left: 1.2rem;
        color: #b0d4ff;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }}
    .important-notes li {{
        margin-bottom: 0.7rem;
    }}

    /* Big button styling */
    .big-button {{
        display: block;
        margin: 3rem auto 0 auto;
        padding: 1.4rem 3.5rem;
        background: linear-gradient(90deg, #0a66c2, #2171c7);
        color: white;
        font-size: 1.8rem;
        font-weight: 900;
        border: none;
        border-radius: 9999px;
        box-shadow: 0 10px 22px rgba(10,102,194,0.7);
        cursor: pointer;
        transition: background 0.35s ease;
        text-align: center;
        width: max-content;
        user-select: none;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        text-decoration: none;
        text-transform: uppercase;
        letter-spacing: 1.1px;
    }}
    .big-button:hover {{
        background: linear-gradient(90deg, #2171c7, #0a66c2);
        box-shadow: 0 15px 28px rgba(33,113,199,0.95);
    }}
    </style>

    <div class="header-box">
      
      <h1>Sky Orbit Visit Dashboard — User Guide & Overview</h1>
      <h3>Your one-stop hub for exploring Oracle visit data with ease!</h3>
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div style='
      padding:32px; 
      background:linear-gradient(90deg, #00c6ff, #0072ff); 
      color:white; 
      border-radius:18px; 
      font-size:1.6em; 
      margin-bottom:2em;
      text-align:center;
      box-shadow:0 4px 24px rgba(0,0,0,0.1);'>
      👋 Welcome back, <b>{st.session_state.username}</b>!<br>
      <span style='font-size:1.1em;'>Thanks for looking after <b>{user['team']}</b>.<br>
      <i>{random.choice(fun_quotes)}</i>
      </span>
    </div>
    """, unsafe_allow_html=True)

    import streamlit as st
    from streamlit_modal import Modal



    # SECTION DATA
    sections = [
        {
            "title": "What is Sky Orbit?",
            "content": "Sky Orbit is an interactive dashboard built especially for our teams that brings together Oracle visit data from multiple teams (VIP North, VIP South, Tier 2 North, Tier 2 South). It helps you explore, analyse, and visualize this data easily with features like secure login, AI-powered chat, detailed KPIs, trends, and forecasting."
        },
        {
            "title": "Main Menu",
            "content": "From here, pick the area you want to explore — like Operational Area, Dashboards, AI Chat (Sky Orbit AI), Suggestions, Forecasts, Sky Retail, Sky Business, and more. Just click the button to jump in."
        },
        {
            "title": "Upload Your Data (Optional)",
            "content": "You can upload your own Excel or CSV file with visit data to explore dynamically. If you don’t upload anything, the app uses default Oracle data combined from the four main teams."
        },
        {
            "title": "Sky Retail Area",
            "content": "View detailed KPIs and trends filtered by stakeholders such as Currys, Sky Retail, and EE. See totals for visits and value, average times, monthly trends, and day-of-week breakdowns. Explore charts and tables by stakeholder, team, or engineer."
        },
        {
            "title": "Sky Business Area",
            "content": "This filters the data to “Sky Business” visit types, showing KPIs for total visits, values, and completion rates. You get activity breakdown charts, monthly trends, sunburst visuals, forecasts by team, and detailed tables with heatmaps."
        },
        {
            "title": "VIP - SB Standby Section",
            "content": "Specialized KPIs and charts for “VIP - SB Standby” visits. Summaries of completed visits, total values, average start/end times, activity status distributions, and monthly counts. Forecasts are included here too. (Note: Time metrics only count rows with valid start/end times.)"
        },
        {
            "title": "SLA Dashboard",
            "content": "Track tickets against SLA buckets like 2h, 4h, 5 day, and 8h targets. View KPIs for total tickets, SLA met/missed counts, and percentages. Visualize ticket volumes by SLA and monthly trends, plus get forecasts for upcoming months."
        },
        {
            "title": "Suggestion Box",
            "content": "Submit or delete suggestions through an Excel-based interface. It uses OneDrive-safe temp files with unique keys to avoid data conflicts."
        },
        {
            "title": "AI Chat Assistant (“Sky Orbit”)",
            "content": "Ask natural language questions about the visit data. You can get answers on total costs, top postcodes, percentage differences, and other stats. The AI also generates charts when relevant. All chats are logged with timestamps and password-protected access."
        }
    ]

        # SESSION STATE FOR EXPANSION
    if "open_card" not in st.session_state:
        st.session_state.open_card = None

    # CSS FOR BOXES
    st.markdown("""
    <style>
    .card-button {
        background: linear-gradient(135deg, #0099ff 30%, #004488 100%);
        color: white;
        border: none;
        border-radius: 20px;
        width: 100%;
        height: 110px;
        font-size: 1.05em;
        font-weight: bold;
        box-shadow: 0 4px 16px rgba(0,0,0,0.07);
        margin-bottom: 12px;
        transition: box-shadow 0.18s, background 0.22s;
        cursor: pointer;
    }
    .card-button:hover {
        box-shadow: 0 8px 30px rgba(0,153,255,0.18);
        background: linear-gradient(135deg, #00c6ff 10%, #004488 90%);
    }
    .card-content {
        background: #151d2c;
        color: #fff;
        border-radius: 18px;
        padding: 16px 20px 10px 20px;
        margin-top: -8px;
        min-height: 85px;
        font-size: 0.99em;
        box-shadow: 0 4px 14px rgba(0,153,255,0.07);
        animation: fadeIn 0.38s;
    }
    @keyframes fadeIn { from { opacity:0; } to { opacity:1; } }
    </style>
    """, unsafe_allow_html=True)

    # GRID LAYOUT: 3 x 3
    for row in range(3):
        cols = st.columns(3)
        for col in range(3):
            idx = row * 3 + col
            if idx < len(sections):
                with cols[col]:
                    # Button triggers expansion
                    if st.button(
                        sections[idx]["title"], 
                        key=f"cardbtn_{idx}",
                        help="Click to expand"
                    ):
                        if st.session_state.open_card == idx:
                            st.session_state.open_card = None
                        else:
                            st.session_state.open_card = idx

                    # Apply custom button style via JS
                    st.markdown(f"""
                    <script>
                    let btn = window.parent.document
                      .querySelector('button[data-testid="baseButton-cardbtn_{idx}"]');
                    if(btn) btn.classList.add("card-button");
                    </script>
                    """, unsafe_allow_html=True)

                    # Expanded content if this card is open
                    if st.session_state.open_card == idx:
                        st.markdown(
                            f'<div class="card-content">{sections[idx]["content"]}</div>', 
                            unsafe_allow_html=True
                        )

    # ——— Darken Important Notes text ———
    st.markdown("""
    <style>
    .section-header {
      font-size: 1.5rem;
      font-weight: 700;
      color: #2c2c2c !important;
      margin-top: 2rem;
      margin-bottom: 0.5rem;
    }
    .important-notes {
      color: #333333 !important;
      font-weight: 500 !important;
      line-height: 1.6 !important;
      padding-left: 1rem;
    }
    .important-notes li {
      margin-bottom: 0.5rem;
    }
    </style>
    """, unsafe_allow_html=True)

    # ——— Important Notes Section ———
    st.markdown('<div class="section-header">Important Notes</div>', unsafe_allow_html=True)
    st.markdown("""
    <ul class="important-notes">
      <li>All data and views refresh on app start and can be filtered dynamically.</li>
      <li>Forecasts use simple linear trends based on recent months.</li>
      <li>Some KPIs or tables depend on available columns—if data is missing, you’ll see warnings.</li>
      <li>Uploading your own file replaces the default Oracle data for analysis.</li>
      <li>Time-related metrics require valid time columns and are calculated carefully.</li>
    </ul>
    """, unsafe_allow_html=True)





    
    st.markdown("""
        <style>
        .stButton > button {
            font-size: 1.2em !important;
            font-weight: bold !important;
            border-radius: 16px !important;
            padding: 0.75em 1.4em !important;
            background: linear-gradient(90deg, #00c6ff 0%, #0072ff 100%) !important;
            color: white !important;
            margin-bottom: 1em;
            box-shadow: 0 3px 16px rgba(0,153,255,0.10);
        }
        .stButton > button:hover {
            background: linear-gradient(90deg, #0072ff 0%, #00c6ff 100%) !important;
            color: #fff !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # 2️⃣ Navigation buttons (side by side + Budget)
    st.markdown("---")
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("🚀 Take me to the App"):
            st.session_state.screen = "area_selection"
            st.rerun()

    with col2:
        if st.button("🔎 Quick Overview Summary"):
            st.session_state.screen = "quick_summary"
            st.rerun()

    with col3:
        if st.button("💰 Go to Budget Page"):
            st.session_state.screen    = "budget"
            st.session_state.user_name = "Dan Homewood"  # or real login logic

elif st.session_state.screen == "budget":
    # 0) Initialize override flag
    if "override_alloc" not in st.session_state:
        st.session_state.override_alloc = False

    # —————— ADMIN AUTH ——————
    ADMIN_PW = "Dan"  # change this to something safe!
    if "is_admin" not in st.session_state:
        st.session_state.is_admin = False

    pw = st.text_input("🔒 Admin password to unlock editing:", type="password")
    if pw and pw == ADMIN_PW:
        st.session_state.is_admin = True
        st.success("🔓 Edit mode unlocked!")
    elif pw:
        st.error("❌ Wrong password")

    # -- Local constants for this screen --
    TOTAL_BUDGET = 280_000

    # -- Header & total budget --
    st.markdown(
        f"<h2 style='text-align:center; color:#fff; font-size:2.5rem;'>"
        f"📅 Total Budget (2025/26): £{TOTAL_BUDGET:,}</h2>",
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # -- Load budgets from SQLite --
    if "budgets_df" not in st.session_state:
        with db.get_conn() as conn:
            budgets_df = pd.read_sql(
                "SELECT area AS Area, allocated AS Allocated FROM budgets",
                conn,
                index_col="Area"
            )
        st.session_state.budgets_df = budgets_df
    budgets_df = st.session_state.budgets_df.copy()

    # -- Load expenses from SQLite --
    with db.get_conn() as conn:
        full_exp = pd.read_sql(
            "SELECT * FROM expenses",
            conn,
            parse_dates=["date"]
        )
    full_exp = full_exp.rename(columns={
        "name":        "Name",
        "date":        "Date",
        "area":        "Area",
        "description": "Description",
        "amount":      "Amount"
    })

    # -- Initialize the "current" display log if needed --
    if "current_exp" not in st.session_state:
        st.session_state.current_exp = full_exp.copy()

    # 1) Compute & display total remaining budget
    exp_sum   = st.session_state.current_exp["Amount"].sum()
    remaining = TOTAL_BUDGET - exp_sum
    st.markdown(
        f"<h2 style='text-align:center; color:#62d2a2; font-size:2.5rem;'>"
        f"💷 Budget Remaining: £{remaining:,.0f}</h2>",
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # 2) Prepare Budget Summary table data
    total_by_area = (
        st.session_state.current_exp
        .groupby("Area")["Amount"].sum()
        .reindex(budgets_df.index)
        .fillna(0)
    )
    summary_df = pd.DataFrame({
        "Total Expense": total_by_area,
        "Allocated":    budgets_df["Allocated"]
    })
    summary_df["Remaining"] = summary_df["Allocated"] - summary_df["Total Expense"]
    summary_df.index.name = "Stakeholder"
    pretty = summary_df.applymap(lambda x: f"£{x:,.2f}")

    # 3) Three tabs: Summary / Adjust / Expenses
    tab_sum, tab_adj, tab_exp = st.tabs([
        "🧾 Budget Summary",
        "🔧 Adjust Allocations",
        "💼 Expenses",
    ])

    # ---- TAB 1: Read-only summary ----
    with tab_sum:
        spent = summary_df["Total Expense"].sum()
        remaining_pct = remaining / TOTAL_BUDGET * 100
        c1, c2, c3 = st.columns(3)
        c1.metric("💷 Spent so far", f"£{spent:,.0f}")
        c2.metric("📊 Remaining budget", f"£{remaining:,.0f}")
        c3.metric("📈 % Remaining", f"{remaining_pct:.1f}%")
        st.markdown("#### 1️⃣ Budget Usage by Stakeholder")
        for stakeholder, row in summary_df.iterrows():
            used_amt  = row["Total Expense"]
            alloc_amt = row["Allocated"]
            pct       = min(int(used_amt / alloc_amt * 100), 100)
            color = "#62d2a2" if pct < 70 else "#f0ad4e" if pct < 90 else "#d9534f"
            st.markdown(f"**{stakeholder}** — £{used_amt:,.2f} / £{alloc_amt:,.2f}")
            st.markdown(
                f"""<progress value="{pct}" max="100"
                    style="width:100%; height:1rem; accent-color: {color};">
                </progress>""",
                unsafe_allow_html=True,
            )
        st.markdown("#### 2️⃣ Allocated vs Spent Chart")
        chart_df = summary_df.reset_index()[["Stakeholder", "Allocated", "Total Expense"]]
        st.bar_chart(chart_df.set_index("Stakeholder"))
        st.markdown("---")
        st.table(pretty)

    # ---- TAB 2: +/- allocators ----
    import matplotlib.pyplot as plt
    with tab_adj:
        # 2.1) Allocation controls
        with st.expander("🔧 Adjust Quarterly Allocations", expanded=False):
            updated = False
            alloc_df = budgets_df.copy()
            for area in alloc_df.index:
                c1, c2, c3, c4 = st.columns([3, 1, 2, 1])
                c1.markdown(f"**{area}**")
                if c2.button("–", key=f"dec_{area}"):
                    alloc_df.at[area, "Allocated"] -= 1_000
                    updated = True
                c3.markdown(f"£{alloc_df.at[area,'Allocated']:,.0f}")
                if c4.button("+", key=f"inc_{area}"):
                    alloc_df.at[area, "Allocated"] += 1_000
                    updated = True
            if updated:
                with db.get_conn() as conn:
                    for area, alloc in alloc_df["Allocated"].items():
                        conn.execute(
                            "UPDATE budgets SET allocated = ? WHERE area = ?",
                            (int(alloc), area)
                        )
                st.session_state.budgets_df = alloc_df
                st.rerun()

        # 2.2) Donut chart
        with st.expander("🍩 Allocation Breakdown", expanded=False):
            fig, ax = plt.subplots(figsize=(4,4), facecolor="none")
            sizes = alloc_df["Allocated"]
            labels = alloc_df.index

            wedges, texts, autotexts = ax.pie(
                sizes,
                labels=None,
                autopct="%1.0f%%",
                startangle=90,
                pctdistance=0.85,
                wedgeprops=dict(width=0.3)
            )

            centre_circle = plt.Circle((0,0), 0.70, fc='white')
            ax.add_artist(centre_circle)

            ax.legend(
                wedges,
                labels,
                title="Stakeholder",
                loc="center left",
                bbox_to_anchor=(1, 0, 0.5, 1)
            )

            ax.set_aspect("equal")
            ax.patch.set_alpha(0)
            st.pyplot(fig)

        # 2.3) Detailed table
        with st.expander("📋 Detailed Budget Table", expanded=True):
            st.table(pretty)

    # ---- TAB 3: Expenses management ----
    with tab_exp:
        # Date filter
        st.markdown("📅 Filter Expenses by Date")
        start_date, end_date = st.date_input(
            "Show from → to:",
            value=(pd.Timestamp.today() - pd.Timedelta(days=90),
                   pd.Timestamp.today())
        )
        cur = st.session_state.current_exp.copy()
        cur["Date"] = pd.to_datetime(cur["Date"], errors="coerce")
        mask = (
            (cur["Date"] >= pd.to_datetime(start_date)) &
            (cur["Date"] <= pd.to_datetime(end_date))
        )
        filtered = cur.loc[mask]

        # Add New Expense
        st.markdown("#### ➕ Add New Expense")
        with st.form("exp_form", clear_on_submit=True):
            name = st.text_input("Name", placeholder="e.g. Invoice #1234")
            d1, d2 = st.columns(2)
            date = d1.date_input("Date", value=pd.Timestamp("today").date())
            stakeholder = d2.selectbox("Stakeholder", budgets_df.index.tolist())
            desc = st.text_input("Description")
            amount = st.number_input("Total Value (£)", 0.0, format="%.2f")
            if st.form_submit_button("Add Entry"):
                new = {"Name": name, "Date": date, "Area": stakeholder,
                       "Description": desc, "Amount": amount}
                db.add_expense(new)
                st.session_state.current_exp = pd.concat(
                    [st.session_state.current_exp, pd.DataFrame([new])],
                    ignore_index=True
                )
                st.success("Entry added!")

        # Display filtered entries
        st.markdown("#### 📑 Current Entries")
        disp = filtered.copy()
        disp["Date"] = disp["Date"].dt.strftime("%d-%m-%Y")
        disp["Total Value"] = disp["Amount"].apply(lambda x: f"£{x:,.2f}")
        st.dataframe(
            disp[["Name","Date","Area","Description","Total Value"]]
            .sort_values("Date", ascending=False),
            use_container_width=True
        )

        # Raw Expense Log
        with st.expander("📚 Raw Expense Log", expanded=False):
            raw = full_exp.copy()
            raw["Date"] = raw["Date"].dt.strftime("%d-%m-%Y")
            raw["Total Value"] = raw["Amount"].apply(lambda x: f"£{x:,.2f}")
            st.dataframe(
                raw[["Name","Date","Area","Description","Total Value"]]
                .sort_values("Date", ascending=False),
                use_container_width=True
            )

        # Control buttons
        btn_clear, btn_reset, btn_reapply, btn_reset_all = st.columns(4)
        if btn_clear.button("🗑️ Clear Current Entries"):
            st.session_state.current_exp = full_exp.iloc[0:0].copy()
            st.success("Current entries cleared.")
            st.rerun()
        if btn_reset.button("🔄 Reset Remaining"):
            st.session_state.override_alloc = True
            st.success("Allocations ignored—remaining reset.")
            st.rerun()
        if st.session_state.override_alloc and btn_reapply.button("↩️ Reapply Allocations"):
            st.session_state.override_alloc = False
            st.success("Allocations reapplied.")
            st.rerun()
        if btn_reset_all.button("⚠️ Reset All Expenses"):
            with db.get_conn() as conn:
                conn.execute("DELETE FROM expenses")
            st.session_state.current_exp = pd.DataFrame(
                columns=["Name","Date","Area","Description","Amount"]
            )
            st.success("All expenses cleared from the database!")
            st.rerun()

        # Back button
        if st.button("⬅️ Back to Dashboard"):
            st.session_state.screen = "dashboard"

















    















      




# ✅ Optimized Dataset Loading Block

import pandas as pd
import streamlit as st
import re

# --- Oracle Files ---
oracle_files = {
    "VIP North": "VIP North Oracle Data.xlsx",
    "VIP South": "VIP South Oracle Data.xlsx",
    "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South": "Tier 2 South Oracle Data.xlsx"
}

@st.cache_data(show_spinner=False)
def load_oracle_files(oracle_files):
    dataframes = []
    for team_name, file in oracle_files.items():
        try:
            df = pd.read_excel(file)
            df["Team"] = team_name
            df.columns = df.columns.str.strip()
            dataframes.append(df)
        except Exception as e:
            st.warning(f"⚠️ Could not load {file}: {e}")
    return dataframes

oracle_dataframes = load_oracle_files(oracle_files)

combined_oracle_df = pd.concat(oracle_dataframes, ignore_index=True) if oracle_dataframes else pd.DataFrame()

# --- Subsets from Oracle ---
vip_north_df = combined_oracle_df[combined_oracle_df["Team"] == "VIP North"]
vip_south_df = combined_oracle_df[combined_oracle_df["Team"] == "VIP South"]
tier2_north_df = combined_oracle_df[combined_oracle_df["Team"] == "Tier 2 North"]
tier2_south_df = combined_oracle_df[combined_oracle_df["Team"] == "Tier 2 South"]

sky_retail_df = combined_oracle_df[
    combined_oracle_df.get("Sky Retail Stakeholder", "").astype(str).str.contains("Sky Retail", case=False, na=False)
] if "Sky Retail Stakeholder" in combined_oracle_df.columns else pd.DataFrame()

highlands_df = combined_oracle_df[
    combined_oracle_df["Visit Type"].astype(str).str.contains("Highlands|Islands", case=False, na=False)
] if "Visit Type" in combined_oracle_df.columns else pd.DataFrame()

# --- Other Datasets ---
@st.cache_data(show_spinner=False)
def load_excel(path):
    return pd.read_excel(path)

try:
    sky_business_df = load_excel("AI Test SB Visits.xlsx")
    sky_business_df.columns = sky_business_df.columns.str.strip()
except Exception as e:
    st.warning(f"⚠️ Could not load Sky Business file: {e}")
    sky_business_df = pd.DataFrame()

try:
    call_log_df = load_excel("Call Log Data.xlsx")
    call_log_df.columns = call_log_df.columns.str.strip()
except Exception as e:
    st.warning(f"⚠️ Could not load Call Log Data.xlsx: {e}")
    call_log_df = pd.DataFrame()

# --- Highlands File ---
try:
    highlands_file = pd.ExcelFile("Highlands Islands.xlsx")
    yearly_sheets = [s for s in highlands_file.sheet_names if "Year" in s]
    company_sheets = [s for s in highlands_file.sheet_names if "Company" in s]
    dfs = {sheet: highlands_file.parse(sheet) for sheet in yearly_sheets + company_sheets}
except Exception as e:
    st.error(f"⚠️ Error loading Highlands & Islands data: {e}")
    highlands_file = None
    yearly_sheets = []
    dfs = {}

# --- Combined Dictionary ---
all_tabs = {
    "VIP North": vip_north_df,
    "VIP South": vip_south_df,
    "Tier 2 North": tier2_north_df,
    "Tier 2 South": tier2_south_df,
    "Sky Retail": sky_retail_df,
    "Sky Business": sky_business_df,
    "Call Log": call_log_df,
    "Highlands & Islands": highlands_df,
}

# --- Time Formatting Functions ---
def format_time_avg(series):
    clean = series.dropna().astype(str)
    clean = clean[~clean.str.contains("00:00:00|00:00|NaT|nan|None", na=False)]
    times = pd.to_timedelta(clean, errors='coerce').dropna()
    return str((times.mean()).components.hours).zfill(2) + ":" + str((times.mean()).components.minutes).zfill(2) if not times.empty else "00:00"

def format_total_time(series):
    times = pd.to_timedelta(series.dropna().astype(str), errors='coerce')
    total = times.sum()
    return f"{int(total.total_seconds() // 3600):02}:{int((total.total_seconds() % 3600) // 60):02}" if not times.empty else "00:00"
# --- Utility for Highlands weighted averages ---
def weighted_avg(df, value_col, weight_col):
    try:
        valid = df[weight_col] > 0
        return (df.loc[valid, value_col] * df.loc[valid, weight_col]).sum() / df.loc[valid, weight_col].sum()
    except Exception:
        return None


if  st.session_state.screen == "quick_summary":
    st.title("🔎 Quick Overview Summary")

    # 1️⃣ Dataset dropdown
    selected_tab = st.selectbox(
        "Choose dataset:",
        list(all_tabs.keys()),
        key="quick_summary_selectbox"
    )
    df = all_tabs[selected_tab]

    # 2️⃣ Navigation buttons (side by side)
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Team Area", key=f"team_area_btn_{selected_tab}"):
            st.session_state.selected_team_area = selected_tab
            # route into the correct screen for detailed charts
            if selected_tab == "Sky Retail":
                st.session_state.screen = "sky_retail"
            elif selected_tab == "Sky Business":
                st.session_state.screen = "sky_business"
            elif selected_tab == "Call Log":
                st.session_state.screen = "dashboard_view"
                st.session_state.selected_dataset = "Call Log Data"
            elif selected_tab == "Highlands & Islands":
                st.session_state.screen = "highlands_islands"
            else:
                st.session_state.screen = "manager_chart_select"
            st.rerun()
    with col2:
        if st.button("🌍 Explore Full App", key=f"explore_app_btn_{selected_tab}"):
            st.session_state.screen = "area_selection"
            st.rerun()

    # 3️⃣ Quick KPIs & Metrics (Sky Business, Highlands & Islands, Call Log, or Generic)
    if df.empty and selected_tab != "Highlands & Islands":
        st.warning(f"No data available for **{selected_tab}**.")
    else:
        if selected_tab == "Sky Business":
            # --- Sky Business custom summary ---
            total_visits     = len(df)
            unique_engineers = df["Business Engineers Name"].nunique() if "Business Engineers Name" in df.columns else 0

            df["Total Value"] = pd.to_numeric(df.get("Total Value", pd.Series(dtype="float")), errors="coerce")
            total_value       = df["Total Value"].sum()
            avg_value         = df["Total Value"].mean()

            r1c1, r1c2 = st.columns(2)
            r2c1, r2c2 = st.columns(2)
            r1c1.metric("Total Visits",        total_visits)
            r1c2.metric("Unique Engineers",    unique_engineers)
            r2c1.metric("Total Value (£)",     f"£{total_value:,.0f}")
            r2c2.metric("Avg Visit Value (£)", f"£{avg_value:,.0f}")

        elif selected_tab == "Highlands & Islands":
            # --- Highlands & Islands custom summary ---
            def safe_sum(sheet, col):
                if col not in dfs[sheet].columns:
                    return 0.0
                return pd.to_numeric(dfs[sheet][col], errors="coerce").sum()

            sheets = yearly_sheets + company_sheets
            total_issued    = sum(safe_sum(s, "Issued Visits")    for s in sheets)
            total_completed = sum(safe_sum(s, "Completed Visits") for s in sheets)
            total_not_done  = sum(safe_sum(s, "Not Done Vists")   for s in sheets)
            total_rev7      = sum(safe_sum(s, "7 Day Revisits")   for s in sheets)
            total_surveys   = sum(safe_sum(s, "Surveys")          for s in sheets)

            completion_pct  = (total_completed / total_issued * 100) if total_issued else 0.0
            not_done_pct    = (total_not_done  / total_issued * 100) if total_issued else 0.0

            concat_df       = pd.concat([dfs[s] for s in sheets], ignore_index=True)
            avg_job_min     = weighted_avg(concat_df, "Average Complete Job Time (Min)", "Completed Visits")
            weighted_nps    = weighted_avg(concat_df, "NPS", "Surveys")

            r1c1, r1c2, r1c3, r1c4 = st.columns(4)
            r1c1.metric("Issued Visits",       f"{int(total_issued):,}")
            r1c2.metric("Completed Visits",    f"{int(total_completed):,}")
            r1c3.metric("Completion Rate (%)", f"{completion_pct:.1f}%")
            r1c4.metric("Not Done (%)",        f"{not_done_pct:.1f}%")

            r2c1, r2c2, r2c3, r2c4 = st.columns(4)
            r2c1.metric("7‑Day Revisits",     f"{int(total_rev7):,}")
            r2c2.metric("Avg Job Time (min)", f"{avg_job_min:.1f}" if avg_job_min is not None else "N/A")
            r2c3.metric("Total Surveys",      f"{int(total_surveys):,}")
            r2c4.metric("Weighted NPS",       f"{weighted_nps:.1f}" if weighted_nps is not None else "N/A")

        elif selected_tab == "Call Log":
            # --- Call Log custom summary ---
            total_calls    = len(df)
            unique_engs    = df["Name of Engineer"].nunique() if "Name of Engineer" in df.columns else 0
            unique_regions = df["Region"].nunique()         if "Region" in df.columns else 0
            unique_opts    = df["Option Selected"].nunique() if "Option Selected" in df.columns else 0
            unique_months  = df["Month"].nunique()          if "Month" in df.columns else 0
            unique_vr      = df["VR Number (If Known)"].nunique()                     if "VR Number (If Known)" in df.columns else 0
            unique_callers = df["Name Of Engineer Who Made The Call"].nunique()        if "Name Of Engineer Who Made The Call" in df.columns else 0
            unique_emails  = df["Engineers email address (who made the call)"].nunique() if "Engineers email address (who made the call)" in df.columns else 0

            # coerce date column before using .dt
            if "Date of Call Taken" in df.columns:
                df["Date of Call Taken"] = pd.to_datetime(
                    df["Date of Call Taken"], errors="coerce"
                )

            # busiest day
            peak_day, peak_count = "N/A", 0
            if "Date of Call Taken" in df.columns:
                calls_per_day = df["Date of Call Taken"].dt.date.value_counts()
                if not calls_per_day.empty:
                    peak_day   = calls_per_day.idxmax().strftime("%d %b %Y")
                    peak_count = calls_per_day.max()

            # layout in a 3×3 grid
            r1c1, r1c2, r1c3 = st.columns(3)
            r2c1, r2c2, r2c3 = st.columns(3)
            r3c1, r3c2, r3c3 = st.columns(3)

            r1c1.metric("Total Calls",          f"{total_calls:,}")
            r1c2.metric("Unique Engineers",     unique_engs)
            r1c3.metric("Unique Regions",       unique_regions)

            r2c1.metric("Options Selected",     unique_opts)
            r2c2.metric("Distinct Months",      unique_months)
            r2c3.metric("Known VR Numbers",     unique_vr)

            r3c1.metric("Caller Engineers",     unique_callers)
            r3c2.metric("Engineer Emails",      unique_emails)
            r3c3.metric("Peak Call Day",        f"{peak_day} ({peak_count})")


        else:
            # --- Generic summary for Oracle / Sky Retail ---
            total_visits = len(df)

            def count_status(col, status):
                return df[col].astype(str).str.contains(status, case=False, na=False).sum() if col in df.columns else 0

            completed   = count_status("Activity Status", "Completed")
            cancelled   = count_status("Activity Status", "Cancelled")
            suspended   = count_status("Activity Status", "Suspended")
            not_done    = count_status("Activity Status", "Not Done")
            pending     = count_status("Activity Status", "Pending")
            started     = count_status("Activity Status", "Started")
            unique_engs = df["Name"].nunique() if "Name" in df.columns else 0

            total_time     = format_total_time(df["Total Time"])           if "Total Time" in df.columns         else "00:00"
            avg_work       = format_time_avg(df["Total Working Time"])     if "Total Working Time" in df.columns else "00:00"
            avg_activate   = format_time_avg(df["Activate"])               if "Activate" in df.columns           else "00:00"
            avg_deactivate = format_time_avg(df["Deactivate"])             if "Deactivate" in df.columns         else "00:00"
            lunch_df       = df[df["Visit Type"].str.contains("Lunch", case=False, na=False)] if "Visit Type" in df.columns else pd.DataFrame()
            avg_lunch      = format_time_avg(lunch_df["Total Time"])       if not lunch_df.empty and "Total Time" in lunch_df.columns else "00:00"

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Total Visits",       total_visits)
            c2.metric("Completed",          completed)
            c3.metric("Cancelled",          cancelled)
            c4.metric("Suspended",          suspended)

            c5, c6, c7, c8 = st.columns(4)
            c5.metric("Not Done",           not_done)
            c6.metric("Pending",            pending)
            c7.metric("Started",            started)
            c8.metric("Unique Engineers",   unique_engs)

            c9, c10, c11, c12 = st.columns(4)
            c9.metric("Total Time",         total_time)
            c10.metric("Avg Working Time",  avg_work)
            c11.metric("Avg Activate",      avg_activate)
            c12.metric("Avg Deactivate",    avg_deactivate)

            st.metric("Avg Lunch (30)",     avg_lunch)







    



if st.session_state.screen == "manager_chart_select":
    if st.button("⬅️ Back to Quick Summary"):
        st.session_state.screen = "quick_summary"
        st.rerun()

    selected_team = st.session_state.get("selected_team_area", None)
    chart_options = {}

    if selected_team in ["VIP North", "VIP South", "Tier 2 North", "Tier 2 South"]:
        st.title(f"👥 {selected_team} — Team Area")
        df = all_tabs[selected_team]
        # … build chart_options …

    elif selected_team == "Sky Retail":
        st.title("👥 Sky Retail — Team Area")
        # … your Sky Retail logic …

    elif selected_team == "Sky Business":
        st.title("👥 Sky Business — Team Area")
        # … your Sky Business logic …

    elif selected_team == "Call Log":
        st.title("📞 Call Log — Team Area")
        # … your Call Log logic …

    elif selected_team == "Highlands & Islands":
        # ← NEW!
        st.session_state.screen = "highlands_islands"
        st.rerun()

    else:
        st.info("Please select a valid team to see options.")
        st.stop()

    # … now your chart_options multiselect and rendering …


        chart_options = {
            "kpis_block": "📈 Team KPIs",
            "summary_kpis": "📋 Summary KPIs",
            "top_visit_types": "📊 Top Visit Types",
            "top_engineers": "👨 Top Engineers by Visits",
            "weekly_trends": "📈 Weekly Visit Trends",
            "status_breakdown": "🟦 Status Breakdown Bar",
            "engineer_pie": "🥧 Engineer Visit Share (Pie)",
            "daily_trend": "📅 Daily Visits Trend",
            "dow_bar": "📊 Visits by Day of Week",
            "value_per_type": "💷 Avg Value per Visit Type",
            "stakeholder": "🏢 Visits by Stakeholder",
            "week_heatmap": "🔥 DOW vs Week Heatmap",
            "monthly_value": "💰 Total Value by Month",
            "sunburst": "🌞 Sunburst Charts (all)",
            "stacked_bar": "📊 Stacked Bar: Visit Type by Month",
            "parallel_categories": "🔗 Parallel Categories: Engineer → Visit Type → Postcode",
            "treemap": "🌲 Drilldown Treemap: Stakeholder → Visit Type → Month",
            "heatmap": "🌡️ Visit Type vs Week Heatmap",
            "visit_type_treemap": "🌳 Treemap: Visit Type by Total Value",
            "pie_chart": "🥧 Visit Type Share (Pie)",
            "table_view": "📋 Full Oracle Visit Table"
        }

    if selected_team == "Sky Retail":
        st.title("👥 Sky Retail — Team Area")
        st.info("Click below to view the full Sky Retail dashboard:")

        if st.button("🔎 Go to Sky Retail Dashboard"):
            st.session_state.screen = "sky_retail"
            st.rerun()

        

    # Only show the dropdown if a valid team is selected
    if chart_options:
        selected_charts = st.multiselect(
            "Select which charts/tables to display:",
            options=list(chart_options.keys()),
            format_func=lambda k: chart_options[k],
            key="team_chart_multiselect"
        )
        for chart_key in selected_charts:
            if chart_key == "kpi":
                # Stakeholder KPIs (put your expander logic here)
                with st.expander("📊 Stakeholder KPIs", expanded=True):
                    # (insert your KPI code block here for the selected stakeholder)
                    pass

            elif chart_key == "monthly_visits":
                # Monthly Visits Table
                st.markdown("**Monthly Visits (last 4 months shown)**")
                show_aggrid(monthly_visits.tail(4))

            elif chart_key == "monthly_value":
                st.markdown("**Monthly Total Value (£) (last 4 months shown)**")
                show_aggrid(monthly_value.tail(4))

            elif chart_key == "visits_by_dow":
                st.markdown("**Visits by Day of Week**")
                show_aggrid(visits_by_day)

            elif chart_key == "value_by_dow":
                st.markdown("**Total Value (£) by Day of Week**")
                show_aggrid(value_by_day)

            elif chart_key == "stakeholder_breakdown":
                with st.expander(f"📋 Breakdown by {stakeholder} Stakeholder", expanded=False):
                    st.dataframe(by_stake, use_container_width=True)

            elif chart_key == "monthly_visit_trends":
                with st.expander(f"📈 Monthly Visit Trends for {stakeholder}", expanded=False):
                    st.plotly_chart(fig_visits, use_container_width=True)

            elif chart_key == "monthly_value_trends":
                with st.expander(f"📈 Monthly Value Trends for {stakeholder}", expanded=False):
                    st.plotly_chart(fig_value, use_container_width=True)

            elif chart_key == "activity_status_pie":
                with st.expander(f"🥧 Visit Activity Status Split for {stakeholder}", expanded=False):
                    st.plotly_chart(fig_pie, use_container_width=True)

            elif chart_key == "team_breakdown":
                with st.expander("📋 Team Breakdown by Stakeholder", expanded=False):
                    st.dataframe(team_pivot, use_container_width=True)

            elif chart_key == "engineer_breakdown":
                with st.expander(f"👤 Engineer Breakdown for {stakeholder}", expanded=False):
                    st.dataframe(display, use_container_width=True)

            elif chart_key == "overall_forecasts":
                with st.expander(f"🔮 Overall {stakeholder} Forecasts", expanded=False):
                    st.plotly_chart(fig_visits_fc, use_container_width=True)
                    st.plotly_chart(fig_value_fc, use_container_width=True)

            elif chart_key == "team_forecasts":
                with st.expander(f"🔮 {stakeholder} Team Forecasts (Visits & Value)", expanded=False):
                    st.plotly_chart(fig_team_visits, use_container_width=True)
                    st.plotly_chart(fig_team_value, use_container_width=True)

            elif chart_key == "mom_change_visits":
                with st.expander("📊 Month-on-Month Change Table (Visits)", expanded=False):
                    st.dataframe(visits_tbl, use_container_width=True)

            elif chart_key == "mom_change_value":
                with st.expander("📊 Month-on-Month Change Table (Value)", expanded=False):
                    st.dataframe(value_tbl, use_container_width=True)

            elif chart_key == "combined_charts":
                with st.expander("📊 Combined Monthly Trend Charts", expanded=False):
                    st.plotly_chart(fig_value_trends, use_container_width=True)
                    st.plotly_chart(fig_visits_trends, use_container_width=True)
                    st.plotly_chart(fig_value_changes, use_container_width=True)
                    st.plotly_chart(fig_visits_changes, use_container_width=True)
                    st.plotly_chart(fig_heatmap_value_max, use_container_width=True)
                    st.plotly_chart(fig_heatmap_visits_max, use_container_width=True)

            elif chart_key == "dow_bar_chart":
                with st.expander("📅 Visits by Day of Week", expanded=False):
                    st.plotly_chart(fig_day, use_container_width=True)

            elif chart_key == "raw_data":
                with st.expander(f"🔎 Show Raw Data for {stakeholder}", expanded=False):
                    st.dataframe(df.dropna(axis=1, how="all"), use_container_width=True)
    
        st.markdown(f"### KPIs for {selected_team}")
    else:
        st.info("Please select a valid team to see options.")

        #st.write(df.head())  # Just shows the data as a table for now

    import datetime, pandas as pd

    def to_timedelta_str(x):
        if pd.isnull(x) or x in ["", "-", "NaT", None, " "]:
            return pd.NaT
        if isinstance(x, (pd.Timedelta, datetime.timedelta)):
            return x
        if isinstance(x, datetime.time):
            return datetime.timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, (float, int)):
            try:
                return datetime.timedelta(seconds=int(float(x) * 86400))
            except Exception:
                return pd.NaT
        if isinstance(x, pd.Timestamp):
            t = x.time()
            return datetime.timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
        if isinstance(x, str):
            s = x.strip()
            if s in ["", "-", "NaT", " "]:
                return pd.NaT
            try:
                h, m, s = map(int, s.split(":")[:3])
                return datetime.timedelta(hours=h, minutes=m, seconds=s)
            except Exception:
                pass
            try:
                return pd.to_timedelta(s)
            except Exception:
                return pd.NaT
        return pd.NaT

    def kpis_block(df):
        if df.empty:
            st.warning("No data for the current selection.")
            return

        local_df = df.copy()
        local_df.columns = local_df.columns.str.strip()
        local_df = local_df.rename(columns={"Name": "Engineer", "Total Value": "Value"})

        # KPIs
        total_visits   = len(local_df)
        unique_engs    = local_df["Engineer"].nunique() if "Engineer" in local_df.columns else "N/A"
        total_value    = f"£{local_df['Value'].sum(skipna=True):,.2f}" if "Value" in local_df.columns else "N/A"

        # Average Activate/Deactivate
        valid_times = local_df.copy()
        if {"Activate", "Deactivate"}.issubset(valid_times.columns):
            valid_times["Activate"]   = valid_times["Activate"].apply(to_timedelta_str)
            valid_times["Deactivate"] = valid_times["Deactivate"].apply(to_timedelta_str)
            valid_times = valid_times[
                valid_times["Activate"].notna() &
                valid_times["Deactivate"].notna() &
                (valid_times["Activate"]   > datetime.timedelta(0)) &
                (valid_times["Deactivate"] > datetime.timedelta(0))
            ]
        else:
            valid_times = pd.DataFrame()

        def avg_time_str(col):
            if col not in valid_times.columns or valid_times.empty:
                return "N/A"
            vals = valid_times[col].dropna()
            if vals.empty:
                return "N/A"
            avg = vals.mean()
            return f"{int(avg.total_seconds()//3600):02}:{int((avg.total_seconds()%3600)//60):02}"

        avg_activate   = avg_time_str("Activate")
        avg_deactivate = avg_time_str("Deactivate")

        # Most Common Visit Type (not lunch)
        most_common_type = "N/A"
        if "Visit Type" in local_df.columns:
            vt = local_df[~local_df["Visit Type"].str.contains("lunch", case=False, na=False)]
            if not vt["Visit Type"].mode().empty:
                most_common_type = vt["Visit Type"].mode()[0]

        # Busiest day
        busiest_day, busiest_count = "N/A", ""
        if "Date" in local_df.columns:
            counts = pd.to_datetime(local_df["Date"], errors="coerce").dt.date.value_counts()
            if not counts.empty:
                busiest_day   = counts.idxmax().strftime("%d %B %Y")
                busiest_count = f"{counts.max()} visits"

        # --- DISPLAY AS BULLET POINTS ---
        
        st.markdown("### 📈 Key Performance Indicators (KPIs)")
        st.markdown(f"""
        - **Total Visits:** {total_visits}
        - **Unique Engineers:** {unique_engs}
        - **Total Value:** {total_value}
        - **Most Common Visit Type:** {most_common_type}
        - **Avg Activate:** {avg_activate}
        - **Avg Deactivate:** {avg_deactivate}
        - **Busiest Day:** {busiest_day} ({busiest_count})
        """)
    if selected_team == "Highlands & Islands":
        st.title("🗺️ Highlands & Islands — Team Area")
        st.markdown("### 📈 Key Performance Indicators (KPIs)")
        st.markdown(f"""
        - **Total Visits:** {total_visits}
        - **Unique Engineers:** {unique_engs}
        - **Total Value:** {total_value}
        - **Most Common Visit Type:** {most_common_type}
        - **Avg Activate:** {avg_activate}
        - **Avg Deactivate:** {avg_deactivate}
        - **Busiest Day:** {busiest_day} ({busiest_count})
        """)

    if st.button("Show Me Charts"):
        st.session_state.selected_team_charts = selected_charts
        st.rerun()
    selected = st.session_state.get("selected_team_charts", [])
    
    for chart_key in selected:
        if chart_key == "summary_kpis":
            with st.expander("📋 Summary KPIs", expanded=True):
                col1, col2, col3, col4 = st.columns(4)
                col1.metric("Total Visits", len(df))
                col2.metric("Unique Engineers", df["Name"].nunique())
                col3.metric("Visit Types", df["Visit Type"].nunique())
                col4.metric("Total Value (£)", f"£{df['Total Value'].dropna().sum():,.2f}")

        elif chart_key == "kpis_block":
            kpis_block(df)

        elif chart_key == "top_visit_types":
            if "Visit Type" in df.columns:
                with st.expander("📊 Top Visit Types"):
                    top_vt = df["Visit Type"].value_counts().head(10).reset_index()
                    top_vt.columns = ["Visit Type", "Count"]
                    st.plotly_chart(px.bar(top_vt, x="Visit Type", y="Count", color="Visit Type",
                                        title="Top Visit Types by Volume"), use_container_width=True)

        elif chart_key == "top_engineers":
            if "Name" in df.columns:
                with st.expander("👨 Top Engineers by Visits"):
                    eng_top = df["Name"].value_counts().head(10).reset_index()
                    eng_top.columns = ["Engineer", "Visits"]
                    st.plotly_chart(px.bar(eng_top, x="Engineer", y="Visits", color="Engineer",
                                        title="Top Engineers by Visit Count"), use_container_width=True)

        elif chart_key == "weekly_trends":
            if "Week" in df.columns:
                with st.expander("📈 Weekly Visit Trends"):
                    weekly = df.groupby("Week").size().reset_index(name="Visits")
                    st.plotly_chart(px.line(weekly, x="Week", y="Visits", title="Visits Over Weeks"),
                                    use_container_width=True)

        elif chart_key == "monthly_value":
            if "Month" in df.columns and "Total Value" in df.columns:
                with st.expander("💰 Total Value by Month"):
                    monthly_val = df.groupby("Month")["Total Value"].sum().reset_index()
                    st.plotly_chart(px.bar(monthly_val, x="Month", y="Total Value", color="Month",
                                        title="Total Value by Month"), use_container_width=True)

        elif chart_key == "sunburst":
            sunburst_configs = [
                ("🌞 Visit Activity Sunburst", ["Visit Type", "Activity Status"], "Visit Type & Activity Status Distribution"),
                ("📍 Sunburst: Visit Type to Postcode", ["Visit Type", "Postcode"], "Visit Type → Postcode Distribution"),
                ("🔀 Sunburst: Engineer → Visit Type → Week", ["Name", "Visit Type", "Week"], "Engineer > Visit Type > Week Breakdown"),
                ("🌀 Sunburst: Visit Type → Week", ["Visit Type", "Week"], "Visit Count by Visit Type and Week"),
                ("🧩 Sunburst: Engineer → Postcode", ["Name", "Postcode"], "Engineer > Postcode Mapping"),
                ("🗓️ Sunburst: Visit Type → Month", ["Visit Type", "Month"], "Visit Type Distribution by Month"),
                ("📅 Sunburst: Visit Type → Date", ["Visit Type", "Date"], "Visit Type Breakdown by Exact Date"),
                ("📋 Sunburst: Visit Type → Day", ["Visit Type", "Day"], "Visit Type by Day of Week"),
                ("📑 Sunburst: Stakeholder → Visit Type", ["Sky Retail Stakeholder", "Visit Type"], "Stakeholder to Visit Type Breakdown")
            ]
            for label, cols, title in sunburst_configs:
                if set(cols).issubset(df.columns):
                    with st.expander(label):
                        sb = df.groupby(cols).size().reset_index(name="Count")
                        fig = px.sunburst(sb, path=cols, values="Count", title=title)
                        st.plotly_chart(fig, use_container_width=True)
        elif chart_key == "status_breakdown":
            if "Activity Status" in df.columns:
                with st.expander("🟦 Status Breakdown Bar"):
                    st.markdown("Volume of visits by status (Completed, Cancelled, Not Done)")
                    bar = df["Activity Status"].value_counts().reset_index()
                    bar.columns = ["Status", "Count"]
                    st.plotly_chart(px.bar(bar, x="Status", y="Count", color="Status",
                                        title="Visit Status Counts"), use_container_width=True)
        elif chart_key == "engineer_pie":
            if "Name" in df.columns:
                with st.expander("🥧 Engineer Visit Share (Pie)"):
                    pie = df["Name"].value_counts().reset_index()
                    pie.columns = ["Engineer", "Visits"]
                    fig = px.pie(pie, names="Engineer", values="Visits", title="Visits by Engineer")
                    st.plotly_chart(fig, use_container_width=True)
                    
        elif chart_key == "daily_trend":
            if "Date" in df.columns:
                with st.expander("📅 Daily Visits Trend"):
                    daily = df.groupby(df["Date"].dt.date).size().reset_index(name="Visits")
                    st.plotly_chart(px.line(daily, x="Date", y="Visits", title="Visits Per Day"), use_container_width=True)

        elif chart_key == "dow_bar":
            if "Date" in df.columns:
                with st.expander("📊 Visits by Day of Week"):
                    dow = df.copy()
                    dow["Day"] = pd.to_datetime(dow["Date"], errors="coerce").dt.day_name()
                    day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
                    chart = (dow["Day"].value_counts().reindex(day_order).reset_index())
                    chart.columns = ["Day", "Visits"]
                    st.plotly_chart(px.bar(chart, x="Day", y="Visits", title="Visits by Day of Week"), use_container_width=True)

        elif chart_key == "value_per_type":
            if "Visit Type" in df.columns and "Value" in df.columns:
                with st.expander("💷 Avg Value per Visit Type"):
                    val_df = df.groupby("Visit Type")["Value"].mean().reset_index()
                    st.plotly_chart(px.bar(val_df, x="Visit Type", y="Value", title="Avg Value by Visit Type"), use_container_width=True)

        elif chart_key == "stakeholder":
            if "Sky Retail Stakeholder" in df.columns:
                with st.expander("🏢 Visits by Stakeholder"):
                    stake = df["Sky Retail Stakeholder"].value_counts().reset_index()
                    stake.columns = ["Stakeholder", "Visits"]
                    st.plotly_chart(px.bar(stake, x="Stakeholder", y="Visits", title="Visits by Stakeholder"), use_container_width=True)


        elif chart_key == "stacked_bar":
            if {"Visit Type", "Month"}.issubset(df.columns):
                with st.expander("📊 Stacked Bar: Visit Type by Month"):
                    bar_df = df.groupby(["Month", "Visit Type"]).size().reset_index(name="Visits")
                    fig = px.bar(bar_df, x="Month", y="Visits", color="Visit Type", title="Monthly Visit Counts by Visit Type",
                                text_auto=True)
                    st.plotly_chart(fig, use_container_width=True)

        elif chart_key == "parallel_categories":
            if {"Name", "Visit Type", "Postcode"}.issubset(df.columns):
                with st.expander("🔗 Parallel Categories: Engineer → Visit Type → Postcode"):
                    pc_df = df[["Name", "Visit Type", "Postcode"]].dropna().astype(str)
                    fig = px.parallel_categories(pc_df, dimensions=["Name", "Visit Type", "Postcode"],
                                                color_continuous_scale=px.colors.sequential.Inferno,
                                                title="Engineer to Visit Type to Postcode Flow")
                    st.plotly_chart(fig, use_container_width=True)

        elif chart_key == "treemap":
            if {"Sky Retail Stakeholder", "Visit Type", "Month", "Total Value"}.issubset(df.columns):
                with st.expander("🌲 Drilldown Treemap: Stakeholder → Visit Type → Month"):
                    tree_df = df.groupby(["Sky Retail Stakeholder", "Visit Type", "Month"])["Total Value"].sum().reset_index()
                    fig = px.treemap(tree_df, path=["Sky Retail Stakeholder", "Visit Type", "Month"],
                                    values="Total Value", title="Value Drilldown by Stakeholder → Visit Type → Month")
                    st.plotly_chart(fig, use_container_width=True)

        elif chart_key == "heatmap":
            if {"Visit Type", "Week"}.issubset(df.columns):
                with st.expander("🌡️ Visit Type vs Week Heatmap"):
                    heat_df = pd.pivot_table(df, index="Visit Type", columns="Week", aggfunc="size", fill_value=0)
                    st.plotly_chart(px.imshow(heat_df, aspect="auto", title="Visit Heatmap: Types by Week"),
                                    use_container_width=True)

        elif chart_key == "visit_type_treemap":
            if {"Visit Type", "Total Value"}.issubset(df.columns):
                with st.expander("🌳 Treemap: Visit Type by Total Value"):
                    tm = df.groupby("Visit Type")["Total Value"].sum().reset_index()
                    fig = px.treemap(tm, path=["Visit Type"], values="Total Value",
                                    title="Total Value by Visit Type")
                    st.plotly_chart(fig, use_container_width=True)

        elif chart_key == "pie_chart":
            if "Visit Type" in df.columns:
                with st.expander("🥧 Visit Type Share (Pie)"):
                    pie = df["Visit Type"].value_counts().reset_index()
                    pie.columns = ["Visit Type", "Count"]
                    fig = px.pie(pie, names="Visit Type", values="Count", title="Visit Type Distribution")
                    st.plotly_chart(fig, use_container_width=True)

        elif chart_key == "week_heatmap":
            if "Date" in df.columns:
                with st.expander("🔥 Day vs Week Heatmap"):
                    temp = df.copy()
                    temp["Week"] = pd.to_datetime(temp["Date"], errors="coerce").dt.isocalendar().week
                    temp["Day"] = pd.to_datetime(temp["Date"], errors="coerce").dt.day_name()
                    pt = pd.pivot_table(temp, index="Day", columns="Week", aggfunc="size", fill_value=0)
                    st.plotly_chart(px.imshow(pt, title="Visits by Day and Week"), use_container_width=True)

        elif chart_key == "table_view":
            with st.expander("📋 Full Oracle Visit Table", expanded=False):
                st.dataframe(df, use_container_width=True)

            # (Add more charts, filters, summaries, etc. here!)
        else:
            st.warning("No team selected! Please go back and pick a dataset.")
 




# --- NUMBER 6 ---#
# --- SECTION: AREA SELECTION MAIN MENU ---
if st.session_state.screen == "area_selection":

    st.markdown(f"""
    <div class="logo">
        <img src='data:image/png;base64,{logo_base64}' />
    </div>
    """, unsafe_allow_html=True)

    # --- TYPING INTRO TEXT ---
    st.markdown("""
    <style>
    .adv-summary {
        font-size: 20px;
        font-weight: 400;
        border-right: 2px solid rgba(255,255,255,0.75);
        white-space: nowrap;
        overflow: hidden;
        width: fit-content;
        margin: 0 auto 30px;
        animation: typing 6s steps(60, end) infinite, blink 0.75s step-end infinite;
    }
    </style>
    <div class='adv-summary'>
    Welcome to the advanced reporting hub use the options below to explore all areas
    </div>
    """, unsafe_allow_html=True)
    if st.button("⬅️ Back to User Guide & Overview"):
        st.session_state.screen = "instruction_guide"
        st.rerun()
def menu_button_with_tooltip(label, tooltip_text, screen_name, key=None): 
    col1, col2 = st.columns([10,1], gap="small")  # wider button, narrow icon
    with col1:
        if st.button(label, use_container_width=True, key=key):
            st.session_state.screen = screen_name
            st.rerun()
    with col2:
        st.markdown(
            f"""
            <div style="position: relative; display: inline-block;">
              <span style="
                cursor: help; 
                font-weight: bold; 
                color: #007bff; 
                font-size: 18px;"
                title="{tooltip_text}">
                ⓘ
              </span>
            </div>
            """,
            unsafe_allow_html=True,
        )

if st.session_state.screen == "area_selection":
    st.markdown("## Choose an area", unsafe_allow_html=True)

    row1_cols = st.columns(3)
    with row1_cols[0]:
        menu_button_with_tooltip("🏢 Operational Area", "Manage field operations and schedules", "operational_area", key="btn1")
    with row1_cols[1]:
        menu_button_with_tooltip("📊 Dashboard Area", "View KPIs and trends", "dashboard", key="btn2")
    with row1_cols[2]:
        menu_button_with_tooltip("🤖 Sky Orbit", "Ask questions and get AI-powered insights", "ai", key="btn3")

    st.write("")

    row2_cols = st.columns(3)
    with row2_cols[0]:
        menu_button_with_tooltip("💡 Suggestion Box", "Submit and view suggestions", "suggestions", key="btn4")
    with row2_cols[1]:
        menu_button_with_tooltip("📈 Forecasts", "View forecasted visits and values", "Forecasts", key="btn5")
    with row2_cols[2]:
        menu_button_with_tooltip("🗺️ Highlands & Islands", "Explore Highlands & Islands area", "highlands_islands", key="btn6")

    row3_cols = st.columns(3)
    with row3_cols[0]:
        menu_button_with_tooltip("🏬 Sky Retail", "Analyze Sky Retail visit data and KPIs", "sky_retail", key="btn7")
    with row3_cols[1]:
        menu_button_with_tooltip("🏢 Sky Business", "Analyze Sky Business visit data and KPIs", "sky_business", key="btn8")
    with row3_cols[2]:
        menu_button_with_tooltip("📁 Sky Orbit File Uploader", "Upload files and query your data", "sky_orbit_file_upload", key="btn9")


if st.session_state.screen == "area_selection":

    # Load the Lottie animation JSON from the URL
    lottie_url = "https://assets2.lottiefiles.com/packages/lf20_jcikwtux.json"  # or your preferred animation URL
    lottie_json = load_lottieurl(lottie_url)
    
    # Show the animation if loaded successfully
    if lottie_json:
        st_lottie(lottie_json, height=280, key="area_selection_anim")
    else:
        st.info("🌀 Animation failed to load — but area selection still works.")
    
    # ... rest of your area_selection code here

import pandas as pd
import streamlit as st
import pdfplumber  # pip install pdfplumber
import docx        # pip install python-docx
from langchain_openai import ChatOpenAI
from langchain.agents.agent_types import AgentType
from langchain_experimental.agents import create_pandas_dataframe_agent

if st.session_state.get("screen") == "sky_orbit_file_upload":

    st.title("📁 Sky Orbit File Uploader")
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_sky_orbit_file_upload"):
        st.session_state.screen = "area_selection"
        st.rerun()
    st.markdown("""
        Upload your Excel, CSV, PDF, TXT, or Word (.docx) files below and ask questions about their contents.
        This AI is specialized in understanding your uploaded files only, but you can still chat normally without uploading.
    """)

    uploaded_file = st.file_uploader(
        "Choose a file to upload",
        type=["xlsx", "xls", "csv", "pdf", "txt", "docx"],
        accept_multiple_files=False
    )

    # Reset uploaded data and chat if new file uploaded
    if uploaded_file is not None:
        # If new file different than previous upload, reset chat & agent
        if uploaded_file.name != st.session_state.get("last_uploaded_file", None):
            st.session_state.file_ai_chat = []
            if "file_df_agent" in st.session_state:
                del st.session_state["file_df_agent"]
            st.session_state.last_uploaded_file = uploaded_file.name
            st.session_state.uploaded_df = None
            st.session_state.uploaded_text = None

        df = None
        extracted_text = None

        if uploaded_file.type in ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                  "application/vnd.ms-excel"]:
            try:
                df = pd.read_excel(uploaded_file)
            except Exception as e:
                st.error(f"Failed to read Excel file: {e}")

        elif uploaded_file.type == "text/csv":
            try:
                df = pd.read_csv(uploaded_file)
            except Exception as e:
                st.error(f"Failed to read CSV file: {e}")

        elif uploaded_file.type == "application/pdf":
            try:
                with pdfplumber.open(uploaded_file) as pdf:
                    extracted_text = ""
                    for page in pdf.pages:
                        extracted_text += page.extract_text() + "\n"
                st.text_area("Preview of extracted PDF text (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to extract PDF text: {e}")

        elif uploaded_file.type == "text/plain":
            try:
                extracted_text = uploaded_file.getvalue().decode("utf-8")
                st.text_area("Preview of uploaded TXT file (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to read TXT file: {e}")

        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(uploaded_file)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""]
                extracted_text = "\n".join(paragraphs)
                st.text_area("Preview of uploaded DOCX file (first 1000 chars):", extracted_text[:1000], height=200)
            except Exception as e:
                st.error(f"Failed to read DOCX file: {e}")

        st.write(f"**Uploaded file:** {uploaded_file.name} ({round(uploaded_file.size / 1024, 2)} KB)")

        st.session_state.uploaded_df = df
        st.session_state.uploaded_text = extracted_text

    else:
        st.info("Upload a file to begin chatting with your file AI or just ask questions below.")
        st.session_state.uploaded_df = None
        st.session_state.uploaded_text = None

    # Initialize chat history if missing
    if "file_ai_chat" not in st.session_state:
        st.session_state.file_ai_chat = []

    # Function to prepare or get existing AI agent
    def get_agent():
        llm_stream = ChatOpenAI(
            api_key=st.secrets["openai"]["api_key"],
            model_name="gpt-4o-mini",
            streaming=True,
        )
        if st.session_state.get("uploaded_df") is not None:
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=st.session_state.uploaded_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )
        elif st.session_state.get("uploaded_text") is not None:
            dummy_df = pd.DataFrame({"Text": [st.session_state.uploaded_text]})
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=dummy_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )
        else:
            # Fallback dummy df for normal conversation without upload
            dummy_df = pd.DataFrame({"Info": ["No file uploaded. Ask me anything or upload a file."]})
            return create_pandas_dataframe_agent(
                llm=llm_stream,
                df=dummy_df,
                verbose=False,
                agent_type=AgentType.OPENAI_FUNCTIONS,
                allow_dangerous_code=True,
            )

    if "file_df_agent" not in st.session_state:
        st.session_state.file_df_agent = get_agent()

    agent = st.session_state.file_df_agent

    # User question input and handling
    question = st.chat_input("Ask your file AI a question:")

    if question:
        with st.spinner("Processing your question..."):
            try:
                response = agent.run(question)
            except Exception as e:
                response = f"⚠️ AI error: {e}"

        st.session_state.file_ai_chat.append({"question": question, "response": response})
        st.rerun()

    # Display conversation bubbles
    for chat in st.session_state.file_ai_chat:
        st.chat_message("user").markdown(chat["question"])
        st.chat_message("assistant").markdown(chat["response"])











import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import re
import os

import streamlit as st
import pandas as pd
import os

# 1. SKY RETAIL PAGE — SHOW BUTTON HERE ONLY
if st.session_state.screen == "sky_retail":
    
    if st.button("📁 View Sky Retail Files"):
        st.session_state.screen = "sky_retail_files"
        st.rerun()

# 2. SKY RETAIL FILE VIEWER PAGE
if st.session_state.screen == "sky_retail_files":
    st.header("📂 Sky Retail File Viewer")

    folder = "Sky Retail"
    excel_files = [f for f in os.listdir(folder) if f.endswith('.xlsx')]

    if not excel_files:
        st.info("No Sky Retail files found in the folder.")
    else:
        selected_file = st.selectbox("Select a Sky Retail file to view:", excel_files)
        if st.button("Open File"):
            file_path = os.path.join(folder, selected_file)
            df = pd.read_excel(file_path)
            st.dataframe(df)

            # --- File summary ---
            df.columns = df.columns.str.strip().str.lower()
            status_col = "store status"
            stakeholder_col = "stakeholder"
            total_open = (df[status_col].str.lower() == "open").sum()
            total_closed = (df[status_col].str.lower() == "closed").sum()
            total_currys = (df[stakeholder_col].str.lower() == "curry's").sum()
            total_sky = (df[stakeholder_col].str.lower() == "sky retail").sum()
            total_ee = (df[stakeholder_col].str.lower() == "ee store").sum()

            st.markdown("### 📝 File Summary")
            st.markdown(f"""
            - **Total Open Stores:** {total_open}
            - **Total Closed Stores:** {total_closed}
            - **Total Currys Stores:** {total_currys}
            - **Total Sky Retail Stores:** {total_sky}
            - **Total EE Stores:** {total_ee}
            """)

    # Add a "Back" button!
    if st.button("⬅️ Back to Sky Retail"):
        st.session_state.screen = "sky_retail"
        st.rerun()





if st.session_state.screen == "sky_retail":
    st.title("Sky Retail")

    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_sky_retail"):
        st.session_state.screen = "area_selection"
        st.rerun()

    col_name = "Sky Retail Stakeholder"
    if col_name not in data.columns:
        st.error(f"Column '{col_name}' not found in Oracle datasets.")
        st.stop()

    # Helper functions
    def time_to_minutes(val):
        import datetime
        if pd.isnull(val):
            return 0
        if isinstance(val, (int, float)):
            return val
        if isinstance(val, datetime.timedelta):
            return val.total_seconds() / 60
        if isinstance(val, datetime.time):
            return val.hour * 60 + val.minute
        if isinstance(val, str):
            try:
                t = pd.to_datetime(val).time()
                return t.hour * 60 + t.minute
            except Exception:
                return 0
        return 0

    def minutes_to_hhmm(minutes):
        hours = int(minutes // 60)
        mins = int(minutes % 60)
        return f"{hours}:{mins:02}"

    def clean_stakeholder(val):
        v = str(val).strip().lower()
        if re.search(r"currys?|curry's", v):
            return "Currys"
        if "ee" in v:
            return "EE"
        if "sky" in v:
            return "Sky Retail"
        return "Other"

    def better_forecast(series, months=6):
        series = series.sort_index()
        n = len(series)
        if n == 0:
            return [0] * months
        if n < 3:
            avg = int(series.mean())
            return [avg] * months
        x = np.arange(n) if n < 4 else np.arange(n-4, n)
        y = series.values if n < 4 else series.values[-4:]
        m, b = np.polyfit(x, y, 1)
        fut_x = np.arange(n, n + months)
        y_pred = (m * fut_x + b)
        recent_avg = max(1, int(series.tail(2).mean()))
        y_pred = np.where(y_pred < recent_avg, recent_avg, y_pred)
        return y_pred.round().astype(int)

    # Prepare cleaned stakeholder column
    data["Sky Retail Stakeholder Clean"] = data[col_name].apply(clean_stakeholder)

    # Define stakeholders and create tabs
    stakeholders = ["Currys", "Sky Retail", "EE"]
    tabs = st.tabs(stakeholders)

    for i, stakeholder in enumerate(stakeholders):
        with tabs[i]:
            df = data[data["Sky Retail Stakeholder Clean"] == stakeholder].copy()
            if df.empty:
                st.info(f"No data for {stakeholder}.")
                continue

            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
            df = df[df["Date"].notna()]
            df["Month"] = df["Date"].dt.to_period("M").astype(str)

            if "Total Time" in df.columns:
                df["Total Time (min)"] = df["Total Time"].apply(time_to_minutes)
            else:
                df["Total Time (min)"] = 0

            # KPIs in expander 
            with st.expander(f"📊 {stakeholder} KPIs", expanded=True):
                count_visits = len(df)
                total_value = df["Total Value"].sum() if "Total Value" in df.columns else 0
                avg_value = df["Total Value"].mean() if "Total Value" in df.columns else 0
                total_time = df["Total Time (min)"].sum()
                avg_time = df["Total Time (min)"].mean()

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.metric(f"Total {stakeholder} Visits", f"{count_visits:,}")

                with col2:
                    st.metric("Total Value (£)", f"£{total_value:,.0f}")

                with col3:
                    st.metric("Avg Value", f"£{avg_value:,.0f}")

                # New row for Total Time and Average Time split into separate columns
                col4, col5, col6 = st.columns(3)

                with col4:
                    st.metric("Total Time", f"{minutes_to_hhmm(total_time)}")

                with col5:
                    st.metric("Average Time", f"{minutes_to_hhmm(avg_time)}")

                with col6:
                    pass  # Empty column for spacing

                
                st.markdown("**Monthly Visits (last 4 months shown)**")
                from st_aggrid import AgGrid, GridOptionsBuilder

                def show_aggrid(df, height=300):
                    gb = GridOptionsBuilder.from_dataframe(df)
                    gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=10)
                    gb.configure_default_column(editable=False, filter=True, sortable=True)
                    grid_options = gb.build()
                    AgGrid(df, gridOptions=grid_options, height=height, fit_columns_on_grid_load=True)
                
                # Monthly Visits table
                monthly_visits = (
                    df.groupby(df["Date"].dt.to_period("M").astype(str))
                      .size()
                      .reset_index(name="Visit Count")
                      .rename(columns={"Date": "Month"})
                )
                st.markdown("**Monthly Visits (last 4 months shown)**")
                show_aggrid(monthly_visits.tail(4))


                # Monthly Total Value table with £ formatting
                monthly_value = (
                    df.groupby(df["Date"].dt.to_period("M").astype(str))["Total Value"]
                      .sum()
                      .reset_index(name="Total Value")
                      .rename(columns={"Date": "Month"})
                )
                monthly_value["Total Value"] = monthly_value["Total Value"].map("£{0:,.0f}".format)
                st.markdown("**Monthly Total Value (£) (last 4 months shown)**")
                show_aggrid(monthly_value.tail(4))

                # Visits by Day of Week table ordered Monday to Sunday
                df["DayOfWeek"] = df["Date"].dt.day_name()
                visits_by_day = (
                    df.groupby("DayOfWeek")
                      .size()
                      .reset_index(name="Visit Count")
                )
                day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
                visits_by_day["DayOfWeek"] = pd.Categorical(visits_by_day["DayOfWeek"], categories=day_order, ordered=True)
                visits_by_day = visits_by_day.sort_values("DayOfWeek")
                st.markdown("**Visits by Day of Week**")
                show_aggrid(visits_by_day)

                # Total Value by Day of Week with £ formatting
                value_by_day = (
                    df.groupby("DayOfWeek")["Total Value"]
                      .sum()
                      .reset_index(name="Total Value")
                )
                value_by_day["DayOfWeek"] = pd.Categorical(value_by_day["DayOfWeek"], categories=day_order, ordered=True)
                value_by_day = value_by_day.sort_values("DayOfWeek")
                value_by_day["Total Value"] = value_by_day["Total Value"].map("£{0:,.0f}".format)
                st.markdown("**Total Value (£) by Day of Week**")
                show_aggrid(value_by_day)

                



            # Breakdown by Stakeholder expander
            with st.expander(f"📋 Breakdown by {stakeholder} Stakeholder", expanded=False):
                by_stake = df.groupby(col_name).agg(
                    Visits=('Name', 'count'),
                    Value=('Total Value', 'sum'),
                    Time=('Total Time (min)', 'sum')
                ).sort_values("Visits", ascending=False)
                by_stake["Time (hh:mm)"] = by_stake["Time"].apply(minutes_to_hhmm)
                by_stake["Value (£)"] = by_stake["Value"].apply(lambda x: f"£{x:,.0f}")
                by_stake = by_stake.drop(columns=["Value", "Time"])
                st.dataframe(by_stake, use_container_width=True)

            # Monthly Visit Trends expander
            with st.expander(f"📈 Monthly Visit Trends for {stakeholder}", expanded=False):
                by_month = df.groupby(["Month", col_name]).size().reset_index(name="Visits")
                fig_visits = px.bar(by_month, x="Month", y="Visits", color=col_name, barmode="group",
                                   title=f"Visits per Month by Stakeholder for {stakeholder}")
                st.plotly_chart(fig_visits, use_container_width=True)

            # Monthly Value Trends expander
            with st.expander(f"📈 Monthly Value Trends for {stakeholder}", expanded=False):
                if "Total Value" in df.columns:
                    by_value = df.groupby(["Month", col_name])["Total Value"].sum().reset_index()
                    fig_value = px.line(by_value, x="Month", y="Total Value", color=col_name,
                                       title=f"Value per Month by Stakeholder for {stakeholder}")
                    st.plotly_chart(fig_value, use_container_width=True)

            # Activity Status Pie Chart expander
            with st.expander(f"📊 Visit Activity Status Split for {stakeholder}", expanded=False):
                if "Activity Status" in df.columns:
                    status_df = df["Activity Status"].value_counts().reset_index()
                    status_df.columns = ["Activity Status", "Count"]
                    fig_pie = px.pie(status_df, names="Activity Status", values="Count")
                    st.plotly_chart(fig_pie, use_container_width=True)

            # Team Breakdown expander
            with st.expander("📋 Team Breakdown by Stakeholder", expanded=False):
                team_pivot = pd.pivot_table(
                    df,
                    index="Team",
                    columns="Sky Retail Stakeholder Clean",
                    values="Name",
                    aggfunc="count",
                    fill_value=0
                )
                st.dataframe(team_pivot, use_container_width=True)

            # Engineer Breakdown expander
            eng_pivot = (
                df.groupby(["Sky Retail Stakeholder Clean", "Name"])
                .agg(Visits=("Name", "count"),
                     Value=("Total Value", "sum"),
                     Time=("Total Time (min)", "sum"))
                .reset_index()
            )
            eng_pivot["Time (hh:mm)"] = eng_pivot["Time"].apply(minutes_to_hhmm)
            eng_pivot["Value (£)"] = eng_pivot["Value"].apply(lambda x: f"£{x:,.0f}")

            for eng_stakeholder in eng_pivot["Sky Retail Stakeholder Clean"].unique():
                with st.expander(f"👤 Engineer Breakdown for {eng_stakeholder}", expanded=False):
                    display = eng_pivot[
                        eng_pivot["Sky Retail Stakeholder Clean"] == eng_stakeholder
                    ][["Name", "Visits", "Value (£)", "Time (hh:mm)"]].sort_values("Visits", ascending=False)
                    st.dataframe(display, use_container_width=True)

            # Overall Forecasts expander
            with st.expander(f"🔮 Overall {stakeholder} Forecasts", expanded=False):
                # Visits Forecast
                st.markdown("**Visits Forecast (based on historic monthly visits)**")
                monthly_visits = df.groupby("Month").size().sort_index()
                if len(monthly_visits) < 2:
                    st.info("Not enough data for an overall visits forecast.")
                else:
                    fc_vals = better_forecast(monthly_visits, months=6)
                    last_month = pd.Period(monthly_visits.index.max(), freq="M")
                    fut_months = [str(last_month + i) for i in range(1, 7)]

                    fc_df = pd.concat([
                        pd.DataFrame({"Month": monthly_visits.index, "Visits": monthly_visits.values, "Type": "Actual"}),
                        pd.DataFrame({"Month": fut_months, "Visits": fc_vals, "Type": "Forecast"})
                    ], ignore_index=True)

                    fig_visits_fc = px.line(
                        fc_df, x="Month", y="Visits", color="Type", markers=True,
                        title=f"{stakeholder} – Overall Visits Forecast"
                    )
                    st.plotly_chart(fig_visits_fc, use_container_width=True)

                # Value Forecast
                st.markdown("**Value Forecast (based on historic monthly value)**")
                if "Total Value" in df.columns:
                    monthly_value = df.groupby("Month")["Total Value"].sum().sort_index()
                    if len(monthly_value) < 2:
                        st.info("Not enough data for a value forecast.")
                    else:
                        fc_vals = better_forecast(monthly_value, months=6)
                        last_month = pd.Period(monthly_value.index.max(), freq="M")
                        fut_months = [str(last_month + i) for i in range(1, 7)]

                        fc_df = pd.concat([
                            pd.DataFrame({"Month": monthly_value.index, "Value": monthly_value.values, "Type": "Actual"}),
                            pd.DataFrame({"Month": fut_months, "Value": fc_vals, "Type": "Forecast"})
                        ], ignore_index=True)

                        fig_value_fc = px.line(
                            fc_df, x="Month", y="Value", color="Type", markers=True,
                            title=f"{stakeholder} – Overall Value Forecast (£)"
                        )
                        st.plotly_chart(fig_value_fc, use_container_width=True)
                else:
                    st.warning("Total Value column not found.")
            
            # Per Team Forecasts expander (fixed to no nested expanders)
            with st.expander(f"🔮 {stakeholder} Team Forecasts (Visits & Value)", expanded=False):
                for team in sorted(df["Team"].unique()):
                    team_df = df[df["Team"] == team]
                    team_monthly = team_df.groupby("Month").size().sort_index()
                    team_monthly_value = team_df.groupby("Month")["Total Value"].sum().sort_index() if "Total Value" in team_df.columns else None

                    st.markdown(f"### 🔮 Forecast: {team} (Visits)")
                    if len(team_monthly) >= 2:
                        fc_team = better_forecast(team_monthly, months=6)
                        last_month = pd.Period(team_monthly.index.max(), freq="M")
                        fut_months = [str(last_month + i) for i in range(1, 7)]

                        team_fc_df = pd.concat([
                            pd.DataFrame({"Month": team_monthly.index, "Visits": team_monthly.values, "Type": "Actual"}),
                            pd.DataFrame({"Month": fut_months, "Visits": fc_team, "Type": "Forecast"})
                        ], ignore_index=True)

                        fig_team_visits = px.line(
                            team_fc_df, x="Month", y="Visits", color="Type", markers=True,
                            title=f"{team} – Visits Forecast"
                        )
                        st.plotly_chart(fig_team_visits, use_container_width=True)

                    st.markdown(f"### 🔮 Forecast: {team} (Value)")
                    if team_monthly_value is not None and len(team_monthly_value) >= 2:
                        fc_team_val = better_forecast(team_monthly_value, months=6)
                        last_month = pd.Period(team_monthly_value.index.max(), freq="M")
                        fut_months = [str(last_month + i) for i in range(1, 7)]

                        team_fc_val_df = pd.concat([
                            pd.DataFrame({"Month": team_monthly_value.index, "Value": team_monthly_value.values, "Type": "Actual"}),
                            pd.DataFrame({"Month": fut_months, "Value": fc_team_val, "Type": "Forecast"})
                        ], ignore_index=True)

                        fig_team_value = px.line(
                            team_fc_val_df, x="Month", y="Value", color="Type", markers=True,
                            title=f"{team} – Value Forecast (£)"
                        )
                        st.plotly_chart(fig_team_value, use_container_width=True)

            # Month-on-Month Change Tables
            pivot_visits = (
                df.groupby(["Month", "Sky Retail Stakeholder Clean"])
                  .size()
                  .unstack(fill_value=0)
                  .reindex(columns=[stakeholder], fill_value=0)
                  .sort_index()
            )
            pivot_value = (
                df.groupby(["Month", "Sky Retail Stakeholder Clean"])["Total Value"]
                  .sum()
                  .unstack(fill_value=0)
                  .reindex(columns=[stakeholder], fill_value=0)
                  .sort_index()
            )

            def make_change_table(pivot):
                df_ = pivot.copy()
                for col in df_.columns:
                    df_[f"{col} Δ"] = df_[col].diff().fillna(0).astype(int)
                    max_val = df_[col].max()
                    min_val = df_[col].min()
                    df_[f"{col} ΔMax"] = df_[col] - max_val
                    df_[f"{col} ΔMin"] = df_[col] - min_val
                return df_

            visits_tbl = make_change_table(pivot_visits)
            value_tbl = make_change_table(pivot_value)

            with st.expander("📊 Month-on-Month Change Table (Visits)", expanded=False):
                st.markdown("**Visits Table**")
                st.dataframe(
                    visits_tbl.style
                    .format({col: "{:,}" for col in visits_tbl.columns if "Δ" not in col})
                    .format({col: "{:+,}" for col in visits_tbl.columns if "Δ" in col}),
                    use_container_width=True
                )

            with st.expander("📊 Month-on-Month Change Table (Value)", expanded=False):
                st.markdown("**Value Table**")
                value_cols = [col for col in value_tbl.columns if "Δ" not in col]
                delta_cols = [col for col in value_tbl.columns if "Δ" in col]

                # Format delta columns with £ and signs:
                for col in delta_cols:
                    value_tbl[col] = value_tbl[col].apply(lambda x: f"£{x:+,}")

                # Round and convert value columns to int
                value_tbl[value_cols] = value_tbl[value_cols].round(0).astype(int)

                # Convert values to strings with £ prefix for display
                for col in value_cols:
                    value_tbl[col] = "£" + value_tbl[col].map("{:,}".format)

                style = value_tbl.style.format({col: "{:+,}" for col in delta_cols})

                st.dataframe(value_tbl, use_container_width=True)

            # Combined Charts (collapsible)
            with st.expander("📊 Combined Monthly Trend Charts", expanded=False):
                # Line chart for total Value per stakeholder
                fig_value_trends = px.line(
                    value_tbl.reset_index(),
                    x='Month',
                    y=[stakeholder],
                    title=f'Monthly Total Value for {stakeholder} (£)',
                    markers=True
                )
                st.plotly_chart(fig_value_trends, use_container_width=True)

                # Line chart for total Visits per stakeholder
                fig_visits_trends = px.line(
                    visits_tbl.reset_index(),
                    x='Month',
                    y=[stakeholder],
                    title=f'Monthly Total Visits for {stakeholder}',
                    markers=True
                )
                st.plotly_chart(fig_visits_trends, use_container_width=True)

                # Select Δ columns only for Value
                value_change_cols = [col for col in value_tbl.columns if 'Δ' in col and 'Max' not in col and 'Min' not in col]
                fig_value_changes = px.bar(
                    value_tbl.reset_index(),
                    x='Month',
                    y=value_change_cols,
                    title=f'Month-on-Month Change in Value for {stakeholder} (£)',
                    barmode='group'
                )
                st.plotly_chart(fig_value_changes, use_container_width=True)

                # Select Δ columns only for Visits
                visits_change_cols = [col for col in visits_tbl.columns if 'Δ' in col and 'Max' not in col and 'Min' not in col]
                fig_visits_changes = px.bar(
                    visits_tbl.reset_index(),
                    x='Month',
                    y=visits_change_cols,
                    title=f'Month-on-Month Change in Visits for {stakeholder}',
                    barmode='group'
                )
                st.plotly_chart(fig_visits_changes, use_container_width=True)

                # Heatmap for Value ΔMax
                fig_heatmap_value_max = go.Figure(
                    go.Heatmap(
                        z=value_tbl[[col for col in value_tbl.columns if 'ΔMax' in col]].values.T,
                        x=value_tbl.index,
                        y=[col.replace(' ΔMax', '') for col in value_tbl.columns if 'ΔMax' in col],
                        colorscale='RdBu',
                        colorbar=dict(title='Difference vs Max'),
                        zmid=0
                    )
                )
                fig_heatmap_value_max.update_layout(title=f'Value Difference vs Max for {stakeholder}')
                st.plotly_chart(fig_heatmap_value_max, use_container_width=True)

                # Heatmap for Visits ΔMax
                fig_heatmap_visits_max = go.Figure(
                    go.Heatmap(
                        z=visits_tbl[[col for col in visits_tbl.columns if 'ΔMax' in col]].values.T,
                        x=visits_tbl.index,
                        y=[col.replace(' ΔMax', '') for col in visits_tbl.columns if 'ΔMax' in col],
                        colorscale='RdBu',
                        colorbar=dict(title='Difference vs Max'),
                        zmid=0
                    )
                )
                fig_heatmap_visits_max.update_layout(title=f'Visits Difference vs Max for {stakeholder}')
                st.plotly_chart(fig_heatmap_visits_max, use_container_width=True)


            # Day of Week Visits Bar Chart
            with st.expander("📅 Visits by Day of Week", expanded=False):
                # Extract day names from the Date column
                df["DayOfWeek"] = df["Date"].dt.day_name()

                # Group by day of week and count visits
                visits_by_day = df.groupby("DayOfWeek").size().reindex([
                    "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"
                ], fill_value=0)

                # Create bar chart with Plotly Express
                fig_day = px.bar(
                    visits_by_day.reset_index(name="Visits"),
                    x="DayOfWeek",
                    y="Visits",
                    title=f"{stakeholder} Visits by Day of Week",
                    labels={"DayOfWeek": "Day of Week", "Visits": "Number of Visits"},
                    color="Visits",
                    color_continuous_scale="Blues"
                )
                st.plotly_chart(fig_day, use_container_width=True)

            # Raw Data expander
            with st.expander(f"🔎 Show Raw Data for {stakeholder}", expanded=False):
                st.dataframe(df.dropna(axis=1, how="all"), use_container_width=True)
         


# BLOCK 24: 📦 Sky Business Area – Combined Oracle Data Dashboard
import pandas as pd
import streamlit as st
import plotly.express as px

if st.session_state.screen == "sky_business":
    # BLOCK 24 stuff


    # Back button to return to area selection menu
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_sky_business"):
        st.session_state.screen = "area_selection"
        st.rerun()

    st.title("Sky Business")

    # You already have 'oracle_all' loaded at the top of your script!
    df_all = data.copy()
    df_all.columns = df_all.columns.str.strip()
    df_all = df_all.dropna(how="all")

    if "Visit Type" not in df_all.columns:
        st.error("Column 'Visit Type' is missing.")
    else:
        df_sky = df_all[df_all["Visit Type"].astype(str).str.contains("Sky Business", case=False, na=False)]
        if df_sky.empty:
            st.info("No rows found for 'Sky Business' in Visit Type.")
        else:
            st.markdown("## 📊 Sky Business Area Dashboard")
            st.caption("Filtered view of all Oracle datasets where `Visit Type` contains 'Sky Business'.")

            # --- Summary KPIs ---
            st.subheader("📌 Summary KPIs")
            col1, col2, col3 = st.columns(3)

            total_visits = len(df_sky)
            total_value = df_sky["Total Value"].sum() if "Total Value" in df_sky.columns else 0

            activity_counts = df_sky["Activity Status"].astype(str).str.lower().value_counts()
            completed = activity_counts.get("completed", 0)
            cancelled = activity_counts.get("cancelled", 0)
            not_done = activity_counts.get("not done", 0)
            failed = cancelled + not_done

            col1.metric("📦 Total Sky Business Visits", total_visits)
            col2.metric("💷 Total Value (£)", f"£{total_value:,.2f}")
            if failed > 0:
                ratio = completed / failed
                col3.markdown(f"🔁 **{ratio:.1f}** visits completed for every **1** cancelled or not done visit")
            else:
                col3.markdown("🔁 No failed visits recorded")

            # --- Activity Breakdown Chart ---
            st.subheader("🧩 Activity Completion Breakdown")
            st.bar_chart(activity_counts)

            # --- Monthly Trends ---
            if "Date" in df_sky.columns:
                df_sky["Month"] = pd.to_datetime(df_sky["Date"], errors="coerce").dt.to_period("M").astype(str)
                by_month = df_sky.groupby("Month").agg({
                    "Visit Type": "count",
                    "Total Value": "sum"
                }).rename(columns={"Visit Type": "Visit Count"})

                st.subheader("📈 Monthly Trends")
                st.plotly_chart(px.line(by_month, y="Visit Count", title="Monthly Visit Count"), use_container_width=True)
                st.plotly_chart(px.line(by_month, y="Total Value", title="Monthly Total Value (£)"), use_container_width=True)
            else:
                st.warning("⚠️ Column 'Date' missing — cannot generate monthly trends.")

            # --- Sunburst: Visit Type → Activity Status ---
            st.subheader("🌞 Visit Type Breakdown by Activity Status")
            if "Activity Status" in df_sky.columns:
                fig = px.sunburst(df_sky, path=["Visit Type", "Activity Status"], title="Sky Business Visit Breakdown")
                st.plotly_chart(fig, use_container_width=True)

            # --- Forecasting (based on 6 months) ---
            st.subheader("🔮 Forecast (based on recent 6 months)")

            if "Month" in df_sky.columns:
                last_6 = by_month.tail(6)
                forecast = round(last_6.mean())
                st.markdown(f"""
                **🗖️ 6-Month Forecast**
                - Avg Monthly Visits: **{forecast['Visit Count']}**
                - Avg Monthly Value: **£{forecast['Total Value']:,.2f}**
                """)
                st.line_chart(last_6)
            else:
                st.info("🗓️ No 'Month' column available for forecasting.")

            st.caption("Data pulled from 4 Oracle sources, filtered to Sky Business only.")


    # --- ADVANCED KPI TABS SECTION ---
    st.markdown("## 📊 Advanced Sky Business KPI Centre")
    st.caption("Everything below is filtered where **Visit Type** contains “Sky Business”.")

    # You already have df_sky, so just use/copy it
    df_all = df_sky.copy()

    # Ensure correct types
    df_all["Date"] = pd.to_datetime(df_all["Date"], errors="coerce")
    df_all.dropna(subset=["Date"], inplace=True)
    df_all["Month"] = df_all["Date"].dt.to_period("M").astype(str)

    file_map = {
        "VIP North":   "VIP North Oracle Data.xlsx",
        "VIP South":   "VIP South Oracle Data.xlsx",
        "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
        "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
    }

    # Assign team name to each row if missing
    if "Team" not in df_all.columns:
        df_all["Team"] = None
    for label in file_map.keys():
        mask = df_all["Source"].str.contains(label.replace(" ", ""), case=False, na=False)
        df_all.loc[mask, "Team"] = label

    # Tab labels
    tab_labels = ["Overall"] + list(file_map.keys())
    tabs       = st.tabs(tab_labels)

    import numpy as np
    import re

    def _simple_forecast(series: pd.Series, periods: int = 6) -> list[int]:
        series = series.sort_index()
        n = len(series)

        if n >= 3:
            y = series.iloc[-4:]
            x = np.arange(len(series) - len(y), len(series))
        elif n == 2:
            y = series
            x = np.arange(n)
        elif n == 1:
            return [int(series.iloc[-1])] * periods
        else:
            return [0] * periods

        m, b = np.polyfit(x, y.values, 1)
        future_x = np.arange(len(series), len(series) + periods)
        return [max(0, int(round(m * xi + b))) for xi in future_x]

    for tab, label in zip(tabs, tab_labels):
        with tab:
            df = df_all if label == "Overall" else df_all[df_all["Team"] == label]
            st.subheader("🌐 Overall" if label == "Overall" else f"📁 {label}")

            if df.empty:
                st.info("No data in this slice.")
                continue

            # Basic KPIs
            with st.expander("🧮 Basic KPIs", expanded=True):
                c1, c2, c3 = st.columns(3)
                c1.metric("Visits",      f"{len(df):,}")
                c2.metric("Value (£)",   f"£{df.get('Total Value', pd.Series(dtype=float)).sum():,.0f}")
                c3.metric("Visit Types", df["Visit Type"].nunique())

            # Historical trends
            with st.expander("📈 Monthly Trend by Visit Type", expanded=False):
                monthly_counts = (
                    df.groupby(["Month", "Visit Type"])
                      .size().reset_index(name="Visits")
                )
                fig_hist = px.line(
                    monthly_counts, x="Month", y="Visits",
                    color="Visit Type", markers=True,
                    title="Historical Visits per Type"
                )
                st.plotly_chart(fig_hist, use_container_width=True,
                                key=f"hist_{re.sub(r'\\W+','_',label)}")

            with st.expander("📊 Monthly Visit Count (Stacked)", expanded=False):
                bar_df = (
                    monthly_counts.pivot(index="Month",
                                         columns="Visit Type",
                                         values="Visits")
                    .fillna(0)
                    .sort_index()
                )
                st.bar_chart(bar_df)

            if "Total Value" in df.columns:
                with st.expander("💷 Monthly Value (£)", expanded=False):
                    value_df = (
                        df.groupby("Month")["Total Value"]
                          .sum()
                          .sort_index()
                    )
                    st.line_chart(value_df)




# ── Detailed Monthly KPI Table by Visit Type & Status ────────────
            with st.expander("📋 KPI Table – Monthly Visit Type x Status", expanded=False):
                if "Activity Status" not in df.columns:
                    st.warning("⚠️ 'Activity Status' column missing.")
                else:
                    kpi_table = (
                        df.groupby(["Month", "Visit Type", "Activity Status"])
                          .size()
                          .reset_index(name="Count")
                    )

                    # Pivot to make Activity Statuses into columns
                    pivot_table = kpi_table.pivot_table(
                        index=["Month", "Visit Type"],
                        columns="Activity Status",
                        values="Count",
                        fill_value=0
                    ).reset_index()

                    # Sort and display
                    pivot_table = pivot_table.sort_values(by=["Month", "Visit Type"])
                    st.dataframe(pivot_table, use_container_width=True)




            # Forecasts -------------------------------------------------
            with st.expander("🔮 Forecasts", expanded=False):

                tab_id = re.sub(r"\W+", "_", label).strip("_")
                key_counter: defaultdict[str, int] = defaultdict(int)

                # 3-A  Overall forecast
                overall_series = df.groupby("Month").size().sort_index()
                fc_vals        = _simple_forecast(overall_series, periods=6)
                last_p         = pd.Period(overall_series.index.max(), freq="M")
                fut_mths       = [str(last_p + i) for i in range(1, 7)]

                overall_df = pd.concat(
                    [
                        pd.DataFrame(
                            {"Month": overall_series.index,
                             "Visits": overall_series.values,
                             "Kind":   "Actual"}
                        ),
                        pd.DataFrame(
                            {"Month": fut_mths,
                             "Visits": fc_vals,
                             "Kind":   "Forecast"}
                        )
                    ],
                    ignore_index=True
                )

                st.plotly_chart(
                    px.line(overall_df, x="Month", y="Visits",
                            line_dash="Kind", markers=True,
                            title="Historical vs Forecast – ALL Visit Types"),
                    use_container_width=True,
                    key=f"fc_overall_{tab_id}"
                )

                # 3-B  Individual visit-type forecasts
                visit_types = df["Visit Type"].dropna().unique()
                for vt in sorted(visit_types):

                    vt_series = (
                        df[df["Visit Type"] == vt]
                          .groupby("Month")
                          .size()
                          .sort_index()
                    )
                    if vt_series.empty:
                        continue

                    vt_fc    = _simple_forecast(vt_series, periods=6)
                    last_p   = pd.Period(vt_series.index.max(), freq="M")
                    fut_mths = [str(last_p + i) for i in range(1, 7)]

                    chart_df = pd.concat(
                        [
                            pd.DataFrame(
                                {"Month": vt_series.index,
                                 "Visits": vt_series.values,
                                 "Kind":   "Actual"}
                            ),
                            pd.DataFrame(
                                {"Month": fut_mths,
                                 "Visits": vt_fc,
                                 "Kind":   "Forecast"}
                            )
                        ],
                        ignore_index=True
                    )

                    clean_vt = re.sub(r"\W+", "_", vt).strip("_")
                    key_counter[clean_vt] += 1
                    safe_key = f"fc_{tab_id}_{clean_vt}_{key_counter[clean_vt]}"

                    st.plotly_chart(
                        px.line(chart_df, x="Month", y="Visits",
                                line_dash="Kind", markers=True,
                                title=f"{vt} – Historical vs Forecast"),
                        use_container_width=True,
                        key=safe_key
                    )

# ── Monthly change summary ───────────────────────────────────
            with st.expander("📊 Month-on-Month Change (Visits)", expanded=False):
                # ➊ Monthly totals for this tab
                monthly_tot = (
                    df.groupby("Month")
                      .size()
                      .sort_index()
                )

                if monthly_tot.empty:
                    st.info("No monthly data available.")
                else:
                    # ➋ Month-over-month deltas
                    delta_abs  = monthly_tot.diff().fillna(0).astype(int)
                    delta_pct  = (monthly_tot.pct_change() * 100).round(1)

                    # ➌ Compare to dataset-wide max / min
                    max_vis = monthly_tot.max()
                    min_vis = monthly_tot.min()

                    summary_df = pd.DataFrame({
                        "Month"        : monthly_tot.index.astype(str),
                        "Visits"       : monthly_tot.values,
                        "Δ vs Prev"    : delta_abs.values,
                        "%Δ vs Prev"   : delta_pct.values,
                        "Δ vs Max"     : (monthly_tot - max_vis).values,
                        "Δ vs Min"     : (monthly_tot - min_vis).values,
                    })

                    # tidy column order
                    summary_df = summary_df[
                        ["Month", "Visits", "Δ vs Prev", "%Δ vs Prev",
                         "Δ vs Max", "Δ vs Min"]
                    ]

                    st.dataframe(summary_df, use_container_width=True)

            # ── Month-on-Month change per *Visit Type* ──────────────────────
            with st.expander("📊 Month-on-Month Change • by Visit Type", expanded=False):
                # ➊ Pivot: rows = Month, cols = Visit Type, values = counts
                pv = (
                    df.groupby(["Month", "Visit Type"])
                      .size()
                      .unstack(fill_value=0)
                      .sort_index()              # chronological
                )

                if pv.empty:
                    st.info("No data available for this slice.")
                else:
                    # ➋ Deltas
                    delta_abs = pv.diff().fillna(0).astype(int)
                    delta_pct = (pv.pct_change() * 100).round(1).fillna(0)

                    # ➌ Build a pretty table
                    tidy_frames = []
                    for vt in pv.columns:
                        tmp = pd.DataFrame({
                            "Month"          : pv.index.astype(str),
                            f"{vt} Visits"   : pv[vt].values,
                            f"{vt} Δ"        : delta_abs[vt].values,
                            f"{vt} %Δ"       : delta_pct[vt].values,
                        })
                        tidy_frames.append(tmp)

                    # ➍ Merge on Month
                    tidy_df = tidy_frames[0]
                    for extra in tidy_frames[1:]:
                        tidy_df = tidy_df.merge(extra, on="Month")

                    # ➎ Show
                    st.dataframe(tidy_df, use_container_width=True)

            # ── KPI Heat-Map • Peaks, Troughs & Growth ──────────────────────
            with st.expander("📊 KPI Dashboard (Peaks • Troughs • Growth)", expanded=False):

                # 1️⃣  Baseline counts ─────────────────────────────────────
                base = (
                    df.groupby("Month")
                      .size()                       # all Sky Business visits
                      .rename("Visits")
                      .sort_index()
                )
                if len(base) < 2:
                    st.info("Need at least 2 months of data for deltas.")
                else:
                    # 2️⃣  Delta vs previous, vs Max, vs Min ─────────────
                    delta_abs = base.diff().fillna(0).astype(int)
                    delta_pct = (base.pct_change() * 100).round(1).fillna(0)

                    max_val   = base.max()
                    min_val   = base.min()

                    kpi_df = pd.DataFrame({
                        "Month"            : base.index.astype(str),
                        "Visits"           : base.values,
                        "Δ Prev Mo"        : delta_abs.values,
                        "%Δ Prev Mo"       : delta_pct.values,
                        "Δ vs Max Peak"    : (base - max_val).values,
                        "Δ vs Min Trough"  : (base - min_val).values,
                    })

                    # 3️⃣  Styling helpers ───────────────────────────────
                    def colour_delta(val):
                        if val > 0:
                            return "background-color:#075E00;color:white"   # green
                        elif val < 0:
                            return "background-color:#8B0000;color:white"   # red
                        else:
                            return "background-color:#444444;color:white"   # grey

                    styled = (
                        kpi_df.style
                              .applymap(colour_delta, subset=["Δ Prev Mo", "%Δ Prev Mo"])
                              .applymap(colour_delta, subset=["Δ vs Max Peak", "Δ vs Min Trough"])
                              .format({"Visits":"{:,}",
                                       "Δ Prev Mo":"{:+,}",
                                       "%Δ Prev Mo":"{:+.1f}%",
                                       "Δ vs Max Peak":"{:+,}",
                                       "Δ vs Min Trough":"{:+,}"})
                    )

            # 📊 EXTRA GRAPH GALLERY  ─  Best-in-class visuals
            with st.expander("📊 Graph Gallery – Visit Trends & Growth", expanded=False):
                import plotly.graph_objects as go

                # 1️⃣ Area Chart – Total Visits
                fig_area = px.area(
                    overall_series.reset_index(name="Visits"),
                    x="Month", y="Visits",
                    title="📈 Area Chart – Total Visits Over Time"
                )
                st.plotly_chart(fig_area, use_container_width=True, key=f"area_{tab_id}")

                # 2️⃣ Line Chart – % Growth Month-over-Month
                pct_change = overall_series.pct_change().fillna(0) * 100
                fig_pct = px.line(
                    pct_change.reset_index(name="% Growth"),
                    x="Month", y="% Growth",
                    title="📊 % Change in Visits (Month-over-Month)"
                )
                st.plotly_chart(fig_pct, use_container_width=True, key=f"pctmo_{tab_id}")

                # 3️⃣ Heatmap – Monthly Visit Counts
                hm_df = overall_series.reset_index(name="Visits")
                hm_df["Month_Num"] = pd.to_datetime(hm_df["Month"]).dt.month
                hm_df["Year"] = pd.to_datetime(hm_df["Month"]).dt.year

                fig_hm = px.density_heatmap(
                    hm_df, x="Month_Num", y="Year", z="Visits",
                    color_continuous_scale="Viridis", nbinsx=12, nbinsy=len(hm_df["Year"].unique()),
                    title="🔥 Monthly Visit Count Heatmap"
                )
                st.plotly_chart(fig_hm, use_container_width=True, key=f"hm_{tab_id}")

                # 4️⃣ Waterfall Chart – Δ Visits from Previous Month
                delta_vals = overall_series.diff().fillna(0).astype(int)
                wf_df = pd.DataFrame({
                    "Month" : overall_series.index.astype(str),
                    "Change": delta_vals.values
                })

                fig_wf = go.Figure(go.Waterfall(
                    x = wf_df["Month"],
                    y = wf_df["Change"],
                    measure = ["absolute"] + ["relative"] * (len(wf_df) - 1),
                    increasing = {"marker": {"color": "green"}},
                    decreasing = {"marker": {"color": "crimson"}},
                    totals = {"marker": {"color": "gray"}},
                    connector = {"line": {"color": "silver"}},
                    textposition="outside"
                ))

                fig_wf.update_layout(
                    title = "🌊 Waterfall – Δ Visits vs Previous Month",
                    showlegend = False,
                    height = 400
                )

                st.plotly_chart(fig_wf, use_container_width=True, key=f"wf_{tab_id}")

    # ── SLA DASHBOARD ─────────────────────────────────────────────────────────
    #  Columns needed in AI Test SB Visits.xlsx …
    #    • "Visit type"    ⇒ SLA bucket (2h / 4h / 5 day / 8h)
    #    • "Date of visit" ⇒ timestamp
    #    • "Met SLA?"      ⇒ optional flag (Y/True/Yes)
    # -------------------------------------------------------------------------
    with st.expander("⏱️ SLA Dashboard – 2 h • 4 h • 5 day • 8 h", expanded=False):

        # 0️⃣ LOAD + CLEAN ----------------------------------------------------
        SLA_FILE   = "AI Test SB Visits.xlsx"
        SLA_COL    = "Visit type"
        DATE_COL   = "Date of visit"
        RESULT_COL = "Met SLA?"          # not in file → created below

        try:
            sla_df = pd.read_excel(SLA_FILE)
        except Exception as e:
            st.error(f"Could not load “{SLA_FILE}”: {e}")
            st.stop()

        sla_df.columns = sla_df.columns.str.strip()

        for col in (SLA_COL, DATE_COL):
            if col not in sla_df.columns:
                st.error(f"Column “{col}” missing – check the sheet header.")
                st.stop()

        sla_df = sla_df.dropna(subset=[SLA_COL, DATE_COL])
        sla_df[DATE_COL] = pd.to_datetime(sla_df[DATE_COL], errors="coerce")
        sla_df = sla_df.dropna(subset=[DATE_COL])

        # 🔎 Filter to the four SLA targets ONLY
        sla_mask = sla_df[SLA_COL].str.lower().str.contains(
            r"\b(2h|2 h|2hr|4h|4 h|4hr|5 ?day|8h|8 h|8hr)\b", regex=True, na=False
        )
        sla_df = sla_df[sla_mask].copy()

        if sla_df.empty:
            st.warning("No rows match 2 h, 4 h, 5 day or 8 h targets.")
            st.stop()

        sla_df["Month"] = sla_df[DATE_COL].dt.to_period("M").astype(str)

        # If “Met SLA?” not present, assume every ticket was met
        if RESULT_COL not in sla_df.columns:
            sla_df[RESULT_COL] = True

        sla_df[RESULT_COL] = (
            sla_df[RESULT_COL]
            .astype(str).str.strip().str.lower()
            .isin(["yes", "y", "true", "1"])
        )

        # 1️⃣ KPI HEADER ------------------------------------------------------
        total_tickets = len(sla_df)
        met_total     = sla_df[RESULT_COL].sum()
        pct_met       = met_total / total_tickets * 100 if total_tickets else 0

        k0, k1, k2, k3 = st.columns(4)
        k0.metric("Total Tickets", f"{total_tickets:,}")
        k1.metric("Met SLA",       f"{met_total:,}")
        k2.metric("Missed SLA",    f"{total_tickets-met_total:,}")
        k3.metric("% Met",         f"{pct_met:.1f}%")

        st.markdown("---")

        # 2️⃣ VOLUME PER SLA BUCKET ------------------------------------------
        vol_df = (sla_df[SLA_COL]
                  .value_counts()
                  .reset_index()
                  .rename(columns={SLA_COL: "SLA Target", "count": "Tickets"}))

        st.plotly_chart(
            px.bar(vol_df, x="SLA Target", y="Tickets", color="SLA Target",
                   title="Ticket Volume by SLA Target"),
            use_container_width=True,
            key="sla_vol"
        )

        # 3️⃣ MONTHLY TREND PER TARGET ---------------------------------------
        trend_df = (sla_df.groupby(["Month", SLA_COL])
                            .size()
                            .reset_index(name="Tickets")
                            .sort_values("Month"))

        st.plotly_chart(
            px.line(trend_df, x="Month", y="Tickets", color=SLA_COL,
                    markers=True, title="Monthly Ticket Trend by SLA Target"),
            use_container_width=True,
            key="sla_trend"
        )


        # 4️⃣ STACKED MET vs MISSED  (only if some misses exist) -------------
        if not sla_df[RESULT_COL].all():
            stack_df = (sla_df.groupby([SLA_COL, RESULT_COL])
                                .size().reset_index(name="Count"))
            stack_df["Status"] = np.where(stack_df[RESULT_COL], "Met", "Missed")

            st.plotly_chart(
                px.bar(stack_df, x=SLA_COL, y="Count", color="Status",
                       title="Met vs Missed SLA (stacked)"),
                use_container_width=True,
                key="sla_stack"
            )

        # 5️⃣ FORECASTS (NEXT 6 MONTHS) --------------------------------------
        st.markdown("### 🔮 Forecasts (next 6 months)")

        def _simple_forecast(series: pd.Series, periods: int = 6) -> list[int]:
            series = series.sort_index()
            n = len(series)
            if n == 0:
                return [0] * periods
            if n == 1:
                return [int(series.iloc[0])] * periods
            if n == 2:
                x = np.arange(2)
                y = series.values
            else:
                x = np.arange(n-4, n)
                y = series.iloc[-4:].values
            m, b = np.polyfit(x, y, 1)
            return [max(0, int(round(m * xi + b)))
                    for xi in range(n, n + periods)]

        for sla_tag in sorted(sla_df[SLA_COL].unique()):
            serie = (sla_df[sla_df[SLA_COL] == sla_tag]
                     .groupby("Month")
                     .size()
                     .sort_index())

            if len(serie) < 2:
                st.info(f"*{sla_tag}* – not enough data for a forecast.")
                continue

            fc_vals  = _simple_forecast(serie)
            last_m   = pd.Period(serie.index.max(), freq="M")
            fut_mths = [str(last_m + i) for i in range(1, 7)]

            plot_df = pd.concat(
                [pd.DataFrame({"Month": serie.index,
                               "Tickets": serie.values,
                               "Kind": "Actual"}),
                 pd.DataFrame({"Month": fut_mths,
                               "Tickets": fc_vals,
                               "Kind": "Forecast"})],
                ignore_index=True
            )

            clean_key = re.sub(r"\W+", "_", sla_tag).strip("_")
            st.plotly_chart(
                px.line(plot_df, x="Month", y="Tickets", line_dash="Kind",
                        markers=True, title=f"{sla_tag} – Actual vs Forecast"),
                use_container_width=True,
                key=f"sla_fc_{clean_key}"
            )

if st.session_state.get("kpi_dataset", (None,))[0] == "Sky Business Area":
    # ── SLA VENUE MATRIX ──────────────────────────────────────────────── 
    st.markdown("## 🏢 Venue SLA Matrix – SLA Counts per Site")

    # 0️⃣  Use previously loaded `sla_df` or load fresh
    if "sla_df" not in locals():
        try:
            sla_df = pd.read_excel("AI Test SB Visits.xlsx")
        except Exception as e:
            st.error(f"Could not reload SLA file: {e}")
            st.stop()

        sla_df.columns = sla_df.columns.str.strip()
        sla_df = sla_df.dropna(subset=["Visit type", "Date of visit"])
        sla_df["Date of visit"] = pd.to_datetime(sla_df["Date of visit"], errors="coerce")
        sla_df = sla_df.dropna(subset=["Date of visit"])
        sla_df["Visit type"] = sla_df["Visit type"].astype(str).str.strip().str.lower()

    # 1️⃣  Filter SLA buckets
    SLA_BUCKETS = ["2hr", "4hr", "5 day", "8h", "8 hr", "8hr"]
    sla_ven = sla_df[
        sla_df["Visit type"].str.lower().isin(SLA_BUCKETS)
    ].copy()

    # 2️⃣  Standardise SLA labels
    sla_ven["SLA"] = sla_ven["Visit type"].str.lower().replace({"8 hr": "8h"})

    # 3️⃣  Pivot: Venue × SLA Counts
    if "Venue Name" in sla_ven.columns and "VR Number" in sla_ven.columns:
        pivot = (
            sla_ven.pivot_table(
                index="Venue Name",
                columns="SLA",
                values="VR Number",
                aggfunc="count",
                fill_value=0
            )
            .assign(Total=lambda d: d.sum(axis=1))
            .sort_values("Total", ascending=False)
        )
        st.dataframe(pivot, use_container_width=True)
    else:
        st.warning("Columns 'Venue Name' or 'VR Number' not found in SLA file.")





    # ── 💾 Table: All Venues ───────────────────────────────────────────
    with st.expander("📋 Full Venue SLA Table", expanded=False):
        st.dataframe(pivot, use_container_width=True)

    # ── 📊 Charts: in tabs (NOT in an expander to avoid nesting) ───────
    tabs = st.tabs(["🏆 Top 20 by Total", "📊 Stacked SLA Mix", "🌡️ Heatmap"])

    # ── Chart A: Top 20 Horizontal Bar
    with tabs[0]:
        top20 = pivot.head(20).reset_index().sort_values("Total")
        fig_top = px.bar(
            top20,
            y="Venue Name", x="Total", orientation="h",
            title="Top 20 Venues by SLA Visits", text="Total"
        )
        st.plotly_chart(fig_top, use_container_width=True)

    # ── Chart B: Stacked SLA Distribution
    with tabs[1]:
        stacked = (
            pivot.head(20)
                .drop(columns="Total")
                .reset_index()
                .melt(id_vars="Venue Name", var_name="SLA", value_name="Tickets")
        )
        fig_stack = px.bar(
            stacked,
            y="Venue Name", x="Tickets", color="SLA", orientation="h",
            title="Top 20 SLA Mix by Venue"
        )
        st.plotly_chart(fig_stack, use_container_width=True)

    # ── Chart C: Heatmap of All Venues × SLA Buckets
    with tabs[2]:
        heat = px.imshow(
            pivot.drop(columns="Total"),
            color_continuous_scale="Blues",
            aspect="auto",
            title="SLA Heatmap – Venue × SLA Bucket"
        )
        st.plotly_chart(heat, use_container_width=True)

    # ── VIP - SB Standby KPI Block (from 4 Oracle sources) ──────────────
    with st.expander("🛡️ VIP - SB Standby Overview (from 4 Oracle sources)", expanded=False):
        # 1⃣  Load & combine Oracle files
        files = {
            "VIP North":   "VIP North Oracle Data.xlsx",
            "VIP South":   "VIP South Oracle Data.xlsx",
            "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
            "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
        }

        standby_frames = []
        for team, path in files.items():
            try:
                tmp = pd.read_excel(path)
                tmp["Team"] = team
                standby_frames.append(tmp)
            except Exception as e:
                st.warning(f"⚠️ {path} could not be loaded – {e}")

        if not standby_frames:
            st.error("❌ None of the Oracle files could be opened – aborting block.")
            st.stop()

        standby_df = pd.concat(standby_frames, ignore_index=True)
        standby_df.columns = standby_df.columns.str.strip()
        standby_df.dropna(how="all", inplace=True)

        # 2⃣  Filter to VIP‑SB Standby rows only
        mask = standby_df["Visit Type"].astype(str).str.contains("VIP - SB Standby", case=False, na=False)
        sb_df_all = standby_df[mask].copy()
        if sb_df_all.empty:
            st.info("No rows found for 'VIP - SB Standby' in the Oracle datasets.")
            st.stop()

        # 3⃣  Basic cleaning & helpers
        sb_df_all["Date"] = pd.to_datetime(sb_df_all["Date"], errors="coerce")
        sb_df_all.dropna(subset=["Date"], inplace=True)
        sb_df_all["Month"] = sb_df_all["Date"].dt.to_period("M").astype(str)

        from datetime import time, timedelta
        def _to_td(val):
            if pd.isna(val):
                return pd.NaT
            if isinstance(val, timedelta):
                return val if val.total_seconds() > 0 else pd.NaT
            if isinstance(val, time):
                return timedelta(hours=val.hour, minutes=val.minute, seconds=val.second) if val != time(0, 0) else pd.NaT
            try:
                h, m, *s = str(val).split(":")
                h, m = int(h), int(m)
                s = int(s[0]) if s else 0
                return pd.NaT if (h == m == s == 0) else timedelta(hours=h, minutes=m, seconds=s)
            except Exception:
                return pd.NaT

        TIME_START_COLS = [c for c in sb_df_all.columns if c.lower() in {"start", "activate", "activate time"}]
        TIME_END_COLS   = [c for c in sb_df_all.columns if c.lower() in {"end", "deactivate", "deactivate time"}]

        chosen_start_col = TIME_START_COLS[0] if TIME_START_COLS else None
        chosen_end_col   = TIME_END_COLS[0]   if TIME_END_COLS   else None

        if chosen_start_col:
            sb_df_all[chosen_start_col] = sb_df_all[chosen_start_col].apply(_to_td)
        if chosen_end_col:
            sb_df_all[chosen_end_col]   = sb_df_all[chosen_end_col].apply(_to_td)

        ACTIVATE_COLS   = [c for c in sb_df_all.columns if "activate"   in c.lower()][:1]
        DEACTIVATE_COLS = [c for c in sb_df_all.columns if "deactivate" in c.lower() and c not in ACTIVATE_COLS][:1]
        chosen_act_col  = ACTIVATE_COLS[0]   if ACTIVATE_COLS   else None
        chosen_dea_col  = DEACTIVATE_COLS[0] if DEACTIVATE_COLS else None

        for col in (chosen_act_col, chosen_dea_col):
            if col:
                sb_df_all[col] = sb_df_all[col].apply(_to_td)

        # 4⃣  Split dataframes for different metric purposes
        if "Activity Status" in sb_df_all.columns:
            completed_mask = sb_df_all["Activity Status"].str.lower() == "completed"
            comp_df = sb_df_all[completed_mask].copy()
            val_df  = sb_df_all[sb_df_all["Activity Status"].str.lower().isin(["completed", "suspended"])]
        else:
            comp_df = sb_df_all.copy()
            val_df  = sb_df_all.copy()

        # 5⃣  KPI Header
        st.markdown("### 📌 Summary KPIs")
        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Total Visits (Completed)", f"{len(comp_df):,}")
        k2.metric("Total Value (£) (Comp + Susp)", f"£{val_df.get('Total Value', pd.Series(dtype=float)).sum():,.0f}")

        def _avg_col_pair(df, primary_col, require_col):
            if not primary_col or not require_col:
                return "–"
            subset = df[df[require_col].notna() & df[primary_col].notna()][primary_col]
            if subset.empty:
                return "–"
            secs = subset.dt.total_seconds().mean()
            return f"{int(secs//3600):02}:{int((secs%3600)//60):02}"

        avg_start = _avg_col_pair(comp_df, chosen_start_col, chosen_act_col if chosen_act_col else chosen_start_col)
        avg_end   = _avg_col_pair(comp_df, chosen_end_col,   chosen_dea_col if chosen_dea_col else chosen_end_col)

        k3.metric("Avg Start (if Activate present)", avg_start)
        k4.metric("Avg End (if Deactivate present)", avg_end)

        st.markdown(
            """
            >⚠️ *Visit count & time metrics only use **Completed** rows **where both primary and corresponding Activate/Deactivate times are present**.*  
            >💰 *Total Value* still aggregates **Completed + Suspended** rows.
            """
        )
        st.markdown("---")

        # 6⃣  Monthly Count (Completed)
        monthly_ct = comp_df.groupby("Month").size().reset_index(name="Visits")
        fig_bar = px.bar(monthly_ct, x="Month", y="Visits", title="Monthly Completed Count – VIP ‑ SB Standby")
        st.plotly_chart(fig_bar, use_container_width=True)

        # 7⃣  Activity Status Pie (all rows)
        if "Activity Status" in sb_df_all.columns:
            pie_df = sb_df_all["Activity Status"].value_counts().reset_index()
            pie_df.columns = ["Activity", "Count"]
            fig_pie = px.pie(pie_df, names="Activity", values="Count", title="Activity Status Distribution")
            st.plotly_chart(fig_pie, use_container_width=True)

        # 8⃣  Sunburst – Team ▸ Month ▸ Activity
        if "Activity Status" in sb_df_all.columns:
            fig_sb = px.sunburst(sb_df_all, path=["Team", "Month", "Activity Status"], title="Team • Month • Activity Breakdown")
            st.plotly_chart(fig_sb, use_container_width=True)

        # 9⃣  6‑Month Forecast (Completed counts)
        def _simple_forecast(series, periods=6):
            import numpy as np
            y = np.array(series.values)
            x = np.arange(len(y))
            if len(x) < 2:
                return np.full(periods, 0)
            coef = np.polyfit(x, y, 1)
            trend = coef[0]
            intercept = coef[1]
            x_future = np.arange(len(y), len(y)+periods)
            y_future = trend * x_future + intercept
            return np.maximum(y_future, 0).round().astype(int)
        
        series = comp_df.groupby("Month").size().sort_index()
        if len(series) >= 2:
            fc_vals = _simple_forecast(series, periods=6)
            last_p  = pd.Period(series.index.max(), freq="M")
            fut_mths = [str(last_p + i) for i in range(1, 7)]

            fc_df = pd.concat([
                pd.DataFrame({"Month": series.index, "Visits": series.values, "Kind": "Actual"}),
                pd.DataFrame({"Month": fut_mths, "Visits": fc_vals, "Kind": "Forecast"})
            ], ignore_index=True)

            fig_fc = px.line(fc_df, x="Month", y="Visits", line_dash="Kind", markers=True, title="VIP ‑ SB Standby – Actual vs Forecast (Completed)")
            st.plotly_chart(fig_fc, use_container_width=True)
        else:
            st.info("Not enough historical points to build a forecast (need ≥2 months).")

        st.caption("All averages use only rows meeting the dual‑time requirement (Start+Activate, End+Deactivate). Total value aggregates Completed + Suspended records across all four Oracle sheets.")



# ──────────────────────────────────────────────────────────────────────
# 💡 SUGGESTION BOX  – Excel‑based (Flow‑friendly) version  V2.1
#     • Keeps Table1 intact for Power Automate
#     • Unique Streamlit keys (no duplicate crash)
#     • Atomic temp‑file save (OneDrive‑safe)
#     • NEW: "🗑️ Delete" button to remove rows on‑screen
# ──────────────────────────────────────────────────────────────────────
import uuid, datetime
from pathlib import Path
import pandas as pd
import streamlit as st
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.utils import get_column_letter

# ▼━━━━━━━━━━━━━━ 1. CONFIG ━━━━━━━━━━━━━━▼
COMMENTS_FILE = Path(r"C:\Users\dah47\OneDrive - Sky\Oracle development\suggestion_comments.xlsx")
SUGG_FILE    = Path(r"C:\Users\dah47\OneDrive - Sky\Oracle development\suggestions log.xlsx")

SUGG_COLS    = ["id","num_id","timestamp","name","tag","idea","status"]
COMM_COLS    = ["id","timestamp","comment"]
NEW_STATUSES = ["Received","Notified"]
TABLE_NAME   = "Table1" ; SHEET_NAME = "Sheet1" ; TMP_SUFFIX = ".tmp.xlsx"
# ▲━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━▲

# ▼━━━━━━━━━━ 2. EXCEL HELPERS ━━━━━━━━━━▼

def _create_table(ws, df):
    n_rows, n_cols = df.shape
    ref = f"A1:{get_column_letter(n_cols)}{n_rows+1}"
    t = Table(displayName=TABLE_NAME, ref=ref)
    t.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2", showRowStripes=True)
    ws.add_table(t)

def _ensure_excel(path: Path, cols):
    if path.exists():
        return
    wb = Workbook(); ws = wb.active; ws.title = SHEET_NAME
    ws.append(cols); _create_table(ws, pd.DataFrame(columns=cols)); wb.save(path)

def _atomic_overwrite(path: Path, df: pd.DataFrame):
    tmp = path.with_suffix(TMP_SUFFIX)
    with pd.ExcelWriter(tmp, engine="openpyxl") as w:
        df.to_excel(w, sheet_name=SHEET_NAME, index=False)
    wb = load_workbook(tmp); ws = wb[SHEET_NAME]
    if TABLE_NAME in ws.tables: del ws.tables[TABLE_NAME]
    _create_table(ws, df); wb.save(tmp); tmp.replace(path)

def _load_df(path: Path, cols): _ensure_excel(path, cols); return pd.read_excel(path, dtype=str).fillna("")

def _save_df(path: Path, df: pd.DataFrame, cols): _ensure_excel(path, cols); _atomic_overwrite(path, df)

load_suggestions = lambda: _load_df(SUGG_FILE, SUGG_COLS)
save_suggestions = lambda d: _save_df(SUGG_FILE, d, SUGG_COLS)
load_comments    = lambda: _load_df(COMMENTS_FILE, COMM_COLS)
save_comments    = lambda d: _save_df(COMMENTS_FILE, d, COMM_COLS)
# ▲━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━▲

# ▼━━━━━━━━ 3. ADD / COMMENT HELPERS ━━━━━▼

def add_suggestion(row):
    df = load_suggestions(); df = pd.concat([df, pd.DataFrame([row])], ignore_index=True)
    save_suggestions(df); st.success("✅ Suggestion submitted – Teams post soon.")

def add_comment(sugg_id, text):
    df = load_comments(); df = pd.concat([df, pd.DataFrame([{"id":sugg_id,"timestamp":datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),"comment":text.strip()}])], ignore_index=True)
    save_comments(df)
# ▲━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━▲
# ▼━━━━━━━━━━ 4. STREAMLIT UI ━━━━━━━━━━━▼
if st.session_state.get("screen") == "suggestions":
    if st.button("⬅️ Back to Main Menu", use_container_width=True):
        st.session_state.screen="area_selection"; st.rerun()

    st.markdown("## 💡 Suggestion Box"); st.caption("Help improve the dashboard by submitting your ideas.")

    with st.form("suggest_form"):
        name = st.text_input("Your name"); tag = st.selectbox("What does your suggestion relate to?",["General","Missing Data","Graphs Required","Sky Stake Holder","Performance"])
        idea = st.text_area("Your suggestion (max 500 chars)", height=150, max_chars=500)
        sent = st.form_submit_button("Submit Suggestion")

    if sent and idea.strip():
        df   = load_suggestions(); next_id = str(len(df)+1)
        add_suggestion({"id":str(uuid.uuid4())[:8],"num_id":next_id,"timestamp":datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),"name":name.strip() or "Anonymous","tag":tag,"idea":idea.strip(),"status":"Received"}); st.rerun()
    elif sent:
        st.warning("Please write something before submitting.")

    tab_new, tab_prog, tab_done = st.tabs(["📥 New","🚧 In Progress","✅ Completed"])
    df_sugg, df_comm = load_suggestions(), load_comments()

    def render_list(df: pd.DataFrame, allowed_status):
        global df_sugg, df_comm
        rows = df[df["status"].isin(allowed_status)]
        if rows.empty:
            st.info("No suggestions here yet."); return
        for idx, (_, r) in enumerate(rows.iterrows()):
            safe_id = r["id"] or r.get("num_id","") or idx
            with st.expander(f"📝 {r['idea']} ({r['name']})"):
                st.caption(f"📅 {r['timestamp']}  |  🏷️ {r['tag']}")

                # ── Comments block ──
                if r["status"] == "In Progress":
                    st.markdown("**Comments**"); cts = df_comm[df_comm["id"]==r["id"]]
                    for _, c in cts.iterrows(): st.markdown(f"- *{c['timestamp']}*: {c['comment']}")
                    new_c = st.text_area("Add new comment", key=f"c_{safe_id}_{idx}")
                    if st.button("Add", key=f"add_{safe_id}_{idx}") and new_c.strip():
                        add_comment(r["id"], new_c); st.success("Added!"); st.rerun()

                # ── Control buttons row ──
                btn_cols = st.columns([1,1,1,5])  # In‑Prog / Complete / Delete / Spacer

                # → Mark In Progress
                if r["status"] in NEW_STATUSES and btn_cols[0].button("Start", key=f"start_{safe_id}_{idx}"):
                    df_sugg.loc[df_sugg["id"]==r["id"],"status"]="In Progress"; save_suggestions(df_sugg); st.rerun()

                # ✓ Completed checkbox
                if r["status"] == "In Progress" and btn_cols[1].checkbox("Done", key=f"done_{safe_id}_{idx}"):
                    df_sugg.loc[df_sugg["id"]==r["id"],"status"]="Completed"; save_suggestions(df_sugg); st.rerun()

                # 🗑️ Delete button
                if btn_cols[2].button("Delete", key=f"del_{safe_id}_{idx}"):
                    df_sugg = df_sugg[df_sugg["id"]!=r["id"]]; save_suggestions(df_sugg); st.success("Deleted."); st.rerun()

    with tab_new:  render_list(df_sugg, NEW_STATUSES)
    with tab_prog: render_list(df_sugg, ["In Progress"])
    with tab_done: render_list(df_sugg, ["Completed"])




# ▲━━━━━━━Highlands_islands━━━━━━━▲

import streamlit as st
import pandas as pd
import plotly.express as px
import numpy as np

# Function to safely convert columns to numeric
def safe_numeric(df, cols):
    for col in cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col].astype(str).str.replace(',', ''), errors='coerce')
    return df

@st.cache_data
def load_data(file_path):
    sheets = [
        "2022", "2023", "2024", "2025",
        "Company 2022", "Company 2023", "Company 2024", "Company 2025"
    ]
    dfs = pd.read_excel(file_path, sheet_name=sheets)
    for sheet, df in dfs.items():
        df.columns = df.columns.str.strip()
        df = safe_numeric(df, ['Issued Visits', 'Completed Visits', 'Average Complete Job Time (Min)',
                              '7 Day Revisits', '7 Day Revisits %', '30 Day Revisits', '30 Day Revisits %',
                              'Surveys', 'NPS%', 'NPS'])
        dfs[sheet] = df
    return dfs

dfs = load_data("Highlands Islands.xlsx")

# Separate sheets for tabs
yearly_sheets = ["2022", "2023", "2024", "2025"]
company_sheets = ["Company 2022", "Company 2023", "Company 2024", "Company 2025"]

# Helper functions at top-level (outside tabs)
def safe_sum(df, col):
    if col in df.columns:
        return pd.to_numeric(df[col], errors='coerce').sum(skipna=True)
    return 0

def weighted_avg(df, val_col, weight_col):
    if val_col in df.columns and weight_col in df.columns:
        df_clean = df[[val_col, weight_col]].dropna()
        if df_clean.empty:
            return None
        total_weight = df_clean[weight_col].sum()
        if total_weight == 0:
            return None
        return (df_clean[val_col] * df_clean[weight_col]).sum() / total_weight
    return None

def generate_summary(df, scope="Yearly"):
    total_issued = safe_sum(df, 'Issued Visits')
    total_completed = safe_sum(df, 'Completed Visits')
    completion_rate = (total_completed / total_issued * 100) if total_issued else 0
    total_not_done = safe_sum(df, 'Not Done Vists')
    not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
    total_7day_revisits = safe_sum(df, '7 Day Revisits')
    total_30day_revisits = safe_sum(df, '30 Day Revisits')
    weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
    total_surveys = safe_sum(df, 'Surveys')
    nps_col = 'NPS' if 'NPS' in df.columns else 'NPS%'
    weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col in df.columns else None

    summary = f"In this {scope} dataset, there were a total of {total_issued:,} issued visits, of which {total_completed:,} were completed, resulting in a completion rate of {completion_rate:.2f}%. "
    summary += f"Not done visits accounted for {total_not_done:,} visits ({not_done_pct:.2f}%). "
    summary += f"Revisit rates included {total_7day_revisits:,} 7-day revisits and {total_30day_revisits:,} 30-day revisits. "
    if weighted_job_time:
        summary += f"The average job completion time was approximately {weighted_job_time:.2f} minutes. "
    if total_surveys > 0 and weighted_nps is not None:
        summary += f"Customer satisfaction measured via NPS had an average score of {weighted_nps:.2f} based on {total_surveys:,} surveys. "
    else:
        summary += "NPS data is currently unavailable. "

    summary += f"This data provides valuable insights to guide operational and customer experience improvements."

    return summary

if st.session_state.get("screen") == "highlands_islands":
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_highlands"):
        st.session_state.screen = "area_selection"
        st.rerun()

    st.title("🗺️ Highlands & Islands Dashboard")

    tab1, tab2 = st.tabs(["Yearly Data", "Company Data"])

    with tab1:
        # default to first year so we can display summary immediately
        year_choice = st.selectbox("Select Year", yearly_sheets)
        df = dfs[year_choice]

        # Show advanced summary paragraph before the dropdown
        if not df.empty:
            st.markdown(f"### Advanced Summary for {year_choice}")
            st.info(generate_summary(df, "Yearly"))
        else:
            st.info("No data available for this year.")

        # Yearly KPIs
        with st.expander("Yearly KPIs", expanded=True):
            total_issued = df['Issued Visits'].sum()
            total_completed = df['Completed Visits'].sum()
            completion_rate = (total_completed / total_issued * 100) if total_issued else 0

            total_not_done = df['Not Done Vists'].sum() if 'Not Done Vists' in df.columns else 0
            not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
            total_7day_revisits = df['7 Day Revisits'].sum() if '7 Day Revisits' in df.columns else 0
            weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
            total_surveys = df['Surveys'].sum() if 'Surveys' in df.columns else 0
            weighted_nps = weighted_avg(df, 'NPS', 'Surveys') if 'NPS' in df.columns else None
            weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')

            cols = st.columns(3)

            cols[0].metric("Total Issued Visits", f"{total_issued:,.0f}")
            cols[1].metric("Total Completed Visits", f"{total_completed:,.0f}")
            cols[2].metric("Completion Rate (%)", f"{completion_rate:.2f}%")

            cols[0].metric("Total Not Done Visits", f"{total_not_done:,.0f}")
            cols[1].metric("Not Done Visits %", f"{not_done_pct:.2f}%")
            cols[2].metric("Total 7 Day Revisits", f"{total_7day_revisits:,.0f}")

            cols[0].metric("Weighted 7 Day Revisits %", f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A")
            cols[1].metric("Total Surveys", f"{total_surveys:,.0f}")
            cols[2].metric("Weighted Average NPS", f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A")

        with st.expander("Sunburst Chart - Issued Visits by Field by FRU"):
            if "Field by FRU" in df.columns:
                df_sb = df.copy()
                df_sb['Field by FRU'] = df_sb['Field by FRU'].fillna("Unknown")
                df_sb['Completion %'] = (df_sb['Completed Visits'] / df_sb['Issued Visits']).fillna(0) * 100

                fig = px.sunburst(df_sb, path=['Field by FRU'], values='Issued Visits',
                                  color='Completion %', color_continuous_scale='RdYlGn',
                                  title="Issued Visits by Field by FRU (colored by Completion %)")
                st.plotly_chart(fig, use_container_width=True, key=f"sunburst_yearly_{year_choice}")
            else:
                st.info("Column 'Field by FRU' not found in this dataset.")

    with tab2:
        comp_choice = st.selectbox("Select Company Year", company_sheets)
        df = dfs[comp_choice]

        if not df.empty:
            st.markdown(f"### Advanced Summary for {comp_choice}")
            st.info(generate_summary(df, "Company"))
        else:
            st.info("No data available for this company dataset.")

        with st.expander("Company KPIs", expanded=True):
            # Group by company and aggregate relevant metrics
            comp_kpis = df.groupby('Company').agg(
                Total_Issued_Visits = ('Issued Visits', 'sum'),
                Total_Completed_Visits = ('Completed Visits', 'sum')
            ).reset_index()

            # Calculate Completion Rate per company
            comp_kpis['Completion Rate (%)'] = (comp_kpis['Total_Completed_Visits'] / comp_kpis['Total_Issued_Visits']) * 100

            # Format numbers for display
            comp_kpis['Total_Issued_Visits'] = comp_kpis['Total_Issued_Visits'].map('{:,.0f}'.format)
            comp_kpis['Total_Completed_Visits'] = comp_kpis['Total_Completed_Visits'].map('{:,.0f}'.format)
            comp_kpis['Completion Rate (%)'] = comp_kpis['Completion Rate (%)'].map('{:.2f}%'.format)

            # Rename columns for nice display
            comp_kpis = comp_kpis.rename(columns={
                'Company': 'Company',
                'Total_Issued_Visits': 'Total Issued Visits',
                'Total_Completed_Visits': 'Total Completed Visits',
                'Completion Rate (%)': 'Completion Rate (%)'
            })

            st.dataframe(comp_kpis)



        with tab1:
            df = dfs[year_choice]

            if not df.empty:
                with st.expander("Section 3: Detailed Metrics & Charts (Yearly)", expanded=False):
                    total_issued = safe_sum(df, 'Issued Visits')
                    total_completed = safe_sum(df, 'Completed Visits')
                    total_not_done = safe_sum(df, 'Not Done Vists')  # spelling as per your header
                    total_7day_revisits = safe_sum(df, '7 Day Revisits')
                    total_30day_revisits = safe_sum(df, '30 Day Revisits')
                    total_surveys = safe_sum(df, 'Surveys')

                    weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
                    weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
                    weighted_30day_revisit_pct = weighted_avg(df, '30 Day Revisits %', 'Completed Visits')

                    nps_col = 'NPS' if 'NPS' in df.columns else None
                    weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col else None

                    # KPI Table
                    kpi_data = {
                        "Metric": [
                            "Total Issued Visits",
                            "Total Completed Visits",
                            "Total Not Done Visits",
                            "Total 7 Day Revisits",
                            "Weighted 7 Day Revisits %",
                            "Total 30 Day Revisits",
                            "Weighted 30 Day Revisits %",
                            "Total Surveys",
                            "Weighted Average NPS",
                            "Weighted Average Job Time (Min)"
                        ],
                        "Value": [
                            f"{total_issued:,.0f}",
                            f"{total_completed:,.0f}",
                            f"{total_not_done:,.0f}",
                            f"{total_7day_revisits:,.0f}",
                            f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A",
                            f"{total_30day_revisits:,.0f}",
                            f"{weighted_30day_revisit_pct:.2f}%" if weighted_30day_revisit_pct is not None else "N/A",
                            f"{total_surveys:,.0f}",
                            f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A",
                            f"{weighted_job_time:.2f}" if weighted_job_time is not None else "N/A",
                        ]
                    }
                    st.table(pd.DataFrame(kpi_data))

                    # Completion vs Not Done Pie Chart
                    pie_df = pd.DataFrame({
                        'Status': ['Completed', 'Not Done'],
                        'Visits': [total_completed, total_not_done]
                    })
                    fig_pie = px.pie(pie_df, names='Status', values='Visits', title="Completion Status Distribution")
                    st.plotly_chart(fig_pie, use_container_width=True, key=f"pie_completion_yearly_{year_choice}")

                    # Revisits Bar Chart (by Field by FRU)
                    revisit_cols = [col for col in ['7 Day Revisits', '30 Day Revisits'] if col in df.columns]
                    if 'Field by FRU' in df.columns and revisit_cols:
                        revisit_df = df.groupby('Field by FRU')[revisit_cols].sum().reset_index()
                        fig_bar = px.bar(revisit_df, x='Field by FRU', y=revisit_cols, barmode='group', title="Revisits by Field by FRU")
                        st.plotly_chart(fig_bar, use_container_width=True, key=f"bar_revisits_yearly_{year_choice}")

        with tab2:
            df = dfs[comp_choice]

            if not df.empty:
                with st.expander("Section 3: Detailed Metrics & Charts (Company)", expanded=False):
                    total_issued = safe_sum(df, 'Issued Visits')
                    total_completed = safe_sum(df, 'Completed Visits')
                    total_not_done = safe_sum(df, 'Not Done Vists')
                    total_7day_revisits = safe_sum(df, '7 Day Revisits')
                    total_30day_revisits = safe_sum(df, '30 Day Revisits')
                    total_surveys = safe_sum(df, 'Surveys')

                    weighted_job_time = weighted_avg(df, 'Average Complete Job Time (Min)', 'Completed Visits')
                    weighted_7day_revisit_pct = weighted_avg(df, '7 Day Revisits %', 'Completed Visits')
                    weighted_30day_revisit_pct = weighted_avg(df, '30 Day Revisits %', 'Completed Visits')

                    nps_col = 'NPS%' if 'NPS%' in df.columns else None
                    weighted_nps = weighted_avg(df, nps_col, 'Surveys') if nps_col else None

                    # KPI Table
                    kpi_data = {
                        "Metric": [
                            "Total Issued Visits",
                            "Total Completed Visits",
                            "Total Not Done Visits",
                            "Total 7 Day Revisits",
                            "Weighted 7 Day Revisits %",
                            "Total 30 Day Revisits",
                            "Weighted 30 Day Revisits %",
                            "Total Surveys",
                            "Weighted Average NPS",
                            "Weighted Average Job Time (Min)"
                        ],
                        "Value": [
                            f"{total_issued:,.0f}",
                            f"{total_completed:,.0f}",
                            f"{total_not_done:,.0f}",
                            f"{total_7day_revisits:,.0f}",
                            f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A",
                            f"{total_30day_revisits:,.0f}",
                            f"{weighted_30day_revisit_pct:.2f}%" if weighted_30day_revisit_pct is not None else "N/A",
                            f"{total_surveys:,.0f}",
                            f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A",
                            f"{weighted_job_time:.2f}" if weighted_job_time is not None else "N/A",
                        ]
                    }
                    st.table(pd.DataFrame(kpi_data))

                    # Completion vs Not Done Pie Chart
                    pie_df = pd.DataFrame({
                        'Status': ['Completed', 'Not Done'],
                        'Visits': [total_completed, total_not_done]
                    })
                    fig_pie = px.pie(pie_df, names='Status', values='Visits', title="Completion Status Distribution")
                    st.plotly_chart(fig_pie, use_container_width=True, key=f"pie_completion_company_{comp_choice}")

                    # Revisits Bar Chart (by Company)
                    revisit_cols = [col for col in ['7 Day Revisits', '30 Day Revisits'] if col in df.columns]
                    if 'Company' in df.columns and revisit_cols:
                        revisit_df = df.groupby('Company')[revisit_cols].sum().reset_index()
                        fig_bar = px.bar(revisit_df, x='Company', y=revisit_cols, barmode='group', title="Revisits by Company")
                        st.plotly_chart(fig_bar, use_container_width=True, key=f"bar_revisits_company_{comp_choice}")

 
#SECTION 4#

        import matplotlib.pyplot as plt
        import seaborn as sns

        with tab1:
            # Yearly Tab
            with st.expander("Section 4: Advanced Analytics (Yearly)", expanded=False):
                # Completion % Trend over Years
                st.subheader("Completion % Trend Over Years")
                year_data = []
                for year in yearly_sheets:
                    d = dfs[year]
                    total_issued = safe_sum(d, 'Issued Visits')
                    total_completed = safe_sum(d, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    year_data.append({"Year": year, "Completion %": comp_pct})
                trend_df = pd.DataFrame(year_data)
                fig_trend = px.line(trend_df, x='Year', y='Completion %', markers=True)
                st.plotly_chart(fig_trend, use_container_width=True, key=f"comp_trend_yearly")

                # Top 5 Field by FRU by Issued Visits for selected year
                st.subheader(f"Top 5 Field by FRU by Issued Visits ({year_choice})")
                if 'Field by FRU' in df.columns:
                    top5 = df.groupby('Field by FRU')['Issued Visits'].sum().nlargest(5).reset_index()
                    fig_top5 = px.bar(top5, x='Field by FRU', y='Issued Visits', title="Top 5 Field by FRU")
                    st.plotly_chart(fig_top5, use_container_width=True, key=f"top5_frus_{year_choice}")
                else:
                    st.info("No 'Field by FRU' data available")

                # Not Done Visits Pie Chart
                st.subheader("Not Done Visits Proportion")
                total_not_done = safe_sum(df, 'Not Done Vists')
                pie_nd_df = pd.DataFrame({
                    'Status': ['Completed', 'Not Done'],
                    'Visits': [safe_sum(df, 'Completed Visits'), total_not_done]
                })
                fig_nd_pie = px.pie(pie_nd_df, names='Status', values='Visits', title="Not Done Visits Proportion")
                st.plotly_chart(fig_nd_pie, use_container_width=True, key=f"nd_pie_{year_choice}")

                # NPS Distribution (if available)
                if 'NPS' in df.columns:
                    st.subheader("NPS Distribution")
                    fig_nps = px.histogram(df, x='NPS', nbins=20, title="NPS Score Distribution")
                    st.plotly_chart(fig_nps, use_container_width=True, key=f"nps_dist_{year_choice}")
                else:
                    st.info("NPS data not available")

        with tab2:
            # Company Tab
            with st.expander("Section 4: Advanced Analytics (Company)", expanded=False):
                # Completion % Trend over Company Years
                st.subheader("Completion % Trend Over Company Years")
                comp_data = []
                for comp_year in company_sheets:
                    d = dfs[comp_year]
                    total_issued = safe_sum(d, 'Issued Visits')
                    total_completed = safe_sum(d, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    comp_data.append({"Year": comp_year, "Completion %": comp_pct})
                comp_trend_df = pd.DataFrame(comp_data)
                fig_comp_trend = px.line(comp_trend_df, x='Year', y='Completion %', markers=True)
                st.plotly_chart(fig_comp_trend, use_container_width=True, key=f"comp_trend_company")

                # Top 5 Companies by Issued Visits for selected company year
                st.subheader(f"Top 5 Companies by Issued Visits ({comp_choice})")
                if 'Company' in df.columns:
                    top5_comp = df.groupby('Company')['Issued Visits'].sum().nlargest(5).reset_index()
                    fig_top5_comp = px.bar(top5_comp, x='Company', y='Issued Visits', title="Top 5 Companies")
                    st.plotly_chart(fig_top5_comp, use_container_width=True, key=f"top5_company_{comp_choice}")
                else:
                    st.info("No 'Company' data available")

                # Not Done Visits Pie Chart
                st.subheader("Not Done Visits Proportion")
                total_not_done = safe_sum(df, 'Not Done Vists')
                pie_nd_df = pd.DataFrame({
                    'Status': ['Completed', 'Not Done'],
                    'Visits': [safe_sum(df, 'Completed Visits'), total_not_done]
                })
                fig_nd_pie = px.pie(pie_nd_df, names='Status', values='Visits', title="Not Done Visits Proportion")
                st.plotly_chart(fig_nd_pie, use_container_width=True, key=f"nd_pie_company_{comp_choice}")

                # NPS% Distribution (if available)
                if 'NPS%' in df.columns:
                    st.subheader("NPS% Distribution")
                    fig_nps = px.histogram(df, x='NPS%', nbins=20, title="NPS% Score Distribution")
                    st.plotly_chart(fig_nps, use_container_width=True, key=f"nps_dist_company_{comp_choice}")
                else:
                    st.info("NPS% data not available")



        with tab1:
            df = dfs[year_choice]

            if not df.empty:
                with st.expander("Section 5: Interactive Filters & Drilldown (Yearly)", expanded=False):
                    # Filter by Field by FRU
                    fru_options = df['Field by FRU'].dropna().unique()
                    selected_fru = st.multiselect(
                        "Filter by Field by FRU",
                        options=fru_options,
                        default=fru_options,
                        key=f"filter_fru_{year_choice}"
                    )

                    # Filter by Completion Status
                    status_options = ['Completed', 'Not Done']
                    selected_status = st.multiselect(
                        "Filter by Completion Status",
                        options=status_options,
                        default=status_options,
                        key=f"filter_status_{year_choice}"
                    )

                    # Filter dataframe based on selections
                    filtered_df = df[df['Field by FRU'].isin(selected_fru)]

                    # Create completion mask
                    if 'Not Done Vists' in filtered_df.columns and 'Completed Visits' in filtered_df.columns:
                        if set(selected_status) == {'Completed', 'Not Done'}:
                            pass
                        elif selected_status == ['Completed']:
                            filtered_df = filtered_df[filtered_df['Completed Visits'] > 0]
                        elif selected_status == ['Not Done']:
                            filtered_df = filtered_df[filtered_df['Not Done Vists'] > 0]
                        else:
                            filtered_df = filtered_df.iloc[0:0]

                    # Show KPIs for filtered data
                    total_issued = safe_sum(filtered_df, 'Issued Visits')
                    total_completed = safe_sum(filtered_df, 'Completed Visits')
                    total_not_done = safe_sum(filtered_df, 'Not Done Vists')

                    st.write("### Filtered KPIs")
                    st.metric("Total Issued Visits", f"{total_issued:,.0f}")
                    st.metric("Total Completed Visits", f"{total_completed:,.0f}")
                    st.metric("Total Not Done Visits", f"{total_not_done:,.0f}")

                    # Optionally show filtered data table
                    if st.checkbox("Show filtered data table", key=f"show_table_{year_choice}"):
                        st.dataframe(filtered_df)

        with tab2:
            df = dfs[comp_choice]

            if not df.empty:
                with st.expander("Section 5: Interactive Filters & Drilldown (Company)", expanded=False):
                    # Filter by Company
                    comp_options = df['Company'].dropna().unique()
                    selected_comp = st.multiselect(
                        "Filter by Company",
                        options=comp_options,
                        default=comp_options,
                        key=f"filter_company_{comp_choice}"
                    )

                    # Filter by Completion Status
                    status_options = ['Completed', 'Not Done']
                    selected_status = st.multiselect(
                        "Filter by Completion Status",
                        options=status_options,
                        default=status_options,
                        key=f"filter_status_company_{comp_choice}"
                    )

                    filtered_df = df[df['Company'].isin(selected_comp)]

                    if 'Not Done Vists' in filtered_df.columns and 'Completed Visits' in filtered_df.columns:
                        if set(selected_status) == {'Completed', 'Not Done'}:
                            pass
                        elif selected_status == ['Completed']:
                            filtered_df = filtered_df[filtered_df['Completed Visits'] > 0]
                        elif selected_status == ['Not Done']:
                            filtered_df = filtered_df[filtered_df['Not Done Vists'] > 0]
                        else:
                            filtered_df = filtered_df.iloc[0:0]

                    # Calculate KPIs for filtered data
                    total_issued = safe_sum(filtered_df, 'Issued Visits')
                    total_completed = safe_sum(filtered_df, 'Completed Visits')
                    total_not_done = safe_sum(filtered_df, 'Not Done Vists')
                    not_done_pct = (total_not_done / total_issued * 100) if total_issued else 0
                    total_7day_revisits = safe_sum(filtered_df, '7 Day Revisits')

                    weighted_7day_revisit_pct = weighted_avg(filtered_df, '7 Day Revisits %', 'Completed Visits')
                    total_surveys = safe_sum(filtered_df, 'Surveys')
                    weighted_nps = weighted_avg(filtered_df, 'NPS%', 'Surveys')

                    completion_rate = (total_completed / total_issued * 100) if total_issued else 0

                    st.write("### Filtered KPIs")
                    cols = st.columns(3)

                    cols[0].metric("Total Issued Visits", f"{total_issued:,.0f}")
                    cols[1].metric("Total Completed Visits", f"{total_completed:,.0f}")
                    cols[2].metric("Completion Rate (%)", f"{completion_rate:.2f}%")

                    cols[0].metric("Total Not Done Visits", f"{total_not_done:,.0f}")
                    cols[1].metric("Not Done Visits %", f"{not_done_pct:.2f}%")
                    cols[2].metric("Total 7 Day Revisits", f"{total_7day_revisits:,.0f}")

                    cols[0].metric("Weighted 7 Day Revisits %", f"{weighted_7day_revisit_pct:.2f}%" if weighted_7day_revisit_pct is not None else "N/A")
                    cols[1].metric("Total Surveys", f"{total_surveys:,.0f}")
                    cols[2].metric("Weighted Average NPS%", f"{weighted_nps:.2f}" if weighted_nps is not None else "N/A")

                    if st.checkbox("Show filtered data table", key=f"show_table_company_{comp_choice}"):
                        st.dataframe(filtered_df)



        from sklearn.linear_model import LinearRegression
        import numpy as np
        import plotly.express as px

        with tab1:
            with st.expander("Section 6: Yearly Trends & Forecasts", expanded=False):
                # Combine all yearly data into one DataFrame
                combined_yearly = []
                for year in yearly_sheets:
                    df_year = dfs[year]
                    total_issued = safe_sum(df_year, 'Issued Visits')
                    total_completed = safe_sum(df_year, 'Completed Visits')
                    comp_pct = (total_completed / total_issued * 100) if total_issued else 0
                    combined_yearly.append({
                        "Year": int(year),
                        "Issued Visits": total_issued,
                        "Completed Visits": total_completed,
                        "Completion %": comp_pct
                    })
                trend_df = pd.DataFrame(combined_yearly).sort_values("Year")

                # KPIs to plot
                kpis = ["Issued Visits", "Completed Visits", "Completion %"]

                for kpi in kpis:
                    st.subheader(f"{kpi} Trend Over Years")
                    fig = px.line(trend_df, x='Year', y=kpi, markers=True, title=f"{kpi} Trend Over Years")

                    # Forecast next year using simple linear regression
                    X = trend_df['Year'].values.reshape(-1,1)
                    y = trend_df[kpi].values

                    model = LinearRegression()
                    model.fit(X, y)

                    next_year = np.array([[trend_df['Year'].max() + 1]])
                    forecast = model.predict(next_year)[0]

                    # Add forecast point to the chart
                    fig.add_scatter(x=[next_year[0][0]], y=[forecast], mode='markers+text',
                                    marker=dict(color='red', size=10),
                                    text=["Forecast"],
                                    textposition="top center",
                                    name="Forecast")

                    st.plotly_chart(fig, use_container_width=True, key=f"trend_forecast_{kpi.replace(' ', '_')}")


# ------- Section 1 Forecasts ------- #

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px

@st.cache_data
def load_file(path):
    try:
        df = pd.read_excel(path)
        if 'Date' in df.columns:
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce')  # Convert Date if it exists
        return df
    except Exception as e:
        st.error(f"Failed to load data from {path}: {e}")
        return pd.DataFrame()


# Your forecast_visit_type function here (unchanged)...
def forecast_visit_type(df, visit_type, selected_activity_statuses, periods=6):
    df_filtered = df[df["Visit Type"] == visit_type].copy()
    df_filtered["Date"] = pd.to_datetime(df_filtered["Date"], errors="coerce")
    df_filtered = df[
        (df["Visit Type"] == visit_type) &
        (df["Activity Status"].isin(selected_activity_statuses))
    ].copy()

    df_filtered.dropna(subset=["Date"], inplace=True)

    # Group by month and count completed visits
    monthly_counts = (
        df_filtered.groupby(pd.Grouper(key="Date", freq="M"))
        .size()
        .reset_index(name="Completed Visits")
    )

    if len(monthly_counts) < 3:
        return None  # Not enough data to forecast

    monthly_counts["Month_Num"] = np.arange(len(monthly_counts))
    X = monthly_counts["Month_Num"].values.reshape(-1, 1)
    y = monthly_counts["Completed Visits"].values

    # Simple linear regression to predict future values
    coeffs = np.polyfit(X.flatten(), y, 1)
    poly = np.poly1d(coeffs)

    # Forecast future months
    future_X = np.arange(len(monthly_counts), len(monthly_counts) + periods)
    future_dates = pd.date_range(
        monthly_counts["Date"].max() + pd.offsets.MonthBegin(),
        periods=periods,
        freq="M"
    )
    forecast_vals = poly(future_X).clip(0).round().astype(int)

    # Build forecast DataFrame
    forecast_df = pd.DataFrame({
        "Month": future_dates.strftime("%b %Y"),
        "Forecasted Completed Visits": forecast_vals
    })

    return monthly_counts, forecast_df


def forecast_ui(df, visit_types, selected_activity_statuses):
    cols_per_row = 8
    n = len(visit_types)

    # Reset selected visit types on page load or if not initialized
    if "selected_visit_types" not in st.session_state or st.session_state.get("reset_forecast", True):
        st.session_state.selected_visit_types = set()
        st.session_state.reset_forecast = False

    # Display historical combined visits line chart (collapsible)
    with st.expander("Historical Completed Visits - All Visit Types", expanded=True):
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        filtered_df = df[df["Activity Status"].isin(selected_activity_statuses)]
        historical_all = (
            filtered_df.groupby(pd.Grouper(key="Date", freq="M"))
            .size()
            .reset_index(name="Completed Visits")
        )
        if not historical_all.empty:
            fig = px.line(historical_all, x="Date", y="Completed Visits", title="All Visit Types Historical Completed Visits")
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No historical data available for selected activity statuses.")

    # Display grid of checkboxes for selection
    st.write("### Select Visit Types to Forecast")
    for i in range(0, n, cols_per_row):
        cols = st.columns(cols_per_row)
        for j, visit_type in enumerate(visit_types[i:i+cols_per_row]):
            checked = visit_type in st.session_state.selected_visit_types
            new_val = cols[j].checkbox(visit_type, value=checked, key=f"vt_{i+j}")
            if new_val and visit_type not in st.session_state.selected_visit_types:
                st.session_state.selected_visit_types.add(visit_type)
            elif not new_val and visit_type in st.session_state.selected_visit_types:
                st.session_state.selected_visit_types.remove(visit_type)

    # Show forecasts only for selected visit types
    if not st.session_state.selected_visit_types:
        st.info("Select one or more visit types above to see forecasts.")
        return

    for vt in st.session_state.selected_visit_types:
        st.markdown(f"## Visit Type: {vt}")
        historical, forecast = forecast_visit_type(df, vt, selected_activity_statuses)
        if historical is None:
            st.info("Not enough data to forecast")
        else:
            st.write("### Historical Completed Visits")
            st.line_chart(historical.set_index("Date")["Completed Visits"])

            st.write("### Forecast for Next 6 Months")
            st.dataframe(forecast)

            fig = px.line(forecast, x="Month", y="Forecasted Completed Visits", markers=True)
            st.plotly_chart(fig, use_container_width=True)



# ▼━━━━━━━━━━ SECTION 2 ━━━━━━━━━━━▼ 
import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

def forecast_visit_type_with_pct_change(df, visit_type, selected_activity_statuses, selected_day='All', periods=6):
    """
    Forecast completed visits per month for a given visit type and calculate percentage change month-on-month.
    Returns historical data and forecast dataframe with % change.
    """
    # Filter for visit type and activity status
    df_filtered = df[
        (df["Visit Type"] == visit_type) & 
        (df["Activity Status"].isin(selected_activity_statuses))
    ].copy()
    
    df_filtered["Date"] = pd.to_datetime(df_filtered["Date"], errors="coerce")
    df_filtered.dropna(subset=["Date"], inplace=True)

    # Filter by day of week if not 'All'
    if selected_day != 'All':
        df_filtered = df_filtered[df_filtered['Date'].dt.day_name() == selected_day]

    # Group by month
    monthly_counts = df_filtered.groupby(pd.Grouper(key="Date", freq="M")).size().reset_index(name="Completed Visits")

    if len(monthly_counts) < 3:
        return None, None  # Not enough data to forecast

    monthly_counts["Month_Num"] = np.arange(len(monthly_counts))
    X = monthly_counts["Month_Num"].values.reshape(-1, 1)
    y = monthly_counts["Completed Visits"].values

    coeffs = np.polyfit(X.flatten(), y, 1)
    poly = np.poly1d(coeffs)

    future_X = np.arange(len(monthly_counts), len(monthly_counts) + periods)
    future_dates = pd.date_range(monthly_counts["Date"].max() + pd.offsets.MonthBegin(), periods=periods, freq="M")
    forecast_vals = poly(future_X).clip(0).round().astype(int)

    forecast_df = pd.DataFrame({
        "Month": future_dates.strftime("%b %Y"),
        "Forecasted Completed Visits": forecast_vals
    })

    forecast_df["Pct Change"] = forecast_df["Forecasted Completed Visits"].pct_change().fillna(0) * 100
    forecast_df["Pct Change"] = forecast_df["Pct Change"].map(lambda x: f"{x:+.2f}%")

    return monthly_counts, forecast_df

def combined_forecast_ui(df, visit_types, selected_activity_statuses, periods=6):
    # Prepare combined data
    combined_data = []

    for visit_type in visit_types:
        historical, forecast = forecast_visit_type_with_pct_change(df, visit_type, selected_activity_statuses, periods)
        if historical is not None:
            # Prepare forecast data for line plot (date + forecasted visits)
            forecast_dates = pd.to_datetime(forecast["Month"], format="%b %Y")
            combined_data.append(pd.DataFrame({
                "Month": forecast_dates,
                "Forecasted Completed Visits": forecast["Forecasted Completed Visits"],
                "Visit Type": visit_type
            }))

    if combined_data:
        combined_df = pd.concat(combined_data)

        fig = px.line(
            combined_df,
            x="Month",
            y="Forecasted Completed Visits",
            color="Visit Type",
            title="Combined Forecast for Selected Visit Types",
            markers=True
        )
        st.plotly_chart(fig, use_container_width=True)

    # Then show individual graphs without tables
    for visit_type in visit_types:
        with st.expander(f"Visit Type: {visit_type}", expanded=False):
            historical, _ = forecast_visit_type_with_pct_change(df, visit_type, selected_activity_statuses, periods)
            if historical is None:
                st.info("Not enough data to forecast for this visit type.")
            else:
                st.write("### Historical Completed Visits")
                st.line_chart(historical.set_index("Date")["Completed Visits"])



def forecast_ui_with_tables(df, visit_types, selected_activity_statuses, selected_day='All'):
    """
    Display forecast charts and tables for each selected visit type, full width.
    """
    for visit_type in visit_types:
        with st.expander(f"Visit Type: {visit_type}", expanded=False):
            historical, forecast = forecast_visit_type_with_pct_change(df, visit_type, selected_activity_statuses, selected_day)
            if historical is None:
                st.info("Not enough data to forecast for this visit type.")
            else:
                st.write("### Historical Completed Visits")
                st.line_chart(historical.set_index("Date")["Completed Visits"])

                st.write("### Forecast for Next 6 Months")
                st.dataframe(forecast)

                combined = pd.concat([
                    historical[["Date", "Completed Visits"]].rename(columns={"Date": "Month"}),
                    pd.DataFrame({
                        "Month": pd.to_datetime(forecast["Month"], format="%b %Y"),
                        "Completed Visits": forecast["Forecasted Completed Visits"]
                    })
                ])

                fig = px.line(
                    combined,
                    x="Month",
                    y="Completed Visits",
                    title=f"Completed Visits Forecast for {visit_type}",
                    markers=True
                )
                st.plotly_chart(fig, use_container_width=True)





# ▼━━━━━━━━━━ SECTION 3 ━━━━━━━━━━━▼  
if st.session_state.get("screen") == "Forecasts":
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_button"):
        st.session_state.screen = "area_selection"
        st.rerun()  # or st.rerun()

    st.title("📊 Forecasts")

    # Load datasets
    vip_south_df = load_file("VIP South Oracle Data.xlsx")
    vip_north_df = load_file("VIP North Oracle Data.xlsx")
    tier2_south_df = load_file("Tier 2 South Oracle Data.xlsx")
    tier2_north_df = load_file("Tier 2 North Oracle Data.xlsx")

    datasets = {
        "All": pd.concat([vip_south_df, vip_north_df, tier2_south_df, tier2_north_df], ignore_index=True),
        "VIP South": vip_south_df,
        "VIP North": vip_north_df,
        "Tier 2 South": tier2_south_df,
        "Tier 2 North": tier2_north_df,
    }

    day_options = ['All', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    selected_day = st.selectbox("Select Day of Week", day_options, key="day_of_week_select")

    selected_dataset_name = st.selectbox("Select dataset", list(datasets.keys()), key="dataset_select")

    if selected_dataset_name:
        selected_df = datasets[selected_dataset_name]

        visit_types = sorted(selected_df['Visit Type'].dropna().unique())
        activity_statuses = sorted(selected_df['Activity Status'].dropna().unique())

        selected_visit_types = st.multiselect("Select Visit Type(s)", visit_types, default=visit_types, key="visit_types_multiselect")

        selected_activity_statuses = st.multiselect("Select Activity Status(es)", activity_statuses, default=activity_statuses, key="activity_status_multiselect")

        # Filter dataframe with visit type and activity status
        filtered_df = selected_df[
            selected_df['Visit Type'].isin(selected_visit_types) &
            selected_df['Activity Status'].isin(selected_activity_statuses)
        ]

        if selected_day != 'All':
            filtered_df = filtered_df[filtered_df['Date'].dt.day_name() == selected_day]

        # Conditionally call forecast functions with all needed parameters
        if len(selected_visit_types) == 1:
            forecast_ui_with_tables(filtered_df, selected_visit_types, selected_activity_statuses, selected_day)
        elif len(selected_visit_types) > 1:
            combined_forecast_ui(filtered_df, selected_visit_types, selected_activity_statuses, selected_day)
        else:
            st.info("Please select at least one Visit Type.")

    # Combine all datasets for any other use
    combined_df = pd.concat([vip_south_df, vip_north_df, tier2_south_df, tier2_north_df], ignore_index=True)
    combined_df['Visit Type'] = combined_df['Visit Type'].astype(str).str.strip()
    visit_types = sorted(combined_df['Visit Type'].dropna().unique())


# --- NUMBER 7 ---#
# --- SECTION: FILE CHOICES / DATA SOURCES ---
file_map = {
    "AI Test SB Visits": "AI Test SB Visits.xlsx",
    "Invoice Data AI": "Invoice Data AI.xlsx",
    "VIP North Oracle Data": "VIP North Oracle Data.xlsx",
    "VIP South Oracle Data": "VIP South Oracle Data.xlsx",
    "Tier 2 North Oracle Data": "Tier 2 North Oracle Data.xlsx",
    "Tier 2 South Oracle Data": "Tier 2 South Oracle Data.xlsx",
    "Call Log Data": "Call Log Data.xlsx",
    "Productivity Report": "Productivity Report.xlsx",
}



#------------------NEW BLOCK 17---------------------#

import streamlit as st
import pandas as pd
import plotly.express as px

@st.cache_data
def load_data():
    df_vip_south = pd.read_excel("VIP South Oracle Data.xlsx")
    df_vip_south["Team"] = "VIP South"
    df_vip_north = pd.read_excel("VIP North Oracle Data.xlsx")
    df_vip_north["Team"] = "VIP North"
    df_t2_south = pd.read_excel("Tier 2 South Oracle Data.xlsx")
    df_t2_south["Team"] = "Tier 2 South"
    df_t2_north = pd.read_excel("Tier 2 North Oracle Data.xlsx")
    df_t2_north["Team"] = "Tier 2 North"
    df_all = pd.concat([df_vip_south, df_vip_north, df_t2_south, df_t2_north], ignore_index=True)
    return df_all, df_vip_south, df_vip_north, df_t2_south, df_t2_north

df_all, df_vip_south, df_vip_north, df_t2_south, df_t2_north = load_data()

def fmt_td(x):
    if pd.isnull(x): return ""
    if isinstance(x, pd.Timedelta):
        s = str(x)
        if "days" in s:
            s = s.split(" ")[-1]
        return s.split(".")[0]
    return str(x)

def fix_time_col(series):
    import datetime
    return series.apply(
        lambda x: x.strftime("%H:%M:%S") if isinstance(x, datetime.time) else x
    )

if st.session_state.get("screen") == "operational_area":
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_operational_area"):
        st.session_state.screen = "area_selection"
        st.rerun()
    # -- Horizontal section buttons --
    col1, col2, col3, col4 = st.columns(4)
    if col1.button("Engineer"):
        st.session_state["op_area_section"] = "engineer"
        st.rerun()
    if col2.button("Time analysis"):
        st.session_state["op_area_section"] = "time"
        st.rerun()
    if col3.button("Visits"):
        st.session_state["op_area_section"] = "visits"
        st.rerun()
    if col4.button("Activity Status"):
        st.session_state["op_area_section"] = "activity_status"
        st.rerun()

    section = st.session_state.get("op_area_section", "engineer")

    # ---- ENGINEER SECTION ----
    from st_aggrid import AgGrid, GridOptionsBuilder

    def show_aggrid(df, height=300):
        gb = GridOptionsBuilder.from_dataframe(df)
        gb.configure_pagination(paginationAutoPageSize=False, paginationPageSize=10)
        gb.configure_default_column(editable=False, filter=True, sortable=True)
        grid_options = gb.build()
        AgGrid(df, gridOptions=grid_options, height=height, fit_columns_on_grid_load=True)
        
    if section == "engineer":
        st.title("Engineer Dashboard (All Oracle Data)")
        st.subheader("All Engineers & Visit Counts")
        engineer_counts = df_all["Name"].value_counts().reset_index()
        engineer_counts.columns = ["Engineer", "Visit Count"]
        st.dataframe(engineer_counts, use_container_width=True)

        with st.expander("📊 Bar Chart: Visits per Engineer"):
            fig = px.bar(engineer_counts, x="Engineer", y="Visit Count", title="Visits per Engineer")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🥧 Pie Chart: Visit Share (Top 10 Engineers)"):
            top10 = engineer_counts.head(10)
            fig = px.pie(top10, names="Engineer", values="Visit Count", title="Top 10 Engineers by Visit Share")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🌞 Sunburst: Engineer → Visit Type"):
            if "Visit Type" in df_all.columns and "Name" in df_all.columns:
                sunburst_df = df_all[["Name", "Visit Type"]].dropna()
                fig = px.sunburst(sunburst_df, path=["Name", "Visit Type"], title="Engineer → Visit Type Breakdown")
                st.plotly_chart(fig, use_container_width=True)

        st.subheader("👷‍♂️ Engineer Visit Summary Stats")
        stats = {
            "Average Visits": int(engineer_counts["Visit Count"].mean()),
            "Min Visits": int(engineer_counts["Visit Count"].min()),
            "Max Visits": int(engineer_counts["Visit Count"].max()),
            "Total Engineers": len(engineer_counts)
        }
        st.table(pd.DataFrame(stats, index=["Value"]).T)

        with st.expander("🏢 Visits per Engineer by Team (Stacked Bar)"):
            if "Team" in df_all.columns and "Name" in df_all.columns:
                grouped = df_all.groupby(["Name", "Team"]).size().unstack(fill_value=0)
                fig = px.bar(
                    grouped,
                    barmode="stack",
                    title="Visits per Engineer by Team"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🔎 Top 10 Engineers with Most Visit Type Diversity"):
            if "Name" in df_all.columns and "Visit Type" in df_all.columns:
                n_types = df_all.groupby("Name")["Visit Type"].nunique().sort_values(ascending=False).head(10)
                st.bar_chart(n_types)
                st.dataframe(n_types)

        with st.expander("🍰 Engineer Visit Status Breakdown (Pie)"):
            if "Name" in df_all.columns and "Activity Status" in df_all.columns:
                engineer_status = df_all.groupby("Name")["Activity Status"].value_counts().unstack(fill_value=0)
                top_eng = engineer_status.sum(axis=1).sort_values(ascending=False).head(5).index
                pie_data = df_all[df_all["Name"].isin(top_eng)]
                fig = px.pie(
                    pie_data, names="Activity Status", title="Status Breakdown (Top 5 Engineers)",
                    color="Activity Status"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📈 Engineer Visits Over Time"):
            if "Date" in df_all.columns and "Name" in df_all.columns:
                df_all["Date"] = pd.to_datetime(df_all["Date"], errors="coerce")
                by_month = df_all.groupby([pd.Grouper(key="Date", freq="M"), "Name"]).size().unstack(fill_value=0)
                top_engs = engineer_counts["Engineer"].head(5)
                fig = px.line(
                    by_month[top_engs], 
                    title="Monthly Visit Count - Top 5 Engineers",
                    labels={"value": "Visits", "Date": "Month"}
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🏷️ Visits per Engineer by Visit Type (Stacked Bar)"):
            if "Visit Type" in df_all.columns and "Name" in df_all.columns:
                grouped = df_all.groupby(["Name", "Visit Type"]).size().unstack(fill_value=0)
                top_engs = engineer_counts["Engineer"].head(5)
                fig = px.bar(
                    grouped.loc[top_engs],
                    barmode="stack",
                    title="Visits per Engineer by Visit Type (Top 5 Engineers)"
                )
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("⏳ Average & Total Time per Engineer"):
            if "Name" in df_all.columns and "Total Time" in df_all.columns:
                df_time = df_all.copy()
                df_time["Total Time"] = pd.to_timedelta(fix_time_col(df_time["Total Time"]), errors="coerce")
                df_time = df_time[~df_time["Total Time"].isna()]
                time_stats = df_time.groupby("Name")["Total Time"].agg(["count", "mean", "sum"])
                time_stats = time_stats.sort_values("count", ascending=False).head(10)
                time_stats["mean"] = time_stats["mean"].apply(fmt_td)
                time_stats["sum"] = time_stats["sum"].apply(fmt_td)
                st.dataframe(time_stats)
                st.bar_chart(time_stats["count"])

        with st.expander("📋 Engineer vs Visit Type Matrix"):
            if "Name" in df_all.columns and "Visit Type" in df_all.columns:
                matrix = pd.crosstab(df_all["Name"], df_all["Visit Type"])
                st.dataframe(matrix)

        with st.expander("👥 Engineer Visits Per Team (Table)"):
            if "Name" in df_all.columns and "Team" in df_all.columns:
                team_table = pd.crosstab(df_all["Name"], df_all["Team"])
                st.dataframe(team_table)




    # ---- VISITS SECTION ----
    elif section == "visits":
        st.title("Visit Dashboard (All Oracle Data)")

        st.subheader("All Visits (Table View)")
        visit_cols = [c for c in ["Visit Type", "Activity Status", "Date", "Team", "Name"] if c in df_all.columns]
        visit_table = df_all[visit_cols].copy()
        st.dataframe(visit_table.head(200), use_container_width=True)

        with st.expander("📊 Bar Chart: Visit Count by Type"):
            vc_type = df_all["Visit Type"].value_counts().reset_index()
            vc_type.columns = ["Visit Type", "Count"]
            fig = px.bar(vc_type, x="Visit Type", y="Count", title="Visit Count by Type")
            st.plotly_chart(fig, use_container_width=True)

        with st.expander("🟢 Visit Status Breakdown (Pie Chart)"):
            if "Activity Status" in df_all.columns:
                fig = px.pie(df_all, names="Activity Status", title="Visit Status Breakdown")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("👥 Visits by Team (Bar Chart)"):
            if "Team" in df_all.columns:
                vc_team = df_all["Team"].value_counts().reset_index()
                vc_team.columns = ["Team", "Visit Count"]
                fig = px.bar(vc_team, x="Team", y="Visit Count", title="Visits by Team")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📈 Visits Over Time (Monthly Line Chart)"):
            if "Date" in df_all.columns:
                df_all["Date"] = pd.to_datetime(df_all["Date"], errors="coerce")
                df_all["Month"] = df_all["Date"].dt.to_period("M").dt.to_timestamp()
                by_month = df_all.groupby("Month").size().reset_index(name="Visit Count")
                fig = px.line(by_month, x="Month", y="Visit Count", markers=True, title="Monthly Visit Trend")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🌞 Sunburst: Team → Visit Type"):
            if "Team" in df_all.columns and "Visit Type" in df_all.columns:
                sunburst_df = df_all[["Team", "Visit Type"]].dropna()
                fig = px.sunburst(sunburst_df, path=["Team", "Visit Type"], title="Visits Breakdown: Team → Visit Type")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("🔗 Parallel Categories: Team, Status, Visit Type"):
            needed_cols = ["Team", "Activity Status", "Visit Type"]
            if all(c in df_all.columns for c in needed_cols):
                pc_df = df_all[needed_cols].dropna().astype(str)
                fig = px.parallel_categories(pc_df, dimensions=needed_cols, title="Team → Status → Visit Type")
                st.plotly_chart(fig, use_container_width=True)

        with st.expander("📅 Visits by Day of Week"):
            if "Date" in df_all.columns:
                df_all["Day of Week"] = df_all["Date"].dt.day_name()
                vc_day = df_all["Day of Week"].value_counts().reindex(
                    ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"], fill_value=0
                )
                st.bar_chart(vc_day)

        with st.expander("🏠 Visits by Postcode (Top 10)"):
            if "Postcode" in df_all.columns:
                vc_postcode = df_all["Postcode"].value_counts().head(10)
                st.bar_chart(vc_postcode)

        with st.expander("👷 Visits per Engineer (Top 10)"):
            if "Name" in df_all.columns:
                vc_engineer = df_all["Name"].value_counts().head(10)
                st.bar_chart(vc_engineer)

        with st.expander("🔥 Visits by Month & Team (Heatmap)"):
            if "Date" in df_all.columns and "Team" in df_all.columns:
                df_all["Month"] = df_all["Date"].dt.to_period("M").dt.strftime('%b %Y')
                pivot = pd.pivot_table(df_all, index="Month", columns="Team", values="Visit Type", aggfunc="count", fill_value=0)
                st.dataframe(pivot)

        with st.expander("📋 Visits by Status & Type"):
            if "Activity Status" in df_all.columns and "Visit Type" in df_all.columns:
                pivot = pd.pivot_table(df_all, index="Activity Status", columns="Visit Type", values="Date", aggfunc="count", fill_value=0)
                st.dataframe(pivot)

    # (You can add a 'time analysis' section next, just like above.)
    elif section == "time":
        import datetime

        def fmt_td(x):
            if pd.isnull(x):
                return ""
            total_seconds = int(x.total_seconds())
            hours = total_seconds // 3600
            minutes = (total_seconds % 3600) // 60
            seconds = total_seconds % 60
            return f"{hours:02}:{minutes:02}:{seconds:02}"

        def clean_times(df, time_cols):
            df = df.copy()
            for col in time_cols:
                if col in df.columns:
                    # Convert any datetime.time to string
                    df[col] = df[col].apply(
                        lambda x: x.strftime("%H:%M:%S") if isinstance(x, datetime.time) else x
                    )
                    # Convert to timedelta
                    df[col] = pd.to_timedelta(df[col], errors="coerce")
                    # Remove zeros, blanks, NaT
                    df = df[~df[col].isna()]
                    df = df[df[col] != pd.Timedelta(0)]
            return df

        # All possible time columns
        time_cols = ["Total Time", "Activate", "Deactivate", "Travel Time", "Total Time (Inc Travel)", "Total Working Time"]


        # 1. Summary Table: Mean/Min/Max by Time Column
        st.subheader("⏳ Summary Table: Time Columns Stats")
        summary_rows = []
        for col in time_cols:
            if col in df_all.columns:
                cleaned = clean_times(df_all[[col]], [col])
                if not cleaned.empty:
                    summary_rows.append({
                        "Column": col,
                        "Mean": cleaned[col].mean(),
                        "Min": cleaned[col].min(),
                        "Max": cleaned[col].max(),
                        "Count": cleaned[col].count()
                    })
        if summary_rows:
            stats_df = pd.DataFrame(summary_rows)
            stats_df["Mean"] = stats_df["Mean"].apply(fmt_td)
            stats_df["Min"] = stats_df["Min"].apply(fmt_td)
            stats_df["Max"] = stats_df["Max"].apply(fmt_td)
            st.dataframe(stats_df)
        else:
            st.info("No time columns found.")

        # 2. Average Total Time per Visit Type
        if "Total Time" in df_all.columns and "Visit Type" in df_all.columns:
            with st.expander("⏱️ Average Total Time per Visit Type"):
                cleaned = clean_times(df_all[["Total Time", "Visit Type"]], ["Total Time"])
                avg_time = cleaned.groupby("Visit Type")["Total Time"].mean().sort_values()
                avg_time = avg_time.apply(fmt_td)
                st.bar_chart(avg_time)
                st.dataframe(avg_time)

        # 3. Activate & Deactivate Time Analysis
        for col in ["Activate", "Deactivate"]:
            if col in df_all.columns:
                with st.expander(f"⚡ {col} Time Analysis"):
                    cleaned = clean_times(df_all[[col, "Visit Type"]], [col])
                    if not cleaned.empty:
                        avg = cleaned.groupby("Visit Type")[col].mean().sort_values()
                        avg = avg.apply(fmt_td)
                        st.bar_chart(avg)
                        st.dataframe(avg)
                        st.markdown(f"**Min:** {fmt_td(cleaned[col].min())} &nbsp;&nbsp; **Max:** {fmt_td(cleaned[col].max())}")
                    else:
                        st.info(f"No data for {col} column.")


        # 4. Total Working Time Analysis
        if "Total Working Time" in df_all.columns:
            with st.expander("🛠️ Total Working Time Analysis"):
                cleaned = clean_times(df_all[["Total Working Time", "Visit Type"]], ["Total Working Time"])
                avg = cleaned.groupby("Visit Type")["Total Working Time"].mean().sort_values()
                avg = avg.apply(fmt_td)
                st.bar_chart(avg)
                st.dataframe(avg)
                st.markdown(f"**Min:** {fmt_td(cleaned['Total Working Time'].min())} &nbsp;&nbsp; **Max:** {fmt_td(cleaned['Total Working Time'].max())}")

        # 5. Boxplot of Total Time per Team
        if "Total Time" in df_all.columns and "Team" in df_all.columns:
            with st.expander("📦 Boxplot: Total Time by Team"):
                cleaned = clean_times(df_all[["Total Time", "Team"]], ["Total Time"])
                import plotly.express as px
                fig = px.box(cleaned, x="Team", y="Total Time", title="Total Time Distribution by Team")
                st.plotly_chart(fig, use_container_width=True)

        # 6. Timeline: Mean Total Time Per Month
        if "Total Time" in df_all.columns and "Date" in df_all.columns:
            with st.expander("📈 Avg Total Time Per Month (Line Chart)"):
                cleaned = clean_times(df_all[["Total Time", "Date"]], ["Total Time"])
                cleaned["Month"] = pd.to_datetime(cleaned["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
                monthly = cleaned.groupby("Month")["Total Time"].mean().dropna()
                monthly_fmt = monthly.apply(fmt_td)
                st.line_chart(monthly_fmt)


        # 7. Table: All Time Columns (first 100 rows, cleaned)
        with st.expander("📋 All Time Columns (Sample)"):
            cols_present = [col for col in time_cols if col in df_all.columns]
            sample = clean_times(df_all[cols_present], cols_present)
            # Format columns
            for c in cols_present:
                sample[c] = sample[c].apply(fmt_td)
            st.dataframe(sample.head(100), use_container_width=True)

        # 8. 🔥 Heatmap: Average Total Time by Team & Visit Type
        if all(col in df_all.columns for col in ["Total Time", "Team", "Visit Type"]):
            with st.expander("🔥 Heatmap: Avg Total Time by Team & Visit Type"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Visit Type"]], ["Total Time"])
                pivot = cleaned.pivot_table(index="Team", columns="Visit Type", values="Total Time", aggfunc="mean")
                # Format for display (HH:MM:SS)
                pivot_fmt = pivot.applymap(fmt_td)
                st.dataframe(pivot_fmt)

        # 9. ⏲️ Distribution of Total Time (Histogram)
        if "Total Time" in df_all.columns:
            with st.expander("⏲️ Distribution of Total Time (Histogram)"):
                cleaned = clean_times(df_all[["Total Time"]], ["Total Time"])
                if not cleaned.empty:
                    import matplotlib.pyplot as plt
                    import numpy as np
                    times = cleaned["Total Time"].dt.total_seconds() / 60  # Minutes
                    fig, ax = plt.subplots()
                    ax.hist(times, bins=30, color='skyblue', edgecolor='black')
                    ax.set_title("Distribution of Total Time (Minutes)")
                    ax.set_xlabel("Minutes")
                    ax.set_ylabel("Frequency")
                    st.pyplot(fig)



        # 10. 🕒 Median Time by Visit Type and Team
        if all(col in df_all.columns for col in ["Total Time", "Team", "Visit Type"]):
            with st.expander("🕒 Median Total Time by Visit Type and Team (Table)"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Visit Type"]], ["Total Time"])
                pivot = cleaned.pivot_table(index="Team", columns="Visit Type", values="Total Time", aggfunc="median")
                pivot_fmt = pivot.applymap(fmt_td)
                st.dataframe(pivot_fmt)

        # 11. 📊 Pie Chart: Proportion of Visits with >1 Hour Total Time
        if "Total Time" in df_all.columns:
            with st.expander("📊 Visits > 1 Hour vs <= 1 Hour (Pie Chart)"):
                cleaned = clean_times(df_all[["Total Time"]], ["Total Time"])
                gt1h = (cleaned["Total Time"] > pd.Timedelta(hours=1)).sum()
                le1h = (cleaned["Total Time"] <= pd.Timedelta(hours=1)).sum()
                pie_df = pd.DataFrame({
                    "Category": ["> 1 Hour", "<= 1 Hour"],
                    "Count": [gt1h, le1h]
                })
                fig = px.pie(pie_df, names="Category", values="Count", title="Proportion of Visits > 1 Hour Total Time")
                st.plotly_chart(fig, use_container_width=True)

        # 12. 🏆 Longest & Shortest Total Times (Per Team)
        if "Total Time" in df_all.columns and "Team" in df_all.columns:
            with st.expander("🏆 Longest & Shortest Total Times per Team"):
                cleaned = clean_times(df_all[["Total Time", "Team", "Name", "Visit Type"]], ["Total Time"])
                idxmax = cleaned.groupby("Team")["Total Time"].idxmax()
                idxmin = cleaned.groupby("Team")["Total Time"].idxmin()
                longest = cleaned.loc[idxmax]
                shortest = cleaned.loc[idxmin]
                longest = longest[["Team", "Name", "Visit Type", "Total Time"]]
                shortest = shortest[["Team", "Name", "Visit Type", "Total Time"]]
                longest["Total Time"] = longest["Total Time"].apply(fmt_td)
                shortest["Total Time"] = shortest["Total Time"].apply(fmt_td)
                st.markdown("#### Longest Total Time per Team")
                st.dataframe(longest)
                st.markdown("#### Shortest Total Time per Team")
                st.dataframe(shortest)

        # 13. ⏱️ Total Working Time > 10:25 (Detailed Summary)
    if section == "time":    
        if "Total Working Time" in df_all.columns:
            with st.expander("⏱️ Total Working Time Over 10:25 Summary", expanded=False):

                def convert_mixed_time(val):
                    import datetime
                    if pd.isnull(val): return 0
                    if isinstance(val, datetime.timedelta): return val.total_seconds() / 60
                    if isinstance(val, datetime.time): return val.hour * 60 + val.minute + val.second / 60
                    try: return pd.to_timedelta(val).total_seconds() / 60
                    except:
                        try:
                            if isinstance(val, str) and ":" in val:
                                return pd.to_timedelta("0 days " + val).total_seconds() / 60
                            if isinstance(val, (int, float)):
                                return float(val) * 24 * 60
                        except: return 0
                    return 0

                def mins_to_hhmm(m):
                    h = int(m // 60)
                    mins = int(m % 60)
                    return f"{h}:{mins:02}"

                valid_df = df_all[df_all["Total Working Time"].notna()].copy()
                valid_df["Total Working Time (min)"] = valid_df["Total Working Time"].apply(convert_mixed_time)
                valid_df["Over Minutes"] = valid_df["Total Working Time (min)"] - 625
                valid_df["Over Minutes"] = valid_df["Over Minutes"].apply(lambda x: x if x > 0 else 0)

                total_minutes = valid_df["Total Working Time (min)"].sum()
                total_over_minutes = valid_df["Over Minutes"].sum()
                total_over_cost = (valid_df["Over Minutes"] / 15).apply(lambda x: round(x)).sum() * 5.50

                overall_summary = {
                    "Total Working Time": mins_to_hhmm(total_minutes),
                    "Total Time Over 10:25": mins_to_hhmm(total_over_minutes),
                    "Total Over Time Cost (£)": f"£{total_over_cost:,.2f}"
                }

                st.markdown("#### 🔢 Overall Summary")
                st.dataframe(pd.DataFrame([overall_summary]), use_container_width=True)

                st.markdown("#### 🧑‍🤝‍🧑 Breakdown by Team")
                team_summary = valid_df.groupby("Team").agg(
                    Total_Working_Minutes=("Total Working Time (min)", "sum"),
                    Over_Minutes=("Over Minutes", "sum"),
                    Over_Cost=("Over Minutes", lambda x: (x / 15).round().sum() * 5.50)
                ).reset_index()
                team_summary["Total Working Time"] = team_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
                team_summary["Time Over 10:25"] = team_summary["Over_Minutes"].apply(mins_to_hhmm)
                team_summary["Over Cost (£)"] = team_summary["Over_Cost"].map("£{:,.2f}".format)
                st.dataframe(team_summary[["Team", "Total Working Time", "Time Over 10:25", "Over Cost (£)"]], use_container_width=True)

                st.markdown("#### 👷 Breakdown by Engineer")
                engineer_summary = valid_df.groupby("Name").agg(
                    Total_Working_Minutes=("Total Working Time (min)", "sum"),
                    Over_Minutes=("Over Minutes", "sum"),
                    Over_Cost=("Over Minutes", lambda x: (x / 15).round().sum() * 5.50)
                ).reset_index()
                engineer_summary["Total Working Time"] = engineer_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
                engineer_summary["Time Over 10:25"] = engineer_summary["Over_Minutes"].apply(mins_to_hhmm)
                engineer_summary["Over Cost (£)"] = engineer_summary["Over_Cost"].map("£{:,.2f}".format)
                st.dataframe(engineer_summary[["Name", "Total Working Time", "Time Over 10:25", "Over Cost (£)"]], use_container_width=True)

       

        # 15. 🗓️ Time Over 10:25: Monthly & Quarterly Breakdown
        with st.expander("📅 Time Over 10:25: Monthly & Quarterly Breakdown"):

            df_time = valid_df.copy()
            df_time["Month"] = pd.to_datetime(df_time["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
            df_time["Quarter"] = pd.to_datetime(df_time["Date"], errors="coerce").dt.to_period("Q").astype(str)

            monthly = df_time.groupby("Month").agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            monthly["% Time Over 10:25"] = (monthly["Over_Minutes"] / monthly["Total_Minutes"] * 100).round(2)
            monthly["Month"] = pd.to_datetime(monthly["Month"]).dt.strftime("%B %Y")
            monthly["Total Working Time (hh:mm)"] = monthly["Total_Minutes"].apply(mins_to_hhmm)
            monthly["Time Over 10:25 (hh:mm)"] = monthly["Over_Minutes"].apply(mins_to_hhmm)
            st.markdown("## 🗓️ Monthly Breakdown")
            st.dataframe(monthly[["Month", "Total Working Time (hh:mm)", "Time Over 10:25 (hh:mm)", "% Time Over 10:25"]], use_container_width=True)

            quarterly = df_time.groupby("Quarter").agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            quarterly["% Time Over 10:25"] = (quarterly["Over_Minutes"] / quarterly["Total_Minutes"] * 100).round(2)
            quarterly["Total Working Time (hh:mm)"] = quarterly["Total_Minutes"].apply(mins_to_hhmm)
            quarterly["Time Over 10:25 (hh:mm)"] = quarterly["Over_Minutes"].apply(mins_to_hhmm)
            st.markdown("## 🗓️ Quarterly Breakdown")
            st.dataframe(quarterly[["Quarter", "Total Working Time (hh:mm)", "Time Over 10:25 (hh:mm)", "% Time Over 10:25"]], use_container_width=True)

            st.markdown("### 📈 % Time Over 10:25 by Team")
            team_monthly = df_time.groupby(["Team", "Month"]).agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            team_monthly["% Time Over 10:25"] = (team_monthly["Over_Minutes"] / team_monthly["Total_Minutes"] * 100).round(2)
            team_monthly["Month"] = pd.to_datetime(team_monthly["Month"]).dt.strftime("%b %Y")
            fig_team = px.line(team_monthly, x="Month", y="% Time Over 10:25", color="Team", markers=True)
            fig_team.update_layout(xaxis_tickangle=45)
            st.plotly_chart(fig_team, use_container_width=True)

            st.markdown("### 📊 % Time Over 10:25 by Engineer")
            eng_monthly = df_time.groupby(["Name", "Month"]).agg(
                Total_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            eng_monthly["% Time Over 10:25"] = (eng_monthly["Over_Minutes"] / eng_monthly["Total_Minutes"] * 100).round(2)
            eng_monthly["Month"] = pd.to_datetime(eng_monthly["Month"]).dt.strftime("%b %Y")
            fig_eng = px.bar(eng_monthly, x="Month", y="% Time Over 10:25", color="Name", barmode="group")
            fig_eng.update_layout(xaxis_tickangle=45)
            st.plotly_chart(fig_eng, use_container_width=True)


        # 16. 📊 Deep Dive: % Time Over 10:25 by Month and Week
        with st.expander("📊 Time Over 10:25: Monthly & Weekly Insight", expanded=False):

            def mins_to_hhmm(m):
                h = int(m // 60)
                mins = int(m % 60)
                return f"{h}:{mins:02}"

            df_over = df_all[df_all["Total Working Time"].notna()].copy()
            df_over["Total Working Time (min)"] = df_over["Total Working Time"].apply(convert_mixed_time)
            df_over["Over Minutes"] = df_over["Total Working Time (min)"] - 625
            df_over["Over Minutes"] = df_over["Over Minutes"].apply(lambda x: x if x > 0 else 0)

            df_over["Month"] = pd.to_datetime(df_over["Date"], errors="coerce").dt.to_period("M").dt.to_timestamp()
            df_over["week"] = df_over["week"].astype(str)

            # Monthly Summary
            monthly_summary = df_over.groupby("Month").agg(
                Total_Working_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            monthly_summary["% Time Over 10:25"] = (monthly_summary["Over_Minutes"] / monthly_summary["Total_Working_Minutes"] * 100).round(2)
            monthly_summary["Total Working Time"] = monthly_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
            monthly_summary["Time Over 10:25"] = monthly_summary["Over_Minutes"].apply(mins_to_hhmm)

            st.markdown("### 📅 Monthly Breakdown")
            st.dataframe(monthly_summary[["Month", "Total Working Time", "Time Over 10:25", "% Time Over 10:25"]], use_container_width=True)

            # Max/Min Monthly Delta
            max_row = monthly_summary.loc[monthly_summary["% Time Over 10:25"].idxmax()]
            min_row = monthly_summary.loc[monthly_summary["% Time Over 10:25"].idxmin()]
            delta = round(max_row["% Time Over 10:25"] - min_row["% Time Over 10:25"], 2)
            delta_table = pd.DataFrame([
                {"Metric": "Max Month", "Month": max_row["Month"].strftime("%B %Y"), "% Over": max_row["% Time Over 10:25"]},
                {"Metric": "Min Month", "Month": min_row["Month"].strftime("%B %Y"), "% Over": min_row["% Time Over 10:25"]},
                {"Metric": "Delta", "Month": f"{max_row['Month'].strftime('%B')} vs {min_row['Month'].strftime('%B')}", "% Over": delta}
            ])

            st.markdown("### 🔺 Max vs Min Monthly Comparison")
            st.dataframe(delta_table, use_container_width=True)

            # Weekly Summary
            weekly_summary = df_over.groupby("week").agg(
                Total_Working_Minutes=("Total Working Time (min)", "sum"),
                Over_Minutes=("Over Minutes", "sum")
            ).reset_index()
            weekly_summary["% Time Over 10:25"] = (weekly_summary["Over_Minutes"] / weekly_summary["Total_Working_Minutes"] * 100).round(2)
            weekly_summary["Total Working Time"] = weekly_summary["Total_Working_Minutes"].apply(mins_to_hhmm)
            weekly_summary["Time Over 10:25"] = weekly_summary["Over_Minutes"].apply(mins_to_hhmm)

            st.markdown("### 📆 Weekly Breakdown")
            st.dataframe(weekly_summary[["week", "Total Working Time", "Time Over 10:25", "% Time Over 10:25"]], use_container_width=True)

        with st.expander("📊 Combined Charts: Time Over 10:25 (Weekly & Monthly)", expanded=False):
            from plotly.subplots import make_subplots
            import plotly.graph_objects as go

            # Convert hh:mm to minutes for line chart
            def hhmm_to_minutes(hhmm_str):
                if isinstance(hhmm_str, str) and ":" in hhmm_str:
                    h, m = hhmm_str.split(":")
                    return int(h) * 60 + int(m)
                return 0

            # Sort week as int for correct order
            weekly_summary["Week_Num"] = weekly_summary["week"].astype(int)
            weekly_summary = weekly_summary.sort_values("Week_Num")
            weekly_summary["Time Over Minutes"] = weekly_summary["Time Over 10:25"].apply(hhmm_to_minutes)

            fig_week = make_subplots(specs=[[{"secondary_y": True}]])

            fig_week.add_trace(
                go.Bar(x=weekly_summary["Week_Num"], y=weekly_summary["% Time Over 10:25"], name="% Time Over 10:25"),
                secondary_y=False
            )

            fig_week.add_trace(
                go.Scatter(x=weekly_summary["Week_Num"], y=weekly_summary["Time Over Minutes"], name="Time Over 10:25 (min)", mode="lines+markers"),
                secondary_y=True
            )

            fig_week.update_layout(
                title="📊 Weekly Time Over 10:25",
                xaxis_title="Week",
                yaxis_title="% Time Over",
                legend_title="Metrics",
                bargap=0.3,
                xaxis=dict(type="category")
            )

            fig_week.update_yaxes(title_text="% Time Over 10:25", secondary_y=False)
            fig_week.update_yaxes(title_text="Time Over (Minutes)", secondary_y=True)

            st.plotly_chart(fig_week, use_container_width=True)

            # Monthly version
            monthly_summary["Month_Str"] = monthly_summary["Month"].dt.strftime("%b %Y")
            monthly_summary = monthly_summary.sort_values("Month")
            monthly_summary["Time Over Minutes"] = monthly_summary["Time Over 10:25"].apply(hhmm_to_minutes)

            fig_month = make_subplots(specs=[[{"secondary_y": True}]])

            fig_month.add_trace(
                go.Bar(x=monthly_summary["Month_Str"], y=monthly_summary["% Time Over 10:25"], name="% Time Over 10:25"),
                secondary_y=False
            )

            fig_month.add_trace(
                go.Scatter(x=monthly_summary["Month_Str"], y=monthly_summary["Time Over Minutes"], name="Time Over 10:25 (min)", mode="lines+markers"),
                secondary_y=True
            )

            fig_month.update_layout(
                title="📅 Monthly Time Over 10:25",
                xaxis_title="Month",
                yaxis_title="% Time Over",
                legend_title="Metrics",
                bargap=0.3,
                xaxis_tickangle=45
            )

            fig_month.update_yaxes(title_text="% Time Over 10:25", secondary_y=False)
            fig_month.update_yaxes(title_text="Time Over (Minutes)", secondary_y=True)

            st.plotly_chart(fig_month, use_container_width=True)


        # 17. 🧠 Last Visit Type When Total Over 10:25
        with st.expander("🧠 Last Visit Type When Total Working Time is Over 10:25", expanded=False):

                st.markdown("This shows the final visit types on days where total working time exceeds 10:25 (625 mins).")

                def to_minutes(t):
                        try:
                                if isinstance(t, str) and ":" in t:
                                        h, m, s = map(int, t.split(":"))
                                        return h * 60 + m + s / 60
                                elif hasattr(t, 'hour'):
                                        return t.hour * 60 + t.minute + t.second / 60
                                return float(t) * 24 * 60  # Excel float
                        except:
                                return 0

                df_check = df_all.copy()
                df_check["Date"] = pd.to_datetime(df_check["Date"], errors="coerce")
                df_check = df_check[df_check["Activity Status"].str.lower() == "completed"]
                df_check = df_check[df_check["Total Working Time"].notna()]
                df_check = df_check[~df_check["Total Working Time"].isin(["00:00", "0:00", "0", 0])]

                df_check["Total Working Time (min)"] = df_check["Total Working Time"].apply(to_minutes)

                # Step 1: Total time per engineer per day
                total_day = df_check.groupby(["Name", "Date"])["Total Working Time (min)"].sum().reset_index()
                total_day = total_day[total_day["Total Working Time (min)"] > 625]
                total_day = total_day.rename(columns={"Total Working Time (min)": "Total Day Working Time (min)"})

                # Step 2: Merge valid days
                df_valid = pd.merge(df_check, total_day, on=["Name", "Date"], how="inner")

                # Step 3: Create minutes from End time for sorting
                df_valid["End_minutes"] = df_valid["End"].apply(to_minutes)

                # Step 4: Sort by End_minutes to get the last visit
                df_valid.sort_values(by=["Name", "Date", "End_minutes"], inplace=True)
                last_visits = df_valid.groupby(["Name", "Date"]).tail(1)

                display_cols = ["Name", "Date", "Visit Type", "Start", "End", "Total Working Time", "Total Day Working Time (min)"]

                if not last_visits.empty and all(col in last_visits.columns for col in display_cols):
                        st.dataframe(last_visits[display_cols], use_container_width=True)
                else:
                        st.warning("⚠️ Required columns for displaying the table are missing or no data matched the criteria.")
        # 18. 📦 Breakdown of Last Visit Types by Month and Week (Total > 10:25)
        with st.expander("📦 Breakdown of Last Visit Types by Month and Week (Total Working Time > 10:25)", expanded=False):

                st.markdown("This shows a breakdown of the last visit types by **month** and **week number** for days where total working time exceeded 10:25 (625 mins).")

                if not last_visits.empty:

                        last_visits["Date"] = pd.to_datetime(last_visits["Date"], errors="coerce")
                        last_visits["Month"] = last_visits["Date"].dt.strftime("%B")
                        last_visits["Week"] = last_visits["Date"].dt.isocalendar().week

                        # Define proper calendar order
                        month_order = ["January", "February", "March", "April", "May", "June", 
                                       "July", "August", "September", "October", "November", "December"]

                        # Monthly breakdown
                        monthly = last_visits.groupby(["Month", "Visit Type"]).size().reset_index(name="Count")
                        monthly["Month"] = pd.Categorical(monthly["Month"], categories=month_order, ordered=True)
                        monthly = monthly.sort_values(by=["Month", "Count"], ascending=[True, False])
                        months = monthly["Month"].dropna().unique().tolist()

                        st.markdown("### 📅 Monthly Breakdown (Select a Month Tab)")

                        month_tabs = st.tabs(months)
                        for i, month in enumerate(months):
                                with month_tabs[i]:
                                        st.dataframe(monthly[monthly["Month"] == month].reset_index(drop=True), use_container_width=True)

                        # Weekly breakdown
                        st.markdown("### 📆 Weekly Breakdown")
                        weekly = last_visits.groupby(["Week", "Visit Type"]).size().reset_index(name="Count")
                        weekly = weekly.sort_values(by=["Week", "Count"], ascending=[True, False])
                        st.dataframe(weekly, use_container_width=True)

                else:
                        st.warning("No data found for last visits over 10:25. Please check upstream filters.")

    # --- NUMBER 7 ----------------------------------------------------------
    # --- SECTION: ACTIVITY STATUS BREAKDOWN --------------------------------
    from prophet import Prophet
    import plotly.express as px
    import plotly.graph_objects as go

    if section == "activity_status":
        st.markdown("## 📊 Activity Status Breakdown")

        # Prep all 4 datasets as (name, df) pairs
        team_datasets = {
            "VIP North": df_vip_north,
            "VIP South": df_vip_south,
            "Tier 2 North": df_t2_north,
            "Tier 2 South": df_t2_south,
        }
        with st.expander("🌍 All Teams Summary", expanded=True):
            # Combine all team dataframes into one for All Teams analysis
            df = pd.concat([df_vip_north, df_vip_south, df_t2_north, df_t2_south], ignore_index=True)

            # Clean data: remove rows with missing or empty Activity Status or Date
            df = df[df["Activity Status"].notna()]
            df = df[df["Activity Status"].str.strip() != ""]
            df = df[df["Date"].notna()]

            # Normalize Activity Status (lowercase + strip)
            df["Activity Status"] = df["Activity Status"].str.lower().str.strip()

            # Convert Date to datetime and add Month period column for sorting
            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
            df = df[df["Date"].notna()]  # Remove rows with invalid dates if any
            df["Month"] = df["Date"].dt.to_period("M")
            df["Month_str"] = df["Month"].astype(str)

            # 1. Overall counts per Activity Status (for bar chart and KPIs)
            overall_counts = df["Activity Status"].value_counts().sort_index()
            overall_counts_df = overall_counts.reset_index()
            overall_counts_df.columns = ["Activity Status", "Count"]  # Fix column names

            st.markdown("### Overall Visit Counts by Activity Status")
            fig_counts = px.bar(
                overall_counts_df,
                x="Activity Status",
                y="Count",
                title="Overall Visit Counts by Activity Status"
            )
            st.plotly_chart(fig_counts, use_container_width=True)

            # 2. Overall completion rate = Completed / (Completed + Cancelled + Not Done)
            completed = overall_counts.get("completed", 0)
            cancelled = overall_counts.get("cancelled", 0)
            not_done = overall_counts.get("not done", 0)
            denom = completed + cancelled + not_done
            completion_rate = (completed / denom * 100) if denom > 0 else 0
            st.markdown(f"**Overall Completion Rate:** {completion_rate:.2f}% (Completed / (Completed + Cancelled + Not Done))")

            # 3. Monthly counts pivot table (Activity Status x Month)
            monthly_pivot = df.pivot_table(
                index="Month",
                columns="Activity Status",
                values="Date",  # Count of non-null Date as proxy for count
                aggfunc="count",
                fill_value=0
            ).sort_index()

            # 4. Monthly Pending % = (pending visits / total visits per month) * 100
            monthly_totals = monthly_pivot.sum(axis=1)
            monthly_pending = monthly_pivot.get("pending", pd.Series(0, index=monthly_pivot.index))
            monthly_pending_pct = (monthly_pending / monthly_totals * 100).fillna(0)

            # 5. Line chart with all Activity Status monthly counts
            monthly_pivot_str = monthly_pivot.copy()
            monthly_pivot_str.index = monthly_pivot_str.index.astype(str)
            fig_monthly = px.line(
                monthly_pivot_str,
                x=monthly_pivot_str.index,
                y=monthly_pivot_str.columns,
                markers=True,
                title="Monthly Activity Status Counts"
            )
            st.plotly_chart(fig_monthly, use_container_width=True)

            # 6. Monthly Pending % line chart
            fig_pending = px.line(
                x=monthly_pending_pct.index.astype(str),
                y=monthly_pending_pct.values,
                labels={"x": "Month", "y": "Pending %"},
                title="Monthly Pending Percentage"
            )
            st.plotly_chart(fig_pending, use_container_width=True)

            # 7. Display Monthly Pending % table sorted by month ascending
            pending_df = monthly_pending_pct.reset_index()
            pending_df.columns = ["Month", "Pending %"]
            pending_df = pending_df.sort_values("Month")
            st.dataframe(pending_df.style.format({"Pending %": "{:.2f}"}), use_container_width=True)


        with st.expander("🌍 All Teams Summary", expanded=True):  # <-- Collapsible expander added here

                # Combine all 4 datasets for the "All Teams" summary
                df_all_teams = pd.concat([df_vip_north, df_vip_south, df_t2_north, df_t2_south], ignore_index=True)

                # 1. Overall totals by Activity Status (all time)
                overall_counts = df_all_teams["Activity Status"].value_counts().sort_index()
                st.markdown("#### Overall Visit Counts by Activity Status")
                st.bar_chart(overall_counts)

                # 2. Monthly Aggregated Summary
                if "Month" in df_all_teams.columns:
                    monthly_counts = df_all_teams.groupby(["Month", "Activity Status"]).size().unstack(fill_value=0)
                    monthly_counts["Total Visits"] = monthly_counts.sum(axis=1)
                    monthly_pct_change = monthly_counts.pct_change().fillna(0) * 100
                    monthly_pct_change = monthly_pct_change.add_suffix(" % Change")
                    max_visits = monthly_counts.max()
                    max_visits.name = "Max Visits"
                    min_visits = monthly_counts.min()
                    min_visits.name = "Min Visits"

                    monthly_summary = monthly_counts.join(monthly_pct_change).join(max_visits).join(min_visits)
                    monthly_summary = monthly_summary.sort_index()

                    st.markdown("#### Monthly Summary KPIs by Activity Status")
                    st.dataframe(monthly_summary.style.format({
                        col: "{:.2f}" for col in monthly_summary.columns if "% Change" in col
                    }))

                    # Monthly totals bar chart for quick visual
                    st.bar_chart(monthly_counts["Total Visits"])

                else:
                    st.warning("Missing 'Month' column for All Teams summary.")

                # 3. Weekly Aggregated Summary
                if "week" in df_all_teams.columns:
                    weekly_counts = df_all_teams.groupby(["week", "Activity Status"]).size().unstack(fill_value=0)
                    weekly_counts["Total Visits"] = weekly_counts.sum(axis=1)
                    weekly_pct_change = weekly_counts.pct_change().fillna(0) * 100
                    weekly_pct_change = weekly_pct_change.add_suffix(" % Change")
                    max_visits = weekly_counts.max()
                    max_visits.name = "Max Visits"
                    min_visits = weekly_counts.min()
                    min_visits.name = "Min Visits"

                    weekly_summary = weekly_counts.join(weekly_pct_change).join(max_visits).join(min_visits)
                    weekly_summary = weekly_summary.sort_index()

                    st.markdown("#### Weekly Summary KPIs by Activity Status")
                    st.dataframe(weekly_summary.style.format({
                        col: "{:.2f}" for col in weekly_summary.columns if "% Change" in col
                    }))

                    # Weekly totals line chart for quick visual
                    st.line_chart(weekly_counts["Total Visits"])

                else:
                    st.warning("Missing 'week' column for All Teams summary.")

        
    # --- Then continue with your existing individual team breakdowns here ---

        for team_name, df in team_datasets.items():
            # Main expander for each team
            with st.expander(f"🔵 {team_name} — Activity Overview", expanded=False):

                # Check for Activity Status column
                if "Activity Status" not in df.columns:
                    st.warning(f"No 'Activity Status' in {team_name}")
                    continue

                # Use tabs inside the expander instead of nested expanders
                tabs = st.tabs([
                    "📅 Monthly Breakdown",
                    "📆 Weekly Breakdown",
                    "📈 Forecast Next 6 Months",
                    "🌞 Sunburst View",
                    "📅 Gantt Chart",
                    "🔎 Drilldown"
                ])

                # Monthly Breakdown Tab
                with tabs[0]:
                    if "Month" in df.columns:
                        # Count visits by month and activity status
                        monthly_counts = df.groupby(["Month", "Activity Status"]).size().unstack(fill_value=0)

                        st.bar_chart(monthly_counts)

                        # Total visits per month (sum across all statuses)
                        monthly_counts["Total Visits"] = monthly_counts.sum(axis=1)

                        # Month-on-Month % Change per status plus total
                        monthly_pct_change = monthly_counts.pct_change().fillna(0) * 100
                        monthly_pct_change = monthly_pct_change.add_suffix(" % Change")

                        # Max visits by status across all months
                        max_visits = monthly_counts.max()
                        max_visits.name = "Max Visits"

                        # Min visits by status across all months
                        min_visits = monthly_counts.min()
                        min_visits.name = "Min Visits"

                        # Combine all KPI info into one table
                        summary_df = monthly_counts.join(monthly_pct_change)
                        summary_df = summary_df.join(max_visits)
                        summary_df = summary_df.join(min_visits)

                        # Sort by Month if possible (you may want to convert to datetime or categorical first)
                        summary_df = summary_df.sort_index()

                        st.markdown("### Monthly Summary KPIs by Activity Status")
                        st.dataframe(summary_df.style.format({
                            col: "{:.2f}" for col in summary_df.columns if "% Change" in col
                        }))
                    else:
                        st.warning("Missing 'Month' column")

                # Weekly Breakdown Tab
                with tabs[1]:
                    if "week" in df.columns:
                        # Count visits by week and activity status
                        weekly_counts = df.groupby(["week", "Activity Status"]).size().unstack(fill_value=0)

                        st.line_chart(weekly_counts)

                        # Total visits per week (sum across all statuses)
                        weekly_counts["Total Visits"] = weekly_counts.sum(axis=1)

                        # Week-on-Week % Change per status plus total
                        weekly_pct_change = weekly_counts.pct_change().fillna(0) * 100
                        weekly_pct_change = weekly_pct_change.add_suffix(" % Change")

                        # Max visits by status across all weeks
                        max_visits = weekly_counts.max()
                        max_visits.name = "Max Visits"

                        # Min visits by status across all weeks
                        min_visits = weekly_counts.min()
                        min_visits.name = "Min Visits"

                        # Combine all KPI info into one table
                        summary_week_df = weekly_counts.join(weekly_pct_change)
                        summary_week_df = summary_week_df.join(max_visits)
                        summary_week_df = summary_week_df.join(min_visits)

                        summary_week_df = summary_week_df.sort_index()

                        st.markdown("### Weekly Summary KPIs by Activity Status")
                        st.dataframe(summary_week_df.style.format({
                            col: "{:.2f}" for col in summary_week_df.columns if "% Change" in col
                        }))
                    else:
                        st.warning("Missing 'week' column")



                # Forecast Tab
                with tabs[2]:
                    try:
                        forecast_data = df[df["Activity Status"].notna()]
                        forecast_df = forecast_data.groupby("Date").size().reset_index(name="y")
                        forecast_df.columns = ["ds", "y"]

                        m = Prophet()
                        m.fit(forecast_df)
                        future = m.make_future_dataframe(periods=180)
                        forecast = m.predict(future)

                        fig = px.line(forecast, x="ds", y="yhat", title="6-Month Forecast")
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.warning(f"Forecasting failed: {e}")

                # Sunburst View Tab
                with tabs[3]:
                    try:
                        df['Month'] = df['Month'].astype(str)
                        fig = px.sunburst(df, path=["Month", "Activity Status"], title="Status Distribution by Month")
                        st.plotly_chart(fig, use_container_width=True)
                    except Exception as e:
                        st.warning(f"Sunburst error: {e}")

                # Gantt Chart Tab
                with tabs[4]:
                    try:
                        if "Date" in df.columns and "Start" in df.columns and "End" in df.columns and "Activity Status" in df.columns:
                            df_gantt = df.copy()

                            # Combine Date with Start/End time into full datetime
                            df_gantt["Start"] = pd.to_datetime(df_gantt["Date"].astype(str) + " " + df_gantt["Start"].astype(str), errors="coerce")
                            df_gantt["End"] = pd.to_datetime(df_gantt["Date"].astype(str) + " " + df_gantt["End"].astype(str), errors="coerce")

                            # Remove rows with missing or invalid times or statuses
                            df_gantt.dropna(subset=["Start", "End", "Activity Status"], inplace=True)

                            fig = px.timeline(
                                df_gantt,
                                x_start="Start",
                                x_end="End",
                                y="Activity Status",
                                color="Activity Status",
                                title=f"Activity Gantt Timeline – {team_name}",
                                height=400
                            )
                            fig.update_yaxes(autorange="reversed")
                            st.plotly_chart(fig)
                    except Exception as e:
                        st.error(f"Gantt chart error: {e}")

                # Drilldown Tab
                with tabs[5]:
                    selected_status = st.selectbox(
                        f"Filter {team_name} by Activity Status",
                        options=sorted(df["Activity Status"].dropna().unique()),
                        key=f"drilldown_{team_name.replace(' ', '_')}"
                    )
                    drill = df[df["Activity Status"] == selected_status]
                    st.dataframe(drill, use_container_width=True)



    





# --- NUMBER 8 ---#
# --- SECTION: DASHBOARD AREA – Dataset Grid Selection ---
from streamlit_lottie import st_lottie
import requests

def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

if st.session_state.screen == "dashboard":

    # Back to Area Selection (Main Menu)
    if st.button("⬅️ Back to Main Menu", use_container_width=True):
        st.session_state.screen = "area_selection"
        st.rerun()

    st.markdown("## 📁 Select a Dataset to Explore", unsafe_allow_html=True)
    st.markdown("Choose one of the available datasets below to enter its dashboard:")

    # Show animation if no dataset selected yet
    if st.session_state.get("selected_dataset") is None:
        lottie_url = "https://assets9.lottiefiles.com/packages/lf20_jcikwtux.json"  # example loading animation
        lottie_json = load_lottieurl(lottie_url)
        if lottie_json:
            st_lottie(lottie_json, height=300, width=300)

    dataset_buttons = {
        "AI Test SB Visits": "📘 AI Test SB Visits",
        "Invoice Data AI": "🧾 Invoice Data AI",
        "Productivity Report": "📈 Productivity Report",
        "Call Log Data": "📞 Call Log Data",
        "VIP North Oracle Data": "🏅 VIP North",
        "VIP South Oracle Data": "🏅 VIP South",
        "Tier 2 North Oracle Data": "🏅 Tier 2 North",
        "Tier 2 South Oracle Data": "🏅 Tier 2 South"
    }

    cols = st.columns(4)
    i = 0
    for dataset_key, label in dataset_buttons.items():
        if i % 4 == 0 and i > 0:
            cols = st.columns(4)
        col = cols[i % 4]
        with col:
            if st.button(label, use_container_width=True, key=f"btn_{dataset_key}"):
                st.session_state.selected_dataset = dataset_key
                st.session_state.screen = "dashboard_view"
                st.rerun()
        i += 1



# --- NUMBER 9 ----------------------------------------------------------
# --- SECTION: SIDEBAR FILTERS & DATA LOADING ---------------------------

# ➊ Skip the entire sidebar when we’re on screens that don’t need it
#    (AI chat assistant *or* the new Operational KPI hub)
# ----------------------------------------------------------------------
if st.session_state.get("screen") in ("ai", "operational_area"):
    # Those screens handle their own layout; no sidebar, no dataset filters.
    pass

else:
    # ------------------------------------------------------------------
    # ORIGINAL SIDEBAR CODE BEGINS HERE (unchanged)
    # ------------------------------------------------------------------

    # ⛔ Only check for dataset if we're NOT on the suggestion screen
  # Only check for dataset if we're NOT on suggestions or Instructions screens
    if st.session_state.get("screen") not in ["suggestions", "instruction_guide"]:
        file_choice = st.session_state.get("selected_dataset")
        if file_choice is None:
            st.info("👈 Pick a dataset first, then filters will appear here.")
            st.stop()



 # 1️⃣ Load the file (only once thanks to @st.cache_data in load_file)
file_choice = st.session_state.get("selected_dataset")
if file_choice:
    file_path = file_map.get(file_choice)
    df = load_file(file_path)
    if df.empty:
        st.warning("❌ No data loaded or file is empty.")
        st.stop()

    filtered_data = df.copy()  # master copy we’ll keep refining

    # 2️⃣ SIDEBAR UI
    with st.sidebar:
        st.image("sky_vip_logo.png", width=180)
        st.markdown(
            """
            <div style='font-size:0.98em;line-height:1.4;margin-bottom:12px;color:#4094D0;'>
                <b>Visit Intelligence Dashboard</b><br>
                Filter the data below and explore insights on the main page.
            </div>
            """,
            unsafe_allow_html=True
        )

        # --- Date --- 
        if (
            st.session_state.selected_dataset not in ["Productivity Report"]
            and "Date" in filtered_data.columns
            and pd.api.types.is_datetime64_any_dtype(filtered_data["Date"])
        ):
            filtered_data["Date"] = pd.to_datetime(filtered_data["Date"], errors="coerce")
            date_opts = ["All"] + sorted(filtered_data["Date"].dt.date.dropna().unique())
            sel_date = st.selectbox("📅 Date", date_opts, index=0)
            if sel_date != "All":
                filtered_data = filtered_data[filtered_data["Date"].dt.date == sel_date]


        # --- Week ---
        if "Week" in filtered_data.columns:
            week_opts = ["All"] + sorted(filtered_data["Week"].dropna().unique())
            sel_week = st.selectbox("🗓️ ISO Week", week_opts, index=0)
            if sel_week != "All":
                filtered_data = filtered_data[filtered_data["Week"] == sel_week]

        # --- Month ---
        if "MonthName" in filtered_data.columns:
            month_opts = ["All"] + sorted(filtered_data["MonthName"].dropna().unique())
            sel_month = st.selectbox("📆 Month", month_opts, index=0)
            if sel_month != "All":
                filtered_data = filtered_data[filtered_data["MonthName"] == sel_month]

        # --- Activity Status ---
        if "Activity Status" in filtered_data.columns:
            act_opts = ["All"] + sorted(filtered_data["Activity Status"].dropna().unique())
            sel_act = st.selectbox("🎯 Activity Status", act_opts, index=0)
            if sel_act != "All":
                filtered_data = filtered_data[filtered_data["Activity Status"] == sel_act]

        # --- Visit Type ---
        if "Visit Type" in filtered_data.columns:
            vt_opts = ["All"] + sorted(filtered_data["Visit Type"].dropna().unique())
            sel_vt = st.selectbox("🛠️ Visit Type", vt_opts, index=0)
            if sel_vt != "All":
                filtered_data = filtered_data[filtered_data["Visit Type"] == sel_vt]

        # --- Free-text search ---
        search_term = st.text_input("🔍 Search all fields", placeholder="Type and hit Enter")
        if search_term:
            filtered_data = filtered_data[
                filtered_data.apply(lambda r: search_term.lower() in str(r).lower(), axis=1)
            ]

    # 3️⃣ Bail if nothing left after filters
    if filtered_data.empty:
        st.warning("No rows match the current filters.")
        st.stop()

    # 4️⃣ Stash for downstream use
    st.session_state.filtered_data = filtered_data

import base64

# --- NUMBER 9 ---#
# --- SECTION: DASHBOARD VIEW – TITLE & ADVANCED SUMMARY ---
if st.session_state.get("screen") == "dashboard_view":

    # Grab the filtered dataframe prepared in Block 8
    selected_dataset = st.session_state.get("selected_dataset", None)
    if selected_dataset is None:
        st.info("Please select a dataset first.")
        st.stop()

    # Load the dataframe fresh for the selected dataset
    file_path = file_map.get(selected_dataset)
    df = load_file(file_path)
    if df.empty:
        st.warning(f"No data found for dataset: {selected_dataset}")
        st.stop()

    # Optionally, you can apply your existing sidebar filters here if needed

    filtered_data = df.copy()

    # ------------- Page title -------------
    st.title("📊 Visit Intelligence Dashboard")

    # ------------- Advanced Summary -------------
    import datetime, pandas as pd

    with st.expander("📢 Advanced Summary", expanded=True):

        if filtered_data.empty:
            st.info("No results found for your selection.")
            st.stop()

        # Use only completed activity rows (if that column exists)
        adv_data = filtered_data.copy()
        if "Activity Status" in adv_data.columns:
            adv_data = adv_data[adv_data["Activity Status"].str.lower() == "completed"]

        # ------- helper: convert messy cells to timedelta -------
        def to_timedelta_str(x):
            if pd.isnull(x) or x in ["", "-", "NaT", None, " "]:
                return pd.NaT
            if isinstance(x, (pd.Timedelta, datetime.timedelta)):
                return x
            if isinstance(x, datetime.time):
                return datetime.timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
            if isinstance(x, (float, int)):
                # Excel stores times as fractions of a day
                try:
                    return datetime.timedelta(seconds=int(float(x) * 86400))
                except Exception:
                    return pd.NaT
            if isinstance(x, pd.Timestamp):
                t = x.time()
                return datetime.timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
            if isinstance(x, str):
                s = x.strip()
                if s in ["", "-", "NaT", " "]:
                    return pd.NaT
                # Try HH:MM:SS
                try:
                    h, m, s = map(int, s.split(":")[:3])
                    return datetime.timedelta(hours=h, minutes=m, seconds=s)
                except Exception:
                    pass
                # Fallback: pandas parser
                try:
                    return pd.to_timedelta(s)
                except Exception:
                    return pd.NaT
            return pd.NaT


        valid_times = adv_data.copy()
        if {"Activate", "Deactivate"}.issubset(valid_times.columns):
            valid_times["Activate"]   = valid_times["Activate"].apply(to_timedelta_str)
            valid_times["Deactivate"] = valid_times["Deactivate"].apply(to_timedelta_str)
            valid_times = valid_times[
                valid_times["Activate"].notna() &
                valid_times["Deactivate"].notna() &
                (valid_times["Activate"]   > datetime.timedelta(0)) &
                (valid_times["Deactivate"] > datetime.timedelta(0))
            ]
        else:
            valid_times = pd.DataFrame()

        # ------- engineer name if single-filtered -------
        name = None
        if "Engineer" in adv_data.columns:
            engs = adv_data["Engineer"].unique()
            if len(engs) == 1:
                name = engs[0]

        visits      = len(adv_data)
        total_value = adv_data["Value"].sum() if "Value" in adv_data.columns else None

        def avg_time_str(col):
            if col not in valid_times.columns or valid_times.empty:
                return "N/A"
            vals = valid_times[col].dropna()
            if vals.empty:
                return "N/A"
            avg = vals.mean()
            return f"{int(avg.total_seconds()//3600):02}:{int((avg.total_seconds()%3600)//60):02}"

        avg_activate   = avg_time_str("Activate")
        avg_deactivate = avg_time_str("Deactivate")

        # Most common visit type (excluding lunch)
        if "Visit Type" in adv_data.columns:
            vt = adv_data[~adv_data["Visit Type"].str.contains("lunch", case=False, na=False)]
            most_common_type = vt["Visit Type"].mode()[0] if not vt["Visit Type"].mode().empty else "N/A"
        else:
            most_common_type = "N/A"

        # Busiest day
        busiest_day, busiest_count = "N/A", ""
        if "Date" in adv_data.columns:
            counts = adv_data["Date"].dt.date.value_counts()
            if not counts.empty:
                busiest_day   = counts.idxmax().strftime("%d %B %Y")
                busiest_count = f"{counts.max()} visits"

        # -------- Build summary sentence --------
        prefix = (
            f"**{name}** completed a total of {visits:,} visits"
            if name else f"Summary of your current selection: {visits:,} completed visits"
        )

        summary = (
            f"{prefix}"
            + (f", generating an overall value of £{total_value:,.2f}" if total_value is not None else "")
            + f". On average, visits began at {avg_activate} and concluded at {avg_deactivate}. "
            f"Excluding lunch, the most frequently performed visit type was '{most_common_type}'. "
            + (f"The busiest day recorded was {busiest_day} with {busiest_count}. " if busiest_day != 'N/A' else "")
        )

        st.markdown(
            f"<div style='font-size:1.05em; line-height:1.65em; margin-bottom:16px;'>{summary}</div>",
            unsafe_allow_html=True,
        )

    @st.cache_data
    def load_file(path):
        try:
            df = pd.read_excel(path)
            df.columns = df.columns.str.strip()  # Strip all column names of whitespace
            #st.write("DEBUG: Columns loaded from", path, df.columns.tolist())

            # Standardise column names by dataset
            if "AI Test SB Visits" in path:
                df = df.rename(columns={
                    'Business Engineers Name': 'Engineer',
                    'Date of visit': 'Date',
                    'Venue Name': 'Venue',
                    'Visit type': 'Visit Type',
                    'Total Value': 'Value'
                })
            elif "Invoice Data AI" in path:
                df = df.rename(columns={
                    'Date of visit': 'Date',
                    'Total Value': 'Value'
                })
            elif "Productivity Report" in path:
                # For Productivity Report, do NOT try to rename or convert date,
                # since it has no date column.
                pass
            elif any(x in path for x in [
                "VIP North Oracle Data", "VIP South Oracle Data",
                "Tier 2 North Oracle Data", "Tier 2 South Oracle Data"
            ]):
                df = df.rename(columns={
                    'Name': 'Engineer',
                    'Date': 'Date',
                    'Visit Type': 'Visit Type',
                    'Total Value': 'Value',
                    'Postcode': 'Venue'
                })

            # Handle date fields for datasets that have a 'Date' column and are NOT Productivity Report
            if 'Date' in df.columns and "Productivity Report" not in path:
                df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True)
                df.dropna(subset=['Date'], inplace=True)
                df['MonthName'] = df['Date'].dt.month_name()
                df['Week'] = df['Date'].dt.isocalendar().week

            return df

        except Exception as e:
            st.error(f"⚠️ Failed to load file: {e}")
            return pd.DataFrame()




# --- NUMBER 10 ---
# --- SECTION: Call Log Data ---

if st.session_state.screen == "dashboard_view" and st.session_state.selected_dataset == "Call Log Data":
    # Back button
    if st.button("⬅️ Back to Dataset Selection", use_container_width=True):
        st.session_state.screen = "dashboard"
        st.session_state.selected_dataset = None
        st.rerun()

    # Convert date column early
    if "Date of Call Taken" in filtered_data.columns:
        filtered_data["Date of Call Taken"] = pd.to_datetime(filtered_data["Date of Call Taken"], errors="coerce")

    st.subheader("📞 Call Log Overview")

    # RAW TABLE
    with st.expander("📋 Raw Call Log Table", expanded=False):
        st.dataframe(filtered_data, use_container_width=True)

    # SUMMARY KPIs
    with st.expander("📋 Summary KPIs", expanded=True):
        cols = st.columns(3)
        col1, col2, col3 = cols[0], cols[1], cols[2]

        # Prepare data metrics
        total_calls = len(filtered_data)
        unique_engineers = filtered_data["Name of Engineer"].nunique() if "Name of Engineer" in filtered_data.columns else "N/A"
        unique_regions = filtered_data["Region"].nunique() if "Region" in filtered_data.columns else "N/A"
        unique_options = filtered_data["Option Selected"].nunique() if "Option Selected" in filtered_data.columns else "N/A"
        unique_months = filtered_data["Month"].nunique() if "Month" in filtered_data.columns else "N/A"
        unique_vr = filtered_data["VR Number (If Known)"].nunique() if "VR Number (If Known)" in filtered_data.columns else "N/A"
        unique_callers = filtered_data["Name Of Engineer Who Made The Call"].nunique() if "Name Of Engineer Who Made The Call" in filtered_data.columns else "N/A"
        unique_emails = filtered_data["Engineers email address (who made the call)"].nunique() if "Engineers email address (who made the call)" in filtered_data.columns else "N/A"

        peak_day = "N/A"
        peak_day_calls = "N/A"
        if "Date of Call Taken" in filtered_data.columns:
            calls_per_day = filtered_data["Date of Call Taken"].dt.date.value_counts()
            if not calls_per_day.empty:
                peak_day = calls_per_day.idxmax().strftime("%d %b %Y")
                peak_day_calls = calls_per_day.max()

        # Display KPIs in a 3x3 grid
        col1.metric("Total Calls", f"{total_calls:,}")
        col2.metric("Unique Engineers", unique_engineers)
        col3.metric("Regions", unique_regions)

        col1.metric("Unique Options Selected", unique_options)
        col2.metric("Unique Months", unique_months)
        col3.metric("Unique VR Numbers", unique_vr)

        col1.metric("Unique Callers", unique_callers)
        col2.metric("Unique Caller Emails", unique_emails)
        col3.metric("Peak Call Day", f"{peak_day} ({peak_day_calls} calls)" if peak_day_calls != "N/A" else peak_day)

    # TOP 5 REGIONS BY CALL VOLUME
    if "Region" in filtered_data.columns:
        with st.expander("🏆 Top 5 Regions by Call Volume"):
            top_regions = (
                filtered_data["Region"].value_counts().head(5)
                .reset_index()
            )
            top_regions.columns = ["Region", "Call Count"]
            st.dataframe(top_regions)
            st.plotly_chart(px.bar(top_regions, x="Region", y="Call Count", color="Region",
                                   title="Top 5 Regions by Call Volume"), use_container_width=True)
            st.plotly_chart(px.pie(top_regions, names="Region", values="Call Count",
                                   title="Region Call Distribution"), use_container_width=True)

    # OPTION SELECTED
    if "Option Selected" in filtered_data.columns:
        with st.expander("📊 Call Volume by Option (Top 10)"):
            option_counts = (
                filtered_data["Option Selected"].value_counts().head(10)
                .reset_index()
            )
            option_counts.columns = ["Option", "Call Count"]
            st.plotly_chart(px.bar(option_counts, x="Option", y="Call Count", color="Option",
                                   title="Top 10 Options by Call Volume"), use_container_width=True)
            st.plotly_chart(px.pie(option_counts, names="Option", values="Call Count",
                                   title="Option Call Distribution"), use_container_width=True)

    # SUNBURST (Region → Option Selected)
    if {"Region", "Option Selected"}.issubset(filtered_data.columns):
        with st.expander("🌞 Region vs Option Sunburst"):
            sunburst_df = filtered_data.groupby(["Region", "Option Selected"]).size().reset_index(name="Count")
            fig = px.sunburst(sunburst_df, path=["Region", "Option Selected"], values="Count",
                              title="Call Distribution: Region → Option")
            st.plotly_chart(fig, use_container_width=True)

    # CALLS OVER TIME
    if "Date of Call Taken" in filtered_data.columns:
        with st.expander("📈 Call Volume Over Time"):
            df_time = filtered_data.copy()
            df_time["Date of Call Taken"] = pd.to_datetime(df_time["Date of Call Taken"], errors="coerce")
            calls_by_day = df_time.groupby("Date of Call Taken").size().reset_index(name="Call Count")
            st.plotly_chart(px.line(calls_by_day, x="Date of Call Taken", y="Call Count",
                                   title="Calls Over Time (Line Chart)"), use_container_width=True)
            st.plotly_chart(px.bar(calls_by_day, x="Date of Call Taken", y="Call Count",
                                   title="Calls Over Time (Bar Chart)"), use_container_width=True)

    # TIME REQUIRED DISTRIBUTION
    if "Time Required Hours" in filtered_data.columns:
        with st.expander("⏱️ Time Required Distribution"):
            df_time = pd.to_numeric(filtered_data["Time Required Hours"], errors="coerce")
            st.plotly_chart(px.histogram(df_time.dropna(), nbins=20,
                                         title="Distribution of Time Required (Hours)"), use_container_width=True)

    # TOP ENGINEERS
    if "Name of Engineer" in filtered_data.columns:
        with st.expander("🧑 Top Engineers by Call Volume"):
            top_eng = (
                filtered_data["Name of Engineer"].value_counts().head(10)
                .reset_index()
            )
            top_eng.columns = ["Engineer", "Call Count"]
            st.plotly_chart(px.bar(top_eng, x="Engineer", y="Call Count", color="Engineer",
                                   title="Top 10 Engineers by Call Volume"), use_container_width=True)
            st.plotly_chart(px.bar(top_eng, y="Engineer", x="Call Count", color="Engineer", orientation="h",
                                   title="Top 10 Engineers by Call Volume (Horizontal)"), use_container_width=True)
    import pandas as pd
    import numpy as np
    import plotly.express as px

    # Additional KPIs expander
    with st.expander("📊 Additional KPIs", expanded=False):
        total_calls = len(filtered_data)
        calls_by_engineer = filtered_data.groupby("Name of Engineer").size()
        calls_by_region = filtered_data.groupby("Region").size()

        avg_calls_per_engineer = calls_by_engineer.mean() if not calls_by_engineer.empty else 0
        avg_calls_per_region = calls_by_region.mean() if not calls_by_region.empty else 0

        most_frequent_option = filtered_data["Option Selected"].mode().iloc[0] if not filtered_data["Option Selected"].mode().empty else "N/A"

        col1, col2, col3 = st.columns(3)
        col1.metric("Avg Calls per Engineer", f"{avg_calls_per_engineer:.2f}")
        col2.metric("Avg Calls per Region", f"{avg_calls_per_region:.2f}")
        col3.metric("Most Frequent Option", most_frequent_option)

    # Calls by Engineer & Region Pivot Table
    with st.expander("📋 Calls by Engineer & Region", expanded=False):
        if {"Name of Engineer", "Region"}.issubset(filtered_data.columns):
            pivot_eng_reg = pd.pivot_table(filtered_data, index="Name of Engineer", columns="Region", values="Date of Call Taken", aggfunc="count", fill_value=0)
            st.dataframe(pivot_eng_reg)

            fig = px.bar(pivot_eng_reg.reset_index(), x="Name of Engineer", y=pivot_eng_reg.columns.tolist(), 
                         title="Calls by Engineer and Region", barmode='stack')
            st.plotly_chart(fig, use_container_width=True)

    # Calls by Option & Month Heatmap
    with st.expander("📅 Calls by Option & Month", expanded=False):
        if {"Option Selected", "Month"}.issubset(filtered_data.columns):
            pivot_opt_month = pd.pivot_table(filtered_data, index="Option Selected", columns="Month", values="Date of Call Taken", aggfunc="count", fill_value=0)
            
            # Sort months in chronological order by parsing month names
            pivot_opt_month = pivot_opt_month.reindex(sorted(pivot_opt_month.columns, key=lambda x: pd.to_datetime(x, format='%B')), axis=1)

            st.dataframe(pivot_opt_month)

            fig = px.imshow(pivot_opt_month, labels=dict(x="Month", y="Option Selected", color="Call Count"), 
                            title="Heatmap of Calls by Option and Month")
            st.plotly_chart(fig, use_container_width=True)

    # Month-on-Month % Change in Calls
    with st.expander("📈 Month-on-Month % Change in Total Calls", expanded=False):
        calls_by_month = filtered_data.groupby(filtered_data["Date of Call Taken"].dt.to_period("M")).size()
        pct_change = calls_by_month.pct_change().fillna(0) * 100
        pct_df = pct_change.reset_index().rename(columns={"Date of Call Taken": "Month", 0: "Percent Change"})
        
        # Rename columns properly
        pct_df.columns = ["Month", "Percent Change"]

        # Convert period to datetime for sorting and format as string
        pct_df["Month_dt"] = pd.to_datetime(pct_df["Month"].astype(str), format='%Y-%m')
        pct_df = pct_df.sort_values("Month_dt")
        pct_df["Month"] = pct_df["Month_dt"].dt.strftime('%b %Y')

        st.dataframe(pct_df.drop(columns=["Month_dt"]))

        fig = px.bar(pct_df, x="Month", y="Percent Change", title="Month-on-Month % Change in Calls")
        st.plotly_chart(fig, use_container_width=True)


    # VR Number Distribution
    with st.expander("🎫 VR Number Distribution", expanded=False):
        if "VR Number (If Known)" in filtered_data.columns:
            vr_counts = filtered_data["VR Number (If Known)"].value_counts().head(15)
            vr_df = vr_counts.reset_index()
            vr_df.columns = ["VR Number", "Call Count"]
            st.dataframe(vr_df)

            fig = px.bar(vr_df, x="VR Number", y="Call Count", title="Top 15 VR Numbers by Call Count")
            st.plotly_chart(fig, use_container_width=True)

    # Region Call Volume & Distribution
    with st.expander("📊 Region Call Volume & Distribution", expanded=False):
        if "Region" in filtered_data.columns:
            calls_by_region = filtered_data["Region"].value_counts().reset_index()
            calls_by_region.columns = ["Region", "Call Count"]
            st.plotly_chart(
                px.bar(calls_by_region, x="Region", y="Call Count", color="Region",
                       title="Call Volume by Region"),
                use_container_width=True,
                key="bar_region_call_volume"
            )

            st.plotly_chart(
                px.pie(calls_by_region, names="Region", values="Call Count",
                       title="Call Volume Distribution by Region"),
                use_container_width=True,
                key="pie_region_call_distribution"
            )

    # Region & Option / Engineer Insights
    with st.expander("📈 Region & Option / Engineer Insights", expanded=False):
        if {"Region", "Option Selected"}.issubset(filtered_data.columns):
            sunburst_df = (
                filtered_data.groupby(["Region", "Option Selected"])
                .size()
                .reset_index(name="Count")
            )
            fig_sunburst = px.sunburst(sunburst_df, path=["Region", "Option Selected"], values="Count",
                                       title="Call Distribution: Region → Option")
            st.plotly_chart(fig_sunburst, use_container_width=True, key="sunburst_region_option")

        if {"Region", "Name of Engineer"}.issubset(filtered_data.columns):
            pivot_eng_reg = pd.pivot_table(
                filtered_data,
                index="Name of Engineer",
                columns="Region",
                values="Date of Call Taken",
                aggfunc="count",
                fill_value=0,
            )
            st.dataframe(pivot_eng_reg)

            fig_stacked_bar = px.bar(
                pivot_eng_reg.reset_index(),
                x="Name of Engineer",
                y=pivot_eng_reg.columns.tolist(),
                title="Calls by Engineer and Region",
                barmode="stack",
            )
            st.plotly_chart(fig_stacked_bar, use_container_width=True, key="stacked_bar_eng_reg")

    # Call Volume by Region & Option Selected (Sunburst or Treemap)
    with st.expander("🌳 Call Volume by Region & Option Selected (Sunburst)", expanded=False):
        if {"Region", "Option Selected"}.issubset(filtered_data.columns):
            region_option = (
                filtered_data.groupby(["Region", "Option Selected"])
                .size()
                .reset_index(name="Count")
            )
            fig_region_option = px.sunburst(region_option,
                                           path=["Region", "Option Selected"],
                                           values="Count",
                                           title="Call Volume by Region & Option Selected")
            st.plotly_chart(fig_region_option, use_container_width=True, key="sunburst_region_option_2")




# --- NUMBER 11 ---
# --- SECTION: Productivity Report ---

if (
    st.session_state.screen == "dashboard_view"
    and st.session_state.selected_dataset == "Productivity Report"
):
    # Back button
    if st.button("⬅️ Back to Dataset Selection", use_container_width=True):
        st.session_state.screen = "dashboard"
        st.session_state.selected_dataset = None
        st.rerun()

    st.subheader("🚀 Productivity Report Overview")

    # ------------------------------------------------------------------
    #   QUICK KPI METRICS
    # ------------------------------------------------------------------
    money_kpi_cols = [
        ("TOTAL REVENUE", "Total Revenue (£)", "sum"),
        ("TARGET REVENUE", "Target Revenue (£)", "sum"),
        ("TARGET REVENUE +/-", "Δ Revenue (£)", "sum"),
    ]
    percent_kpi_cols = [
        ("TARGET REVENUE % +/-", "Δ Revenue %", "mean"),
        ("TOTAL COMPLETION RATE % Overall", "Total Completion %", "mean"),
    ]

    with st.expander("📋 KPI Summary", expanded=True):
        k1, k2, k3, k4, k5 = st.columns(5)
        # Money KPIs
        if "TOTAL REVENUE" in filtered_data.columns:
            k1.metric(
                "Total Revenue (£)",
                f"£{filtered_data['TOTAL REVENUE'].sum():,.0f}",
            )
        if {c[0] for c in money_kpi_cols}.issubset(filtered_data.columns):
            delta = (
                filtered_data["TOTAL REVENUE"].sum()
                - filtered_data["TARGET REVENUE"].sum()
            )
            k2.metric(
                "Δ vs Target (£)",
                f"£{delta:,.0f}",
                delta_color="inverse" if delta < 0 else "normal",
            )
        # Percentage KPIs
        if "TARGET REVENUE % +/-" in filtered_data.columns:
            pct = filtered_data["TARGET REVENUE % +/-"].mean() * 100
            k3.metric("Δ Revenue %", f"{pct:+.1f}%")
        if "TOTAL COMPLETION RATE % Overall" in filtered_data.columns:
            comp_pct = filtered_data["TOTAL COMPLETION RATE % Overall"].mean() * 100
            k4.metric("Completion %", f"{comp_pct:.1f}%")
        if "TOTAL VISITS COMPLETED" in filtered_data.columns:
            total_visits = filtered_data["TOTAL VISITS COMPLETED"].sum()
            k5.metric("Visits Completed", f"{total_visits:,}")

    # ------------------------------------------------------------------
    #   CLEAN & GROUP COLUMNS FOR CHART GALLERY
    # ------------------------------------------------------------------
    money_cols = [
        "TOTAL REVENUE",
        "TARGET REVENUE",
        "TARGET REVENUE +/-",
        "Overtime Average",
        "Total OT for Month",
    ]
    percent_cols = [
        "TARGET REVENUE % +/-",
        "Invoice Completeion Rate",
        "Total Percentage Productivity",
        "TOTAL COMPLETION RATE % Overall",
        "Average Daily Completion Rate",
        "% ABOVE OR BELOW TARGET",
        "TOTAL Capactity FOR THE MONTH",
    ]
    visit_cols = [
        "TOTAL VISITS ISSUED",
        "TOTAL VISITS COMPLETED",
        "TOTAL VISITS PENDING (NOT INCLUDING LUNCH)",
        "TOTAL VISITS CANCELLED",
        "TOTAL VISITS STARTED NOT COMPLETED",
        "TOTAL VISITS NOT DONE",
        "ESTIMATED VISITS FOR THE MONTH",
    ]

    # Filter to existing cols
    money_cols   = [c for c in money_cols if c in filtered_data.columns]
    percent_cols = [c for c in percent_cols if c in filtered_data.columns]
    visit_cols   = [c for c in visit_cols if c in filtered_data.columns]

    # ------------------------------------------------------------------
    #   CHART GALLERY PER METRIC
    # ------------------------------------------------------------------
    chart_columns = money_cols + percent_cols + visit_cols
    for col in chart_columns:
        if col in money_cols:
            value_col = col
            display_col = col.replace("_", " ")
        elif col in percent_cols:
            value_col = col
            display_col = col.replace("_", " ")
            # convert to percent (0-1 ➜ 0-100) if needed
            if filtered_data[value_col].max() <= 1.01:
                filtered_data[value_col] = filtered_data[value_col] * 100
        else:  # visits
            value_col = col
            display_col = col.replace("_", " ")

        with st.expander(f"📊 {display_col} Charts"):
            left, right = st.columns(2)

            # Vertical Bar
            with left:
                st.plotly_chart(
                    px.bar(
                        filtered_data,
                        x="Team",       # Using Team grouping as valid column
                        y=value_col,
                        color="Team",
                        title=f"{display_col} by Team (Bar)",
                        labels={value_col: display_col, "Team": "Team"},
                    ),
                    use_container_width=True,
                )
                st.plotly_chart(
                    px.pie(
                        filtered_data,
                        names="Team",
                        values=value_col,
                        title=f"{display_col} Share by Team (Donut)",
                        hole=0.45,
                    ),
                    use_container_width=True,
                )

            # Horizontal + Pie
            with right:
                st.plotly_chart(
                    px.bar(
                        filtered_data,
                        y="Team",
                        x=value_col,
                        color="Team",
                        orientation="h",
                        title=f"{display_col} by Team (Horizontal)",
                        labels={value_col: display_col, "Team": "Team"},
                    ),
                    use_container_width=True,
                )
                st.plotly_chart(
                    px.pie(
                        filtered_data,
                        names="Team",
                        values=value_col,
                        title=f"{display_col} Share by Team (Pie)",
                    ),
                    use_container_width=True,
                )

    # ------------------------------------------------------------------
    #   HEATMAP – Team vs Metrics (money_cols only)
    # ------------------------------------------------------------------
    if len(money_cols) >= 2:
        with st.expander("🌡️ Revenue Metrics Heatmap"):
            pivot = filtered_data.pivot_table(index="Team", values=money_cols, aggfunc="sum")
            heat = px.imshow(pivot, text_auto=True, aspect="auto",
                             title="Revenue Metrics by Team")
            st.plotly_chart(heat, use_container_width=True)

    # ------------------------------------------------------------------
    #   RAW TABLE
    # ------------------------------------------------------------------
    with st.expander("📋 Full Productivity Data"):
        st.dataframe(filtered_data, use_container_width=True)


# --- NUMBER 12 ---
# --- SECTION: Invoice Data AI ---

if st.session_state.screen == "dashboard_view" and st.session_state.selected_dataset == "Invoice Data AI":
    # Back button
    if st.button("⬅️ Back to Dataset Selection", use_container_width=True):
        st.session_state.screen = "dashboard"
        st.rerun()

    st.subheader("📄 Invoice Data AI Overview")

    # Clean data
    df_invoice = filtered_data.copy()
    df_invoice = df_invoice.replace(["", " ", "00:00", "00:00:00", 0, "0", None], pd.NA)
    df_invoice.dropna(how='all', inplace=True)

    # Fix column naming and add YearMonth for monthly grouping
    if "Date of visit" in df_invoice.columns:
        df_invoice["Date of visit"] = pd.to_datetime(df_invoice["Date of visit"], errors='coerce')
        df_invoice["YearMonth"] = df_invoice["Date of visit"].dt.to_period("M").astype(str)
        df_invoice["Week"] = df_invoice["Date of visit"].dt.isocalendar().week

    # Ensure numeric conversions for relevant columns
    for col in ["Total Value", "Labour Value", "Time On-Site"]:
        if col in df_invoice.columns:
            df_invoice[col] = pd.to_numeric(df_invoice[col], errors='coerce')

    # KPI Grid (3x3) with additional detailed metrics
    with st.expander("📋 Summary KPIs", expanded=True):
        col1, col2, col3 = st.columns(3)

        total_value = df_invoice["Total Value"].sum() if "Total Value" in df_invoice.columns else 0
        avg_invoice = df_invoice["Total Value"].mean() if "Total Value" in df_invoice.columns else 0
        max_invoice = df_invoice["Total Value"].max() if "Total Value" in df_invoice.columns else 0
        min_invoice = df_invoice["Total Value"].min() if "Total Value" in df_invoice.columns else 0
        median_invoice = df_invoice["Total Value"].median() if "Total Value" in df_invoice.columns else 0

        total_labour = df_invoice["Labour Value"].sum() if "Labour Value" in df_invoice.columns else 0
        avg_time_onsite = df_invoice["Time On-Site"].mean() if "Time On-Site" in df_invoice.columns else 0

        total_invoices = len(df_invoice)
        unique_types = df_invoice["Visit Type"].nunique() if "Visit Type" in df_invoice.columns else 0

        internal_counts = df_invoice["Internal Or External"].value_counts() if "Internal Or External" in df_invoice.columns else pd.Series()

        col1.metric("Total Invoice Value (£)", f"£{total_value:,.2f}")
        col2.metric("Average Invoice Value (£)", f"£{avg_invoice:,.2f}")
        col3.metric("Median Invoice Value (£)", f"£{median_invoice:,.2f}")

        col1.metric("Max Invoice Value (£)", f"£{max_invoice:,.2f}")
        col2.metric("Min Invoice Value (£)", f"£{min_invoice:,.2f}")
        

        
        col2.metric("Total Invoices", f"{total_invoices:,}")
        col3.metric("Unique Invoice Types", f"{unique_types}")

    import pandas as pd

    import pandas as pd

    if "Visit Type" in df_invoice.columns:
        visit_types = df_invoice["Visit Type"].unique()

        st.markdown("### 📊 KPIs by Visit Type")

        for vt in visit_types:
            vt_df = df_invoice[df_invoice["Visit Type"] == vt]

            # Exclude unwanted columns
            exclude_cols = ["Time On-Site", "Week"]
            numeric_cols = [col for col in vt_df.select_dtypes(include="number").columns if col not in exclude_cols]

            # Calculate KPIs for this Visit Type
            kpi_data = {}
            for col in numeric_cols:
                if any(keyword in col for keyword in ["Value", "Cost", "Total", "Amount", "Count"]):
                    kpi_data[col] = vt_df[col].sum()
                else:
                    kpi_data[col] = vt_df[col].mean()

            if not kpi_data:
                st.info(f"No numeric KPIs found for Visit Type: {vt}")
                continue

            kpi_df = pd.DataFrame.from_dict(kpi_data, orient="index", columns=["Value"])

            with st.expander(f"Visit Type: {vt} KPIs", expanded=False):
                st.dataframe(kpi_df.style.format("{:,.2f}"), use_container_width=True)



    # Raw data table
    with st.expander("🧾 Raw Invoice Table", expanded=False):
        st.dataframe(df_invoice, use_container_width=True)

    # Box plot: Invoice Value distribution by Visit Type
    if "Visit Type" in df_invoice.columns and "Total Value" in df_invoice.columns:
        with st.expander("📦 Invoice Value Distribution by Visit Type"):
            fig_box = px.box(df_invoice, x="Visit Type", y="Total Value",
                             title="Invoice Value Distribution by Visit Type")
            st.plotly_chart(fig_box, use_container_width=True)

    # Histogram: Time On-Site distribution
    if "Time On-Site" in df_invoice.columns:
        with st.expander("⏱️ Time On-Site Distribution"):
            fig_hist = px.histogram(df_invoice, x="Time On-Site", nbins=30, title="Time On-Site Distribution (mins)")
            st.plotly_chart(fig_hist, use_container_width=True)

    # Bar chart: Total Labour Value by Team
    if "Team" in df_invoice.columns and "Labour Value" in df_invoice.columns:
        with st.expander("💼 Total Labour Value by Team"):
            labour_team = df_invoice.groupby("Team")["Labour Value"].sum().reset_index()
            fig_bar_labour = px.bar(labour_team, x="Team", y="Labour Value", title="Total Labour Value by Team")
            st.plotly_chart(fig_bar_labour, use_container_width=True)

    # Pie chart: Internal vs External
    if "Internal Or External" in df_invoice.columns:
        with st.expander("🔍 Internal vs External Visits"):
            internal_counts = df_invoice["Internal Or External"].value_counts()
            fig_pie_int_ext = px.pie(names=internal_counts.index, values=internal_counts.values, title="Internal vs External Visits")
            st.plotly_chart(fig_pie_int_ext, use_container_width=True)

    # Line chart: Monthly Invoice Value with Min, Max, Median
    if "YearMonth" in df_invoice.columns and "Total Value" in df_invoice.columns:
        with st.expander("📈 Monthly Invoice Value with Min, Max, Median"):
            monthly_stats = df_invoice.groupby("YearMonth")["Total Value"].agg(["min", "max", "median"]).reset_index()
            fig_line_stats = px.line(monthly_stats, x="YearMonth", y=["min", "max", "median"],
                                    title="Monthly Invoice Value: Min, Max, Median", markers=True)
            st.plotly_chart(fig_line_stats, use_container_width=True)

    # Keep your existing charts too
    if "Visit Type" in df_invoice.columns and "Value" in df_invoice.columns:
        with st.expander("💰 Top 5 Visit Types by Value"):
            top_value = df_invoice.groupby("Visit Type")["Value"].sum().sort_values(ascending=False).head(5).reset_index()
            st.plotly_chart(px.bar(top_value, x="Visit Type", y="Value", color="Visit Type",
                                   title="Top 5 Visit Types by Value (£)"), use_container_width=True)

    if "Visit Type" in df_invoice.columns:
        with st.expander("📊 Top 5 Visit Types by Count"):
            top_count = df_invoice["Visit Type"].value_counts().head(5).reset_index()
            top_count.columns = ["Visit Type", "Count"]
            st.plotly_chart(px.bar(top_count, x="Visit Type", y="Count", color="Visit Type",
                                   title="Top 5 Visit Types by Volume"), use_container_width=True)

    if {"Visit Type", "Week"}.issubset(df_invoice.columns):
        with st.expander("🌞 Visit Type → Week Sunburst"):
            sun_df = df_invoice.groupby(["Visit Type", "Week"]).size().reset_index(name="Count")
            fig = px.sunburst(sun_df, path=["Visit Type", "Week"], values="Count",
                              title="Visit Type Breakdown by Week")
            st.plotly_chart(fig, use_container_width=True)

    if {"Week", "Value"}.issubset(df_invoice.columns):
        with st.expander("📈 Visit Trends by Week"):
            week_value = df_invoice.groupby("Week")["Value"].sum().reset_index()
            week_count = df_invoice.groupby("Week").size().reset_index(name="Count")
            st.plotly_chart(px.line(week_value, x="Week", y="Value", title="Total Invoice Value by Week"), use_container_width=True)
            st.plotly_chart(px.bar(week_count, x="Week", y="Count", title="Invoice Volume by Week"), use_container_width=True)

    if {"Visit Type", "Week"}.issubset(df_invoice.columns):
        with st.expander("🌡️ Visit Type vs Week Heatmap"):
            heat_df = pd.pivot_table(df_invoice, index="Visit Type", columns="Week", aggfunc="size", fill_value=0)
            st.plotly_chart(px.imshow(heat_df, aspect="auto", title="Visit Heatmap: Types by Week"), use_container_width=True)

    with st.expander("📋 Full Invoice Table", expanded=False):
        st.dataframe(df_invoice, use_container_width=True)




# --- NUMBER 13 ---
# --- SECTION: AI Test SB Visits ---

if st.session_state.screen == "dashboard_view" and st.session_state.selected_dataset == "AI Test SB Visits":
    # Back button
    if st.button("⬅️ Back to Dataset Selection", use_container_width=True):
        st.session_state.screen = "dashboard"
        st.session_state.selected_dataset = None
        st.rerun()

    st.subheader("🧪 AI Test SB Visits Overview")

    df_sb = filtered_data.copy()
    df_sb = df_sb.replace(["", " ", "00:00", "00:00:00", 0, "0", None], pd.NA)
    df_sb.dropna(how="all", inplace=True)

    # Rename for easier access
    df_sb.rename(columns={
        "Business Engineers Name": "Engineer",
        "Date of visit": "Date",
        "Visit type": "Visit Type",
        "Total Value": "Value"
    }, inplace=True)

    # Convert Date and extract Week and Month
    df_sb["Date"] = pd.to_datetime(df_sb["Date"], errors='coerce')
    df_sb["Week"] = df_sb["Date"].dt.isocalendar().week
    df_sb["Month"] = df_sb["Date"].dt.to_period('M').astype(str)

    # Ensure Value column is numeric for aggregation
    df_sb["Value"] = pd.to_numeric(df_sb["Value"], errors="coerce")

    # --- Summary KPIs per Visit Type (no nested expanders) ---
    with st.expander("📋 Summary KPIs per Visit Type", expanded=False):
        visit_types = df_sb["Visit Type"].dropna().unique()
        for vt in visit_types:
            sub_df = df_sb[df_sb["Visit Type"] == vt]
            total_visits = len(sub_df)
            unique_engineers = sub_df["Engineer"].nunique()
            total_value = sub_df["Value"].sum()
            avg_value = sub_df["Value"].mean()

            st.markdown(f"### Visit Type: {vt}")
            kpi1, kpi2, kpi3 = st.columns(3)
            kpi1.metric("Total Visits", f"{total_visits:,}")
            kpi2.metric("Unique Engineers", f"{unique_engineers:,}")
            kpi3.metric("Total Invoice Value (£)", f"£{total_value:,.2f}")
            st.write(f"Average Invoice Value (£): £{avg_value:,.2f}")

    # --- Overall KPIs (3x3 grid) ---
    with st.expander("📊 Overall KPIs", expanded=True):
        col1, col2, col3 = st.columns(3)
        col4, col5, col6 = st.columns(3)
        col7, col8, col9 = st.columns(3)

        total_visits = len(df_sb)
        unique_engineers = df_sb["Engineer"].nunique()
        unique_visit_types = df_sb["Visit Type"].nunique()
        total_value = df_sb["Value"].sum()
        avg_value = df_sb["Value"].mean()
        min_value = df_sb["Value"].min()
        max_value = df_sb["Value"].max()
        median_value = df_sb["Value"].median()

        col1.metric("Total Visits", f"{total_visits:,}")
        col2.metric("Unique Engineers", f"{unique_engineers:,}")
        col3.metric("Unique Visit Types", f"{unique_visit_types:,}")

        col4.metric("Total Invoice Value (£)", f"£{total_value:,.2f}")
        col5.metric("Average Invoice Value (£)", f"£{avg_value:,.2f}")
        col6.metric("Median Invoice Value (£)", f"£{median_value:,.2f}")

        col7.metric("Minimum Invoice Value (£)", f"£{min_value:,.2f}")
        col8.metric("Maximum Invoice Value (£)", f"£{max_value:,.2f}")
        col9.metric("Data Range (Months)", f"{df_sb['Month'].nunique()}")

    # (rest of your code unchanged...)


    # --- Top Visit Types by Volume and Value ---
    if "Visit Type" in df_sb.columns:
        with st.expander("📊 Top 5 Visit Types by Volume"):
            top_visits = df_sb["Visit Type"].value_counts().head(5).reset_index()
            top_visits.columns = ["Visit Type", "Count"]
            st.plotly_chart(px.bar(top_visits, x="Visit Type", y="Count", color="Visit Type",
                                   title="Top 5 Visit Types by Volume"), use_container_width=True)

        with st.expander("💰 Top 5 Visit Types by Invoice Value"):
            value_sum = df_sb.groupby("Visit Type")["Value"].sum().sort_values(ascending=False).head(5).reset_index()
            st.plotly_chart(px.bar(value_sum, x="Visit Type", y="Value", color="Visit Type",
                                   title="Top 5 Visit Types by Invoice Value"), use_container_width=True)

    # --- Top Engineers by Visits and Value ---
    if "Engineer" in df_sb.columns:
        with st.expander("🧑 Top 5 Engineers by Visit Volume"):
            eng_counts = df_sb["Engineer"].value_counts().head(5).reset_index()
            eng_counts.columns = ["Engineer", "Count"]
            st.plotly_chart(px.bar(eng_counts, x="Engineer", y="Count", color="Engineer",
                                   title="Top 5 Engineers by Visits"), use_container_width=True)

        with st.expander("💰 Top 5 Engineers by Invoice Value"):
            value_sum = df_sb.groupby("Engineer")["Value"].sum().nlargest(5).reset_index()
            st.plotly_chart(px.bar(value_sum, x="Engineer", y="Value", color="Engineer",
                                   title="Top 5 Engineers by Invoice Value"), use_container_width=True)

    # --- Visits Over Time by Month ---
    if "Month" in df_sb.columns:
        with st.expander("📈 Visit Volume Over Time (Monthly)"):
            monthly_visits = df_sb.groupby("Month").size().reset_index(name="Visits")
            st.plotly_chart(px.line(monthly_visits, x="Month", y="Visits", title="Monthly Visit Volume"), use_container_width=True)

    # --- Pie Charts ---
    if "Visit Type" in df_sb.columns:
        with st.expander("🥧 Visit Type Share"):
            visit_type_counts = df_sb["Visit Type"].value_counts().reset_index()
            visit_type_counts.columns = ["Visit Type", "Count"]
            st.plotly_chart(px.pie(visit_type_counts, values="Count", names="Visit Type", title="Visit Type Share"), use_container_width=True)

    if "Engineer" in df_sb.columns:
        with st.expander("👥 Engineer Contribution Share"):
            engineer_counts = df_sb["Engineer"].value_counts().reset_index()
            engineer_counts.columns = ["Engineer", "Count"]
            st.plotly_chart(px.pie(engineer_counts, values="Count", names="Engineer", title="Engineer Contribution Share"), use_container_width=True)

    # --- Heatmap: Visit Type vs Week ---
    if {"Visit Type", "Week"}.issubset(df_sb.columns):
        with st.expander("🌡️ Visit Type vs Week Heatmap"):
            heat_df = pd.pivot_table(df_sb, index="Visit Type", columns="Week", aggfunc="size", fill_value=0)
            st.plotly_chart(px.imshow(heat_df, aspect="auto", title="Visit Heatmap: Types by Week"),
                            use_container_width=True)

    # --- Sunburst: Visit Type → Engineer ---
    if {"Visit Type", "Engineer"}.issubset(df_sb.columns):
        with st.expander("🌞 Visit Type → Engineer Sunburst"):
            sb_counts = df_sb.groupby(["Visit Type", "Engineer"]).size().reset_index(name="Count")
            fig = px.sunburst(sb_counts, path=["Visit Type", "Engineer"], values="Count",
                              title="Visits by Type and Engineer")
            st.plotly_chart(fig, use_container_width=True)

    # --- Full Table ---
    with st.expander("📋 Full AI Test SB Visit Table", expanded=False):
        # Drop 'Week' and 'Time On-Site' columns if they exist
        show_df = df_sb.drop(columns=[col for col in ["Week", "Time On-Site"] if col in df_sb.columns])
        st.dataframe(show_df, use_container_width=True)

# --- NUMBER 14 ---
# --- SECTION: Oracle Team – Advanced Summary ---

import datetime
import os
import pandas as pd
import streamlit as st

ORACLE_TEAMS = [
    "VIP North Oracle Data",
    "VIP South Oracle Data",
    "Tier 2 North Oracle Data",
    "Tier 2 South Oracle Data",
]

if (
    st.session_state.screen == "dashboard_view"
    and st.session_state.selected_dataset in ORACLE_TEAMS
):

    # Back button
    if st.button("⬅️ Back to Dataset Selection", use_container_width=True, key="back_oracle"):
        st.session_state.screen = "dashboard"
        st.session_state.selected_dataset = None
        st.rerun()

    # Local function to load Oracle data only for this block
    @st.cache_data(show_spinner=False)
    def load_all_data():
        def clean_df(df: pd.DataFrame) -> pd.DataFrame:
            # Replace unwanted values with NaN
            df.replace({"0": pd.NA, "": pd.NA, " ": pd.NA}, inplace=True)

            # Strip whitespace from all object/string columns
            for c in df.select_dtypes("object").columns:
                df[c] = df[c].astype(str).str.strip()

            # Replace additional unwanted string variants with NaN
            df.replace({"": pd.NA, "nan": pd.NA, "None": pd.NA}, inplace=True)

            # Convert Date column to datetime, drop rows with invalid dates
            if "Date" in df.columns:
                df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
                df.dropna(subset=["Date"], inplace=True)

            # Convert certain columns to timedelta if they exist
            for td in ("Activate", "Deactivate", "Total Time"):
                if td in df.columns:
                    df[td] = pd.to_timedelta(df[td].astype(str), errors="coerce")

            return df

        # Dictionary of Oracle dataset filenames keyed by team name
        oracle_files = {
            "VIP South": "VIP South Oracle Data.xlsx",
            "VIP North": "VIP North Oracle Data.xlsx",
            "Tier 2 South": "Tier 2 South Oracle Data.xlsx",
            "Tier 2 North": "Tier 2 North Oracle Data.xlsx",
        }

        oracle_frames, missing = [], []

        # Load each Oracle dataset file if it exists
        for team, path in oracle_files.items():
            if os.path.exists(path):
                tmp = pd.read_excel(path)
                tmp["Team"] = team  # Add team column for identification
                oracle_frames.append(tmp)
            else:
                missing.append(path)

        # Warn if any files are missing
        if missing:
            st.warning("Missing Oracle files: " + ", ".join(missing))

        # Concatenate all dataframes if any loaded, else return empty DataFrame
        if oracle_frames:
            df_oracle = clean_df(pd.concat(oracle_frames, ignore_index=True))
        else:
            df_oracle = pd.DataFrame()

        return df_oracle

    # Load data only here inside block 14
    with st.spinner("Loading Oracle data, please wait..."):
        df_oracle = load_all_data()

    selected_dataset = st.session_state.get("selected_dataset")

# Map full dataset name to team label used in df_oracle["Team"]
    dataset_to_team = {
        "VIP South Oracle Data": "VIP South",
        "VIP North Oracle Data": "VIP North",
        "Tier 2 South Oracle Data": "Tier 2 South",
        "Tier 2 North Oracle Data": "Tier 2 North",
    }

    if selected_dataset in dataset_to_team:
        team_name = dataset_to_team[selected_dataset]
        df_oracle = df_oracle[df_oracle["Team"] == team_name]



    # Now you can safely use df_oracle below in Block 14 only

    # ---------------- Load & Clean ----------------
    df = df_oracle.copy()
    df.columns = df.columns.str.strip()
    df.replace(["", " ", "00:00", "00:00:00", 0, "0", None], pd.NA, inplace=True)

    # Standard names we rely on
    df = df.rename(columns={
        "Name": "Engineer",
        "Total Value": "Value",
    })

    if df.empty:
        st.warning("No data for the selected filters.")
        st.stop()

    # -------------- Helper functions --------------
    def to_timedelta_str(x):
        if pd.isnull(x) or x in ["", "-", "NaT", None, " "]:
            return pd.NaT
        if isinstance(x, (pd.Timedelta, datetime.timedelta)):
            return x
        if isinstance(x, datetime.time):
            return datetime.timedelta(hours=x.hour, minutes=x.minute, seconds=x.second)
        if isinstance(x, (float, int)):
            try:
                return datetime.timedelta(seconds=int(float(x) * 86400))
            except Exception:
                return pd.NaT
        if isinstance(x, pd.Timestamp):
            t = x.time()
            return datetime.timedelta(hours=t.hour, minutes=t.minute, seconds=t.second)
        if isinstance(x, str):
            s = x.strip()
            if s in ["", "-", "NaT", " "]:
                return pd.NaT
            try:
                h, m, s = map(int, s.split(":")[:3])
                return datetime.timedelta(hours=h, minutes=m, seconds=s)
            except Exception:
                pass
            try:
                return pd.to_timedelta(s)
            except Exception:
                return pd.NaT
        return pd.NaT

    # --- Replace old avg logic ---
    df["Activate"] = df["Activate"].apply(to_timedelta_str)
    df["Deactivate"] = df["Deactivate"].apply(to_timedelta_str)

    valid_times = df[
        df["Activate"].notna() &
        df["Deactivate"].notna() &
        (df["Activate"] > datetime.timedelta(0)) &
        (df["Deactivate"] > datetime.timedelta(0))
    ]

    avg_activate_time = avg_deactivate_time = "N/A"
    if not valid_times.empty:
        avg_act = valid_times["Activate"].mean()
        avg_deact = valid_times["Deactivate"].mean()
        avg_activate_time = f"{int(avg_act.total_seconds()//3600):02}:{int((avg_act.total_seconds()%3600)//60):02}"
        avg_deactivate_time = f"{int(avg_deact.total_seconds()//3600):02}:{int((avg_deact.total_seconds()%3600)//60):02}"


    def td_to_str(td):
        if isinstance(td, (pd.Timedelta, datetime.timedelta)):
            s = int(td.total_seconds())
            return f"{s//3600:02}:{(s%3600)//60:02}:{s%60:02}"
        return "N/A"

    # -------------- Value columns -----------------
    if "Value" in df.columns:
        df["Value"] = pd.to_numeric(df["Value"], errors="coerce")
        total_value = f"£{df['Value'].sum(skipna=True):,.2f}"
        avg_value = f"£{df['Value'].mean(skipna=True):,.2f}"
    else:
        total_value = avg_value = "N/A"

    # -------------- Date range --------------------
    if "Date" in df.columns:
        earliest = df["Date"].min().strftime("%d %b %Y")
        latest = df["Date"].max().strftime("%d %b %Y")
    else:
        earliest = latest = "N/A"



    # -------------- Lunch duration ----------------
    lunch_col = next((c for c in df.columns if c.lower().startswith("total time")), None)
    avg_lunch_str = "N/A"
    if lunch_col:
        def to_td(val):
            if pd.isnull(val):
                return pd.NaT
            if isinstance(val, datetime.time):
                return datetime.timedelta(hours=val.hour, minutes=val.minute, seconds=val.second)
            try:
                return pd.to_timedelta(str(val))
            except Exception:
                return pd.NaT
        lunches = (
            df.loc[df["Visit Type"] == "Lunch (30)", lunch_col]
              .apply(to_td)
              .dropna()
        )
        if not lunches.empty:
            avg_lunch_str = td_to_str(lunches.mean())

    # -------------- Common visit type -------------
    data_no_lunch = df[df["Visit Type"] != "Lunch (30)"] if "Visit Type" in df.columns else df.copy()
    if "Visit Type" in data_no_lunch.columns and not data_no_lunch["Visit Type"].mode().empty:
        common_type = data_no_lunch["Visit Type"].mode()[0]
    else:
        common_type = "N/A"

    # ----------- Advanced Summary output ----------
    st.subheader("📊 Oracle Team Visit Overview")
    st.dataframe(df.head())

    st.markdown(
        f"""
        <div style='background:#1f2937;padding:16px 20px;border-radius:10px;color:#e0e0e0;font-size:1.03em;line-height:1.6em'>
        <b>Advanced Summary:</b><br><br>
        Across <b>{len(df):,}</b> rows, engineers completed <b>{len(data_no_lunch):,}</b> visits (excluding lunch),
        generating <b>{total_value}</b> in total value and averaging <b>{avg_value}</b> per visit.<br>
        The most common visit type was <b>{common_type}</b>.<br><br>
        Shifts typically started at <b>{avg_activate_time}</b> and ended by <b>{avg_deactivate_time}</b>.<br>
        Average lunch duration was <b>{avg_lunch_str}</b>.<br><br>
        Data covers <b>{earliest}</b> to <b>{latest}</b>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -------------- Bullet list -------------------
    st.markdown(
        f"""
        - **Total Rows:** {len(df):,}
        - **Unique Engineers:** {df['Engineer'].nunique() if 'Engineer' in df.columns else 'N/A'}
        - **Unique Visit Types:** {df['Visit Type'].nunique() if 'Visit Type' in df.columns else 'N/A'}
        - **Date Range:** {earliest} → {latest}
        - **Total Value:** {total_value}
        - **Average Value per Visit:** {avg_value}
        - **Average Activate Time:** {avg_activate_time}
        - **Average Deactivate Time:** {avg_deactivate_time}
        - **Most Common Visit Type:** {common_type}
        """
    )

    # Section 3 of Block 14: Oracle Team Visit Data (All Regions)

    # Copy and clean filtered_data into df_oracle for detailed analysis
    df_oracle = filtered_data.copy()
    df_oracle.columns = df_oracle.columns.str.strip()
    columns_to_clean = [col for col in df_oracle.columns if col != "Activity Status"]

    df_oracle[columns_to_clean] = df_oracle[columns_to_clean].replace(
        ["", " ", "00:00", "00:00:00", 0, "0", None], pd.NA
    )
    df_oracle.dropna(how="all", inplace=True)

    # Parse date columns
    if "Date" in df_oracle.columns:
        df_oracle["Date"] = pd.to_datetime(df_oracle["Date"], errors="coerce")
        df_oracle["Week"] = df_oracle["Date"].dt.isocalendar().week
        df_oracle["Month"] = df_oracle["Date"].dt.strftime("%B")

    # Load full raw Oracle data for completion breakdown (adjust loading method if needed)
    vip_south_df = load_file("VIP South Oracle Data.xlsx")
    vip_north_df = load_file("VIP North Oracle Data.xlsx")
    tier2_south_df = load_file("Tier 2 South Oracle Data.xlsx")
    tier2_north_df = load_file("Tier 2 North Oracle Data.xlsx")

    df_all = pd.concat([vip_south_df, vip_north_df, tier2_south_df, tier2_north_df], ignore_index=True)

    # Normalize status column
    status = (
        df_all["Activity Status"]
            .astype(str)
            .str.strip()
            .str.casefold()
    )

    # Count status values
    vc = status.value_counts()

    completed  = vc.get("completed", 0)
    not_done   = vc.get("not done", 0)
    cancelled  = status.str.contains("cancel", na=False).sum()

    known      = completed + cancelled + not_done
    total      = int(vc.sum())
    other      = total - known

    # Calculate metrics
    completion_rate_pct       = (completed / known * 100) if known else 0
    completion_vs_failed_ratio = (completed / (cancelled + not_done)) if (cancelled + not_done) > 0 else float("inf")

    # Display Activity Completion Breakdown
    with st.expander("🧩 Activity Completion Breakdown", expanded=False):
        st.markdown(f"""
        ✅ **Completed**: {completed:,} ({completed / total:.1%})  
        ❌ **Cancelled**: {cancelled:,} ({cancelled / total:.1%})  
        🚫 **Not Done**:  {not_done:,} ({not_done / total:.1%})  
        ❓ **Other/Unknown**: {other:,} ({other / total:.1%})
        """)

        col1, col2, col3 = st.columns(3)
        col1.metric("✔ Completion Rate", f"{completion_rate_pct:.1f}%")
        col2.metric("🔁 Completed : Failed", f"{completion_vs_failed_ratio:.1f} ×")
        col3.markdown(
            f"🔁 **{completion_vs_failed_ratio:.1f}** visits completed for every **1** cancelled or not done visit"
        )

        st.bar_chart(vc)

    # Optional Debug Output (commented)
    # with st.expander("📊 Unique statuses in data", expanded=False):
    #     st.dataframe(
    #         pd.DataFrame(vc).reset_index().rename(
    #             columns={"index": "Activity Status", 0: "Count"}
    #         )
    #     )

    # --- KPIs ---
    with st.expander("📋 Summary KPIs", expanded=True):
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Total Visits", len(df_oracle))
        col2.metric("Unique Engineers", df_oracle["Name"].nunique())
        col3.metric("Visit Types", df_oracle["Visit Type"].nunique())
        col4.metric("Total Value (£)", f"£{df_oracle['Total Value'].dropna().sum():,.2f}")

    # --- Top Visit Types ---
    if "Visit Type" in df_oracle.columns:
        with st.expander("📊 Top Visit Types"):
            top_vt = df_oracle["Visit Type"].value_counts().head(10).reset_index()
            top_vt.columns = ["Visit Type", "Count"]
            st.plotly_chart(px.bar(top_vt, x="Visit Type", y="Count", color="Visit Type",
                                   title="Top Visit Types by Volume"), use_container_width=True)

    # --- Top Engineers ---
    if "Name" in df_oracle.columns:
        with st.expander("👨 Top Engineers by Visits"):
            eng_top = df_oracle["Name"].value_counts().head(10).reset_index()
            eng_top.columns = ["Engineer", "Visits"]
            st.plotly_chart(px.bar(eng_top, x="Engineer", y="Visits", color="Engineer",
                                   title="Top Engineers by Visit Count"), use_container_width=True)

    # --- Weekly Trends ---
    if "Week" in df_oracle.columns:
        with st.expander("📈 Weekly Visit Trends"):
            weekly = df_oracle.groupby("Week").size().reset_index(name="Visits")
            st.plotly_chart(px.line(weekly, x="Week", y="Visits", title="Visits Over Weeks"),
                            use_container_width=True)

    # --- Monthly Value Breakdown ---
    if "Month" in df_oracle.columns and "Total Value" in df_oracle.columns:
        with st.expander("💰 Total Value by Month"):
            monthly_val = df_oracle.groupby("Month")["Total Value"].sum().reset_index()
            st.plotly_chart(px.bar(monthly_val, x="Month", y="Total Value", color="Month",
                                   title="Total Value by Month"), use_container_width=True)

    # --- Sunburst Charts ---
    sunburst_configs = [
        ("🌞 Visit Activity Sunburst", ["Visit Type", "Activity Status"], "Visit Type & Activity Status Distribution"),
        ("📍 Sunburst: Visit Type to Postcode", ["Visit Type", "Postcode"], "Visit Type → Postcode Distribution"),
        ("🔀 Sunburst: Engineer → Visit Type → Week", ["Name", "Visit Type", "Week"], "Engineer > Visit Type > Week Breakdown"),
        ("🌀 Sunburst: Visit Type → Week", ["Visit Type", "Week"], "Visit Count by Visit Type and Week"),
        ("🧩 Sunburst: Engineer → Postcode", ["Name", "Postcode"], "Engineer > Postcode Mapping"),
        ("🗓️ Sunburst: Visit Type → Month", ["Visit Type", "Month"], "Visit Type Distribution by Month"),
        ("📅 Sunburst: Visit Type → Date", ["Visit Type", "Date"], "Visit Type Breakdown by Exact Date"),
        ("📋 Sunburst: Visit Type → Day", ["Visit Type", "Day"], "Visit Type by Day of Week"),
        ("📑 Sunburst: Stakeholder → Visit Type", ["Sky Retail Stakeholder", "Visit Type"], "Stakeholder to Visit Type Breakdown")
    ]

    for label, cols, title in sunburst_configs:
        if set(cols).issubset(df_oracle.columns):
            with st.expander(label):
                sb = df_oracle.groupby(cols).size().reset_index(name="Count")
                fig = px.sunburst(sb, path=cols, values="Count", title=title)
                st.plotly_chart(fig, use_container_width=True)

    # --- Stacked Bar: Visit Type by Month ---
    if {"Visit Type", "Month"}.issubset(df_oracle.columns):
        with st.expander("📊 Stacked Bar: Visit Type by Month"):
            bar_df = df_oracle.groupby(["Month", "Visit Type"]).size().reset_index(name="Visits")
            fig = px.bar(bar_df, x="Month", y="Visits", color="Visit Type", title="Monthly Visit Counts by Visit Type",
                         text_auto=True)
            st.plotly_chart(fig, use_container_width=True)

    # --- Parallel Categories: Engineer → Visit Type → Postcode ---
    if {"Name", "Visit Type", "Postcode"}.issubset(df_oracle.columns):
        with st.expander("🔗 Parallel Categories: Engineer → Visit Type → Postcode"):
            pc_df = df_oracle[["Name", "Visit Type", "Postcode"]].dropna().astype(str)
            fig = px.parallel_categories(pc_df, dimensions=["Name", "Visit Type", "Postcode"],
                                         color_continuous_scale=px.colors.sequential.Inferno,
                                         title="Engineer to Visit Type to Postcode Flow")
            st.plotly_chart(fig, use_container_width=True)

    # --- Drilldown Treemap: Stakeholder → Visit Type → Month ---
    if {"Sky Retail Stakeholder", "Visit Type", "Month", "Total Value"}.issubset(df_oracle.columns):
        with st.expander("🌲 Drilldown Treemap: Stakeholder → Visit Type → Month"):
            tree_df = df_oracle.groupby(["Sky Retail Stakeholder", "Visit Type", "Month"])["Total Value"].sum().reset_index()
            fig = px.treemap(tree_df, path=["Sky Retail Stakeholder", "Visit Type", "Month"],
                             values="Total Value", title="Value Drilldown by Stakeholder → Visit Type → Month")
            st.plotly_chart(fig, use_container_width=True)

    # --- Heatmap: Visit Type vs Week ---
    if {"Visit Type", "Week"}.issubset(df_oracle.columns):
        with st.expander("🌡️ Visit Type vs Week Heatmap"):
            heat_df = pd.pivot_table(df_oracle, index="Visit Type", columns="Week", aggfunc="size", fill_value=0)
            st.plotly_chart(px.imshow(heat_df, aspect="auto", title="Visit Heatmap: Types by Week"),
                            use_container_width=True)

    # --- Treemap: Visit Type by Value ---
    if {"Visit Type", "Total Value"}.issubset(df_oracle.columns):
        with st.expander("🌳 Treemap: Visit Type by Total Value"):
            tm = df_oracle.groupby("Visit Type")["Total Value"].sum().reset_index()
            fig = px.treemap(tm, path=["Visit Type"], values="Total Value",
                             title="Total Value by Visit Type")
            st.plotly_chart(fig, use_container_width=True)

    # --- Pie Chart: Visit Type Share ---
    if "Visit Type" in df_oracle.columns:
        with st.expander("🥧 Visit Type Share (Pie)"):
            pie = df_oracle["Visit Type"].value_counts().reset_index()
            pie.columns = ["Visit Type", "Count"]
            fig = px.pie(pie, names="Visit Type", values="Count", title="Visit Type Distribution")
            st.plotly_chart(fig, use_container_width=True)

    # --- Table View ---
    with st.expander("📋 Full Oracle Visit Table", expanded=False):
        st.dataframe(df_oracle, use_container_width=True)






import datetime
import os
import pandas as pd
import streamlit as st

# ── 1. Horizontal uploader and chat input (inline layout) ─────
if st.session_state.get("screen") == "ai":

    # Inject combined CSS for uploader layout and chat bubbles
    st.markdown(
        """
        <style>
        .upload-chat-row {
            display: flex;
            gap: 1rem;
            align-items: center;
            margin-bottom: 1rem;
        }
        .upload-chat-row > div {
            flex: 1;
        }

        /* Chat bubble styles */
        .user-msg {
            background-color: #4c6ef5;
            color: white;
            padding: 12px 20px;
            border-radius: 20px 20px 0 20px;
            max-width: 70%;
            margin-left: auto;
            margin-bottom: 10px;
            font-size: 1rem;
            font-weight: 600;
        }
        .assistant-msg {
            background-color: #e9ecef;
            color: #343a40;
            padding: 12px 20px;
            border-radius: 20px 20px 20px 0;
            max-width: 70%;
            margin-right: auto;
            margin-bottom: 10px;
            font-size: 1rem;
            font-weight: 400;
        }
        </style>
        """,
        unsafe_allow_html=True
    )




# Section 3 of Block 16 #
if st.session_state.get("screen") == "ai":
    # ---------- Navigation ----------
    if st.button("⬅️ Back to Main Menu", use_container_width=True, key="back_ai_16"):
        st.session_state.screen = "area_selection"
        st.rerun()

    st.markdown("## 🤖 Sky Orbit")
    st.markdown(
        "**SKY ORBIT** stands for: Sky Business focus, Knowledge & Intelligence (AI-powered), Your Data Unified, Oracle integration, Reporting & Results, Business Insights, Interactive Visualizations, Tracking visits & forecasts."
    )
    user_df = st.session_state.get("user_df", None)

    # Button to show logs
    if st.button("Show AI Chat Logs"):
        st.session_state.show_logs = True

    # Button to clear chat screen
    if st.button("Clear Chat Screen"):
        st.session_state.ai_chat = []  # Clears visible chat messages only
        st.rerun()  # Refresh the page to clear UI

    # Render chat messages from session state with timestamps
    for msg in st.session_state.get("ai_chat", []):
        role = msg.get("role", "")
        content = msg.get("content", "")
        timestamp = msg.get("timestamp", None)

        if timestamp:
            try:
                ts_str = datetime.datetime.fromisoformat(timestamp).strftime("%Y-%m-%d %H:%M")
            except Exception:
                ts_str = ""
        else:
            ts_str = ""

        if role == "user":
            st.markdown(f'''
                <div class="user-msg">
                    👤 {content}
                    <div style="font-size: 0.7rem; color: #ccc; text-align: right;">{ts_str}</div>
                </div>
            ''', unsafe_allow_html=True)
        else:
            st.markdown(f'''
                <div class="assistant-msg">
                    🤖 {content}
                    <div style="font-size: 0.7rem; color: #aaa; text-align: left;">{ts_str}</div>
                </div>
            ''', unsafe_allow_html=True)

    # Show chat logs if toggled
    if st.session_state.get("show_logs", False):
        password = st.text_input("Enter admin password to view logs:", type="password")
        
        if password:
            if password == "AI Chat":  # Change to your secure password
                st.success("Access granted to chat logs 📊")
                log_file_path = os.path.join(os.getcwd(), "chat_logs.csv")
                if os.path.exists(log_file_path):
                    logs_df = pd.read_csv(log_file_path, parse_dates=["timestamp"], on_bad_lines='skip')
                    if logs_df.empty:
                        st.warning("No chat logs found!")
                    else:
                        st.dataframe(logs_df)
                else:
                    st.warning("No chat logs file found!")
            else:
                st.error("Incorrect password. Access denied.")




# --- Section 4 of Block 16: Data Loading & Agent Setup ---

from langchain_openai import ChatOpenAI
from langchain.agents.agent_types import AgentType
from langchain_experimental.agents import create_pandas_dataframe_agent

@st.cache_data(show_spinner=False)
def load_all_data():
    def clean_df(df: pd.DataFrame) -> pd.DataFrame:
        df.replace({"0": pd.NA, "": pd.NA, " ": pd.NA}, inplace=True)
        for c in df.select_dtypes("object").columns:
            df[c] = df[c].astype(str).str.strip()
        df.replace({"": pd.NA, "nan": pd.NA, "None": pd.NA}, inplace=True)
        if "Date" in df.columns:
            df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
            df.dropna(subset=["Date"], inplace=True)
        for td in ("Activate", "Deactivate", "Total Time"):
            if td in df.columns:
                df[td] = pd.to_timedelta(df[td].astype(str), errors="coerce")
        return df

    oracle_files = {
        "VIP South":   "VIP South Oracle Data.xlsx",
        "VIP North":   "VIP North Oracle Data.xlsx",
        "Tier 2 South":"Tier 2 South Oracle Data.xlsx",
        "Tier 2 North":"Tier 2 North Oracle Data.xlsx",
    }

    oracle_frames, missing = [], []
    for team, path in oracle_files.items():
        if os.path.exists(path):
            tmp = pd.read_excel(path)
            tmp["Team"] = team
            oracle_frames.append(tmp)
        else:
            missing.append(path)

    if missing:
        st.warning("Missing Oracle files: " + ", ".join(missing))

    if oracle_frames:
        df_oracle = clean_df(pd.concat(oracle_frames, ignore_index=True))
    else:
        df_oracle = pd.DataFrame()

    return df_oracle

# Load Oracle data (cache-enabled)
with st.spinner("Loading Oracle data, please wait..."):
    df_oracle = load_all_data()


if df_oracle.empty:
    st.warning("No Oracle data loaded.")
    st.stop()

# Use uploaded file if available, else default to df_oracle
user_df = st.session_state.get("user_df", None)
default_df = user_df if user_df is not None else df_oracle

# Setup LangChain streaming LLM agent once per session
if "df_agent" not in st.session_state:
    llm_stream = ChatOpenAI(
        api_key=st.secrets["openai"]["api_key"],
        model_name="gpt-4o-mini",
        streaming=True,
    )
    st.session_state.df_agent = create_pandas_dataframe_agent(
        llm=llm_stream,
        df=default_df,
        verbose=False,
        agent_type=AgentType.OPENAI_FUNCTIONS,
        allow_dangerous_code=True,
    )

df_agent = st.session_state.df_agent

# Optional debugging info (uncomment to enable)
# df_map = {"df_oracle": df_oracle}
# if user_df is not None:
#     df_map["uploaded"] = user_df
# combined_schema = "\n".join(f"{name}: {list(df.columns)}" for name, df in df_map.items())
# with st.expander("🧾 Loaded DataFrames (debug)", expanded=False):
#     st.markdown(
#         f"""
#         <div style='font-size:0.9rem;color:grey;'>
#         <strong>📂 Loaded DataFrames:</strong><br>{combined_schema}
#         </div>
#         """,
#         unsafe_allow_html=True,
#     )




# --- Section 5 of Block 16: Data Filtering Helper ---

def filter_zero_values(df: pd.DataFrame, query: str) -> pd.DataFrame:
    qlc = query.lower()

    # If user explicitly asks about zeros or similar, return full dataframe without filtering
    if any(keyword in qlc for keyword in ["0 entries", "zero counts", "show zero", "count zero", "how many zero"]):
        return df

    df_filtered = df.copy()

    # Exclude invalid postcodes like "0", empty, or nan values
    if "Postcode" in df_filtered.columns:
        df_filtered = df_filtered[
            ~df_filtered["Postcode"].astype(str).str.strip().isin(["0", "", "nan", "None"])
        ]

    # Columns to filter out zero or missing numeric values (customize as needed)
    numeric_cols = ["Visit Count", "Total Time", "Total Value", "Total Cost Inc Travel"]
    time_cols = ["Activate", "Deactivate", "Total Time", "Travel Time", "Total Time (Inc Travel)"]

    for col in numeric_cols:
        if col in df_filtered.columns:
            df_filtered = df_filtered[
                ~((df_filtered[col] == 0) | (df_filtered[col].isna()))
            ]

    for col in time_cols:
        if col in df_filtered.columns:
            df_filtered = df_filtered[
                ~((df_filtered[col] == pd.Timedelta(0)) | df_filtered[col].astype(str).str.startswith("00:00"))
            ]

    return df_filtered


# --- Section 6 of Block 16 -----------------------

# Define aliases to map user-friendly terms to actual dataframe columns
ALIASES = {
    "activate time": "Activate",
    "activate": "Activate",
    "deactivate time": "Deactivate",
    "deactivate": "Deactivate",
    "visit type": "Visit Type",
    "total £": "Total Value",
    "total value": "Total Value",
    "total cost": "Total Cost Inc Travel",
    "total time": "Total Time",
    "total working time": "Total Working Time",
    "travel time": "Travel Time",
    "stakeholder": "Sky Retail Stakeholder",
}

# Quick keyword to canonical column mapping
KEYWORDS = {
    "stakeholder": "Sky Retail Stakeholder",
    "engineer": "Name" if "Name" in df_oracle.columns else "Engineer",
    "postcode": "Postcode",
    "visit type": "Visit Type",
    "status": "Activity Status",
    "team": "Team",
    "month": "Month",
    "week": "Week",
    "day": "Day",
}

def alias(text: str) -> str:
    """Replace friendly phrases in the user prompt with canonical column names."""
    t = text.lower()
    for k, v in ALIASES.items():
        t = t.replace(k, v)
    return t

# Setup LangChain streaming OpenAI model
llm_stream = ChatOpenAI(
    api_key=st.secrets["openai"]["api_key"],
    model_name="gpt-4o-mini",
    streaming=True,
)

# Create LangChain pandas dataframe agent
df_agent = create_pandas_dataframe_agent(
    llm=llm_stream,
    df=default_df,
    verbose=False,
    agent_type=AgentType.OPENAI_FUNCTIONS,
    allow_dangerous_code=True,
)

# Fallback OpenAI client (non-LangChain streaming)
fallback_client = OpenAI(api_key=st.secrets["openai"]["api_key"]).chat.completions




# ──────────────────────────────────────────────────────────────────────
# 4⃣  TYPING EFFECT HELPER
# ──────────────────────────────────────────────────────────────────────
def slow_stream(token_iter, placeholder, delay: float = 0.03):
    buf = ""
    for chunk in token_iter:
        buf += chunk
        placeholder.markdown(
            f"<div class='streaming-token'>{buf}</div>",
            unsafe_allow_html=True,
        )
        time.sleep(delay)
    placeholder.markdown(buf, unsafe_allow_html=True)
    return buf

# ── 0. Remember any file the user has already dropped ──────────
if "user_df" not in st.session_state:
    st.session_state.user_df = None         # holds the DataFrame
if "user_file_name" not in st.session_state:
    st.session_state.user_file_name = None  # keeps the name for display

# ──────────────────────────────────────────────────────────────────────
# 5⃣  CHAT HISTORY UI
# ──────────────────────────────────────────────────────────────────────
import datetime

# Initialize chat state if not present
if "ai_chat" not in st.session_state:
    st.session_state.ai_chat = []
if "ai" not in st.session_state:
    st.session_state.ai = True

# Inject CSS for chat bubbles (you can put this in your main CSS section too)
st.markdown(
    """
    <style>
    .user-msg {
        background-color: #4c6ef5;
        color: white;
        padding: 12px 20px;
        border-radius: 20px 20px 0 20px;
        max-width: 70%;
        margin-left: auto;
        margin-bottom: 10px;
        font-size: 1rem;
        font-weight: 600;
    }
    .assistant-msg {
        background-color: #e9ecef;
        color: #343a40;
        padding: 12px 20px;
        border-radius: 20px 20px 20px 0;
        max-width: 70%;
        margin-right: auto;
        margin-bottom: 10px;
        font-size: 1rem;
        font-weight: 400;
    }
    .timestamp {
        font-size: 0.7rem;
        color: #888;
        margin-top: 4px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Render chat messages with timestamps and bubbles
for msg in st.session_state.ai_chat:
    role = msg.get("role", "assistant")
    content = msg.get("content", "")
    timestamp = msg.get("timestamp", None)
    if timestamp:
        ts_str = datetime.datetime.fromisoformat(timestamp).strftime("%Y-%m-%d %H:%M")
    else:
        ts_str = ""

    if role == "user":
        st.markdown(
            f'''
            <div class="user-msg">
                👤 {content}
                <div class="timestamp" style="text-align: right;">{ts_str}</div>
            </div>
            ''',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f'''
            <div class="assistant-msg">
                🤖 {content}
                <div class="timestamp" style="text-align: left;">{ts_str}</div>
            </div>
            ''',
            unsafe_allow_html=True,
        )


# ──────────────────────────────────────────────────────────────────────
# 6⃣  CHART-TYPE PICKER (RULE-BASED)  ← ***FIXED***
#     • Only returns a chart type when the question clearly implies one.
#     • If the user doesn’t hint at a visual, returns None.
# ──────────────────────────────────────────────────────────────────────
def pick_chart_type(q: str) -> str | None:
    q = q.lower()
    if any(k in q for k in ("sunburst", "treemap", "parallel")):
        return "sunburst"
    if any(k in q for k in ("corr", "correlation", "matrix")):
        return "corr"
    if any(k in q for k in ("parallel", "flow")):
        return "parallel"
    if any(k in q for k in ("kpi", "dashboard")):
        return "kpi"
    if any(k in q for k in ("trend", "over time", "line")):
        return "line"
    if any(k in q for k in ("share", "proportion", "percentage", "pie")):
        return "pie"
    if any(k in q for k in ("bar chart", "histogram", "distribution", 
                            "grouped", "split", "vs", "versus", "by ")):
        return "bar"
    return None  # ← default is now *no* chart


def render_chart(chart_type: str, df: pd.DataFrame, query: str = ""):
    query_lc = query.lower()

    # ----- SPECIAL CASE: Including vs Excluding Travel ----------------
    if "including travel" in query_lc and "excluding travel" in query_lc:
        # Use df_oracle, ensure it's defined in global scope before calling
        global df_oracle
        df = df_oracle.copy()
        df["Total Time"] = pd.to_timedelta(df["Total Time"], errors="coerce")
        df["Travel Time"] = pd.to_timedelta(df["Travel Time"], errors="coerce")
        df["Excluding Travel"] = df["Total Time"] - df["Travel Time"]
        df["Including Travel"] = df["Total Time"]
        summary = (
            df.groupby("Team")[["Including Travel", "Excluding Travel"]]
              .mean()
              .reset_index()
        )
        melt = summary.melt(id_vars="Team", var_name="Type", value_name="Avg Time")
        melt["Seconds"] = melt["Avg Time"].dt.total_seconds()
        fig = px.bar(
            melt, x="Team", y="Seconds", color="Type", barmode="group",
            title="Avg Time by Team (Including vs Excluding Travel)",
        )
        st.plotly_chart(fig, use_container_width=True)
        return  # don’t draw anything else for this query

    # ---------- Helpers ----------
    def pick_col(cols):
        for kw, canon in KEYWORDS.items():
            if kw in query_lc and canon in cols:
                return canon
        for c in cols:
            if c.lower() in query_lc:
                return c
        return next(
            (c for c in cols if df[c].dtype == "object" and df[c].nunique() < 30),
            None,
        )

    def filter_invalid_postcodes(df: pd.DataFrame):
        df_clean = df.copy()
        # Remove rows where 'Total Value' is 0 or missing
        df_clean = df_clean[df_clean['Total Value'].fillna(0) > 0]
        # Exclude numeric-looking postcodes
        df_clean = df_clean[~df_clean['Postcode'].str.strip().str.isnumeric()]
        return df_clean

    def get_top_postcodes_by_value(df, top_n=5):
        filtered = df[df["Total Value"].fillna(0) > 0].copy()
        filtered = filtered[~filtered["Postcode"].astype(str).str.strip().isin(["0", "", "nan", "None"])]
        grouped = filtered.groupby("Postcode")["Total Value"].sum().reset_index()
        top_postcodes = grouped.sort_values("Total Value", ascending=False).head(top_n)
        lines = []
        for i, row in enumerate(top_postcodes.itertuples(), 1):
            postcode = row.Postcode
            value = f"£{row._2:,.2f}"  # format as currency
            lines.append(f"{i}. Postcode: {postcode} - Total Value: {value}")
        return "\n".join(lines)

if st.session_state.get("screen") == "ai":
	# ──────────── Helpers (Used by Main Q&A Loop) ────────────

	def get_total_cost_completed_visits(df, team=None):
		"""Calculate total cost for completed visits, optionally filtered by team."""
		if "Activity Status" not in df.columns or "Total Cost Inc Travel" not in df.columns:
			return None
		filtered = df[df["Activity Status"].str.lower() == "completed"]
		if team:
			filtered = filtered[filtered["Team"] == team]
		total_cost = filtered["Total Cost Inc Travel"].sum()
		return total_cost

	def calculate_percentage_diff_avg(df, group_col="Team", value_col="Total Value"):
		"""Calculate percentage difference from average total value per group."""
		totals = df.groupby(group_col)[value_col].sum().reset_index()
		baseline = totals[value_col].mean()
		totals["Pct Difference"] = ((totals[value_col] - baseline) / baseline * 100).round(2)
		totals["Total Value"] = totals[value_col].map("£{:,.2f}".format)
		totals["Pct Difference"] = totals["Pct Difference"].map(lambda x: f"{x:.2f}%")

		output_lines = []
		for _, row in totals.iterrows():
			output_lines.append(f"{row[group_col]}: {row['Total Value']} ({row['Pct Difference']})")

		return output_lines

	def clean_postcodes(df):
		"""Clean postcode column: uppercase, strip, and exclude invalid postcodes."""
		df = df.copy()
		df['Postcode'] = df['Postcode'].astype(str).str.strip().str.upper()

		invalids = {"", "0", "NAN", "NONE", "NULL"}

		df = df[~df['Postcode'].isin(invalids)]
		return df

	def filter_zero_values(df: pd.DataFrame, query: str) -> pd.DataFrame:
		"""
		Remove rows with zero or missing values in important numeric/time columns,
		unless user specifically requests to see zeros.
		"""
		qlc = query.lower()
		if any(keyword in qlc for keyword in ["0 entries", "zero counts", "show zero", "count zero", "how many zero"]):
			return df  # user explicitly wants zeros

		df_filtered = df.copy()

		if "Postcode" in df_filtered.columns:
			df_filtered = df_filtered[
				~df_filtered["Postcode"].astype(str).str.strip().isin(["0", "", "nan", "None"])
			]

		numeric_cols = ["Visit Count", "Total Value", "Total Cost Inc Travel"]
		time_cols = ["Activate", "Deactivate", "Total Time", "Travel Time", "Total Time (Inc Travel)"]

		for col in numeric_cols:
			if col in df_filtered.columns:
				df_filtered = df_filtered[~((df_filtered[col] == 0) | (df_filtered[col].isna()))]

		for col in time_cols:
			if col in df_filtered.columns:
				df_filtered = df_filtered[
					~((df_filtered[col] == pd.Timedelta(0)) | df_filtered[col].astype(str).str.startswith("00:00"))
				]

		return df_filtered


if st.session_state.get("screen") == "ai":
    # ──────────────────────────────────────────────────────────────────────
    # 7b️⃣  FORECAST RENDERER  (NOW TOP-LEVEL, NOT NESTED)
    # ──────────────────────────────────────────────────────────────────────
    from sklearn.linear_model import LinearRegression
    import numpy as np

    def render_forecast(query: str, df: pd.DataFrame):
        qlc = query.lower()
        if "forecast" not in qlc and "projection" not in qlc:
            return
        if "completed" not in qlc:
            st.info("Only 'completed visits' forecasting is supported currently.")
            return
        if "Date" not in df.columns or "Visit Type" not in df.columns:
            st.warning("Data is missing Date or Visit Type columns.")
            return

        df = df.copy()
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        df.dropna(subset=["Date"], inplace=True)

        # ✅ UPDATED FILTERING HERE
        df = df[df["Activity Status"].str.lower() == "completed"]
        if df.empty:
            st.warning("No completed visits found in the dataset.")
            return

        df["Month"] = df["Date"].dt.to_period("M").dt.to_timestamp()
        monthly = df.groupby("Month").size().reset_index(name="Completed Visits")
        if len(monthly) < 3:
            st.warning("Not enough data to forecast.")
            return

        monthly["Month_Num"] = np.arange(len(monthly))
        X = monthly["Month_Num"].values.reshape(-1, 1)
        y = monthly["Completed Visits"].values
        model = LinearRegression().fit(X, y)

        fut = 6
        future_X = np.arange(len(monthly), len(monthly) + fut).reshape(-1, 1)
        future_dates = pd.date_range(
            start=monthly["Month"].max() + pd.DateOffset(months=1),
            periods=fut, freq="MS"
        )
        forecast_vals = model.predict(future_X).round().astype(int)

        forecast_df = pd.DataFrame({
            "Month": future_dates.strftime("%B %Y"),
            "Forecasted Completed Visits": forecast_vals
        })
        st.subheader("📈 Forecasted Completed Visits (Next 6 Months)")
        st.dataframe(forecast_df)

        full = pd.concat([
            monthly[["Month", "Completed Visits"]],
            pd.DataFrame({"Month": future_dates, "Completed Visits": forecast_vals})
        ], ignore_index=True)
        full["Type"] = ["Historical"] * len(monthly) + ["Forecast"] * fut

        fig = px.line(full, x="Month", y="Completed Visits", color="Type",
                      title="Completed Visits Forecast (Historical + Next 6 Months)",
                      markers=True)
        st.plotly_chart(fig, use_container_width=True)


    # ──────────────────────────────────────────────────────────────────────
    # 7c️⃣  FORECAST TOTAL VALUE (£)  – helper runs for “forecast value …”
    # ──────────────────────────────────────────────────────────────────────
    import pandas as pd, plotly.express as px
    from statsmodels.tsa.arima.model import ARIMA

    def render_value_forecast(query: str, df: pd.DataFrame):
        """Forecast total £ value for each Visit Type (ex-Lunch(30)) next 6 months."""
        qlc = query.lower()
        if "forecast" not in qlc or "value" not in qlc or "visit type" not in qlc:
            return  # user didn’t ask for this

        needed = {"Date", "Visit Type", "Total Value"}
        if not needed.issubset(df.columns):
            st.warning(f"Missing columns: {needed}")
            return

        # ── clean & filter ────────────────────────────────────────────────
        df = df.copy()
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        df["Total Value"] = pd.to_numeric(df["Total Value"], errors="coerce")
        df.dropna(subset=["Date", "Total Value"], inplace=True)
        df = df[~df["Visit Type"].str.contains("Lunch(30)", case=False, na=False)]
        df = df[df["Visit Type"].str.strip().ne("")]

        if df.empty:
            st.warning("No rows after excluding Lunch(30) / blanks.")
            return

        # ── monthly totals ───────────────────────────────────────────────
        df["Month"] = df["Date"].dt.to_period("M").dt.to_timestamp()
        monthly = (df.groupby(["Month", "Visit Type"])["Total Value"]
                   .sum().reset_index())

        fut = 6
        rows = []
        for vtype, grp in monthly.groupby("Visit Type"):
            if len(grp) < 3:
                continue  # not enough history
            grp = grp.sort_values("Month")
            grp["n"] = np.arange(len(grp))
            # simple linear trend
            coeffs = np.polyfit(grp["n"], grp["Total Value"], 1)
            future_n = np.arange(len(grp), len(grp)+fut)
            preds = np.poly1d(coeffs)(future_n).clip(0).round(2)
            future_dates = pd.date_range(grp["Month"].max()+pd.DateOffset(months=1),
                                         periods=fut, freq="MS")
            rows.append(pd.DataFrame({"Month": future_dates,
                                      "Visit Type": vtype,
                                      "Forecasted Total Value": preds}))
        if not rows:
            st.warning("Nothing to forecast.")
            return

        fc = pd.concat(rows, ignore_index=True)

        # ── table (formatted £) ──────────────────────────────────────────
        table = fc.copy()
        table["Month"] = table["Month"].dt.strftime("%B %Y")
        table["Forecasted Total Value (£)"] = table["Forecasted Total Value"].map("£{:,.2f}".format)

        st.subheader("💰 Forecasted Total Value by Visit Type (Next 6 Months)")
        st.dataframe(table[["Month", "Visit Type", "Forecasted Total Value (£)"]])

        # ── chart ────────────────────────────────────────────────────────
        st.plotly_chart(
            px.line(fc, x="Month", y="Forecasted Total Value",
                    color="Visit Type", markers=True),
            use_container_width=True
        )

    # ──────────── Helpers (Used by Main Q&A Loop) ────────────

    def get_total_cost_completed_visits(df, team=None):
        """Calculate total cost for completed visits, optionally filtered by team."""
        if "Activity Status" not in df.columns or "Total Cost Inc Travel" not in df.columns:
            return None
        filtered = df[df["Activity Status"].str.lower() == "completed"]
        if team:
            filtered = filtered[filtered["Team"] == team]
        total_cost = filtered["Total Cost Inc Travel"].sum()
        return total_cost

    def calculate_percentage_diff_avg(df, group_col="Team", value_col="Total Value"):
        """Calculate percentage difference from average total value per group."""
        totals = df.groupby(group_col)[value_col].sum().reset_index()
        baseline = totals[value_col].mean()
        totals["Pct Difference"] = ((totals[value_col] - baseline) / baseline * 100).round(2)
        totals["Total Value"] = totals[value_col].map("£{:,.2f}".format)
        totals["Pct Difference"] = totals["Pct Difference"].map(lambda x: f"{x:.2f}%")

        output_lines = []
        for _, row in totals.iterrows():
            output_lines.append(f"{row[group_col]}: {row['Total Value']} ({row['Pct Difference']})")

        return output_lines

    def clean_postcodes(df):
        """Clean postcode column: uppercase, strip, and exclude invalid postcodes."""
        df = df.copy()
        df['Postcode'] = df['Postcode'].astype(str).str.strip().str.upper()
        invalids = {"", "0", "NAN", "NONE", "NULL"}
        df = df[~df['Postcode'].isin(invalids)]
        return df

    def filter_zero_values(df: pd.DataFrame, query: str) -> pd.DataFrame:
        """
        Remove rows with zero or missing values in important numeric/time columns,
        unless user specifically requests to see zeros.
        """
        qlc = query.lower()
        if any(keyword in qlc for keyword in ["0 entries", "zero counts", "show zero", "count zero", "how many zero"]):
            return df  # user explicitly wants zeros

        df_filtered = df.copy()

        if "Postcode" in df_filtered.columns:
            df_filtered = df_filtered[
                ~df_filtered["Postcode"].astype(str).str.strip().isin(["0", "", "nan", "None"])
            ]

        numeric_cols = ["Visit Count", "Total Value", "Total Cost Inc Travel"]
        time_cols = ["Activate", "Deactivate", "Total Time", "Travel Time", "Total Time (Inc Travel)"]

        for col in numeric_cols:
            if col in df_filtered.columns:
                df_filtered = df_filtered[~((df_filtered[col] == 0) | (df_filtered[col].isna()))]

        for col in time_cols:
            if col in df_filtered.columns:
                df_filtered = df_filtered[
                    ~((df_filtered[col] == pd.Timedelta(0)) | df_filtered[col].astype(str).str.startswith("00:00"))
                ]

        return df_filtered

    # ────────────── Main Q&A Loop ──────────────

    user_q = st.chat_input(
        "Ask me anything … e.g. 'Can you write me an email'",
        key="oracle_ai"   # always provide a unique key!
    )
    user_q = (user_q or "").strip()
    if not user_q:
        st.stop()

    chart_type = pick_chart_type(user_q)

    # Function to get top postcodes by Total Value
    def get_top_postcodes_by_value(df, top_n=5):
        df_clean = clean_postcodes(df)
        grouped = (
            df_clean.groupby("Postcode")["Total Value"]
            .sum()
            .reset_index()
            .sort_values("Total Value", ascending=False)
            .head(top_n)
        )
        return grouped

    # Use uploaded data if available, else df_oracle
    df_for_ai_raw = st.session_state.user_df if st.session_state.user_df is not None else df_oracle
    df_for_ai = filter_zero_values(df_for_ai_raw, user_q)
    default_df = df_for_ai

    import datetime

    # Log user question with timestamp
    st.session_state.ai_chat.append({
        "role": "user",
        "content": user_q,
        "timestamp": datetime.datetime.now().isoformat()
    })
    with st.chat_message("user"):
        st.markdown(user_q)

    handled_manually = False
    answer = ""

    # Handle specific queries manually for quick replies
    if "total cost" in user_q.lower() and "completed visits" in user_q.lower():
        team = None
        for t in ["VIP South", "VIP North", "Tier 2 South", "Tier 2 North"]:
            if t.lower() in user_q.lower():
                team = t
                break
        total_cost = get_total_cost_completed_visits(default_df, team)
        if total_cost is not None:
            answer = f"Total cost including travel for completed visits{f' in {team}' if team else ''} is £{total_cost:,.2f}."
        else:
            answer = "Sorry, unable to calculate total cost for that query."
        with st.chat_message("assistant"):
            st.markdown(answer)
        st.session_state.ai_chat.append({
            "role": "assistant",
            "content": answer,
            "timestamp": datetime.datetime.now().isoformat()
        })
        handled_manually = True


    elif "top postcodes by total value" in user_q.lower():
        filtered_df = filter_zero_values(default_df, user_q)
        filtered_df["Postcode"] = filtered_df["Postcode"].astype(str).str.strip()
        filtered_df = filtered_df[~filtered_df["Postcode"].isin(["", "0", "nan", "None"])]
        grouped = (
            filtered_df.groupby("Postcode")["Total Value"]
            .sum()
            .reset_index()
            .sort_values("Total Value", ascending=False)
            .head(5)
        )
        lines = []
        for i, row in enumerate(grouped.itertuples(index=False), 1):
            postcode = row.Postcode if row.Postcode and str(row.Postcode).strip() else "(blank)"
            formatted_val = f"£{row._2:,.2f}"
            lines.append(f"{i}. Postcode: {postcode} - Total Value: {formatted_val}")
        answer = "The top 5 postcodes by Total Value are as follows:\n\n" + "\n".join(lines)
        with st.chat_message("assistant"):
            st.markdown(answer)
        st.session_state.ai_chat.append({"role": "assistant", "content": answer})
        handled_manually = True

    elif "percentage difference" in user_q.lower():
        filtered_df = filter_zero_values(default_df, user_q)
        lines = calculate_percentage_diff_avg(filtered_df)
        answer = "The percentage difference in total value between the teams is as follows:\n\n" + "\n".join(lines)
        with st.chat_message("assistant"):
            st.markdown(answer)
        st.session_state.ai_chat.append({"role": "assistant", "content": answer})
        handled_manually = True

    # If none of the manual handlers match, fallback to AI agent
    if not handled_manually:
        with st.chat_message("assistant"):
            placeholder = st.empty()
            try:
                answer = slow_stream(
                    (chunk if isinstance(chunk, str) else chunk.get("output", "")
                        for chunk in df_agent.stream(alias(user_q))),
                    placeholder,
                )
            except Exception:
                stream = fallback_client.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": user_q}],
                    stream=True,
                )
                answer = slow_stream(
                    (p.choices[0].delta.content or "" for p in stream),
                    placeholder
                )
        st.session_state.ai_chat.append({"role": "assistant", "content": answer})

    # Save Q&A to chat logs CSV
    import csv
    import datetime
    with open("chat_logs.csv", "a", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        if f.tell() == 0:
            w.writerow(["timestamp", "question", "answer", "feedback"])
        last_feedback = st.session_state.ai_chat[-1].get("feedback", "")
        w.writerow([datetime.datetime.now().isoformat(timespec="seconds"),
                    user_q, answer, last_feedback])

    # Logs viewer toggle & password protection
    show_logs = st.checkbox("Show AI Chat Logs", key="show_logs_checkbox")
    if show_logs:
        password = st.text_input("Enter admin password to view logs:", type="password", key="logs_password")
        if password == "AI Chat":  # Change your password here
            st.success("Access granted to chat logs 📊")
            import pandas as pd
            def load_logs(file_path):
                df = pd.read_csv(file_path, parse_dates=["timestamp"])
                return df
            log_file = "chat_logs.csv"
            logs_df = load_logs(log_file)
            if logs_df.empty:
                st.warning("No chat logs found!")
            else:
                st.markdown(f"### Total Questions Asked: {len(logs_df)}")
                min_date = logs_df['timestamp'].min().date()
                max_date = logs_df['timestamp'].max().date()
                date_range = st.date_input("Filter by date range:", [min_date, max_date], key="log_date_filter")
                filtered_logs = logs_df[
                    (logs_df['timestamp'].dt.date >= date_range[0]) &
                    (logs_df['timestamp'].dt.date <= date_range[1])
                ]
                keyword = st.text_input("Search questions or answers:", key="log_keyword").strip().lower()
                if keyword:
                    filtered_logs = filtered_logs[
                        filtered_logs['question'].str.lower().str.contains(keyword) |
                        filtered_logs['answer'].str.lower().str.contains(keyword)
                    ]
                st.markdown(f"### Showing {len(filtered_logs)} chat logs")
                st.dataframe(filtered_logs)
                freq_data = filtered_logs.groupby(filtered_logs['timestamp'].dt.date).size().reset_index(name='Count')
                st.markdown("### Questions Over Time")
                st.line_chart(freq_data.rename(columns={'timestamp': 'Date'}).set_index('timestamp'))
                if "feedback" in logs_df.columns and logs_df["feedback"].dropna().any():
                    st.markdown("### Feedback Summary")
                    feedback_counts = filtered_logs["feedback"].fillna("No Feedback").value_counts()
                    st.bar_chart(feedback_counts)
                st.markdown("---")
                st.caption("Chat logs loaded from `chat_logs.csv`")
        elif password:
            st.error("Incorrect password. Access denied.")

    # Auto-generate chart if answer suggests it
    # After the assistant gives a reply...
    # After the assistant reply, show a chart if the query or answer asks for it
    chart_keywords = [
        "chart", "graph", "distribution", "plot", "visual", "pie", "bar", "line", "trend", "correlation", "sunburst", "parallel"
    ]
    if any(k in user_q.lower() for k in chart_keywords) or any(k in answer.lower() for k in chart_keywords):
        # Guess a sensible column to plot
        col_candidates = [
            "Visit Type", "Activity Status", "Team", "Name"
        ]
        chart_col = None
        for c in col_candidates:
            if c.lower() in user_q.lower() or c.lower() in answer.lower():
                if c in df_for_ai.columns:
                    chart_col = c
                    break
        # Fallback to Visit Type if not found
        if not chart_col and "Visit Type" in df_for_ai.columns:
            chart_col = "Visit Type"
        # Render a bar chart if column found
        if chart_col:
            st.subheader(f"Chart: {chart_col} distribution")
            vc = (
                df_for_ai[chart_col]
                .astype(str)
                .str.strip()
                .replace({"0": pd.NA, "nan": pd.NA})
                .dropna()
                .value_counts()
                .head(12)
            )
            st.bar_chart(vc)
        else:
            st.info("Sorry, couldn't find a suitable column for charting.")





    # Quick summary by two columns if requested
    import re
    by_two = re.search(r"\bby ([\w ]+?) and ([\w ]+?)\b", user_q.lower())
    if by_two:
        a, b = by_two.group(1).strip(), by_two.group(2).strip()
        a_col = KEYWORDS.get(a, a.title())
        b_col = KEYWORDS.get(b, b.title())
        if {a_col, b_col}.issubset(df_oracle.columns):
            summary = df_oracle.groupby([a_col, b_col]).size().reset_index(name="Visits")
            st.dataframe(summary.head(200))
            try:
                fig = px.sunburst(summary, path=[a_col, b_col],
                                values="Visits",
                                title=f"Visits by {a_col} → {b_col}")
                st.plotly_chart(fig, use_container_width=True)
            except Exception:
                st.bar_chart(summary.set_index(a_col)["Visits"])

    # Always run forecast helper
    render_value_forecast(user_q, df_for_ai)


elif st.session_state.screen == "budget":
    import pandas as pd  # move this to the top if you prefer

    # — Budget Header —
    st.markdown(
        """
        <div style="text-align:center; margin-top:2em;">
          <span style="
            font-size:3em; font-weight:bold; color:#62d2a2;
            background:#f2f2f2; border-radius:20px;
            padding:30px 60px; box-shadow:0 4px 24px #62d2a244;">
            💷 Budget Remaining: <span style="color:#1d3557">£275,000</span>
          </span>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("<br>", unsafe_allow_html=True)

    # — Section breakdown cards —
    sections = [
        ("Sky Retail",   67000),
        ("Sky Business", 51000),
        ("Sky VIP",      42000),
        ("Tier 2",       35000),
    ]
    pastel = ["#a0c4ff", "#b9fbc0", "#ffd6a5", "#ffd6e0"]
    cols = st.columns(len(sections))
    for i, (name, spent) in enumerate(sections):
        with cols[i]:
            st.markdown(f"""
                <div style="
                  background:{pastel[i]}; border-radius:15px;
                  padding:20px; text-align:center;
                  box-shadow:0 2px 12px {pastel[i]}77;">
                  <span style="font-size:1.5em; font-weight:bold;">
                    {name}
                  </span><br>
                  <span style="font-size:2em; color:#222;">
                    £{spent:,.0f}
                  </span>
                </div>
            """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # — Back button —
    if st.button("⬅️ Back to Dashboard"):
        st.session_state.screen = "dashboard"

        



























        





























    

      


























        






























